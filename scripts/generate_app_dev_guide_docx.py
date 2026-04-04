#!/usr/bin/env python3
"""
hospital_app_development_guide_v2.md → Word(.docx) 変換スクリプト

Markdownファイルを読み込み、構造化されたWordドキュメントを生成する。
出力先: docs/admin/hospital_app_development_guide_v2.docx
"""

import os
import re
from datetime import date

from docx import Document
from docx.shared import Pt, Cm, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

# === パス設定 ===
SCRIPT_DIR = os.path.dirname(os.path.abspath(__file__))
PROJECT_DIR = os.path.dirname(SCRIPT_DIR)
INPUT_PATH = os.path.join(PROJECT_DIR, "docs", "admin", "hospital_app_development_guide_v2.md")
OUTPUT_PATH = os.path.join(PROJECT_DIR, "docs", "admin", "hospital_app_development_guide_v2.docx")
FIGURES_DIR = os.path.join(PROJECT_DIR, "docs", "admin", "figures")

# === フォント設定 ===
JP_FONT = "游ゴシック"
EN_FONT = "Calibri"
CODE_FONT = "Consolas"
FONT_SIZE_BODY = Pt(10.5)
FONT_SIZE_SMALL = Pt(9)
FONT_SIZE_CODE = Pt(9)
FONT_SIZE_H1 = Pt(18)
FONT_SIZE_H2 = Pt(14)
FONT_SIZE_H3 = Pt(12)

# === 施設情報 ===
FACILITY_NAME = "おもろまちメディカルセンター"


# ============================================================
# ヘルパー関数
# ============================================================

def set_font(run, name=JP_FONT, size=FONT_SIZE_BODY, bold=False, color=None, italic=False):
    """フォントを設定するヘルパー関数"""
    run.font.name = name
    run.font.size = size
    run.font.bold = bold
    run.font.italic = italic
    # 日本語フォント設定（eastAsia）
    r = run._element
    rPr = r.find(qn("w:rPr"))
    if rPr is None:
        rPr = parse_xml(f"<w:rPr {nsdecls('w')}/>")
        r.insert(0, rPr)
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = parse_xml(
            f'<w:rFonts {nsdecls("w")} w:ascii="{EN_FONT}" w:hAnsi="{EN_FONT}" w:eastAsia="{JP_FONT}"/>'
        )
        rPr.insert(0, rFonts)
    else:
        rFonts.set(qn("w:eastAsia"), JP_FONT)
        rFonts.set(qn("w:ascii"), EN_FONT)
        rFonts.set(qn("w:hAnsi"), EN_FONT)
    if color:
        run.font.color.rgb = color


def set_code_font(run, size=FONT_SIZE_CODE):
    """コード用フォントを設定"""
    run.font.name = CODE_FONT
    run.font.size = size
    r = run._element
    rPr = r.find(qn("w:rPr"))
    if rPr is None:
        rPr = parse_xml(f"<w:rPr {nsdecls('w')}/>")
        r.insert(0, rPr)
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = parse_xml(
            f'<w:rFonts {nsdecls("w")} w:ascii="{CODE_FONT}" w:hAnsi="{CODE_FONT}" w:eastAsia="{CODE_FONT}"/>'
        )
        rPr.insert(0, rFonts)
    else:
        rFonts.set(qn("w:ascii"), CODE_FONT)
        rFonts.set(qn("w:hAnsi"), CODE_FONT)
        rFonts.set(qn("w:eastAsia"), CODE_FONT)


def add_paragraph(doc, text, font_name=JP_FONT, font_size=FONT_SIZE_BODY, bold=False,
                  alignment=WD_ALIGN_PARAGRAPH.LEFT, space_after=Pt(6), space_before=Pt(0),
                  color=None):
    """段落を追加するヘルパー関数"""
    p = doc.add_paragraph()
    p.alignment = alignment
    p.paragraph_format.space_after = space_after
    p.paragraph_format.space_before = space_before
    run = p.add_run(text)
    set_font(run, name=font_name, size=font_size, bold=bold, color=color)
    return p


def add_heading_styled(doc, text, level=1):
    """見出しを追加してフォント設定"""
    size_map = {1: FONT_SIZE_H1, 2: FONT_SIZE_H2, 3: FONT_SIZE_H3}
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        set_font(run, name=JP_FONT, size=size_map.get(level, FONT_SIZE_H3), bold=True)
    return h


def add_code_block(doc, text):
    """コードブロック（グレー背景・等幅フォント）を追加"""
    for line in text.split("\n"):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.left_indent = Cm(1)
        # グレー背景を設定
        pPr = p._element.get_or_add_pPr()
        shading = parse_xml(f'<w:shd {nsdecls("w")} w:val="clear" w:color="auto" w:fill="EDEDED"/>')
        pPr.append(shading)
        run = p.add_run(line if line else " ")
        set_code_font(run)
    # コードブロック後のスペース
    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_before = Pt(0)
    spacer.paragraph_format.space_after = Pt(4)


def add_mermaid_placeholder(doc):
    """Mermaid図のプレースホルダーを追加"""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(8)
    # 枠線付きの注記
    pPr = p._element.get_or_add_pPr()
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:val="clear" w:color="auto" w:fill="FFF8E1"/>')
    pPr.append(shading)
    run = p.add_run("※ 図はMarkdown版を参照してください")
    set_font(run, size=FONT_SIZE_SMALL, italic=True, color=RGBColor(0x99, 0x66, 0x00))


def add_page_numbers(doc):
    """フッターにページ番号を追加"""
    for section in doc.sections:
        footer = section.footer
        footer.is_linked_to_previous = False
        p = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        fldChar1 = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="begin"/>')
        run._element.append(fldChar1)
        run2 = p.add_run()
        instrText = parse_xml(
            f'<w:instrText {nsdecls("w")} xml:space="preserve"> PAGE </w:instrText>'
        )
        run2._element.append(instrText)
        run3 = p.add_run()
        fldChar2 = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="end"/>')
        run3._element.append(fldChar2)


def add_toc(doc):
    """目次ページを追加（Wordフィールドコードによる自動目次）"""
    add_heading_styled(doc, "目次", level=1)

    p = doc.add_paragraph()
    run = p.add_run()
    fldChar1 = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="begin"/>')
    run._element.append(fldChar1)

    run2 = p.add_run()
    instrText = parse_xml(
        f'<w:instrText {nsdecls("w")} xml:space="preserve">'
        f' TOC \\o "1-3" \\h \\z \\u </w:instrText>'
    )
    run2._element.append(instrText)

    run3 = p.add_run()
    fldChar2 = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="separate"/>')
    run3._element.append(fldChar2)

    run4 = p.add_run("（Wordで開いた後、目次を右クリック →「フィールドの更新」で目次が生成されます）")
    set_font(run4, size=FONT_SIZE_SMALL, color=RGBColor(0x88, 0x88, 0x88), italic=True)

    run5 = p.add_run()
    fldChar3 = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="end"/>')
    run5._element.append(fldChar3)

    doc.add_page_break()


# ============================================================
# Markdownパーサー
# ============================================================

def parse_markdown(filepath):
    """
    Markdownファイルを解析してブロック要素のリストを返す。

    返却するブロック要素の種類:
    - ("heading", level, text)
    - ("paragraph", text)
    - ("bullet", text)
    - ("checkbox", checked, text)
    - ("code", language, text)
    - ("mermaid", text)
    - ("image", alt_text, image_path)
    - ("table", headers, rows)
    - ("hr",)
    """
    with open(filepath, "r", encoding="utf-8") as f:
        lines = f.readlines()

    blocks = []
    i = 0
    n = len(lines)

    while i < n:
        line = lines[i]
        stripped = line.rstrip("\n")

        # 空行はスキップ
        if not stripped.strip():
            i += 1
            continue

        # 見出し
        heading_match = re.match(r'^(#{1,6})\s+(.+)$', stripped)
        if heading_match:
            level = len(heading_match.group(1))
            text = heading_match.group(2).strip()
            blocks.append(("heading", level, text))
            i += 1
            continue

        # 画像 ![alt](path)
        image_match = re.match(r'^!\[([^\]]*)\]\(([^)]+)\)\s*$', stripped.strip())
        if image_match:
            alt_text = image_match.group(1)
            image_path = image_match.group(2)
            blocks.append(("image", alt_text, image_path))
            i += 1
            continue

        # コードブロック（```で始まる）
        code_match = re.match(r'^```(\w*)', stripped)
        if code_match:
            lang = code_match.group(1).lower()
            i += 1
            code_lines = []
            while i < n and not lines[i].rstrip("\n").startswith("```"):
                code_lines.append(lines[i].rstrip("\n"))
                i += 1
            i += 1  # 閉じる```をスキップ
            code_text = "\n".join(code_lines)
            if lang == "mermaid":
                blocks.append(("mermaid", code_text))
            else:
                blocks.append(("code", lang, code_text))
            continue

        # 表（|で始まる行）
        if stripped.startswith("|"):
            table_lines = []
            while i < n and lines[i].rstrip("\n").strip().startswith("|"):
                table_lines.append(lines[i].rstrip("\n"))
                i += 1
            # 表をパース
            if len(table_lines) >= 2:
                headers = [c.strip() for c in table_lines[0].split("|")[1:-1]]
                # 2行目はセパレーターなのでスキップ
                rows = []
                for tl in table_lines[2:]:
                    cells = [c.strip() for c in tl.split("|")[1:-1]]
                    rows.append(cells)
                blocks.append(("table", headers, rows))
            continue

        # チェックリスト（- [ ] / - [x]）
        checkbox_match = re.match(r'^[-*]\s+\[([ xX])\]\s+(.+)$', stripped.strip())
        if checkbox_match:
            checked = checkbox_match.group(1).lower() == "x"
            text = checkbox_match.group(2)
            blocks.append(("checkbox", checked, text))
            i += 1
            continue

        # 箇条書き（- / * で始まる）
        bullet_match = re.match(r'^(\s*)[-*]\s+(.+)$', stripped)
        if bullet_match:
            text = bullet_match.group(2)
            blocks.append(("bullet", text))
            i += 1
            continue

        # 番号付きリスト
        numbered_match = re.match(r'^\s*\d+[.)]\s+(.+)$', stripped)
        if numbered_match:
            text = numbered_match.group(1)
            blocks.append(("bullet", text))
            i += 1
            continue

        # 水平線
        if re.match(r'^[-*_]{3,}\s*$', stripped):
            blocks.append(("hr",))
            i += 1
            continue

        # 通常の段落（連続する非空行をまとめる）
        para_lines = []
        while i < n:
            curr = lines[i].rstrip("\n")
            if not curr.strip():
                break
            if re.match(r'^#{1,6}\s', curr):
                break
            if curr.startswith("```"):
                break
            if curr.strip().startswith("|"):
                break
            if re.match(r'^[-*]\s', curr.strip()):
                break
            if re.match(r'^\s*\d+[.)]\s', curr):
                break
            if re.match(r'^[-*_]{3,}\s*$', curr):
                break
            if re.match(r'^!\[', curr.strip()):
                break
            para_lines.append(curr.strip())
            i += 1
        if para_lines:
            blocks.append(("paragraph", " ".join(para_lines)))
        continue

    return blocks


def extract_title_from_blocks(blocks):
    """ブロックリストから最初のH1見出しをタイトルとして抽出する"""
    for block in blocks:
        if block[0] == "heading" and block[1] == 1:
            return block[2]
    return "院内アプリ開発ガイド"


# ============================================================
# インラインMarkdown処理
# ============================================================

def add_rich_text(paragraph, text):
    """
    インラインMarkdown（**太字**、`コード`、[リンク](url)等）を処理して
    paragraphにrunを追加する。
    """
    # パターン: **bold**, `code`, *italic*
    pattern = re.compile(
        r'(\*\*(.+?)\*\*)'   # **bold**
        r'|(`(.+?)`)'         # `code`
        r'|(\*(.+?)\*)'       # *italic*
        r'|\[([^\]]+)\]\([^\)]+\)'  # [text](url) → textのみ表示
    )

    pos = 0
    for m in pattern.finditer(text):
        # マッチ前のテキスト
        if m.start() > pos:
            run = paragraph.add_run(text[pos:m.start()])
            set_font(run)

        if m.group(2):  # **bold**
            run = paragraph.add_run(m.group(2))
            set_font(run, bold=True)
        elif m.group(4):  # `code`
            run = paragraph.add_run(m.group(4))
            set_code_font(run, size=FONT_SIZE_BODY)
        elif m.group(6):  # *italic*
            run = paragraph.add_run(m.group(6))
            set_font(run, italic=True)
        elif m.group(7):  # [text](url)
            run = paragraph.add_run(m.group(7))
            set_font(run, color=RGBColor(0x05, 0x63, 0xC1))

        pos = m.end()

    # 残りのテキスト
    if pos < len(text):
        run = paragraph.add_run(text[pos:])
        set_font(run)


# ============================================================
# メイン変換処理
# ============================================================

def create_cover_page(doc, title):
    """表紙ページを生成"""
    for _ in range(6):
        add_paragraph(doc, "", space_after=Pt(20))

    add_paragraph(doc, title,
                  font_size=Pt(24), bold=True,
                  alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=Pt(40))

    add_paragraph(doc, f"作成日：{date.today().strftime('%Y年%m月%d日')}",
                  font_size=Pt(12),
                  alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=Pt(8))

    add_paragraph(doc, FACILITY_NAME,
                  font_size=Pt(14), bold=True,
                  alignment=WD_ALIGN_PARAGRAPH.CENTER)

    doc.add_page_break()


def render_blocks(doc, blocks):
    """パースしたブロック要素をWordドキュメントに描画する"""
    # 最初のH1は表紙で使用済みなのでスキップ
    skip_first_h1 = True

    for block in blocks:
        btype = block[0]

        if btype == "heading":
            level = block[1]
            text = block[2]
            if level == 1 and skip_first_h1:
                skip_first_h1 = False
                continue
            # H1-H3のみサポート、H4以降はH3扱い
            doc_level = min(level, 3)
            add_heading_styled(doc, text, level=doc_level)

        elif btype == "paragraph":
            text = block[1]
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(6)
            add_rich_text(p, text)

        elif btype == "bullet":
            text = block[1]
            p = doc.add_paragraph(style="List Bullet")
            p.text = ""
            add_rich_text(p, text)

        elif btype == "checkbox":
            checked = block[1]
            text = block[2]
            mark = "☑" if checked else "□"
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Cm(1)
            p.paragraph_format.space_after = Pt(2)
            run = p.add_run(f"{mark} {text}")
            set_font(run, size=FONT_SIZE_BODY)

        elif btype == "code":
            _lang = block[1]
            text = block[2]
            add_code_block(doc, text)

        elif btype == "image":
            alt_text = block[1]
            rel_path = block[2]
            # figuresディレクトリからの相対パスを解決
            abs_path = os.path.join(os.path.dirname(INPUT_PATH), rel_path)
            if os.path.exists(abs_path):
                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = p.add_run()
                run.add_picture(abs_path, width=Inches(5.5))
                # キャプション
                cap = doc.add_paragraph()
                cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
                cap.paragraph_format.space_after = Pt(12)
                cap_run = cap.add_run(alt_text)
                set_font(cap_run, size=FONT_SIZE_SMALL, italic=True,
                         color=RGBColor(0x66, 0x66, 0x66))
            else:
                p = doc.add_paragraph()
                run = p.add_run(f"[図: {alt_text}]（画像ファイルが見つかりません: {rel_path}）")
                set_font(run, size=FONT_SIZE_SMALL, italic=True,
                         color=RGBColor(0xFF, 0x00, 0x00))

        elif btype == "mermaid":
            add_mermaid_placeholder(doc)

        elif btype == "table":
            headers = block[1]
            rows = block[2]
            n_cols = len(headers)
            table = doc.add_table(rows=1 + len(rows), cols=n_cols)
            table.style = "Light List Accent 1"
            table.alignment = WD_TABLE_ALIGNMENT.CENTER

            # ヘッダー行
            for ci, h in enumerate(headers):
                cell = table.rows[0].cells[ci]
                cell.text = ""
                run = cell.paragraphs[0].add_run(h)
                set_font(run, bold=True, size=FONT_SIZE_BODY)

            # データ行
            for ri, row_data in enumerate(rows, start=1):
                for ci in range(n_cols):
                    cell_text = row_data[ci] if ci < len(row_data) else ""
                    cell = table.rows[ri].cells[ci]
                    cell.text = ""
                    run = cell.paragraphs[0].add_run(cell_text)
                    set_font(run, size=FONT_SIZE_BODY)

            # 表後のスペース
            add_paragraph(doc, "", space_after=Pt(4))

        elif btype == "hr":
            # 水平線 → 改ページとして扱う（または薄い線の段落）
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(6)
            p.paragraph_format.space_after = Pt(6)
            pPr = p._element.get_or_add_pPr()
            pBdr = parse_xml(
                f'<w:pBdr {nsdecls("w")}>'
                f'  <w:bottom w:val="single" w:sz="4" w:space="1" w:color="CCCCCC"/>'
                f'</w:pBdr>'
            )
            pPr.append(pBdr)


def generate_docx():
    """メイン処理：Markdownを読み込み、Wordドキュメントを生成する"""
    if not os.path.exists(INPUT_PATH):
        print(f"エラー: 入力ファイルが見つかりません: {INPUT_PATH}")
        return

    # Markdownをパース
    blocks = parse_markdown(INPUT_PATH)
    title = extract_title_from_blocks(blocks)

    # Wordドキュメント作成
    doc = Document()

    # ページ設定（A4）
    section = doc.sections[0]
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

    # デフォルトフォント設定
    style = doc.styles["Normal"]
    style.font.name = EN_FONT
    style.font.size = FONT_SIZE_BODY
    rPr = style.element.find(qn("w:rPr"))
    if rPr is None:
        rPr = parse_xml(f"<w:rPr {nsdecls('w')}/>")
        style.element.append(rPr)
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = parse_xml(
            f'<w:rFonts {nsdecls("w")} w:ascii="{EN_FONT}" w:hAnsi="{EN_FONT}" w:eastAsia="{JP_FONT}"/>'
        )
        rPr.insert(0, rFonts)
    else:
        rFonts.set(qn("w:eastAsia"), JP_FONT)
        rFonts.set(qn("w:ascii"), EN_FONT)
        rFonts.set(qn("w:hAnsi"), EN_FONT)

    # 表紙
    create_cover_page(doc, title)

    # 目次
    add_toc(doc)

    # 本文
    render_blocks(doc, blocks)

    # フッターにページ番号
    add_page_numbers(doc)

    # 保存
    os.makedirs(os.path.dirname(OUTPUT_PATH), exist_ok=True)
    doc.save(OUTPUT_PATH)
    print(f"Wordドキュメントを生成しました: {OUTPUT_PATH}")
    print(f"  入力: {INPUT_PATH}")
    print(f"  出力: {OUTPUT_PATH}")


if __name__ == "__main__":
    generate_docx()
