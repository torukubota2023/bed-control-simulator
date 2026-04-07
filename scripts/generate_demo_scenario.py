#!/usr/bin/env python3
"""ベッドコントロールシミュレーター 幹部会議 実演シナリオ Word文書生成"""

from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
import os

OUTPUT = "/Users/torukubota/ai-management/docs/admin/bed_control_demo_scenario.docx"

doc = Document()

# --- Page setup ---
section = doc.sections[0]
section.page_width = Cm(21.0)
section.page_height = Cm(29.7)
section.top_margin = Cm(2.5)
section.bottom_margin = Cm(2.5)
section.left_margin = Cm(2.5)
section.right_margin = Cm(2.5)

# --- Font helpers ---
JP_FONT = "游ゴシック"
JP_FONT_FALLBACK = "Hiragino Sans"

def set_font(run, name=JP_FONT, size=Pt(11), color=None, bold=False):
    run.font.name = name
    run.font.size = size
    run.font.bold = bold
    r = run._element
    r.rPr.rFonts.set(qn('w:eastAsia'), name)
    if color:
        run.font.color.rgb = color

def add_paragraph_with_style(text, size=Pt(11), color=None, bold=False, align=None, space_after=Pt(6), space_before=Pt(0), font_name=JP_FONT):
    p = doc.add_paragraph()
    if align is not None:
        p.alignment = align
    pf = p.paragraph_format
    pf.space_after = space_after
    pf.space_before = space_before
    run = p.add_run(text)
    set_font(run, name=font_name, size=size, color=color, bold=bold)
    return p, run

def add_empty_line():
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.space_before = Pt(0)
    run = p.add_run("")
    run.font.size = Pt(6)

def add_separator():
    """Add a horizontal rule using bottom border on an empty paragraph."""
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(12)
    p.paragraph_format.space_before = Pt(12)
    pPr = p._element.get_or_add_pPr()
    pBdr = parse_xml(
        '<w:pBdr %s>'
        '  <w:bottom w:val="single" w:sz="6" w:space="1" w:color="AAAAAA"/>'
        '</w:pBdr>' % nsdecls('w')
    )
    pPr.append(pBdr)

def add_scene_heading(text):
    """Scene heading: 13pt, navy, bold, underline."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(8)
    run = p.add_run(text)
    set_font(run, size=Pt(13), color=RGBColor(0x1B, 0x4F, 0x72), bold=True)
    run.underline = True

def add_operation(text):
    """Operation line (▶): green, bold."""
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.space_before = Pt(4)
    run = p.add_run(text)
    set_font(run, size=Pt(11), color=RGBColor(0x1E, 0x84, 0x49), bold=True)

def add_serif(text, size=Pt(11), color=RGBColor(0x2C, 0x3E, 0x50), bold=False):
    """Serif line (speech): dark blue."""
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.left_indent = Cm(0.5)
    run = p.add_run(text)
    set_font(run, size=size, color=color, bold=bold)
    return p

def add_serif_with_highlight(parts):
    """Speech line with mixed formatting. parts = [(text, {kwargs}), ...]"""
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.left_indent = Cm(0.5)
    for text, kwargs in parts:
        run = p.add_run(text)
        size = kwargs.get('size', Pt(11))
        color = kwargs.get('color', RGBColor(0x2C, 0x3E, 0x50))
        bold = kwargs.get('bold', False)
        set_font(run, size=size, color=color, bold=bold)
    return p

# Color constants
NAVY = RGBColor(0x1B, 0x4F, 0x72)
GREY = RGBColor(0x66, 0x66, 0x66)
PURPLE = RGBColor(0x8E, 0x44, 0xAD)
GREEN = RGBColor(0x1E, 0x84, 0x49)
DARK_BLUE = RGBColor(0x2C, 0x3E, 0x50)
RED = RGBColor(0xE7, 0x4C, 0x3C)

# ===== TITLE =====
add_paragraph_with_style(
    "ベッドコントロールシミュレーター\n幹部会議 実演シナリオ（約18分）",
    size=Pt(16), color=NAVY, bold=True,
    align=WD_ALIGN_PARAGRAPH.CENTER, space_after=Pt(8), space_before=Pt(24)
)

add_paragraph_with_style(
    "2026年4月　副院長　久保田 徹",
    size=Pt(12), color=GREY,
    align=WD_ALIGN_PARAGRAPH.CENTER, space_after=Pt(8)
)

add_paragraph_with_style(
    "【改訂版 v2.0 — 平均在院日数クリア計画・稼働率ギャップ分析を追加】",
    size=Pt(11), color=PURPLE, bold=True,
    align=WD_ALIGN_PARAGRAPH.CENTER, space_after=Pt(16)
)

add_separator()

# ===== 事前準備 =====
add_paragraph_with_style("■ 事前準備", size=Pt(13), color=NAVY, bold=True, space_after=Pt(8))

# Table with light blue background
table = doc.add_table(rows=1, cols=1)
table.alignment = WD_TABLE_ALIGNMENT.CENTER
cell = table.cell(0, 0)
# Set light blue background
shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="DCE6F1" w:val="clear"/>')
cell._element.get_or_add_tcPr().append(shading)

# Set table width
tbl = table._tbl
tblPr = tbl.tblPr if tbl.tblPr is not None else parse_xml(f'<w:tblPr {nsdecls("w")}/>')
tblW = parse_xml(f'<w:tblW {nsdecls("w")} w:type="pct" w:w="5000"/>')
tblPr.append(tblW)

# Remove borders or set light ones
borders = parse_xml(
    '<w:tblBorders %s>'
    '  <w:top w:val="single" w:sz="4" w:color="8DB4E2"/>'
    '  <w:left w:val="single" w:sz="4" w:color="8DB4E2"/>'
    '  <w:bottom w:val="single" w:sz="4" w:color="8DB4E2"/>'
    '  <w:right w:val="single" w:sz="4" w:color="8DB4E2"/>'
    '</w:tblBorders>' % nsdecls('w')
)
tblPr.append(borders)

prep_items = [
    "会議室モニターにノートPCを接続し、ブラウザでアプリを開いておく",
    "サイドバーで「シミュレーション実行」を押し、データが表示された状態にしておく\n　→ 月の日数（経過日数）: 20、カレンダー月日数: 30 を確認",
    "病棟セレクターは「全体（94床）」にしておく",
    "診療報酬改定プリセット：2024年度（令和6年度）",
    "「85歳以上が20%以上（在院日数+1日緩和）」にチェック済み",
    "運営分析タブを開いた状態でスタート",
]

# Clear default paragraph
cell.paragraphs[0].clear()
for i, item in enumerate(prep_items):
    if i == 0:
        p = cell.paragraphs[0]
    else:
        p = cell.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.space_before = Pt(2)
    run = p.add_run(f"・{item}")
    set_font(run, size=Pt(11))

add_empty_line()

# Red bold point
p = doc.add_paragraph()
p.paragraph_format.space_after = Pt(12)
run = p.add_run("＊ポイント：「今日は4月20日、月の半ばです。残り10日で何ができるかをお見せします」から入る")
set_font(run, size=Pt(11), color=RED, bold=True)

add_separator()

# ===== Scene 1 =====
add_scene_heading("■ Scene 1　「今日は4月20日。あと10日で、いくら取り戻せますか？」　(2分)")

add_operation("▶ 運営分析タブの稼働率ギャップパネルを見せる")
add_serif("💬 画面をご覧ください。今日4月20日時点の成績表です。")
add_operation("▶ パネルを指しながら")

add_serif_with_highlight([
    ("💬 月平均稼働率 ", {}),
    ("89.5%", {'color': RED, 'bold': True}),
    ("。目標の", {}),
    ("90%", {'color': RED, 'bold': True}),
    ("に対して", {}),
    ("マイナス0.5ポイント", {'color': RED, 'bold': True}),
    ("。これが今の私たちの現在地です。", {}),
])

add_serif_with_highlight([
    ("💬 この下の表を見てください。現ペース月末予測の運営貢献額が", {}),
    ("7,158万円", {'color': RED, 'bold': True}),
    ("。目標値（稼働率90%）なら", {}),
    ("7,274万円", {'color': RED, 'bold': True}),
    ("。その差額が", {}),
    ("約116万円", {'color': RED, 'bold': True}),
    ("です。", {}),
])

add_serif_with_highlight([
    ("💬 つまり、このまま現ペースを続けると、目標との差", {}),
    ("116万円", {'color': RED, 'bold': True}),
    ("が取り戻せない。逆に言えば、残り10日の動き方次第で、", {}),
    ("116万円", {'color': RED, 'bold': True}),
    ("を上乗せできる可能性がある。", {}),
])

add_operation("▶ 左サイドバーの「稼働率1%（≈1名の入院）の価値：年間1,199万円」を指す")

add_serif_with_highlight([
    ("💬 稼働率1%、つまりあと1名の入院で", {}),
    ("年間約1,200万円", {'color': RED, 'bold': True}),
    ("。常勤医師1名分の手取り年収に相当します。この数字を常に画面上に表示しています。感覚ではなく、金額で判断する文化をつくるためです。", {}),
])

add_separator()

# ===== Scene 2 =====
add_scene_heading("■ Scene 2　稼働率の波を「見る」— 週末の谷に隠れたコスト　(2分)")

add_operation("▶ 「日次推移」タブをクリック")
add_serif("💬 過去20日間の稼働率の推移です。黄色い帯が目標レンジの90〜95%。グレーの背景が土日。")
add_operation("▶ グラフの週末の谷を指しながら")
add_serif("💬 注目してほしいのは、この繰り返しのパターンです。金曜日に退院が集中して、土日に稼働率が落ちる。そして月曜に入院が入って回復する。毎週、同じ波形を繰り返しています。")

add_serif_with_highlight([
    ("💬 この「谷」の1つ1つが、空床コストです。空床1床は1日約", {}),
    ("約2万9千円", {'color': RED, 'bold': True}),
    ("。週末に5床空けば、2日で", {}),
    ("約29万円", {'color': RED, 'bold': True}),
    ("のロス。月4回で", {}),
    ("約116万円", {'color': RED, 'bold': True}),
    ("。パターンが見えれば、対策が打てます。", {}),
])

add_serif("💬 退院の曜日を平準化する。木曜にも退院を入れる。それだけで谷が浅くなり、月間で数十万円の改善になります。")

add_separator()

# ===== Scene 3 =====
add_scene_heading("■ Scene 3　「平均在院日数 21.4日」— 算定基準を超えています　(3分)")

add_operation("▶ 「運営分析」タブに戻る")
add_operation("▶ 下段のメトリクスカード「平均在院日数 21.4日」の赤いdelta表示を見せる")

add_serif_with_highlight([
    ("💬 ここを見てください。平均在院日数 ", {}),
    ("21.4日", {'color': RED, 'bold': True}),
    ("。赤い矢印で「基準超過 ", {}),
    ("+0.4日", {'color': RED, 'bold': True}),
    ("」と出ています。", {}),
])

add_serif_with_highlight([
    ("💬 地域包括医療病棟の算定基準は、平均在院日数", {}),
    ("21日以内", {'color': RED, 'bold': True}),
    ("です。2026年改定では", {}),
    ("20日以内", {'color': RED, 'bold': True}),
    ("に短縮されますが、当院は85歳以上が20%を超えているため+1日の緩和措置で21日以内。いずれにしても、今の", {}),
    ("21.4日は超過", {'color': RED, 'bold': True}),
    ("しています。", {}),
])

add_operation("▶ その下の赤いアラートを見せる")

add_serif("💬 アプリが自動で検知して、赤いアラートを出しています。「施設基準を満たさなくなるリスク」と、具体的な4つの対策。C群——在院15日以上の患者さんから退院調整を進める必要があります。")

add_serif("💬 でも、「何人退院させればいいのか」が分からなければ動けません。その答えが、次のタブにあります。")

add_separator()

# ===== Scene 4 =====
add_scene_heading("■ Scene 4　クリア計画 —「C群からあと何名退院させれば基準クリアか」　(3分)")

# Highlight badge
p = doc.add_paragraph()
p.paragraph_format.space_after = Pt(8)
run = p.add_run("★ 本日のハイライト")
set_font(run, size=Pt(12), color=RED, bold=True)

add_operation("▶ 「運営改善アラート」タブをクリック")
add_operation("▶ 下にスクロールして「平均在院日数 21日以内クリア計画」を見せる")

add_serif("💬 これが今日のハイライトです。")

add_operation("▶ 赤いパネルを指しながら")

add_serif_with_highlight([
    ("💬 現状：平均在院日数 ", {}),
    ("21.4日", {'color': RED, 'bold': True}),
    ("、基準を", {}),
    ("+0.4日超過", {'color': RED, 'bold': True}),
    ("。このままのペースだと月末には約", {}),
    ("21.4日", {'color': RED, 'bold': True}),
    ("——基準超過が継続します。必要な対策：残り10日でC群から追加", {}),
    ("3名", {'color': RED, 'bold': True}),
    ("の退院で基準クリア見込み。", {}),
])

add_serif_with_highlight([
    ("💬 「", {}),
    ("3名", {'color': RED, 'bold': True}),
    ("」。この数字があるだけで、行動が変わります。漠然と「在院日数を短くしましょう」ではなく、「C群から3名、今週中に退院調整を進めましょう」と具体的に動ける。", {}),
])

add_operation("▶ 左側のシミュレーションテーブルを見せる")

add_serif_with_highlight([
    ("💬 ここに退院計画シミュレーションがあります。C群追加退院0名なら", {}),
    ("21.4日で超過", {'color': RED, 'bold': True}),
    ("、2名なら", {}),
    ("21.1日でまだ超過", {'color': RED, 'bold': True}),
    ("、3名退院させると", {}),
    ("21.0日でクリア", {'color': RGBColor(0x1E, 0x84, 0x49), 'bold': True}),
    ("。", {}),
])

add_serif("💬 数字の根拠があるから、連携室にも「あと3名の退院先確保をお願いします」と明確に依頼できます。看護師長にも「C群の中で退院可能な方を3名リストアップしてください」と具体的に相談できます。")

add_operation("▶ 右側の「退院候補の優先順位」を見せる")

add_serif("💬 在院日数の長い方から順に退院調整。同時に新規入院の受入れも有効です。分母が増えれば平均在院日数は下がります。入院を創出してくださる先生方の貢献が、ここでも活きてきます。")

add_separator()

# ===== Scene 5 =====
add_scene_heading("■ Scene 5　病棟別に見る — 5Fと6Fでは課題が違う　(2分)")

add_operation("▶ 病棟セレクターで「5F（47床）」に切り替える")

add_serif_with_highlight([
    ("💬 5Fに切り替えます。外科・整形系の病棟です。稼働率が", {}),
    ("80%台", {'color': RED, 'bold': True}),
    ("で推移していて、空床が常に4〜8床。", {}),
])

add_serif_with_highlight([
    ("💬 5Fの課題は明確です。「", {}),
    ("入院を増やすこと", {'color': RED, 'bold': True}),
    ("」。稼働率を上げないと、空床が毎日コストを生み続けます。入院を創出してくれる先生方との連携強化が鍵です。", {}),
])

add_operation("▶ 病棟セレクターで「6F（47床）」に切り替える")

add_serif_with_highlight([
    ("💬 一方、6Fは内科・ペイン系。グラフを見てください。稼働率が常に", {}),
    ("90%以上", {'color': RED, 'bold': True}),
    ("で推移し、満床や入院断りも発生しています。", {}),
])

add_serif_with_highlight([
    ("💬 6Fの課題は逆です。「", {}),
    ("長期入院の方の退院調整", {'color': RED, 'bold': True}),
    ("」。ベッドが空かないから新しい患者さんを受けられない。C群——在院15日以上の長期患者さんの退院を計画的に進めることで、回転を生む必要があります。", {}),
])

add_serif("💬 同じ病院でも、5Fは「入院を増やす」、6Fは「退院を進める」。打つべき手が正反対です。全体の平均だけ見ていては、この違いは見えません。病棟別に見るからこそ、的確な手が打てるのです。")

add_operation("▶ 「全体（94床）」に戻す")

add_separator()

# ===== Scene 6 =====
add_scene_heading("■ Scene 6　2026年度改定への備え　(1分)")

add_operation("▶ サイドバーの「診療報酬改定プリセット」を「2026年度（令和8年度）」に切り替える")

add_serif("💬 2026年度の診療報酬改定にも対応しています。プリセットを切り替えるだけで、入院料1のイ・ロ・ハ区分に基づく新しい報酬で自動再計算されます。")

add_operation("▶ サイドバーの「平均在院日数上限: 21日以内（通常20日+1日緩和）」を指す")

add_serif_with_highlight([
    ("💬 注目してほしいのはここです。2026年改定では基準が", {}),
    ("20日", {'color': RED, 'bold': True}),
    ("に短縮されます。ただし当院は85歳以上の割合が20%を超えているため、+1日の緩和で", {}),
    ("21日以内", {'color': RED, 'bold': True}),
    ("。このチェックボックス一つで切り替えられます。改定前後の影響を事前にシミュレーションして、今から準備できます。", {}),
])

add_separator()

# ===== Scene 7 =====
add_scene_heading("■ Scene 7　「稼働率が上がると、私たちの賞与はどうなるか」　(2分)")

# Highlight badge
p = doc.add_paragraph()
p.paragraph_format.space_after = Pt(8)
run = p.add_run("★ ここから全職員に関わる話です")
set_font(run, size=Pt(12), color=RED, bold=True)

add_serif("💬 少し話を変えます。ここからは皆さん一人ひとりに関わる話です。")

add_serif_with_highlight([
    ("💬 私は事務の専門家ではありませんので、AIと一緒に試算した概算です。正確な数字は事務部門に確認が必要ですが、大きな方向性としてお聞きください。", {}),
])

add_empty_line()

add_serif_with_highlight([
    ("💬 まず、2026年度の診療報酬改定。先ほどプリセットを切り替えてお見せしましたが、同じ稼働率でも入院料が上がるため、運営貢献額が月あたり約", {}),
    ("600万円", {'color': RED, 'bold': True}),
    ("増えます。年間にすると", {}),
    ("約7,200万円", {'color': RED, 'bold': True}),
    ("の増収です。", {}),
])

add_serif_with_highlight([
    ("💬 当院の人件費率は58%です。この増収分を人件費率で按分すると、職員290名に対して一人あたり年間約", {}),
    ("14万円", {'color': RED, 'bold': True}),
    ("。これは賞与に換算すると、", {}),
    ("約10%の増額", {'color': RED, 'bold': True}),
    ("に相当します。改定の恩恵だけで、です。", {}),
])

add_empty_line()

add_serif_with_highlight([
    ("💬 さらに。今の月平均稼働率89.5%を目標の", {}),
    ("94%", {'color': RED, 'bold': True}),
    ("まで引き上げたらどうなるか。稼働率1%の年間価値は約1,200万円ですから、4.5ポイント改善で年間約", {}),
    ("5,400万円", {'color': RED, 'bold': True}),
    ("の増収。人件費按分で一人あたり年間約", {}),
    ("11万円", {'color': RED, 'bold': True}),
    ("。賞与に換算して、さらに", {}),
    ("約8%の増額", {'color': RED, 'bold': True}),
    ("です。", {}),
])

add_empty_line()

# Table for summary
table = doc.add_table(rows=4, cols=4)
table.alignment = WD_TABLE_ALIGNMENT.CENTER
# Set table width
tbl2 = table._tbl
tblPr2 = tbl2.tblPr if tbl2.tblPr is not None else parse_xml(f'<w:tblPr {nsdecls("w")}/>')
tblW2 = parse_xml(f'<w:tblW {nsdecls("w")} w:type="pct" w:w="4500"/>')
tblPr2.append(tblW2)

_bonus_data = [
    ["", "年間増収", "一人あたり", "賞与影響"],
    ["改定効果のみ", "+7,200万円", "+14万円", "約+10%"],
    ["稼働率94%達成", "+5,400万円", "+11万円", "約+8%"],
    ["合計", "+1億2,600万円", "+25万円", "約+18%"],
]
for r_idx, row_data in enumerate(_bonus_data):
    for c_idx, cell_text in enumerate(row_data):
        cell = table.cell(r_idx, c_idx)
        cell.text = ""
        run = cell.paragraphs[0].add_run(cell_text)
        is_header = (r_idx == 0)
        is_total = (r_idx == 3)
        _r_color = NAVY if is_header else RED if is_total else DARK_BLUE
        _r_bold = is_header or is_total
        set_font(run, size=Pt(10), color=_r_color, bold=_r_bold)
        if is_header:
            shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="DCE6F1" w:val="clear"/>')
            cell._element.get_or_add_tcPr().append(shading)
        if c_idx >= 1:
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT

add_empty_line()

add_serif_with_highlight([
    ("💬 つまり、改定と稼働率改善を合わせれば、", {}),
    ("賞与が約2割増える可能性", {'color': RED, 'bold': True}),
    ("がある。繰り返しますが、これは事務専門でない副院長がAIと計算した概算です。ただ、方向性は間違っていないと思います。", {}),
])

add_serif_with_highlight([
    ("💬 「今より1日あたり4〜5名多く入院がいる状態を維持する」。それだけで、皆さん一人ひとりの年間賞与が", {}),
    ("約25万円増える", {'color': RED, 'bold': True}),
    ("計算になります。", {}),
])

add_separator()

# ===== Scene 8 =====
add_scene_heading("■ Scene 8　均等努力方式（Equal Effort） —「同じだけ頑張る」から始める　(3分)")

add_serif("💬 ここまで病棟別に課題を見てきました。5Fは稼働率が低め、6Fは満床に近い。普通ならこう思いますよね——「5Fはもっと頑張れ」「6Fは問題ない」と。")

add_empty_line()

add_serif("💬 でも、「5Fだけ頑張れ」と言われたら、5Fのスタッフはどう感じるでしょうか？　不公平ですよね。")

add_empty_line()

add_serif_with_highlight([
    ("💬 当院の全体目標は94床全体で", {}),
    ("稼働率90%", {'color': RED, 'bold': True}),
    ("。5Fだけ、6Fだけで達成する必要はありません。", {}),
])

add_serif("💬 このアプリには「均等努力方式（Equal Effort）」という考え方を組み込みました。キーワードは「平等感」です。")

add_empty_line()

# Concept explanation
add_serif_with_highlight([
    ("💬 考え方はシンプルです。全体目標を達成するために必要な上昇幅Δを計算し、", {}),
    ("両方の病棟が同じだけ稼働率を上げる", {'color': RED, 'bold': True}),
    ("。片方だけに負担を押し付けない。", {}),
])

add_empty_line()

add_serif_with_highlight([
    ("💬 例えば今月、5Fの平均が", {}),
    ("87%", {'color': RED, 'bold': True}),
    ("、6Fの平均が", {}),
    ("92%", {'color': NAVY, 'bold': True}),
    ("だとします。全体を90%にするには、両病棟がそれぞれ", {}),
    ("+1.6ポイント", {'color': RED, 'bold': True}),
    ("ずつ上げればいい。5Fは88.6%へ、6Fは93.5%へ。全体で約90%達成です。", {}),
])

add_empty_line()

add_serif("💬 「同じだけ頑張る」。これなら誰も文句は言えません。")

add_empty_line()

# Demo instruction
p = doc.add_paragraph()
p.paragraph_format.space_after = Pt(8)
run = p.add_run("【操作】病棟セレクターで5Fまたは6Fを選択 → 均等努力テーブルが表示される")
set_font(run, size=Pt(11), color=NAVY, bold=True)

add_empty_line()

add_serif("💬 アプリでは、この均等努力テーブルが自動計算されます。上昇幅Δを変えた複数のシナリオが一覧で見えます。")

add_empty_line()

# Equal effort table example
_holistic_table = doc.add_table(rows=6, cols=5)
_holistic_table.alignment = WD_TABLE_ALIGNMENT.CENTER
_ht_tbl = _holistic_table._tbl
_htPr = _ht_tbl.tblPr if _ht_tbl.tblPr is not None else parse_xml(f'<w:tblPr {nsdecls("w")}/>')
_htW = parse_xml(f'<w:tblW {nsdecls("w")} w:type="pct" w:w="4500"/>')
_htPr.append(_htW)

_holistic_data = [
    ["上昇幅Δ", "5F目標", "6F目標", "全体", "達成"],
    ["+0.5pt", "87.5%", "92.4%", "89.6%", "❌"],
    ["+1.0pt", "88.0%", "92.9%", "89.8%", "❌"],
    ["+1.6pt", "88.6%", "93.5%", "90.0%", "✅"],
    ["+2.0pt", "89.0%", "93.9%", "90.1%", "✅"],
    ["+4.0pt", "91.0%", "95.9%", "90.8%", "✅"],
]
for r_idx, row_data in enumerate(_holistic_data):
    for c_idx, cell_text in enumerate(row_data):
        cell = _holistic_table.cell(r_idx, c_idx)
        cell.text = ""
        run = cell.paragraphs[0].add_run(cell_text)
        is_header = (r_idx == 0)
        is_target = (r_idx == 3)
        _ht_color = NAVY if is_header else GREEN if is_target else DARK_BLUE
        _ht_bold = is_header or is_target
        set_font(run, size=Pt(10), color=_ht_color, bold=_ht_bold)
        if is_header:
            shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="DCE6F1" w:val="clear"/>')
            cell._element.get_or_add_tcPr().append(shading)
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

add_empty_line()

add_serif_with_highlight([
    ("💬 +1.6ポイントずつ——これが今月の「均等努力ライン」です。テーブルの", {}),
    ("緑の行", {'color': GREEN, 'bold': True}),
    ("がちょうど全体90%を達成するポイント。両病棟が同じ幅だけ改善すればいい。", {}),
])

add_empty_line()

add_serif("💬 下の行を見てください。Δが大きすぎると、片方が96%に近づきます。+4.1ポイントで96%上限に達してしまう。アプリでは96%上限を超えるケースに警告マークが付きます。「そこまでは求めない」というラインが見えるわけです。")

add_empty_line()

add_serif_with_highlight([
    ("💬 これが「均等努力方式」の考え方です。「うちの病棟さえ良ければいい」ではなく、", {}),
    ("「同じだけ努力して、全体で目標を達成する」", {'color': RED, 'bold': True}),
    ("。5Fも6Fも同じ上昇幅を目指す。平等だからこそ、チームとして動ける。", {}),
])

add_empty_line()

add_serif("💬 月によって各病棟の実績は変わりますから、必要なΔも毎月自動で再計算されます。先月は5Fが良くて6Fが苦しかったかもしれない。今月は逆かもしれない。どちらの場合でも、「同じ幅だけ頑張ろう」——この公平さが、病棟間の協力を自然に生みます。")

add_separator()

# ===== Scene 9 =====
add_scene_heading("■ Scene 9　まとめ —「見える」が「動ける」に変わる　(1分)")

add_serif("💬 今日お見せしたことをまとめます。")
add_empty_line()
add_serif("💬 1つ目。月平均稼働率の現在地と、残り日数でいくら取り戻せるかが分かる。")
add_serif("💬 2つ目。平均在院日数が基準を超えそうなとき、「C群からあと何名」と具体的に分かる。")
add_serif("💬 3つ目。病棟ごとに課題が違うことが一目で分かり、的確に手が打てる。")
add_serif("💬 4つ目。均等努力方式。「同じだけ頑張る」という公平なルールで、病棟間の協力体制をデータで支える。")
add_serif("💬 5つ目。稼働率の改善が、皆さんの賞与にどうつながるかが見える。")
add_empty_line()
add_serif("💬 このアプリに必要なのは、毎日5分のデータ入力だけです。電子カルテとの連携なし。追加コストゼロ。")
add_empty_line()

# Special emphasized paragraphs (12pt, navy)
add_serif(
    "💬 入院を創出してくださる先生方。外来で患者さんを診て、「この方は入院が必要だ」と判断してくださるその一つひとつの意思決定が、年間1,200万円の価値を生んでいます。このアプリは、その貢献を数字で可視化します。",
    size=Pt(12), color=NAVY, bold=False
)

add_serif(
    "💬 そして病棟で患者さんを受けてくださる看護スタッフの皆さん。「また入院が来た」「ベッドが足りない」——そう感じる日もあるでしょう。でもこのデータが示しているのは、皆さんが1名を受けるたびに、病院に年間1,200万円の価値が生まれ、それが皆さんの賞与に返ってくるということです。",
    size=Pt(12), color=NAVY, bold=False
)

add_empty_line()

add_serif("💬 目的は、誰かを責めることではありません。「見えなかったものを見えるようにする」こと。見えれば、自分で考えて、自分で動ける。そしてその頑張りは、数字を通じてちゃんと自分に返ってくる。そういう文化を、このアプリで一緒につくっていきたいと思います。")

add_serif("💬 まずは1ヶ月、試してみませんか。")

add_empty_line()

add_serif("💬 最後に、今日お見せしきれなかった機能を簡単にご紹介します。", size=Pt(12), color=NAVY, bold=True)

add_empty_line()

# Feature list table
feat_table = doc.add_table(rows=12, cols=2)
feat_table.alignment = WD_TABLE_ALIGNMENT.CENTER
# Set table width
_ft = feat_table._tbl
_ftPr = _ft.tblPr if _ft.tblPr is not None else parse_xml(f'<w:tblPr {nsdecls("w")}/>')
_ftW = parse_xml(f'<w:tblW {nsdecls("w")} w:type="pct" w:w="5000"/>')
_ftPr.append(_ftW)

_features = [
    ["機能", "できること"],
    ["📊 日次推移", "稼働率・入退院数の日別グラフ。曜日パターンや週末の谷が一目で分かる"],
    ["🔄 フェーズ構成", "A群/B群/C群の構成比を可視化。理想構成比との乖離を表示"],
    ["💰 運営分析", "診療報酬・コスト・運営貢献額を実績と月末予測で表示。改定前後の比較も可能"],
    ["🚨 運営改善アラート", "在院日数超過の自動検知とC群退院クリア計画。「あと何名」を算出"],
    ["🎯 意思決定ダッシュボード", "5日間の稼働率予測、最適在院日数レンジの算出、推奨アクション一覧"],
    ["🔮 What-if分析", "「もし金曜退院を平準化したら？」「入院が5名増えたら？」を即シミュレーション"],
    ["👨\u200d⚕️ 医師別分析", "入院創出医・担当医ごとの貢献度を可視化。フェーズ別の分析も可能"],
    ["💡 改善のヒント", "金曜退院の平準化、月曜入院の強化など、改善策を金額換算で提示"],
    ["🤝 均等努力方式", "病棟間の公平な改善目標を自動算出。同じΔだけ稼働率を上げるシナリオ表示（96%上限警告付き）"],
    ["📨 HOPE送信", "電子カルテのToDo一斉送信用メッセージを自動生成（400文字制限対応）"],
    ["📋 日次データ入力", "5F/6F病棟ごとに在院患者数・入院数・退院患者の在院日数をスライダーで入力。毎日5分で完了"],
]
for r_idx, row_data in enumerate(_features):
    for c_idx, cell_text in enumerate(row_data):
        cell = feat_table.cell(r_idx, c_idx)
        cell.text = ""
        run = cell.paragraphs[0].add_run(cell_text)
        is_header = (r_idx == 0)
        _f_color = NAVY if is_header else DARK_BLUE
        _f_bold = is_header or c_idx == 0
        _f_size = Pt(10) if c_idx == 1 and not is_header else Pt(11)
        set_font(run, size=_f_size, color=_f_color, bold=_f_bold)
        if is_header:
            shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="DCE6F1" w:val="clear"/>')
            cell._element.get_or_add_tcPr().append(shading)

add_empty_line()

add_serif("💬 すべてブラウザで動きます。インストール不要。追加コストゼロ。ご質問はありますか？")

add_separator()

# ===== Q&A =====
add_paragraph_with_style("■ 想定Q&A", size=Pt(13), color=NAVY, bold=True, space_after=Pt(12))

qa_pairs = [
    ("Q. データ入力の手間はどのくらいか？",
     "A. 1日5分程度です。在院患者数・入院数を入力し、退院患者はスライダーで在院日数を選ぶだけ。すべてマウスのみで操作可能です。"),
    ("Q. 平均在院日数の「C群から3名」はどうやって計算しているのか？",
     "A. 厚労省公式の平均在院日数（在院患者延日数÷半回転数）を使い、残り日数での入退院ペースを現ペースで推計した上で、C群退院の追加による在院患者延日数の減少と退院数の増加を反映しています。退院が早いほど効果が大きいため、退院1名あたり残り日数の6割分の患者日数を節約する前提です。"),
    ("Q. 数字の精度は信頼できるのか？",
     "A. 診療報酬のイ/ロ/ハ区分による誤差は±8%程度です。「稼働率を上げるべきか」「退院の偏りを減らすべきか」という方向性の判断には影響しません。カーナビの到着予想時刻のようなもので、30分か3時間かの判断には十分な精度です。"),
    ("Q. 効果が出なかったらどうなるか？",
     "A. 導入コストがゼロですので、金銭的な損失はありません。むしろ、PoCで得られた稼働データ自体が今後の病棟運営判断に活用できます。「やってみて損はない」のが最大の強みです。"),
    ("Q. 個人の評価に使われないか？",
     "A. 使いません。このツールは病棟全体の傾向を把握し、改善の方向性を示すものです。将来的に医師別分析機能もありますが、目的は個人攻撃ではなく、貢献の可視化とフィードバックです。"),
    ("Q. 2026年改定で平均在院日数が20日になったらどうするのか？",
     "A. アプリのプリセットを切り替えるだけで対応できます。当院は85歳以上が20%を超えているため+1日の緩和で21日以内ですが、もし患者構成が変わった場合はチェックボックスを外せば20日基準に切り替わります。今から改定後の影響をシミュレーションして備えることが重要です。"),
    ("Q. 賞与の試算は本当に正確なのか？",
     "A. あくまでAIと計算した概算です。人件費率58%で按分し、増収分が賞与に反映されると仮定した試算ですので、実際の配分は事務部門・経営判断によります。ただ「稼働率が上がれば病院の収入が増え、それが職員に還元される」という方向性は間違いありません。大事なのは正確な数字より、自分たちの仕事がどう病院経営につながっているかを理解することです。"),
]

for q, a in qa_pairs:
    # Question: bold, navy
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.space_before = Pt(8)
    run = p.add_run(q)
    set_font(run, size=Pt(11), color=NAVY, bold=True)

    # Answer: normal
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.left_indent = Cm(0.5)
    run = p.add_run(a)
    set_font(run, size=Pt(11))

# --- Set default font for document ---
style = doc.styles['Normal']
font = style.font
font.name = JP_FONT
font.size = Pt(11)
style.element.rPr.rFonts.set(qn('w:eastAsia'), JP_FONT)

doc.save(OUTPUT)
print(f"Saved: {OUTPUT}")
