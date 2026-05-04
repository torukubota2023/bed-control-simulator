"""医師別 病棟運営パターン分析 — 理事会・運営会議 説明資料 DOCX 生成（4 層分析 抜本改訂版）.

副院長指示 (2026-05-04 / Rev. 4):
  - Codex 批判的レビューを受け、ランキング型 → 4 層分析へ全面再構成
  - Layer 1: 病床利用貢献（既存）
  - Layer 2: 出来高上乗せ可能性（手術あり症例比率を新規追加）
  - Layer 3: 制度適合貢献（救急車入院 proxy・在宅復帰・予定外を新規追加）
  - Layer 4: 経営上の解釈（5+1 類型による医師運営パターン分類）
  - 「医師別利益ランキング」表現を排除、「病棟運営構造の理解資料」へ位置づけ直し
"""
from __future__ import annotations

import argparse
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

from generate_doctor_kpi_charts import (
    OUT_DIR as CHART_DIR,
    PROFIT_A, PROFIT_B, PROFIT_C,
    PROFIT_A_LOW, PROFIT_A_HIGH,
    PROFIT_B_LOW, PROFIT_B_HIGH,
    PROFIT_C_LOW, PROFIT_C_HIGH,
    REV_A, REV_B, REV_C,
    COST_A, COST_B, COST_C,
    COST_UNCERTAINTY,
    EMERGENCY_RATIO_THRESHOLD,
    HOME_DISCHARGE_THRESHOLD,
    TYPE_ORDER,
    TYPE_DESCRIPTIONS,
    load_and_compute,
    make_anonymous_map,
)

REPO_ROOT = Path(__file__).resolve().parent.parent
JP_FONT = "Hiragino Mincho ProN"

C_TITLE = RGBColor(0x1F, 0x29, 0x37)
C_ACCENT = RGBColor(0x25, 0x63, 0xEB)
C_DANGER = RGBColor(0xDC, 0x26, 0x26)
C_OK = RGBColor(0x10, 0xB9, 0x81)
C_GOLD = RGBColor(0xD9, 0x77, 0x06)
C_MUTED = RGBColor(0x6B, 0x72, 0x80)
C_FAINT = RGBColor(0x9C, 0xA3, 0xAF)


def _set_jp_font(run, size_pt=11, bold=False, color=None):
    run.font.name = JP_FONT
    run.font.size = Pt(size_pt)
    run.font.bold = bold
    if color is not None:
        run.font.color.rgb = color
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.insert(0, rFonts)
    rFonts.set(qn("w:eastAsia"), JP_FONT)
    rFonts.set(qn("w:ascii"), JP_FONT)
    rFonts.set(qn("w:hAnsi"), JP_FONT)


def add_para(doc, text, size=11, bold=False, align=None, color=None):
    p = doc.add_paragraph()
    if align is not None:
        p.alignment = align
    run = p.add_run(text)
    _set_jp_font(run, size_pt=size, bold=bold, color=color)
    return p


def add_heading(doc, text, level=1):
    sizes = {0: 24, 1: 18, 2: 15, 3: 13}
    h = doc.add_heading(level=level)
    run = h.add_run(text)
    _set_jp_font(run, size_pt=sizes.get(level, 12), bold=True, color=C_TITLE)


def add_callout(doc, label, body, color=None):
    if color is None:
        color = C_ACCENT
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.5)
    run_label = p.add_run(label + " ")
    _set_jp_font(run_label, size_pt=11, bold=True, color=color)
    run_body = p.add_run(body)
    _set_jp_font(run_body, size_pt=11)


def add_image(doc, path, width_cm=16):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(str(path), width=Cm(width_cm))


def add_table(doc, header, rows, col_widths_cm=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(header))
    table.style = "Light Grid Accent 1"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    if col_widths_cm:
        for col, w in zip(table.columns, col_widths_cm):
            for cell in col.cells:
                cell.width = Cm(w)
    hdr = table.rows[0].cells
    for i, label in enumerate(header):
        hdr[i].text = ""
        run = hdr[i].paragraphs[0].add_run(label)
        _set_jp_font(run, size_pt=10, bold=True, color=C_TITLE)
    for r_i, row_data in enumerate(rows, start=1):
        cells = table.rows[r_i].cells
        for c_i, val in enumerate(row_data):
            cells[c_i].text = ""
            run = cells[c_i].paragraphs[0].add_run(str(val))
            _set_jp_font(run, size_pt=10)


def build(mode: str):
    df, agg, ward_total_days, total_hospital_days = load_and_compute()
    anon_map = make_anonymous_map(agg)
    label_map = (lambda x: x) if mode == 'named' else (lambda x: anon_map.get(x, x))
    title_audience = "理事会" if mode == 'named' else "運営会議"

    # 集計値
    overall_surgery = agg['手術あり件数'].sum() / agg['入院件数'].sum() * 100
    overall_emerg = agg['救急搬送後件数'].sum() / agg['入院件数'].sum() * 100
    overall_home = agg['在宅復帰件数'].sum() / agg['入院件数'].sum() * 100
    overall_unp = agg['予定外件数'].sum() / agg['入院件数'].sum() * 100

    # 分母情報（CSV 全体 vs 分析対象）
    total_admissions = len(df)
    analysis_admissions = int(agg['入院件数'].sum())
    excluded_admissions = total_admissions - analysis_admissions

    type_counts = {t: int((agg['タイプ分類'] == t).sum()) for t in TYPE_ORDER}

    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = JP_FONT
    style.font.size = Pt(11)
    for section in doc.sections:
        section.top_margin = Cm(2.0)
        section.bottom_margin = Cm(2.0)
        section.left_margin = Cm(2.2)
        section.right_margin = Cm(2.2)

    # ====================================================================
    # 表紙
    # ====================================================================
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(70)
    run = p.add_run("医師別 病棟運営パターン分析")
    _set_jp_font(run, size_pt=26, bold=True, color=C_TITLE)

    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run2 = p2.add_run(f"{title_audience}説明資料")
    _set_jp_font(run2, size_pt=18, bold=True, color=C_ACCENT)

    p3 = doc.add_paragraph()
    p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p3.paragraph_format.space_before = Pt(30)
    run3 = p3.add_run("〜 病床利用 + 出来高上乗せ可能性 + 制度適合 + 経営上の解釈 の 4 層分析 〜")
    _set_jp_font(run3, size_pt=13, color=C_ACCENT, bold=True)

    p3b = doc.add_paragraph()
    p3b.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p3b.paragraph_format.space_before = Pt(8)
    run3b = p3b.add_run("（医師個人の評価ではなく、地域包括医療病棟の運営構造を理解するための資料）")
    _set_jp_font(run3b, size_pt=11, color=C_MUTED)

    p4 = doc.add_paragraph()
    p4.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p4.paragraph_format.space_before = Pt(50)
    run4 = p4.add_run("おもろまちメディカルセンター")
    _set_jp_font(run4, size_pt=14)

    p5 = doc.add_paragraph()
    p5.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run5 = p5.add_run("副院長 久保田 透")
    _set_jp_font(run5, size_pt=12)

    p6 = doc.add_paragraph()
    p6.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p6.paragraph_format.space_before = Pt(8)
    run6 = p6.add_run("作成日：2026 年 5 月 4 日 / Rev. 4（4 層分析 抜本改訂版）")
    _set_jp_font(run6, size_pt=11, color=C_FAINT)

    if mode == 'anonymous':
        p7 = doc.add_paragraph()
        p7.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p7.paragraph_format.space_before = Pt(20)
        run7 = p7.add_run("※ 個人名はすべて A 医師, B 医師 ... と匿名化しています")
        _set_jp_font(run7, size_pt=10, color=C_DANGER)

    doc.add_page_break()

    # ====================================================================
    # 1. 目的・本資料の位置づけ
    # ====================================================================
    add_heading(doc, "1. 目的・本資料の位置づけ", level=1)

    add_para(
        doc,
        "病棟稼働率は施設基準（地域包括医療病棟は 90% 目標）を満たすための"
        "経営的・制度的な最重要指標です。実態は、各主治医の入院・在院・退院"
        "判断の総和が病棟稼働率を作っています。",
    )
    add_para(doc, "")
    add_callout(
        doc,
        "📌 大前提",
        "本資料は「医師個人の評価」ではなく「病棟運営の構造把握」を目的とします。"
        "個別の優劣判定ではなく、診療科特性・患者背景・退院調整難度を踏まえた上で、"
        "病棟全体の運営最適化に向けた議論材料として提示します。",
        color=C_DANGER,
    )
    add_para(doc, "")

    add_heading(doc, "1.1 4 層分析の枠組み（本資料の見方）", level=2)
    add_para(
        doc,
        "従来「病床日数」と「ベッド粗利単価」だけで医師別の経営貢献を語ると、"
        "次のような歪みが起こります。",
    )
    add_para(doc, "  ・手術症例を多く担当する医師の収益寄与を過小評価する")
    add_para(doc, "  ・リハビリ提供量の多い症例の経営上の意味を過小評価する")
    add_para(doc, "  ・短期でも処置・検査・麻酔等を伴う症例を軽く見てしまう")
    add_para(doc, "  ・「病棟を埋めていること」と「医療資源投入を伴う収益貢献」を混同する")
    add_para(doc, "")
    add_para(
        doc,
        "そこで本資料は、医師別の貢献を以下の 4 層に分けて読み解きます。",
        bold=True,
    )
    add_para(doc, "")
    add_table(
        doc,
        ["層", "観点", "本資料での扱い"],
        [
            ["Layer 1", "病床利用貢献",
             "稼働率寄与率・病床日数シェア・ベッド粗利単価（仮定）"],
            ["Layer 2", "出来高上乗せ可能性",
             "手術あり症例比率（入口指標）／点数は未反映"],
            ["Layer 3", "制度適合貢献",
             "救急車入院比率（参考 proxy）・在宅復帰率・予定外比率（施設基準と対比）"],
            ["Layer 4", "経営上の解釈",
             "5+1 類型による医師運営パターン分類"],
        ],
        col_widths_cm=[2, 4, 10],
    )
    add_para(doc, "")
    add_callout(
        doc,
        "💡 読み方",
        "Layer 1 だけを見れば「ベッドを埋める量と単価」しか分かりません。Layer 2/3 を併せて見て、"
        "初めて「この医師は地域包括医療病棟をどのような形で支えているか」が立体的に見えます。",
        color=C_OK,
    )
    doc.add_page_break()

    # ====================================================================
    # 2. データ概要・限界
    # ====================================================================
    add_heading(doc, "2. データ概要・限界", level=1)

    add_heading(doc, "2.1 データ概要", level=2)
    add_table(
        doc,
        ["項目", "内容"],
        [
            ["対象期間", "2025-04-01 〜 2026-04-25（約 13 ヶ月）"],
            ["対象データ（CSV 全体）",
             f"全入院 {total_admissions:,} 件（事務提供 CSV、患者氏名なし、医師コード化）"],
            ["分析対象（医師別集計）",
             f"年間入院 10 件以上の {len(agg)} 医師 / 合計 {analysis_admissions:,} 件"
             f"（CSV 全体の {analysis_admissions/total_admissions*100:.1f}%）"],
            ["対象病棟", "5F 病棟（47 床）+ 6F 病棟（47 床）= 計 94 床"],
            ["除外",
             f"件数 10 未満の医師（4 名、{excluded_admissions} 件）は統計的に弱いため医師別集計から除外"],
            ["利用可能な変数",
             "患者番号 / 病棟 / 救急車 / 緊急 / 入経路 / 入院日 / 退院日 / 退経路 / 日数 / 診療科 / 医師 / 手術"],
        ],
        col_widths_cm=[4, 12],
    )
    add_para(doc, "")

    add_heading(doc, "2.2 本資料の限界（必読）", level=2)
    add_para(
        doc,
        "本資料は CSV から計算可能な指標で構成しており、以下は **未反映** です。"
        "経営判断する際は必ずこれらの限界を踏まえてください。",
        bold=True,
    )
    add_para(doc, "")
    add_para(doc, "【未反映 ① コスト精緻化】", bold=True)
    add_para(doc, "  ・ベッド粗利単価のコスト見積もり（A 12,000 / B 6,000 / C 4,500 円）は業界目安")
    add_para(doc, f"  ・当院 DPC データでの直接検証は未実施。±{COST_UNCERTAINTY*100:.0f}% の不確実性")
    add_para(doc, "  ・コスト ±20% の幅でも順位（B > C > A）は揺らがないが、絶対値は精度が限定的")
    add_para(doc, "")
    add_para(doc, "【未反映 ② 出来高上乗せの正確な点数】", bold=True)
    add_para(doc, "  ・手術料・麻酔料の点数（CSV には手術 ○/× フラグのみ）")
    add_para(doc, "  ・リハビリ単位数・リハビリ料（PT / OT / ST）")
    add_para(doc, "  ・処置加算（中心静脈、人工呼吸器、ドレーン等）")
    add_para(doc, "  ・検査加算（内視鏡、心カテ、画像診断の出来高分）")
    add_para(doc, "  ・1,000 点以上の処置、高額薬剤、特定保険医療材料")
    add_para(doc, "")
    add_para(doc, "【未反映 ③ 看護必要度（医師別集計）】", bold=True)
    add_para(doc, "  ・別集計（病棟単位）では実施しているが、医師別の重症度評価は本資料に含まず")
    add_para(doc, "")
    add_callout(
        doc,
        "⚠️ 結論への影響",
        "本資料は「総合収益分析」ではなく「病床利用構造の基礎分析」です。"
        "Layer 1（粗利貢献額・順位）は CSV から計算した管理用スコアであり、"
        "確定的な会計利益・人事評価指標ではありません。Layer 2/3 を併せて読んで初めて"
        "実態に近い貢献の姿が見えます。Stage 2 で事務（医事課・経理）から追加データを"
        "取得し、当院実測値での再計算を予定しています。",
        color=C_DANGER,
    )
    doc.add_page_break()

    # ====================================================================
    # 3. Layer 1：病床利用貢献
    # ====================================================================
    add_heading(doc, "3. Layer 1：病床利用貢献", level=1)
    add_para(
        doc,
        "各医師が病棟稼働率（量）と病床日数あたりの管理用スコア（質の仮定）に"
        "どう関わっているかを見ます。",
    )
    add_para(doc, "")

    add_heading(doc, "3.1 ベッド粗利単価の考え方（参考）", level=2)
    add_para(
        doc,
        "本資料では、地域包括医療病棟入院料の在院日数フェーズ別に以下の管理用スコアを使います"
        "（コストは業界目安 ±20%、§2.2 参照）。",
    )
    add_para(doc, "")
    add_table(
        doc,
        ["フェーズ", "在院期間", "① 報酬", "② 変動費", "③ 粗利仮定（① − ②）"],
        [
            ["A 群（急性期）", "1 〜 5 日",
             f"{REV_A:,} 円", f"{COST_A:,} 円", f"{PROFIT_A:,} 円"],
            ["B 群（回復期）", "6 〜 14 日",
             f"{REV_B:,} 円", f"{COST_B:,} 円", f"{PROFIT_B:,} 円"],
            ["C 群（退院準備）", "15 日以降",
             f"{REV_C:,} 円", f"{COST_C:,} 円", f"{PROFIT_C:,} 円"],
        ],
        col_widths_cm=[3, 2.5, 3, 3, 4],
    )
    add_para(doc, "")
    add_callout(
        doc,
        "💡 注意",
        "この粗利仮定は **管理用スコア** であり、会計上の確定利益ではありません。"
        "「短期回転 = 高単価」という思い込みを否定する基礎としてのみ使い、"
        "個人評価指標として使わないでください。",
        color=C_GOLD,
    )
    add_para(doc, "")
    add_image(doc, CHART_DIR / f"01_concept_{mode}.png", width_cm=15.5)
    doc.add_page_break()

    add_heading(doc, "3.2 医師別 ポジショニング（散布図）", level=2)
    add_image(doc, CHART_DIR / f"02_scatter_{mode}.png", width_cm=15.5)
    add_para(doc, "")
    add_callout(
        doc,
        "🔍 読み方",
        "右側に位置するほど稼働率寄与（量）が大きく、上側に位置するほど病床日数あたりの"
        "管理用スコア（仮定）が高い医師です。エラーバーはコスト ±20% に伴う不確実性を表します。",
        color=C_ACCENT,
    )
    doc.add_page_break()

    add_heading(doc, "3.3 病棟別 稼働率寄与率", level=2)
    add_image(doc, CHART_DIR / f"03_ward_ranking_{mode}.png", width_cm=16)
    add_para(doc, "")

    sub_5f = agg[agg['5F_寄与率'] > 0].sort_values('5F_寄与率', ascending=False)
    sub_6f = agg[agg['6F_寄与率'] > 0].sort_values('6F_寄与率', ascending=False)
    okuk_label = label_map('OKUK')

    add_para(doc, "5F 病棟 寄与率上位 5 名（タイプ分類併記）:", bold=True)
    rows_5f = [
        [str(i+1), label_map(r['医師']), f"{r['5F_寄与率']:.1f}%",
         f"{int(r['5F_延日数']):,} 床日", f"{r['平均在院日数']:.1f} 日",
         r['タイプ分類']]
        for i, (_, r) in enumerate(sub_5f.head(5).iterrows())
    ]
    add_table(doc, ["順位", "医師", "5F 寄与率", "延日数", "平均在院日数", "タイプ"],
              rows_5f, col_widths_cm=[1.3, 2.2, 2.5, 2.5, 2.7, 3.5])
    add_para(doc, "")
    add_para(doc, "6F 病棟 寄与率上位 5 名（タイプ分類併記）:", bold=True)
    rows_6f = [
        [str(i+1), label_map(r['医師']), f"{r['6F_寄与率']:.1f}%",
         f"{int(r['6F_延日数']):,} 床日", f"{r['平均在院日数']:.1f} 日",
         r['タイプ分類']]
        for i, (_, r) in enumerate(sub_6f.head(5).iterrows())
    ]
    add_table(doc, ["順位", "医師", "6F 寄与率", "延日数", "平均在院日数", "タイプ"],
              rows_6f, col_widths_cm=[1.3, 2.2, 2.5, 2.5, 2.7, 3.5])
    doc.add_page_break()

    add_heading(doc, "3.4 病床日数シェアと管理用スコアの関係", level=2)
    add_image(doc, CHART_DIR / f"05_correlation_{mode}.png", width_cm=15.5)
    add_para(doc, "")
    add_callout(
        doc,
        "⚠️ 解釈の注意",
        "稼働率寄与率と年間粗利スコアは r=0.96 と高い相関を示しますが、これは"
        "**寄与率を使って粗利スコアを計算しているため統計的に独立ではなく**、"
        "「単価差より病床日数シェアが管理用スコアの大半を規定する」という説明として読みます。"
        "医師別貢献が独立に統計的証明されたものではありません。",
        color=C_DANGER,
    )
    doc.add_page_break()

    # ====================================================================
    # 4. Layer 2：出来高上乗せ収益の可能性
    # ====================================================================
    add_heading(doc, "4. Layer 2：出来高上乗せ収益の可能性", level=1)
    add_para(
        doc,
        "地域包括医療病棟入院料の包括範囲外として、手術・麻酔・リハビリ・一部の処置・検査が"
        "収益に影響します（出典：厚生労働省「令和 6 年度診療報酬改定の概要 入院 I」）。"
        "Layer 1 の病床日数だけでは、これらの出来高貢献を捕捉できません。",
    )
    add_para(doc, "")
    add_callout(
        doc,
        "📌 本層の位置づけ",
        "本資料の Layer 2 は **入口指標のみ**（手術あり症例比率）です。"
        "手術料・麻酔料・リハ単位の正確な点数は CSV に未収録のため、絶対金額化は行いません。"
        "「単価が低い = 経営貢献が小さい」と解釈する前に、この層を必ず併読してください。",
        color=C_DANGER,
    )
    add_para(doc, "")

    add_heading(doc, "4.1 医師別 手術あり症例比率", level=2)
    add_image(doc, CHART_DIR / f"08_layer2_surgery_{mode}.png", width_cm=15.5)
    add_para(doc, "")

    surgery_top = agg.sort_values('手術あり比率', ascending=False).head(5)
    rows_surg = [
        [label_map(r['医師']),
         f"{r['手術あり比率']:.1f}%",
         f"{int(r['手術あり件数'])} / {int(r['入院件数'])} 件",
         f"{r['平均在院日数']:.1f} 日",
         r['主たる診療科'],
         r['タイプ分類']]
        for _, r in surgery_top.iterrows()
    ]
    add_para(doc, "手術あり症例比率 上位 5 名（出来高上乗せ可能性が大きいと推定される医師）:", bold=True)
    add_table(doc, ["医師", "手術あり比率", "件数", "平均在院日数", "主科", "タイプ"],
              rows_surg, col_widths_cm=[2.0, 2.5, 3.0, 2.5, 2.0, 3.5])
    add_para(doc, "")
    add_callout(
        doc,
        f"🔍 分析対象 {len(agg)} 医師 {analysis_admissions:,} 件中の平均: "
        f"{overall_surgery:.1f}%（{int(agg['手術あり件数'].sum())} 件）",
        "分析対象の約 1/4 が手術あり症例です。手術あり比率が 40% を超える医師は「手術・処置型」に分類され、"
        "Layer 1 の管理用スコアでは捕捉できない出来高上乗せ収益の可能性があります。"
        f"（CSV 全体 {total_admissions:,} 件のうち、件数 10 未満医師の {excluded_admissions} 件は分析対象外）",
        color=C_ACCENT,
    )
    doc.add_page_break()

    # ====================================================================
    # 5. Layer 3：制度適合貢献
    # ====================================================================
    add_heading(doc, "5. Layer 3：制度適合貢献", level=1)
    add_para(
        doc,
        "地域包括医療病棟は単純収益だけでなく、施設基準維持が運営継続に直結します。"
        "本層では救急車入院比率（参考 proxy）・在宅復帰率・予定外入院比率の 3 指標で、各医師が"
        "制度適合にどう寄与しているかを見ます。",
    )
    add_para(doc, "")

    add_heading(doc, "5.1 医師別 制度適合 3 指標", level=2)
    add_image(doc, CHART_DIR / f"09_layer3_compliance_{mode}.png", width_cm=16)
    add_para(doc, "")

    add_table(
        doc,
        ["指標", "全体平均", "制度・施設基準", "意味"],
        [
            ["救急車入院比率（参考 proxy）",
             f"{overall_emerg:.1f}%",
             f"{EMERGENCY_RATIO_THRESHOLD:.0f}%（病棟別 rolling 3 ヶ月）",
             "制度判定は病棟単位、本指標は医師別 proxy で別物"],
            ["在宅復帰率（自宅+居住系）",
             f"{overall_home:.1f}%",
             f"{HOME_DISCHARGE_THRESHOLD:.0f}% 以上",
             "地域包括医療病棟の施設基準"],
            ["予定外入院比率",
             f"{overall_unp:.1f}%",
             "（参考値）",
             "緊急入院全般。救急受入機能の厚みを示す"],
        ],
        col_widths_cm=[4.5, 2.5, 4, 5],
    )
    add_para(doc, "")
    add_callout(
        doc,
        "💡 読み方",
        "「Layer 1 でスコアが低い医師でも、Layer 3 で制度適合（救急受入・在宅復帰）に貢献している」"
        "ケースが少なくありません。逆に「Layer 1 でスコア高くても、Layer 3 で施設基準を悪化させる」"
        "症例も理屈上はあります。本層は両者を区別する材料です。",
        color=C_OK,
    )
    doc.add_page_break()

    # ====================================================================
    # 6. Layer 4：経営上の解釈（タイプ分類）
    # ====================================================================
    add_heading(doc, "6. Layer 4：経営上の解釈（5+1 類型）", level=1)
    add_para(
        doc,
        "Layer 1〜3 を総合し、各医師を「病棟をどのような形で支えているか」の 5+1 類型に分類します。"
        "個人の優劣判定ではなく、病棟構成のバランスを見るためのフレームです。",
    )
    add_para(doc, "")

    add_heading(doc, "6.1 5+1 類型の定義", level=2)
    add_table(
        doc,
        ["類型", "判定基準（優先順位順）", "経営上の意味"],
        [
            ["手術・処置型", "手術あり比率 ≥ 40%", TYPE_DESCRIPTIONS["手術・処置型"]],
            ["救急・予定外受入基盤型",
             "救急車入院比率（参考 proxy） ≥ 25% かつ 手術あり比率 < 30%",
             TYPE_DESCRIPTIONS["救急・予定外受入基盤型"]],
            ["長期安定型", "平均在院日数 ≥ 14 日", TYPE_DESCRIPTIONS["長期安定型"]],
            ["リハ・在宅復帰型",
             "在宅復帰率 ≥ 95% かつ 平均在院日数 8〜13 日",
             TYPE_DESCRIPTIONS["リハ・在宅復帰型"]],
            ["回転供給型", "平均在院日数 ≤ 8 日", TYPE_DESCRIPTIONS["回転供給型"]],
            ["混合型", "上記いずれにも該当せず", TYPE_DESCRIPTIONS["混合型"]],
        ],
        col_widths_cm=[3.0, 5.5, 7.5],
    )
    add_para(doc, "")

    add_heading(doc, "6.2 当院の医師別 タイプ分類（散布図）", level=2)
    add_image(doc, CHART_DIR / f"10_layer4_typing_{mode}.png", width_cm=16)
    add_para(doc, "")

    add_heading(doc, "6.3 当院の構成内訳", level=2)
    rows_type = [
        [t, f"{type_counts.get(t, 0)} 名", TYPE_DESCRIPTIONS[t]]
        for t in TYPE_ORDER if type_counts.get(t, 0) > 0
    ]
    add_table(doc, ["類型", "人数", "病棟運営上の意味"], rows_type,
              col_widths_cm=[3.5, 1.5, 11])
    add_para(doc, "")
    add_callout(
        doc,
        "💡 観察",
        "当院は手術・処置型と救急・予定外受入基盤型が中心となり、長期安定型がそれを支える構成になっています。"
        "「短期回転型」を称する医師は実は手術・処置型として出来高を支えており、"
        "Layer 1 の管理用スコアだけでは見えなかった貢献構造が浮かび上がります。",
        color=C_OK,
    )
    add_para(doc, "")
    add_callout(
        doc,
        "⚠️ 0 名類型の解釈",
        "「リハ・在宅復帰型」「回転供給型」は当院で 0 名ですが、これは以下の理由により"
        "**「専門化が進んでいない」と断定できません**：(1) この分類ルールが優先順位"
        "ベースの heuristic であり閾値依存、(2) リハ単位データが CSV に未収録のため"
        "「リハ・在宅復帰型」の精度が限定的、(3) 同じ医師が複数類型に当てはまる場合に"
        "優先順位の高い類型へ吸収される。**この分類ルールでは該当なし、追加データで再評価が必要**"
        "という解釈にとどめます。",
        color=C_DANGER,
    )
    doc.add_page_break()

    # ====================================================================
    # 7. 主な観察（順位ではなく問いかけ型）
    # ====================================================================
    add_heading(doc, "7. 主な観察 — 経営会議で問うべきこと", level=1)
    add_para(
        doc,
        "本資料は「順位」を見るためのものではありません。次の問いに使います。",
    )
    add_para(doc, "")

    questions = [
        ("当院の地域包括医療病棟は、どの類型の医師群で稼働を支えているか？",
         "→ Layer 4 の構成内訳が示す。当院は手術・処置型 + 救急・予定外受入基盤型 + 長期安定型のバランス"),
        ("「Layer 1 でスコア低い医師」は、本当に経営貢献が小さいのか？",
         "→ Layer 2/3 を併読すべき。手術あり比率や救急車入院比率（proxy）が高ければ、別軸の貢献が見える"),
        ("5F 病棟は単一主治医に病床日数の約 1/4 が集中する構造になっている。これは脆弱性か、強みか？",
         f"→ {okuk_label} 1 名で 23.3% を占有（5F のタイプ：救急・予定外受入基盤型）。"
         "不在時の代替体制をどう組むかが構造課題"),
        ("施設基準（救急 15% / 在宅復帰 70%）に対して、誰がどう貢献しているか？",
         "→ Layer 3 で症例構成の厚みを見る。救急車入院比率（proxy）25% 超の医師は病棟全体の救急受入を支える一群"),
        ("「手術・処置型」5 名が運営する出来高貢献を、現資料は反映できているか？",
         "→ できていない（Layer 2 は入口指標のみ）。Stage 2 で点数データを取得して精緻化"),
        ("回転供給型・リハ・在宅復帰型が当院で 0 名なのは、何を意味するか？",
         "→ 「専門化が進んでいない」と断定はできない。閾値依存の heuristic であること、"
         "リハ単位データが CSV 未収録であること、優先順位による吸収の可能性を踏まえ、"
         "「この分類ルールでは該当なし、追加データで再評価が必要」と解釈する"),
        ("Layer 1〜3 で見える指標に矛盾はないか？（高単価かつ施設基準悪化など）",
         "→ 当院では現状目立った矛盾なし。継続監視が必要"),
    ]
    for i, (q, hint) in enumerate(questions, 1):
        add_para(doc, f"問 {i}: {q}", bold=True)
        add_para(doc, f"   {hint}", color=C_MUTED, size=10)
        add_para(doc, "")
    doc.add_page_break()

    # ====================================================================
    # 8. 提言（4 軸）
    # ====================================================================
    add_heading(doc, "8. 提言", level=1)

    add_heading(doc, "8.1 提言 1：病棟運営構造の継続監視", level=2)
    add_para(
        doc,
        f"5F 病棟は寄与率トップの医師（{okuk_label}、救急・予定外受入基盤型）に病床日数の 23.3% が集中しており、"
        "個人攻撃ではなく構造リスクとして次の対応を提言します。",
    )
    add_para(doc, "  ・5F に主たる病棟を持つ医師を中期的に増員する選択肢を検討")
    add_para(doc, f"  ・{okuk_label} の患者バックアップ体制（指示・引継ぎ）の整備")
    add_para(doc, "  ・5F 主治医の月次寄与率を可視化し、依存度を継続監視")
    add_para(doc, "")
    add_para(
        doc,
        "なお 6F 病棟は上位 3 名で寄与率が分散しており、"
        "現時点で特定医師依存リスクは相対的に小さい構造です。",
        size=10, color=C_MUTED,
    )
    add_para(doc, "")

    add_heading(doc, "8.2 提言 2：制度適合の維持（Layer 3 の月次監視）", level=2)
    add_para(
        doc,
        "2026-06-01 から救急搬送後 15% は本則完全適用（病棟単位 rolling 3 ヶ月）。"
        "在宅復帰 70% は地域包括医療病棟の継続要件。次の対応を提言します。",
    )
    add_para(doc, "  ・病棟別 Layer 3 指標の月次ダッシュボード化（既存アプリへ組込み済）")
    add_para(doc, "  ・施設基準下振れ徴候があれば、救急受入比率と在宅復帰調整の介入を検討")
    add_para(doc, "  ・ALOS の階段関数（短手 3 Day 5/6 境界）の遵守を病床コントロールで継続")
    add_para(doc, "")

    add_heading(doc, "8.3 提言 3：必要データの整備（Stage 2）", level=2)
    add_para(
        doc,
        "本資料は CSV から計算可能な範囲に留まっています。出来高上乗せの正確な評価には"
        "追加データが必要ですが、当院事務部門は患者別コスト計算ができない前提のため、"
        "**事務に重い負担を強いる順序ではなく、請求収入側の集計可能性確認を先行**する"
        "提言とします。",
    )
    add_para(doc, "")
    add_para(doc, "【優先度 1】請求収入側の月次集計可能性の確認", bold=True)
    add_para(doc, "  ・手術料・麻酔料の患者別 月次集計が可能か確認")
    add_para(doc, "  ・リハビリ単位数（PT / OT / ST）と リハビリ料の月次集計が可能か確認")
    add_para(doc, "  ・処置加算（中心静脈、人工呼吸器、ドレーン管理）と検査加算の月次集計可能性")
    add_para(doc, "  ・1,000 点以上の処置・高額薬剤・特定保険医療材料の月次集計可能性")
    add_para(doc, "  → 既存の請求データ・レセプトの集計範囲で見える出来高情報を整理")
    add_para(doc, "")
    add_para(doc, "【優先度 2】コスト見積もりの精緻化（社内分析）", bold=True)
    add_para(doc, "  ・優先度 1 が見えてから、当院 DPC データを活用したコスト精緻化を検討")
    add_para(doc, "  ・事務部門への精密コスト計算依頼ではなく、社内分析として段階的に整備")
    add_para(doc, "")
    add_callout(
        doc,
        "📌 注意",
        "事務部門に患者別コスト計算の精密化を求めるものではありません。"
        "まずは既存の請求データ集計範囲内で、Layer 2 の点数化が可能かを確認する位置づけです。",
        color=C_GOLD,
    )
    add_para(doc, "")

    add_heading(doc, "8.4 提言 4：KPI 月次運用 — 4 層併読の文化", level=2)
    add_para(
        doc,
        "「稼働率寄与率（Layer 1）」だけを月次会議で見ると、Layer 2/3 で見える"
        "出来高貢献・制度適合貢献が抜け落ちます。次の運用を提言します。",
    )
    add_para(doc, "  ・月次会議では Layer 1〜3 の指標を併記して提示")
    add_para(doc, "  ・順位ではなく Layer 4 散布図で「タイプ構成」として議論")
    add_para(doc, "  ・既存ベッドコントロールシミュレーターに月次自動更新機能を組込み（次フェーズ）")
    add_para(doc, "  ・四半期に 1 回、Stage 2 データの取り込み状況をレビュー")
    add_para(doc, "")
    add_callout(
        doc,
        "📌 倫理的配慮",
        "本資料の指標は患者選別・適応外入院・在院延長の誘因にしてはいけません。"
        "「適応のある患者を、適切な期間で治療し、自宅復帰を支援する」という"
        "医療の質を犠牲にしないことを大前提とします。",
        color=C_DANGER,
    )
    doc.add_page_break()

    # ====================================================================
    # 9. 参考：全医師 集計表（Layer 1+2+3 拡張）
    # ====================================================================
    add_heading(doc, "9. 参考：全医師 集計表（Layer 1+2+3 拡張）", level=1)
    add_para(doc, "Layer 1〜3 の主要指標を全医師について併記します。", size=10, color=C_MUTED)
    add_para(doc, "")

    rows_all = []
    for _, r in agg.iterrows():
        rows_all.append([
            label_map(r['医師']),
            f"{int(r['入院件数'])}",
            f"{r['平均在院日数']:.1f}",
            f"{r['病院全体寄与率']:.2f}%",
            f"{r['手術あり比率']:.1f}%",
            f"{r['救急搬送後比率']:.1f}%",
            f"{r['在宅復帰率']:.1f}%",
            r['タイプ分類'],
            r['主たる診療科'],
        ])
    add_table(
        doc,
        ["医師", "件数", "平均在院日数(日)", "病院全体寄与率",
         "手術%", "救急車%", "在宅復帰%", "タイプ", "主科"],
        rows_all,
        col_widths_cm=[1.8, 1.2, 2.1, 2.4, 1.6, 1.8, 2.2, 2.5, 1.3],
    )
    add_para(doc, "")
    add_para(
        doc,
        f"参考：全体年間粗利仮定（管理用スコア合計）{agg['総粗利'].sum()/1e8:.2f} 億円　／　"
        f"全体平均ベッド粗利単価仮定 {agg['総粗利'].sum()/agg['総延日数'].sum():,.0f} 円/床日",
        size=10, color=C_MUTED,
    )
    add_para(
        doc,
        "※ 上記の絶対額は会計上の確定利益ではなく、コスト ±20% の業界目安に基づく管理用スコアの集計値。"
        "経営判断時は Layer 2/3 を併読し、Stage 2 の請求収入側集計と DPC コスト精緻化を経て、より妥当な推計として再評価する。",
        size=9, color=C_MUTED,
    )

    suffix = '理事会説明資料' if mode == 'named' else '運営会議説明資料'
    out_path = REPO_ROOT / 'docs' / 'admin' / f'医師別KPI解析_{suffix}_2026-05-04.docx'
    doc.save(out_path)
    print(f"✅ Generated: {out_path}")
    return out_path


if __name__ == '__main__':
    parser = argparse.ArgumentParser()
    parser.add_argument('--mode', choices=['named', 'anonymous'], default='named')
    args = parser.parse_args()
    build(args.mode)
