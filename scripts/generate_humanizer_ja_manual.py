"""日本語 Humanizer JA 導入マニュアル DOCX 生成スクリプト.

副院長指示 (2026-04-28): 初心者向けに手取り足取り、ヒラギノ明朝で。
Skills として配置するのが正解（公式の方法、~/.claude/skills/humanizer-ja/ に置くだけ）。

参考リポジトリ:
    https://github.com/gonta223/humanizer-ja

使い方:
    .venv/bin/python scripts/generate_humanizer_ja_manual.py
    → docs/admin/日本語Humanizer_JA_導入マニュアル_2026-04-28.docx
"""
from __future__ import annotations

from io import BytesIO
from pathlib import Path

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH


JP_FONT = "Hiragino Mincho ProN"
MONO_FONT = "Menlo"


def _set_jp_font(run, size_pt: int = 11, bold: bool = False, color=None):
    run.font.name = JP_FONT
    run.font.size = Pt(size_pt)
    run.font.bold = bold
    if color is not None:
        run.font.color.rgb = color
    # 日本語フォント明示（Word 互換）
    rPr = run._element.get_or_add_rPr()
    from docx.oxml.ns import qn
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        from docx.oxml import OxmlElement
        rFonts = OxmlElement("w:rFonts")
        rPr.insert(0, rFonts)
    rFonts.set(qn("w:eastAsia"), JP_FONT)
    rFonts.set(qn("w:ascii"), JP_FONT)
    rFonts.set(qn("w:hAnsi"), JP_FONT)


def _set_mono_font(run, size_pt: int = 10):
    run.font.name = MONO_FONT
    run.font.size = Pt(size_pt)
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.insert(0, rFonts)
    rFonts.set(qn("w:ascii"), MONO_FONT)
    rFonts.set(qn("w:hAnsi"), MONO_FONT)


def add_para(doc, text: str, size: int = 11, bold: bool = False,
             align: int | None = None, color=None):
    """通常段落（ヒラギノ明朝）."""
    p = doc.add_paragraph()
    if align is not None:
        p.alignment = align
    run = p.add_run(text)
    _set_jp_font(run, size_pt=size, bold=bold, color=color)
    return p


def add_heading(doc, text: str, level: int = 1):
    """見出し（ヒラギノ明朝・太字）."""
    sizes = {0: 26, 1: 20, 2: 16, 3: 14}
    h = doc.add_heading(level=level)
    run = h.add_run(text)
    _set_jp_font(run, size_pt=sizes.get(level, 12), bold=True,
                  color=RGBColor(0x1F, 0x29, 0x37))


def add_code(doc, code: str):
    """コードブロック（等幅フォント、薄い背景）."""
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(1.0)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(code)
    _set_mono_font(run, size_pt=10)


def add_callout(doc, label: str, body: str, color=None):
    """コールアウト（💡 ヒント / ⚠️ 注意 等）."""
    if color is None:
        color = RGBColor(0x25, 0x63, 0xEB)
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.5)
    run_label = p.add_run(label + " ")
    _set_jp_font(run_label, size_pt=11, bold=True, color=color)
    run_body = p.add_run(body)
    _set_jp_font(run_body, size_pt=11)


def add_step(doc, num: int, title: str, body_lines: list[str], code: str | None = None):
    """ステップ番号付きセクション."""
    p = doc.add_paragraph()
    run_num = p.add_run(f"ステップ {num}：")
    _set_jp_font(run_num, size_pt=14, bold=True, color=RGBColor(0xDC, 0x26, 0x26))
    run_title = p.add_run(title)
    _set_jp_font(run_title, size_pt=14, bold=True)
    for line in body_lines:
        add_para(doc, line, size=11)
    if code:
        add_code(doc, code)


def build():
    doc = Document()

    # 既定スタイル設定
    style = doc.styles["Normal"]
    style.font.name = JP_FONT
    style.font.size = Pt(11)

    # ページ余白 A4 縦
    for section in doc.sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    # =====================================================================
    # 表紙
    # =====================================================================
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(80)
    run = p.add_run("日本語 Humanizer JA")
    _set_jp_font(run, size_pt=36, bold=True, color=RGBColor(0x1F, 0x29, 0x37))

    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run2 = p2.add_run("導入マニュアル")
    _set_jp_font(run2, size_pt=28, bold=True, color=RGBColor(0x1F, 0x29, 0x37))

    p3 = doc.add_paragraph()
    p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p3.paragraph_format.space_before = Pt(40)
    run3 = p3.add_run("〜 Claude Code 初心者向け 手取り足取り版 〜")
    _set_jp_font(run3, size_pt=16, color=RGBColor(0x6B, 0x72, 0x80))

    p4 = doc.add_paragraph()
    p4.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p4.paragraph_format.space_before = Pt(120)
    run4 = p4.add_run("おもろまちメディカルセンター")
    _set_jp_font(run4, size_pt=14)

    p5 = doc.add_paragraph()
    p5.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run5 = p5.add_run("作成日：2026年4月28日")
    _set_jp_font(run5, size_pt=12, color=RGBColor(0x9C, 0xA3, 0xAF))

    doc.add_page_break()

    # =====================================================================
    # 目次
    # =====================================================================
    add_heading(doc, "目次", level=1)
    toc_items = [
        "1. このマニュアルについて",
        "2. 日本語 Humanizer JA とは（5 分で理解）",
        "3. 配置方法の選択：Skills が正解（新フォルダは不要）",
        "4. Claude Code 用 導入手順（5 ステップ）",
        "5. Codex デスクトップアプリ用 導入手順（4 ステップ）",
        "6. 両方使う場合の最適手順（1 つのコマンドで一括）",
        "7. 動作確認：最初の 1 回",
        "8. 使い方 3 パターン",
        "9. うまくいかないとき（トラブルシューティング）",
        "10. アップデート方法（GitHub に新版が出たとき）",
        "11. アンインストール方法",
        "12. 参考リンク",
    ]
    for item in toc_items:
        add_para(doc, item, size=12)

    doc.add_page_break()

    # =====================================================================
    # 1. このマニュアルについて
    # =====================================================================
    add_heading(doc, "1. このマニュアルについて", level=1)

    add_para(doc, "このマニュアルは、Claude Code を使い始めたばかりの方に向けて、"
                  "日本語 Humanizer JA というツール（スキル）を導入する手順を、"
                  "コマンド 1 行ずつ、つまずきポイントまで含めて解説するものです。")

    add_callout(doc, "💡 対象読者：",
                "Claude Code を 1 度でも起動したことがあれば OK。"
                "ターミナル（黒い画面）に少しでも触れたことがあれば理解できる構成です。")

    add_para(doc, "")
    add_para(doc, "■ このマニュアルで分かること", size=12, bold=True)
    add_para(doc, "・日本語 Humanizer JA が何をしてくれるツールか")
    add_para(doc, "・どこに置けばいいか（結論：~/.claude/skills/ の中）")
    add_para(doc, "・どうやってインストールするか（コマンド 3 行）")
    add_para(doc, "・どう使うか（3 つの使い方）")
    add_para(doc, "・うまくいかないときどうするか")

    doc.add_page_break()

    # =====================================================================
    # 2. 日本語 Humanizer JA とは
    # =====================================================================
    add_heading(doc, "2. 日本語 Humanizer JA とは（5 分で理解）", level=1)

    add_heading(doc, "ひと言で言うと", level=2)
    add_para(doc, "「AI が書いた感じの日本語」を、「人間が書いた自然な日本語」に直してくれる Claude Code 用のスキルです。",
             size=13)

    add_heading(doc, "どんな問題を解決してくれるか", level=2)
    add_para(doc, "AI（Claude / ChatGPT / Gemini など）に文章を作らせると、"
                  "便利な一方で「いかにも AI が書いた」感じが残ってしまうことがよくあります。")

    add_para(doc, "")
    add_para(doc, "■ AI が書いた感じが出る典型例（20 パターンのうち抜粋）", size=12, bold=True)

    table = doc.add_table(rows=6, cols=2)
    table.style = "Light Grid"
    headers = ["AI が書いた感じが残る表現", "なぜ気になるか"]
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ""
        run = cell.paragraphs[0].add_run(h)
        _set_jp_font(run, size_pt=11, bold=True, color=RGBColor(0xFF, 0xFF, 0xFF))
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement
        tcPr = cell._tc.get_or_add_tcPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:fill"), "374151")
        tcPr.append(shd)

    examples = [
        ("「〜することができます」", "回りくどい。普通は「〜できます」"),
        ("「今後の展開が注目されます」", "中身がない決まり文句"),
        ("「浮き彫りにしており」", "新聞記事っぽい大げさな表現"),
        ("太字 + コロン + 箇条書きの多用", "Markdown ぽさが残る"),
        ("「〜と言えるでしょう」を多用", "結論を曖昧にする逃げ表現"),
    ]
    for i, (ex, why) in enumerate(examples, start=1):
        for j, val in enumerate([ex, why]):
            cell = table.rows[i].cells[j]
            cell.text = ""
            run = cell.paragraphs[0].add_run(val)
            _set_jp_font(run, size_pt=11)

    add_para(doc, "")
    add_para(doc, "Humanizer JA は、こうした「20 パターン」のチェックリストを内蔵していて、"
                  "AI が生成した文章をこれらの観点で書き直してくれます。")

    add_heading(doc, "副院長の業務にどう役立つか", level=2)
    add_para(doc, "・ 病院運営文書（経営会議資料、提案書、ブリーフィング）の AI ぽさ除去")
    add_para(doc, "・ 教育資料（レジデント向け、看護師向け）を読みやすい自然な日本語に")
    add_para(doc, "・ 院内お知らせ・SNS 紹介文の自然な言い回し化")
    add_para(doc, "・ 医療倫理的な文書で「決まり文句」を避け、誠実な印象に")

    doc.add_page_break()

    # =====================================================================
    # 3. 配置方法の選択
    # =====================================================================
    add_heading(doc, "3. 配置方法の選択：Skills が正解（新フォルダは不要）", level=1)

    add_heading(doc, "結論", level=2)
    add_callout(doc, "✅ 結論：",
                "Claude Code の「Skills」として配置してください。"
                "新しいフォルダーを別途作る必要はありません。",
                color=RGBColor(0x10, 0xB9, 0x81))

    add_heading(doc, "なぜ Skills なのか", level=2)
    add_para(doc, "Humanizer JA は、配布元（GitHub の gonta223/humanizer-ja）が"
                  "「Claude Code Skill」として作成しているためです。"
                  "公式の置き場所に置くだけで、Claude Code が自動的に認識してくれます。")

    add_heading(doc, "Claude Code と Codex で配置先が違う", level=2)
    add_para(doc, "重要な前提：Claude Code（Anthropic）と Codex（OpenAI）は別の AI ツールで、"
                  "それぞれ別のフォルダから Skills を読み込みます。")

    table2 = doc.add_table(rows=4, cols=3)
    table2.style = "Light Grid"
    h2 = ["ツール / スコープ", "配置先フォルダ", "推奨度"]
    for i, h in enumerate(h2):
        cell = table2.rows[0].cells[i]
        cell.text = ""
        run = cell.paragraphs[0].add_run(h)
        _set_jp_font(run, size_pt=11, bold=True, color=RGBColor(0xFF, 0xFF, 0xFF))
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement
        tcPr = cell._tc.get_or_add_tcPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:fill"), "374151")
        tcPr.append(shd)

    rows = [
        ("Claude Code（ユーザー全体）", "~/.claude/skills/humanizer-ja/", "★ 推奨"),
        ("Codex（ユーザー全体）", "~/.agents/skills/humanizer-ja/", "★ 推奨（Codex 利用時）"),
        ("プロジェクト個別", ".claude/skills/ または .agents/skills/", "△ 限定的"),
    ]
    for i, (where, path, rec) in enumerate(rows, start=1):
        for j, val in enumerate([where, path, rec]):
            cell = table2.rows[i].cells[j]
            cell.text = ""
            run = cell.paragraphs[0].add_run(val)
            if j == 1:
                _set_mono_font(run, size_pt=10)
            else:
                _set_jp_font(run, size_pt=11)

    add_para(doc, "")
    add_callout(doc, "💡 副院長の用途では：",
                "病院運営文書も教育資料も、いろんなフォルダで作業します。"
                "「ユーザー全体（~/.claude/skills/ や ~/.agents/skills/）」に置けば、"
                "どのフォルダから起動しても呼び出せます。")

    add_callout(doc, "⚠️ Claude Code と Codex の両方を使う場合：",
                "両方のフォルダにインストールが必要です。後述の「6. 両方使う場合の最適手順」"
                "で 1 つのコマンドで一括導入する方法を案内します。",
                color=RGBColor(0xF5, 0x9E, 0x0B))

    add_heading(doc, "Codex 互換性についての注意", level=2)
    add_para(doc, "humanizer-ja は配布元（gonta223）が「Claude Code Skill」として作成したものです。"
                  "ただし、Codex も SKILL.md ベースの似たスキル形式を採用しているため、"
                  "実用上は Codex でも同様に動作することが多いです。")
    add_para(doc, "もし Codex で「スキルが認識されない」場合は、後述の「9. トラブルシューティング」"
                  "を参照してください。")

    doc.add_page_break()

    # =====================================================================
    # 4. Claude Code 用 導入手順
    # =====================================================================
    add_heading(doc, "4. Claude Code 用 導入手順（5 ステップ）", level=1)

    add_callout(doc, "📝 所要時間：",
                "5 分以内。コマンドは 3 行だけです。")

    # ステップ 1
    add_step(doc, 1, "ターミナルを開く", [
        "・「Launchpad」または「アプリケーション → ユーティリティ」から「ターミナル」を起動。",
        "・もしくは「Spotlight 検索（⌘+スペース）」で「ターミナル」と入力。",
        "・黒い画面（または白い画面）に文字が表示されればOK。",
    ])

    # ステップ 2
    add_step(doc, 2, "スキル置き場フォルダを作る", [
        "・以下のコマンドをコピーして、ターミナルに貼り付けて Enter キーを押してください。",
        "・このコマンドは「~/.claude/skills/ という場所がなければ作る」という意味です。",
        "・既に存在していてもエラーになりません（安全）。",
    ], code="mkdir -p ~/.claude/skills")

    # ステップ 3
    add_step(doc, 3, "GitHub からダウンロード", [
        "・以下のコマンドをコピーして、ターミナルに貼り付けて Enter。",
        "・git clone とは、「インターネット上のリポジトリ（保管庫）を、自分のパソコンに丸ごとコピーする」コマンドです。",
        "・ダウンロード完了後、ファイルは自動的に ~/.claude/skills/humanizer-ja/ に配置されます。",
    ], code="git clone https://github.com/gonta223/humanizer-ja.git ~/.claude/skills/humanizer-ja")

    # ステップ 4
    add_step(doc, 4, "ちゃんと配置されたか確認", [
        "・以下のコマンドで、フォルダの中身を確認できます。",
        "・SKILL.md / README.md / LICENSE など複数のファイルが表示されればOK。",
    ], code="ls ~/.claude/skills/humanizer-ja/")

    # ステップ 5
    add_step(doc, 5, "Claude Code を再起動", [
        "・既に Claude Code が開いていたら、いったん終了して、もう一度開き直してください。",
        "・新しいスキルは、Claude Code 起動時に読み込まれます。",
        "・再起動後に「日本語 Humanizer JA」が使えるようになります。",
    ])

    doc.add_page_break()

    # =====================================================================
    # 5. Codex デスクトップアプリ用 導入手順
    # =====================================================================
    add_heading(doc, "5. Codex デスクトップアプリ用 導入手順（4 ステップ）", level=1)

    add_callout(doc, "📝 所要時間：",
                "Claude Code 用と同じく 5 分以内。配置先フォルダが違うだけです。")

    add_callout(doc, "💡 ポイント：",
                "Codex の正式な Skills フォルダは `~/.agents/skills/` です（`~/.codex/skills/` ではありません）。"
                "ここがよく間違われる箇所なので注意。")

    # ステップ 1
    add_step(doc, 1, "ターミナルを開く", [
        "・Claude Code 用と同じ手順で、ターミナルを起動してください。",
        "・既にターミナルが開いていれば、それを使って OK。",
    ])

    # ステップ 2
    add_step(doc, 2, "Codex 用スキル置き場フォルダを作る", [
        "・以下のコマンドをコピーして、ターミナルに貼り付けて Enter。",
        "・~/.agents/skills/ が Codex の正式な Skills 配置先です。",
        "・既に存在していてもエラーになりません（安全）。",
    ], code="mkdir -p ~/.agents/skills")

    # ステップ 3
    add_step(doc, 3, "GitHub からダウンロード", [
        "・以下のコマンドをコピーして、ターミナルに貼り付けて Enter。",
        "・URL は Claude Code 用と同じ（同じ humanizer-ja を Codex 用フォルダにコピー）。",
    ], code="git clone https://github.com/gonta223/humanizer-ja.git ~/.agents/skills/humanizer-ja")

    # ステップ 4
    add_step(doc, 4, "Codex デスクトップアプリを再起動", [
        "・既に Codex が起動していたら、いったん終了して、もう一度開き直してください。",
        "・再起動後に humanizer-ja スキルが認識されます。",
    ])

    add_callout(doc, "🔍 Codex で動作確認の方法：",
                "Codex で適当なフォルダを開いた状態で「humanizer-jaスキルは使える？」"
                "と聞いてみてください。Codex が「はい、使えます」「○○について書き直します」"
                "といった反応をすれば成功。",
                color=RGBColor(0x10, 0xB9, 0x81))

    doc.add_page_break()

    # =====================================================================
    # 6. 両方使う場合の最適手順
    # =====================================================================
    add_heading(doc, "6. 両方使う場合の最適手順（1 つのコマンドで一括）", level=1)

    add_para(doc, "Claude Code と Codex の両方を使うなら、以下のように "
                  "1 行のコマンドで一括設定できます。")

    add_heading(doc, "一括導入コマンド（コピペ 3 行）", level=2)

    add_code(doc, "# 1) フォルダを両方作る\n"
                  "mkdir -p ~/.claude/skills ~/.agents/skills\n\n"
                  "# 2) Claude Code 用に clone\n"
                  "git clone https://github.com/gonta223/humanizer-ja.git ~/.claude/skills/humanizer-ja\n\n"
                  "# 3) Codex 用にも clone（同じ URL から別フォルダに）\n"
                  "git clone https://github.com/gonta223/humanizer-ja.git ~/.agents/skills/humanizer-ja")

    add_callout(doc, "💡 「同じものを 2 回ダウンロードするのは無駄では？」と思った方へ：",
                "それぞれのツールは独立して Skills を読むので、別々に置くのが最もシンプル。"
                "更新も `git pull` を 2 回するだけ（後述の「10. アップデート方法」参照）。")

    add_heading(doc, "上級者向け：シンボリックリンクで 1 か所に統一", level=2)
    add_para(doc, "「同じファイルを 2 か所に置くのは嫌、1 か所で管理したい」という方は、"
                  "以下のようにシンボリックリンクを使えます（任意・上級向け）。")

    add_code(doc, "# 1) Claude Code 用に clone\n"
                  "mkdir -p ~/.claude/skills ~/.agents/skills\n"
                  "git clone https://github.com/gonta223/humanizer-ja.git ~/.claude/skills/humanizer-ja\n\n"
                  "# 2) Codex 側はシンボリックリンクで参照\n"
                  "ln -s ~/.claude/skills/humanizer-ja ~/.agents/skills/humanizer-ja")

    add_para(doc, "この場合、`git pull` は 1 回（~/.claude/skills/humanizer-ja のみ）"
                  "で両方に反映されます。ただしリンクが壊れるリスクもあるため、"
                  "初心者は前述の「2 回 clone する方法」が無難です。")

    doc.add_page_break()

    # =====================================================================
    # 7. 動作確認
    # =====================================================================
    add_heading(doc, "7. 動作確認：最初の 1 回", level=1)

    add_para(doc, "うまくインストールできたかを確認するため、簡単な例で試してみましょう。")

    add_heading(doc, "試し方", level=2)
    add_para(doc, "Claude Code を開いた状態で、以下の文を入力してみてください。")

    add_code(doc,
             "以下の文章を人間っぽくして：\n\n"
             "本研究は、地域医療における看護必要度の重要性を浮き彫りにしており、\n"
             "今後の展開が大いに注目されると言えるでしょう。")

    add_heading(doc, "うまくいけば", level=2)
    add_para(doc, "Claude が humanizer-ja スキルを呼び出して、たとえば次のように書き直してくれます。")

    add_code(doc,
             "今回の研究で、地域医療における看護必要度の意味がはっきり見えてきました。\n"
             "今後の動きから目が離せません。")

    add_callout(doc, "💡 ポイント：",
                "「〜することができます」が「〜できます」に、「浮き彫りにしており」が普通の言い方に変わります。"
                "20 パターンを内蔵しているので、AI ぽい表現を網羅的に直してくれます。")

    doc.add_page_break()

    # =====================================================================
    # 8. 使い方 3 パターン
    # =====================================================================
    add_heading(doc, "8. 使い方 3 パターン", level=1)

    add_heading(doc, "パターン A：すでに書いた文章を直してもらう", level=2)
    add_para(doc, "AI が作った（または自分で書いた）文章を、自然な日本語に書き直してもらう使い方です。")
    add_code(doc, "この文章を人間っぽくして：\n\n[AI が作った文章を貼り付け]")

    add_heading(doc, "パターン B：最初から自然な文体で書かせる", level=2)
    add_para(doc, "新しい文章を Claude に作らせるとき、最初から humanizer-ja を意識させる使い方です。")
    add_code(doc, "humanizer-ja スキルを使って、地域包括医療病棟の経営会議向け提案書を作って。")

    add_heading(doc, "パターン C：チェックだけしてもらう", level=2)
    add_para(doc, "書いた文章のどこが「AI ぽい」かを指摘してもらう使い方です。")
    add_code(doc,
             "humanizer-ja の 20 パターンで、この文章をチェックして。\n"
             "どこが AI ぽいか指摘だけして、書き換えはしないで。")

    add_callout(doc, "💡 副院長の業務でのおすすめ：",
                "経営会議資料・教育資料の最終チェックに「パターン A」、"
                "新規ドラフト作成に「パターン B」、"
                "学会発表原稿の推敲に「パターン C」が効果的です。")

    doc.add_page_break()

    # =====================================================================
    # 9. トラブルシューティング
    # =====================================================================
    add_heading(doc, "9. うまくいかないとき（トラブルシューティング）", level=1)

    add_heading(doc, "症状 1：git clone でエラーが出る", level=2)
    add_para(doc, "■ 考えられる原因：インターネット接続、または git コマンドが入っていない")
    add_para(doc, "■ 対処：")
    add_para(doc, "1. インターネットに繋がっているか確認")
    add_para(doc, "2. git が入っているか確認するため、以下を実行")
    add_code(doc, "git --version")
    add_para(doc, "  → バージョン番号が出れば OK。出なければ Mac の場合は次のコマンドで案内に従う")
    add_code(doc, "xcode-select --install")

    add_heading(doc, "症状 2：Claude Code で再起動しても humanizer-ja が呼び出せない", level=2)
    add_para(doc, "■ 考えられる原因：配置フォルダが間違っている")
    add_para(doc, "■ 対処：")
    add_para(doc, "以下のコマンドで、SKILL.md がちゃんと存在するか確認")
    add_code(doc, "ls ~/.claude/skills/humanizer-ja/SKILL.md")
    add_para(doc, "  → 「No such file or directory」と出る場合は、Claude Code 用のステップ 3 のダウンロードからやり直し")

    add_heading(doc, "症状 3：Codex で再起動しても humanizer-ja が認識されない", level=2)
    add_para(doc, "■ 考えられる原因：")
    add_para(doc, "  ① 配置フォルダが ~/.codex/skills/ になっている（正しくは ~/.agents/skills/）")
    add_para(doc, "  ② Codex のバージョンが古く、Skills 機能未対応")
    add_para(doc, "  ③ humanizer-ja の SKILL.md 形式が Codex の期待と若干違う")
    add_para(doc, "■ 対処：")
    add_para(doc, "1. 以下のコマンドで Codex 用の正しいフォルダにあるか確認")
    add_code(doc, "ls ~/.agents/skills/humanizer-ja/SKILL.md")
    add_para(doc, "2. なければ、Codex 用のステップ 3 のダウンロードからやり直し")
    add_para(doc, "3. それでも動かない場合は、Codex を最新版にアップデート")
    add_para(doc, "4. 最悪、SKILL.md の冒頭に Codex 形式（name / description のメタデータ）が"
                  "あるか確認。なければ手動追加（上級向け）")

    add_heading(doc, "症状 4：書き換えが期待ほど自然でない", level=2)
    add_para(doc, "■ 考えられる原因：humanizer-ja は20パターンチェックリストベースなので、文脈次第で完璧ではない")
    add_para(doc, "■ 対処：")
    add_para(doc, "・ もう一度書き直してもらう（「もっと自然にして」と追加で指示）")
    add_para(doc, "・ 具体的に「医師宛のメール調で」「学会発表のスライド向けに」など文脈を追加")
    add_para(doc, "・ 元の文章が短すぎる場合は精度が落ちる傾向あり、ある程度のまとまり（数百字）で投入")

    doc.add_page_break()

    # =====================================================================
    # 10. アップデート方法
    # =====================================================================
    add_heading(doc, "10. アップデート方法（GitHub に新版が出たとき）", level=1)

    add_para(doc, "humanizer-ja は GitHub で更新されることがあります（新パターン追加、バグ修正など）。"
                  "最新版を取り込むには以下のコマンドを実行します。")

    add_heading(doc, "Claude Code 用フォルダの更新", level=2)
    add_code(doc, "cd ~/.claude/skills/humanizer-ja && git pull")

    add_heading(doc, "Codex 用フォルダの更新", level=2)
    add_code(doc, "cd ~/.agents/skills/humanizer-ja && git pull")

    add_heading(doc, "両方一気に更新（推奨）", level=2)
    add_code(doc, "cd ~/.claude/skills/humanizer-ja && git pull && \\\n"
                  "cd ~/.agents/skills/humanizer-ja && git pull")

    add_callout(doc, "💡 ヒント：",
                "更新後は Claude Code / Codex を再起動して、新版が読み込まれるようにしてください。")

    doc.add_page_break()

    # =====================================================================
    # 11. アンインストール
    # =====================================================================
    add_heading(doc, "11. アンインストール方法", level=1)

    add_para(doc, "もし不要になった場合、以下のコマンドで完全に削除できます。")

    add_heading(doc, "Claude Code 用のみアンインストール", level=2)
    add_code(doc, "rm -rf ~/.claude/skills/humanizer-ja")

    add_heading(doc, "Codex 用のみアンインストール", level=2)
    add_code(doc, "rm -rf ~/.agents/skills/humanizer-ja")

    add_heading(doc, "両方一気にアンインストール", level=2)
    add_code(doc, "rm -rf ~/.claude/skills/humanizer-ja ~/.agents/skills/humanizer-ja")

    add_callout(doc, "⚠️ 注意：",
                "rm -rf は「中身ごと完全に削除」するコマンドで、ゴミ箱には入りません。"
                "対象フォルダ名（~/.claude/skills/humanizer-ja および "
                "~/.agents/skills/humanizer-ja）を間違えないようにしてください。",
                color=RGBColor(0xF5, 0x9E, 0x0B))

    add_para(doc, "")
    add_para(doc, "削除後、Claude Code / Codex を再起動すれば、スキルが認識されなくなります。")

    doc.add_page_break()

    # =====================================================================
    # 12. 参考リンク
    # =====================================================================
    add_heading(doc, "12. 参考リンク", level=1)

    add_para(doc, "■ 配布元 GitHub リポジトリ", size=12, bold=True)
    p = doc.add_paragraph()
    run = p.add_run("https://github.com/gonta223/humanizer-ja")
    _set_mono_font(run, size_pt=11)

    add_para(doc, "")
    add_para(doc, "■ Claude Code 公式ドキュメント", size=12, bold=True)
    add_para(doc, "Skills の仕様について詳しく知りたい場合：")
    p = doc.add_paragraph()
    run = p.add_run("https://docs.claude.com/en/docs/claude-code")
    _set_mono_font(run, size_pt=11)

    add_para(doc, "")
    add_para(doc, "■ Codex 公式 Skills ドキュメント", size=12, bold=True)
    p = doc.add_paragraph()
    run = p.add_run("https://developers.openai.com/codex/skills")
    _set_mono_font(run, size_pt=11)
    add_para(doc, "")
    p = doc.add_paragraph()
    run = p.add_run("https://github.com/openai/skills")
    _set_mono_font(run, size_pt=11)

    add_para(doc, "")
    add_para(doc, "■ AGENTS.md（Codex のカスタム指示）", size=12, bold=True)
    p = doc.add_paragraph()
    run = p.add_run("https://agents.md/")
    _set_mono_font(run, size_pt=11)

    add_para(doc, "")
    add_para(doc, "■ ライセンス", size=12, bold=True)
    add_para(doc, "MIT ライセンス（自由に使用・修正・配布可能）")

    # 締め
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(40)
    run = p.add_run("— マニュアル ここまで —")
    _set_jp_font(run, size_pt=12, color=RGBColor(0x9C, 0xA3, 0xAF))

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("おもろまちメディカルセンター 副院長 久保田 徹  /  作成: 2026-04-28")
    _set_jp_font(run, size_pt=10, color=RGBColor(0x9C, 0xA3, 0xAF))

    return doc


def main():
    doc = build()
    out = Path("docs/admin/日本語Humanizer_JA_導入マニュアル_2026-04-28.docx")
    out.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(out))
    print(f"✅ 生成完了: {out}")
    print(f"   サイズ: {out.stat().st_size / 1024:.1f} KB")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
