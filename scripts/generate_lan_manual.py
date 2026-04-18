#!/usr/bin/env python3
"""
ベッドコントロールシミュレーター 社内LAN共有マニュアル Word文書生成スクリプト
出力先: docs/admin/streamlit_lan_sharing_manual.docx
"""

import os
from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

# === パス設定 ===
SCRIPT_DIR = os.path.dirname(os.path.abspath(__file__))
PROJECT_DIR = os.path.dirname(SCRIPT_DIR)
OUTPUT_PATH = os.path.join(PROJECT_DIR, "docs", "admin", "streamlit_lan_sharing_manual.docx")

# === フォント設定 ===
JP_FONT = "游ゴシック"
MONO_FONT = "Courier New"
FONT_SIZE_BODY = Pt(10.5)
FONT_SIZE_SMALL = Pt(9)
FONT_SIZE_CODE = Pt(9)


def set_font(run, name=JP_FONT, size=FONT_SIZE_BODY, bold=False, color=None, italic=False):
    """フォントを設定するヘルパー関数"""
    run.font.name = name
    run.font.size = size
    run.font.bold = bold
    run.font.italic = italic
    # 日本語フォント設定
    r = run._element
    rPr = r.find(qn("w:rPr"))
    if rPr is None:
        rPr = parse_xml(f"<w:rPr {nsdecls('w')}/>")
        r.insert(0, rPr)
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = parse_xml(f'<w:rFonts {nsdecls("w")} w:eastAsia="{name}"/>')
        rPr.insert(0, rFonts)
    else:
        rFonts.set(qn("w:eastAsia"), name)
    if color:
        run.font.color.rgb = color


def add_paragraph(doc, text, font_name=JP_FONT, font_size=FONT_SIZE_BODY, bold=False,
                  alignment=WD_ALIGN_PARAGRAPH.LEFT, space_after=Pt(6), space_before=Pt(0),
                  color=None):
    """段落を追加するヘルパー関数"""
    p = doc.add_paragraph()
    p.alignment = alignment
    p.paragraph_format.space_after = space_after
    p.paragraph_format.space_before = space_before
    run = p.add_run(text)
    set_font(run, name=font_name, size=font_size, bold=bold, color=color)
    return p


def add_code_block(doc, text):
    """コマンドブロック（グレー背景・等幅フォント）を追加"""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    # グレー背景を設定
    pPr = p._element.get_or_add_pPr()
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:val="clear" w:color="auto" w:fill="F0F0F0"/>')
    pPr.append(shading)
    # 左インデント
    p.paragraph_format.left_indent = Cm(1)
    run = p.add_run(text)
    set_font(run, name=MONO_FONT, size=FONT_SIZE_CODE, bold=False)
    return p


def add_code_with_description(doc, command, description=None):
    """コマンド+説明の組み合わせ"""
    add_code_block(doc, command)
    if description:
        add_paragraph(doc, description, font_size=FONT_SIZE_SMALL,
                      color=RGBColor(0x55, 0x55, 0x55))


def add_heading_styled(doc, text, level=1):
    """見出しを追加してフォント設定"""
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        set_font(run, name=JP_FONT, size=Pt(16 if level == 1 else 13 if level == 2 else 11),
                 bold=True)
    return h


def add_page_numbers(doc):
    """フッターにページ番号を追加"""
    for section in doc.sections:
        footer = section.footer
        footer.is_linked_to_previous = False
        p = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        # ページ番号フィールド
        run = p.add_run()
        fldChar1 = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="begin"/>')
        run._element.append(fldChar1)
        run2 = p.add_run()
        instrText = parse_xml(f'<w:instrText {nsdecls("w")} xml:space="preserve"> PAGE </w:instrText>')
        run2._element.append(instrText)
        run3 = p.add_run()
        fldChar2 = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="end"/>')
        run3._element.append(fldChar2)


def add_checklist_item(doc, text):
    """チェックリスト項目"""
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(1)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run("□ " + text)
    set_font(run, size=FONT_SIZE_BODY)
    return p


def create_manual():
    doc = Document()

    # ページ設定
    section = doc.sections[0]
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

    # ============================================================
    # 表紙ページ
    # ============================================================
    # 空行で位置調整
    for _ in range(6):
        add_paragraph(doc, "", space_after=Pt(20))

    add_paragraph(doc, "ベッドコントロールシミュレーター",
                  font_size=Pt(24), bold=True,
                  alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=Pt(8))
    add_paragraph(doc, "社内LAN共有マニュアル",
                  font_size=Pt(22), bold=True,
                  alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=Pt(30))
    add_paragraph(doc, "院内のどのPCからでもブラウザでアクセスする方法",
                  font_size=Pt(14),
                  alignment=WD_ALIGN_PARAGRAPH.CENTER,
                  color=RGBColor(0x55, 0x55, 0x55), space_after=Pt(60))
    add_paragraph(doc, "作成日：2026年3月25日",
                  font_size=Pt(12),
                  alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=Pt(8))
    add_paragraph(doc, "おもろまちメディカルセンター",
                  font_size=Pt(14), bold=True,
                  alignment=WD_ALIGN_PARAGRAPH.CENTER)

    # 改ページ
    doc.add_page_break()

    # ============================================================
    # 1. はじめに
    # ============================================================
    add_heading_styled(doc, "1. はじめに（このマニュアルでできること）", level=1)

    add_paragraph(doc,
        "このマニュアルに従うと、病院内のWi-Fiまたは有線LANに接続された任意のPC・タブレット・"
        "スマホから、ブラウザでベッドコントロールシミュレーターにアクセスできるようになります。")

    add_paragraph(doc, "前提条件：", bold=True, space_before=Pt(10))

    preconditions = [
        "サーバー役のPC（以下「ホストPC」）を1台決めます",
        "ホストPCでアプリを起動している間だけ、他のPCからアクセスできます",
        "ホストPCの電源を切る、またはアプリを終了すると、アクセスできなくなります",
        "外部（インターネット）からはアクセスできません。院内ネットワーク限定です",
    ]
    for item in preconditions:
        p = doc.add_paragraph(item, style="List Bullet")
        for run in p.runs:
            set_font(run, size=FONT_SIZE_BODY)

    # ============================================================
    # 2. 必要なもの
    # ============================================================
    add_heading_styled(doc, "2. 必要なもの", level=1)

    # テーブルで表示
    table = doc.add_table(rows=4, cols=2)
    table.style = "Light List Accent 1"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    headers = ["項目", "詳細"]
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ""
        run = cell.paragraphs[0].add_run(h)
        set_font(run, bold=True, size=FONT_SIZE_BODY)

    data = [
        ("ホストPC", "Mac 1台（このマニュアルではmacOS前提で説明します）"),
        ("ネットワーク", "院内Wi-FiまたはLANへの接続"),
        ("アクセス端末", "PC、タブレット、スマホ（ブラウザがあればOS不問）"),
    ]
    for row_idx, (col1, col2) in enumerate(data, start=1):
        for col_idx, text in enumerate([col1, col2]):
            cell = table.rows[row_idx].cells[col_idx]
            cell.text = ""
            run = cell.paragraphs[0].add_run(text)
            set_font(run, size=FONT_SIZE_BODY)

    add_paragraph(doc, "", space_after=Pt(6))

    # ============================================================
    # 3. ホストPCの初期セットアップ（初回のみ）
    # ============================================================
    add_heading_styled(doc, "3. ホストPCの初期セットアップ（初回のみ）", level=1)

    # 3-1
    add_heading_styled(doc, "3-1. ターミナルの開き方", level=2)
    steps_terminal = [
        "キーボードで ⌘（コマンド）キー + スペースキー を同時に押す",
        "「Spotlight検索」が表示される",
        "「ターミナル」と入力してEnterキーを押す",
        "黒い（または白い）ウィンドウが開く。これが「ターミナル」です",
    ]
    for i, step in enumerate(steps_terminal, 1):
        add_paragraph(doc, f"{i}. {step}")

    # 3-2
    add_heading_styled(doc, "3-2. Pythonの確認", level=2)
    add_paragraph(doc, "ターミナルに以下を入力してEnter：")
    add_code_block(doc, "python3 --version")
    add_paragraph(doc,
        "「Python 3.xx.x」と表示されればOK。表示されない場合はIT部門に相談してください。")

    # 3-3
    add_heading_styled(doc, "3-3. 必要なソフトのインストール", level=2)
    add_paragraph(doc, "以下のコマンドを1行ずつ入力してEnter：")

    add_code_with_description(doc, "cd ~/ai-management",
                              "意味：ai-managementフォルダに移動する")
    add_code_with_description(doc, "python3 -m venv .venv",
                              "意味：仮想環境を作る（他のソフトに影響しない安全な箱のようなもの）")
    add_code_with_description(doc, "source .venv/bin/activate",
                              "意味：仮想環境に入る（ターミナルの行頭に(.venv)と表示される）")
    add_code_with_description(doc, "pip install -r requirements.txt",
                              "意味：シミュレーターに必要なソフトを自動インストールする")

    p = add_paragraph(doc, "「Successfully installed ...」と表示されれば完了です。", bold=True)

    # 3-4
    add_heading_styled(doc, "3-4. IPアドレスの確認", level=2)
    add_paragraph(doc,
        "IPアドレスとは、ネットワーク上でのこのPCの住所のようなものです。")

    add_heading_styled(doc, "【方法A】設定画面で確認（簡単）", level=3)
    steps_ip_a = [
        "画面左上のAppleマーク（🍎）をクリック",
        "「システム設定」をクリック",
        "左メニューの「Wi-Fi」をクリック",
        "接続中のネットワーク名の横にある「詳細...」をクリック",
        "「TCP/IP」タブをクリック",
        "「IPアドレス」の欄に表示される数字（例：192.168.0.11）をメモ",
    ]
    for i, step in enumerate(steps_ip_a, 1):
        add_paragraph(doc, f"{i}. {step}")

    add_heading_styled(doc, "【方法B】ターミナルで確認", level=3)
    add_code_block(doc, "ipconfig getifaddr en0")
    add_paragraph(doc,
        "表示された数字（例：192.168.0.11）をメモ。これが「ホストPCのアドレス」です。")

    # ============================================================
    # 4. アプリの起動方法（毎回の手順）
    # ============================================================
    add_heading_styled(doc, "4. アプリの起動方法（毎回の手順）", level=1)

    add_heading_styled(doc, "4-1. ターミナルを開く", level=2)
    add_paragraph(doc, "⌘ + スペース →「ターミナル」と入力 → Enter")

    add_heading_styled(doc, "4-2. 以下の3行を1行ずつ入力", level=2)
    add_code_block(doc, "cd ~/ai-management")
    add_code_block(doc, "source .venv/bin/activate")
    add_code_block(doc, "streamlit run scripts/bed_control_simulator_app.py")

    add_heading_styled(doc, "4-3. 起動の確認", level=2)
    add_paragraph(doc, "以下のような表示が出れば成功です：")
    add_code_block(doc,
        "You can now view your Streamlit app in your browser.\n"
        "Local URL: http://localhost:8501\n"
        "Network URL: http://192.168.0.11:8501")
    add_paragraph(doc,
        "「Network URL」の行に表示されているアドレスが、他のPCからアクセスするためのURLです。",
        bold=True)

    add_heading_styled(doc, "4-4. 動作確認", level=2)
    add_paragraph(doc,
        "ホストPC自身のブラウザ（Safari/Chrome）で http://localhost:8501 にアクセスして、"
        "シミュレーターが表示されることを確認してください。")

    # ============================================================
    # 5. 他のPCからアクセスする方法
    # ============================================================
    add_heading_styled(doc, "5. 他のPCからアクセスする方法", level=1)

    add_heading_styled(doc, "5-1. 同じネットワークに接続", level=2)
    add_paragraph(doc,
        "アクセスする側のPC/タブレットが、ホストPCと同じWi-Fi（同じネットワーク名）に"
        "接続されていることを確認します。")

    add_heading_styled(doc, "5-2. ブラウザでURLを入力", level=2)
    add_paragraph(doc,
        "ブラウザ（Chrome、Safari、Edge等なんでもOK）のアドレスバーに以下を入力してEnter：")
    add_code_block(doc, "http://192.168.0.11:8501")

    p = add_paragraph(doc, "")
    run = p.add_run("※「192.168.0.11」の部分は、3-4でメモしたホストPCのIPアドレスに置き換えてください。")
    set_font(run, bold=True, color=RGBColor(0xCC, 0x00, 0x00))
    p2 = add_paragraph(doc, "")
    run2 = p2.add_run("※「:8501」を忘れずに入力してください。")
    set_font(run2, bold=True, color=RGBColor(0xCC, 0x00, 0x00))

    add_heading_styled(doc, "5-3. 完了", level=2)
    add_paragraph(doc,
        "シミュレーターの画面が表示されれば成功です！\n"
        "左サイドバーでパラメータを変更し、「シミュレーション実行」ボタンを押して使ってください。")

    # ============================================================
    # 6. アクセスできないときの対処法
    # ============================================================
    add_heading_styled(doc, "6. アクセスできないときの対処法", level=1)

    add_heading_styled(doc, "症状1: ページが表示されない・読み込みが終わらない", level=2)
    add_paragraph(doc, "確認すること：", bold=True)
    add_checklist_item(doc, "ホストPCでアプリが起動しているか（ターミナルにStreamlitの表示が出ているか）")
    add_checklist_item(doc, "同じWi-Fiに接続されているか")
    add_checklist_item(doc, "IPアドレスが正しいか（ホストPCで再確認：ipconfig getifaddr en0）")
    add_checklist_item(doc, "URLの末尾に :8501 を付けているか")

    add_heading_styled(doc, "症状2: ファイアウォールでブロックされている場合", level=2)
    fw_steps = [
        "Appleメニュー（🍎）→ システム設定 → ネットワーク → ファイアウォール",
        "ファイアウォールが「オン」の場合 →「オプション」をクリック",
        "「外部からの接続をすべてブロック」が オフ になっていることを確認",
        "または「+」ボタンで「Python」を許可リストに追加",
    ]
    for i, step in enumerate(fw_steps, 1):
        add_paragraph(doc, f"{i}. {step}")

    add_heading_styled(doc, "症状3: IPアドレスが変わってしまった", level=2)
    add_paragraph(doc,
        "Wi-Fiの再接続やPC再起動でIPアドレスが変わることがあります。\n"
        "ホストPCで再度 ipconfig getifaddr en0 を実行して、新しいアドレスを他のスタッフに伝えてください。")

    # ============================================================
    # 7. アプリの終了方法
    # ============================================================
    add_heading_styled(doc, "7. アプリの終了方法", level=1)
    p = add_paragraph(doc, "")
    run = p.add_run("ターミナルで Control + C を押す")
    set_font(run, bold=True)
    run2 = p.add_run("（Controlキーを押しながらCキー）。")
    set_font(run2)
    add_paragraph(doc,
        "これでアプリが停止し、他のPCからもアクセスできなくなります。\n"
        "ターミナルウィンドウは閉じてOKです。")

    # ============================================================
    # 8. 毎朝の起動をワンクリックにする方法
    # ============================================================
    add_heading_styled(doc, "8. 毎朝の起動をワンクリックにする方法", level=1)
    add_paragraph(doc,
        "デスクトップにショートカットファイルを作っておくと、ダブルクリックだけで起動できます。")

    add_heading_styled(doc, "作成手順", level=2)
    shortcut_steps = [
        "Finderで「アプリケーション」→「テキストエディット」を開く",
        "「フォーマット」メニュー →「標準テキストにする」を選択",
        "以下の4行を入力（コピー＆ペースト推奨）：",
    ]
    for i, step in enumerate(shortcut_steps, 1):
        add_paragraph(doc, f"{i}. {step}")

    add_code_block(doc,
        "#!/bin/bash\n"
        "cd ~/ai-management\n"
        "source .venv/bin/activate\n"
        "streamlit run scripts/bed_control_simulator_app.py")

    remaining_steps = [
        ("4", "「ファイル」→「保存」→ 保存先を「デスクトップ」に変更"),
        ("5", "ファイル名を「シミュレーター起動.command」にして保存"),
        ("6", "ターミナルを開いて以下を入力："),
    ]
    for num, step in remaining_steps:
        add_paragraph(doc, f"{num}. {step}")

    add_code_block(doc, "chmod +x ~/Desktop/シミュレーター起動.command")

    add_paragraph(doc,
        "7. 以後、デスクトップの「シミュレーター起動.command」をダブルクリックするだけで起動できます")

    # ============================================================
    # 9. 運用のヒント
    # ============================================================
    add_heading_styled(doc, "9. 運用のヒント", level=1)

    hints = [
        "ホストPCがスリープすると他のPCからアクセスできなくなります。日中はスリープしない設定にすると"
        "便利です（システム設定 → ディスプレイ → スリープを「しない」に設定）",
        "IPアドレスが毎回変わって不便な場合は、IT部門に「IPアドレスの固定」を依頼してください",
        "複数人が同時にアクセスしても問題ありません（ただし10人以上になると動作が遅くなる場合があります）",
        "各ユーザーのパラメータ変更は他のユーザーに影響しません（セッションは独立しています）",
    ]
    for hint in hints:
        p = doc.add_paragraph(hint, style="List Bullet")
        for run in p.runs:
            set_font(run, size=FONT_SIZE_BODY)

    # ============================================================
    # 10. よくある質問（FAQ）
    # ============================================================
    add_heading_styled(doc, "10. よくある質問（FAQ）", level=1)

    faqs = [
        ("iPadやスマホからもアクセスできますか？",
         "はい。同じWi-Fiに接続し、ブラウザでURLにアクセスすればOKです。"),
        ("外出先からアクセスできますか？",
         "できません。院内ネットワーク限定です。外出先からも使いたい場合は「Streamlit Community Cloud "
         "デプロイマニュアル」を参照してください。"),
        ("ホストPCのスペックはどのくらい必要ですか？",
         "メモリ4GB以上のMacであれば十分です。"),
        ("Windowsでもホストになれますか？",
         "はい。Python3とpipが使えれば同じ手順で可能です。ターミナルの代わりにコマンドプロンプトまたは"
         "PowerShellを使います。"),
        ("アプリのアップデートはどうすればいいですか？",
         "scripts/フォルダのファイルを差し替えて、アプリを再起動（Ctrl+C で停止 → 再度 streamlit run ...）"
         "するだけです。"),
    ]

    # FAQテーブル
    faq_table = doc.add_table(rows=len(faqs) + 1, cols=2)
    faq_table.style = "Light List Accent 1"
    faq_table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # ヘッダー
    for i, h in enumerate(["質問", "回答"]):
        cell = faq_table.rows[0].cells[i]
        cell.text = ""
        run = cell.paragraphs[0].add_run(h)
        set_font(run, bold=True, size=FONT_SIZE_BODY)

    # 列幅設定
    for row in faq_table.rows:
        row.cells[0].width = Cm(6)
        row.cells[1].width = Cm(10)

    for row_idx, (q, a) in enumerate(faqs, start=1):
        cell_q = faq_table.rows[row_idx].cells[0]
        cell_q.text = ""
        run_q = cell_q.paragraphs[0].add_run(f"Q: {q}")
        set_font(run_q, bold=True, size=FONT_SIZE_BODY)

        cell_a = faq_table.rows[row_idx].cells[1]
        cell_a.text = ""
        run_a = cell_a.paragraphs[0].add_run(f"A: {a}")
        set_font(run_a, size=FONT_SIZE_BODY)

    # ページ番号追加
    add_page_numbers(doc)

    # 保存
    os.makedirs(os.path.dirname(OUTPUT_PATH), exist_ok=True)
    doc.save(OUTPUT_PATH)
    print(f"マニュアルを生成しました: {OUTPUT_PATH}")


if __name__ == "__main__":
    create_manual()
