#!/usr/bin/env python3
"""
Claude Code 移行マニュアル DOCX生成スクリプト
おもろまちメディカルセンター AI ワークスペースを別PCに移行するための手順書
"""

from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
import os

OUTPUT_PATH = os.path.join(
    os.path.dirname(os.path.dirname(os.path.abspath(__file__))),
    "docs", "admin", "claude_code_migration_manual.docx"
)


def setup_styles(doc):
    """ドキュメントのスタイルを設定"""
    style = doc.styles["Normal"]
    font = style.font
    font.name = "Yu Gothic"
    font.size = Pt(10.5)
    font.color.rgb = RGBColor(0x33, 0x33, 0x33)
    style.element.rPr.rFonts.set(qn("w:eastAsia"), "Yu Gothic")
    pf = style.paragraph_format
    pf.space_after = Pt(6)
    pf.line_spacing = 1.4

    for level in range(1, 4):
        h = doc.styles[f"Heading {level}"]
        h.font.name = "Yu Gothic"
        h.element.rPr.rFonts.set(qn("w:eastAsia"), "Yu Gothic")
        h.font.color.rgb = RGBColor(0x1A, 0x47, 0x8A)
        h.font.bold = True
        if level == 1:
            h.font.size = Pt(20)
            h.paragraph_format.space_before = Pt(24)
            h.paragraph_format.space_after = Pt(12)
        elif level == 2:
            h.font.size = Pt(16)
            h.paragraph_format.space_before = Pt(18)
            h.paragraph_format.space_after = Pt(8)
        else:
            h.font.size = Pt(13)
            h.paragraph_format.space_before = Pt(12)
            h.paragraph_format.space_after = Pt(6)

    # Code style
    if "Code" not in [s.name for s in doc.styles]:
        code_style = doc.styles.add_style("Code", WD_STYLE_TYPE.PARAGRAPH)
        code_style.font.name = "Consolas"
        code_style.font.size = Pt(9.5)
        code_style.font.color.rgb = RGBColor(0x1E, 0x1E, 0x1E)
        code_style.paragraph_format.space_before = Pt(4)
        code_style.paragraph_format.space_after = Pt(4)
        code_style.paragraph_format.line_spacing = 1.2
        code_style.paragraph_format.left_indent = Cm(0.5)

    # Code Char style for inline
    if "Code Char" not in [s.name for s in doc.styles]:
        code_char = doc.styles.add_style("Code Char", WD_STYLE_TYPE.CHARACTER)
        code_char.font.name = "Consolas"
        code_char.font.size = Pt(9.5)
        code_char.font.color.rgb = RGBColor(0x1E, 0x1E, 0x1E)

    # List Bullet
    lb = doc.styles["List Bullet"]
    lb.font.name = "Yu Gothic"
    lb.element.rPr.rFonts.set(qn("w:eastAsia"), "Yu Gothic")
    lb.font.size = Pt(10.5)
    lb.paragraph_format.space_after = Pt(3)


def add_code_block(doc, code_text):
    """コードブロックを追加（薄いグレー背景付き）"""
    lines = code_text.strip().split("\n")
    for line in lines:
        p = doc.add_paragraph(line, style="Code")
        # Add gray shading
        shading = parse_xml(
            f'<w:shd {nsdecls("w")} w:fill="F0F0F0" w:val="clear"/>'
        )
        p.paragraph_format.element.get_or_add_pPr().append(shading)


def add_bullet(doc, text, bold_prefix=None):
    """箇条書きを追加"""
    p = doc.add_paragraph(style="List Bullet")
    if bold_prefix:
        run = p.add_run(bold_prefix)
        run.bold = True
        p.add_run(text)
    else:
        p.add_run(text)
    return p


def add_note(doc, text):
    """注意書きを追加"""
    p = doc.add_paragraph()
    run = p.add_run(f"  {text}")
    run.font.size = Pt(9.5)
    run.font.color.rgb = RGBColor(0xCC, 0x66, 0x00)
    run.font.name = "Yu Gothic"
    run.element.rPr.rFonts.set(qn("w:eastAsia"), "Yu Gothic")
    return p


def add_warning(doc, text):
    """警告を追加"""
    p = doc.add_paragraph()
    run = p.add_run(f">>> {text}")
    run.bold = True
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0xCC, 0x33, 0x00)
    return p


def add_page_numbers(doc):
    """ページ番号をフッターに追加"""
    section = doc.sections[0]
    footer = section.footer
    footer.is_linked_to_previous = False
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    fld_char_begin = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="begin"/>')
    run._r.append(fld_char_begin)
    run2 = p.add_run(" PAGE ")
    fld_code = parse_xml(f'<w:instrText {nsdecls("w")} xml:space="preserve"> PAGE </w:instrText>')
    run2._r.append(fld_code)
    run3 = p.add_run()
    fld_char_end = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="end"/>')
    run3._r.append(fld_char_end)
    p.runs[0].font.size = Pt(9)


def add_separator(doc):
    """区切り線を追加"""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(8)
    pPr = p.paragraph_format.element.get_or_add_pPr()
    bottom = parse_xml(
        f'<w:pBdr {nsdecls("w")}>'
        f'<w:bottom w:val="single" w:sz="4" w:space="1" w:color="CCCCCC"/>'
        f'</w:pBdr>'
    )
    pPr.append(bottom)


def create_manual():
    doc = Document()
    setup_styles(doc)

    # Page setup
    section = doc.sections[0]
    section.page_height = Cm(29.7)
    section.page_width = Cm(21.0)
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

    add_page_numbers(doc)

    # ========== Title Page ==========
    for _ in range(4):
        doc.add_paragraph()

    title = doc.add_heading("Claude Code 移行マニュアル", level=1)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in title.runs:
        run.font.size = Pt(28)
        run.font.color.rgb = RGBColor(0x1A, 0x47, 0x8A)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run("― 別のPCで同じ環境を再現する ―")
    run.font.size = Pt(16)
    run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
    run.font.name = "Yu Gothic"
    run.element.rPr.rFonts.set(qn("w:eastAsia"), "Yu Gothic")

    doc.add_paragraph()
    doc.add_paragraph()

    info = doc.add_paragraph()
    info.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = info.add_run("対象: おもろまちメディカルセンター AI ワークスペース")
    run.font.size = Pt(12)
    run.font.name = "Yu Gothic"
    run.element.rPr.rFonts.set(qn("w:eastAsia"), "Yu Gothic")

    date_p = doc.add_paragraph()
    date_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = date_p.add_run("作成日: 2026年3月27日")
    run.font.size = Pt(12)
    run.font.name = "Yu Gothic"
    run.element.rPr.rFonts.set(qn("w:eastAsia"), "Yu Gothic")

    doc.add_page_break()

    # ========== Table of Contents ==========
    doc.add_heading("目次", level=1)
    toc_items = [
        ("はじめに", ""),
        ("前提条件", ""),
        ("Step 1:", "Claude Codeデスクトップアプリのインストール"),
        ("Step 2:", "Homebrewのインストール"),
        ("Step 3:", "必要なツールのインストール"),
        ("Step 4:", "プロジェクトのクローン（ダウンロード）"),
        ("Step 5:", "Python仮想環境のセットアップ"),
        ("Step 6:", "Node.js依存パッケージのインストール"),
        ("Step 7:", "MCP（PubMed検索）の設定"),
        ("Step 8:", "Claude Codeの設定"),
        ("Step 9:", "プロジェクトフォルダをClaude Codeで開く"),
        ("Step 10:", "動作確認"),
        ("Step 11:", "Streamlit Cloud（オプション）"),
        ("トラブルシューティング", ""),
        ("フォルダ構成（参考）", ""),
        ("重要なポイント", ""),
    ]
    for prefix, desc in toc_items:
        p = doc.add_paragraph()
        run = p.add_run(f"  {prefix}")
        run.bold = True
        run.font.size = Pt(11)
        if desc:
            run2 = p.add_run(f" {desc}")
            run2.font.size = Pt(11)
        p.paragraph_format.space_after = Pt(2)

    doc.add_page_break()

    # ========== はじめに ==========
    doc.add_heading("はじめに", level=1)
    add_bullet(doc, "このマニュアルでは、現在のClaude Codeプロジェクトを別のMac PCに移行し、同じ環境で作業を継続するための手順を説明します")
    add_bullet(doc, "全てコピペで実行できるようにしています")
    add_bullet(doc, "所要時間: 約30〜45分")

    add_separator(doc)

    # ========== 前提条件 ==========
    doc.add_heading("前提条件", level=1)
    add_bullet(doc, "移行先: macOS搭載のMac")
    add_bullet(doc, "インターネット接続あり")
    add_bullet(doc, "Anthropicアカウント（Claude Pro/Team）を持っていること")

    add_separator(doc)

    # ========== Step 1 ==========
    doc.add_heading("Step 1: Claude Codeデスクトップアプリのインストール", level=1)

    doc.add_paragraph("Claude for Desktop をダウンロードします。")
    doc.add_paragraph()
    p = doc.add_paragraph()
    run = p.add_run("ダウンロードURL: ")
    run.bold = True
    p.add_run("https://claude.ai/download")

    doc.add_paragraph()
    doc.add_paragraph("手順:")
    p = doc.add_paragraph("1. 上記URLにアクセスし、Mac版をダウンロード", style="List Bullet")
    p = doc.add_paragraph("2. ダウンロードした .dmg ファイルを開く", style="List Bullet")
    p = doc.add_paragraph("3. Claude アプリを「アプリケーション」フォルダにドラッグ", style="List Bullet")
    p = doc.add_paragraph("4. アプリを開き、Anthropicアカウントでログイン", style="List Bullet")

    doc.add_paragraph()
    p = doc.add_paragraph()
    run = p.add_run("Claude for Desktop は、チャットインターフェースとターミナル（Claude Code）が統合されたデスクトップアプリです。以前は別途 claude コマンドのインストールが必要でしたが、現在はデスクトップアプリに統合されています。")
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

    add_separator(doc)

    # ========== Step 2 ==========
    doc.add_heading("Step 2: Homebrewのインストール（未インストールの場合）", level=1)

    doc.add_paragraph("HomebrewはmacOS用のパッケージマネージャーです。開発ツールを簡単にインストールできます。")
    doc.add_paragraph()
    doc.add_paragraph("ターミナルを開いて以下のコマンドをコピペして実行してください:")

    add_code_block(doc, '/bin/bash -c "$(curl -fsSL https://raw.githubusercontent.com/Homebrew/install/HEAD/install.sh)"')

    doc.add_paragraph()
    doc.add_paragraph("インストール完了後、画面に表示されるPATH設定コマンドを実行してください。通常は以下のようなコマンドです:")

    add_code_block(doc, """echo 'eval "$(/opt/homebrew/bin/brew shellenv)"' >> ~/.zprofile
eval "$(/opt/homebrew/bin/brew shellenv)\"""")

    add_note(doc, "注意: Apple Silicon Mac（M1/M2/M3/M4）の場合、上記コマンドが表示されます。Intel Macの場合はパスが異なります。画面の指示に従ってください。")

    doc.add_paragraph()
    doc.add_paragraph("インストール確認:")
    add_code_block(doc, "brew --version")

    add_separator(doc)

    # ========== Step 3 ==========
    doc.add_heading("Step 3: 必要なツールのインストール", level=1)

    doc.add_paragraph("以下のコマンドで必要なツールを一括インストールします:")
    add_code_block(doc, "brew install git node python@3.12")

    doc.add_paragraph()
    doc.add_paragraph("各ツールの用途:")
    add_bullet(doc, "ソースコード管理。GitHubからプロジェクトをダウンロードするのに必要", bold_prefix="git: ")
    add_bullet(doc, "JavaScript実行環境。文献のMarkdown→Word変換スクリプトで使用", bold_prefix="node: ")
    add_bullet(doc, "シミュレーター・分析ツールの実行に必要", bold_prefix="python@3.12: ")

    doc.add_paragraph()
    doc.add_paragraph("インストール確認:")
    add_code_block(doc, """git --version
node --version
python3.12 --version""")

    add_separator(doc)

    # ========== Step 4 ==========
    doc.add_heading("Step 4: プロジェクトのクローン（ダウンロード）", level=1)

    doc.add_paragraph("GitHubからプロジェクトをダウンロードします。")
    doc.add_paragraph()
    doc.add_paragraph("方法A: フォルダを先に作成してからダウンロード")
    add_code_block(doc, """mkdir ~/ai-management
cd ~/ai-management
git clone https://github.com/torukubota2023/bed-control-simulator.git .""")

    add_note(doc, '注意: 末尾の「.」は「現在のフォルダにダウンロードする」という意味です。忘れずに入力してください。')

    doc.add_paragraph()
    doc.add_paragraph("方法B: ワンコマンドでダウンロード")
    add_code_block(doc, """git clone https://github.com/torukubota2023/bed-control-simulator.git ~/ai-management
cd ~/ai-management""")

    add_note(doc, "注意: フォルダ名は「ai-management」でなくても構いません。好みの名前に変更してOKです。")

    add_separator(doc)

    # ========== Step 5 ==========
    doc.add_heading("Step 5: Python仮想環境のセットアップ", level=1)

    doc.add_paragraph("プロジェクト専用のPython環境を作成します。仮想環境を使うことで、システム全体のPythonに影響を与えずにパッケージを管理できます。")
    doc.add_paragraph()

    add_code_block(doc, """cd ~/ai-management
python3.12 -m venv .venv
source .venv/bin/activate
pip install -r requirements.txt""")

    doc.add_paragraph()
    add_warning(doc, "重要: source .venv/bin/activate は、ターミナルを開くたびに毎回実行する必要があります。")
    doc.add_paragraph()
    p = doc.add_paragraph()
    run = p.add_run("仮想環境が有効になっているときは、ターミナルのプロンプトの先頭に (.venv) と表示されます。")
    run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

    add_separator(doc)

    # ========== Step 6 ==========
    doc.add_heading("Step 6: Node.js依存パッケージのインストール", level=1)

    doc.add_paragraph("文献のMarkdown→Word変換に必要なパッケージをインストールします:")
    add_code_block(doc, """cd ~/ai-management
npm install docx""")

    add_note(doc, "注意: npm install を実行すると node_modules フォルダが作成されます。これは正常です。")

    add_separator(doc)

    # ========== Step 7 ==========
    doc.add_heading("Step 7: MCP（PubMed検索）の設定", level=1)

    doc.add_paragraph("MCP（Model Context Protocol）は、Claude CodeからPubMed文献検索を行うための接続設定です。この設定ファイルはGitHubには含まれていないため、手動で作成する必要があります。")
    doc.add_paragraph()
    doc.add_paragraph("以下のコマンドをターミナルにコピペして実行してください:")

    add_code_block(doc, """cd ~/ai-management
cat > .mcp.json << 'EOF'
{
  "mcpServers": {
    "pubmed-mcp": {
      "type": "stdio",
      "command": "npx",
      "args": [
        "-y",
        "@ncukondo/pubmed-mcp",
        "--email",
        "t.kubochan@gmail.com",
        "--cache-dir",
        "./pubmed-cache"
      ],
      "env": {}
    }
  }
}
EOF""")

    doc.add_paragraph()
    add_note(doc, '注意: "t.kubochan@gmail.com" の部分は、ご自身のメールアドレスに変更しても構いません（PubMed APIの利用者識別用）。')
    add_bullet(doc, "pubmed-cache フォルダは検索実行時に自動作成されます")

    add_separator(doc)

    # ========== Step 8 ==========
    doc.add_heading("Step 8: Claude Codeの設定", level=1)

    # 8-1
    doc.add_heading("8-1: グローバル設定", level=2)

    doc.add_paragraph("Claude Codeの全体設定ファイルを作成します。以下のコマンドを実行してください:")

    add_code_block(doc, """mkdir -p ~/.claude
cat > ~/.claude/settings.json << 'EOF'
{
  "env": {
    "CLAUDE_CODE_EXPERIMENTAL_AGENT_TEAMS": "1"
  },
  "teammateMode": "auto",
  "autoApprove": {
    "enabled": true
  }
}
EOF""")

    doc.add_paragraph()
    add_bullet(doc, "Agent Teams機能を有効化します（複数のAIエージェントが協力して作業する機能）", bold_prefix="CLAUDE_CODE_EXPERIMENTAL_AGENT_TEAMS: ")
    add_bullet(doc, "チームメイトモードを自動に設定", bold_prefix="teammateMode: ")
    add_bullet(doc, "ツール実行の自動承認を有効化", bold_prefix="autoApprove: ")

    # 8-2
    doc.add_heading("8-2: カスタムコマンドの設定", level=2)

    doc.add_paragraph("カスタムコマンド（/pico-search, /evidence-update, /clinical-docx）はGitHubには含まれていないため、手動で作成する必要があります。")
    doc.add_paragraph()
    doc.add_paragraph("まず、コマンドフォルダを作成します:")
    add_code_block(doc, """cd ~/ai-management
mkdir -p .claude/commands""")

    doc.add_paragraph()
    add_warning(doc, "重要: 以下のファイル内容に含まれる /Users/torukubota/ の部分は、ご自身のホームディレクトリパスに変更してください。")
    doc.add_paragraph()
    doc.add_paragraph("自分のホームディレクトリパスを確認するコマンド:")
    add_code_block(doc, "echo $HOME")
    doc.add_paragraph("（例: /Users/yamada のように表示されます）")

    doc.add_paragraph()
    add_separator(doc)

    # Command 1: pico-search.md
    doc.add_heading("コマンド1: pico-search.md（臨床疑問→PubMed文献検索）", level=3)
    doc.add_paragraph("以下のコマンドでファイルを作成します:")

    pico_content = r'''cd ~/ai-management
cat > .claude/commands/pico-search.md << 'CMDEOF'
# 臨床疑問 → PubMed文献検索スキル

あなたはオーケストレーターです。以下のワークフローを**全てsubagent/Taskに委託**して実行してください。自分では実装しないこと。

## 入力された臨床疑問
$ARGUMENTS

---

## ワークフロー

### Step 1: PICO構造化（subagentに委託）

以下のPICO形式に分解するよう subagent に指示する:
- **P** (Patient/Problem): 患者背景・問題
- **I** (Intervention/Exposure): 介入・暴露・検査・所見
- **C** (Comparison): 比較対象（ある場合）
- **O** (Outcome): 知りたいアウトカム（診断・治療・予後）

PICO分解の結果から、以下の**3種類の検索クエリ**を英語で生成させる:
1. 主診断/疾患名の検索クエリ
2. 鑑別疾患の検索クエリ
3. 治療・マネジメントの検索クエリ

### Step 2: PubMed検索（subagentに委託、3並列）

各クエリについて別々のsubagentを立て、PubMed MCPで検索:
- 検索結果から**上位10件**を選定
- 各文献について: タイトル・著者・雑誌・年・PMID・DOIを記録
- 可能な限りアブストラクトを取得
- 重要度順にランキング（メタ解析 > RCT > コホート > ケースシリーズ > 症例報告）

### Step 3: エビデンス統合・日本語要約（subagentに委託）

取得した文献を統合して以下の構造でまとめる:

```
# 臨床疑問: [疑問を端的に表現]

**検索日:** YYYY-MM-DD
**PICO:**
- P: ...
- I/E: ...
- C: ...
- O: ...

---

## 1. [主要診断/疾患名]

[日本語要約: 病態・診断基準・臨床的特徴]

### エビデンスレベルサマリー
| エビデンスレベル | 内容 | 参考文献 |
|---|---|---|
| Level I (SR/Meta) | ... | PMID:XXXXXXXX |
| Level II (RCT) | ... | PMID:XXXXXXXX |

### 診断精度（利用可能な場合）
| 検査/所見 | 感度 | 特異度 | +LR | -LR | 参考文献 |
|---|---|---|---|---|---|

**参考文献:**
- [著者] [タイトル]. [雑誌] [年];[巻(号)]:[ページ]. PMID:XXXXXXXX [DOI](https://doi.org/...)

---

## 2. [鑑別疾患]
...

---

## 3. 治療・マネジメント
...

---

## 4. 推奨される臨床アプローチ

1. ...
2. ...

---

## まとめ

[3-5行の臨床的インプリケーション]

**検索クエリ記録:**
- Query 1: ...
- Query 2: ...
- Query 3: ...
```

### Step 4: ファイル保存（subagentに委託）

- ファイル名: `YYYY-MM-DD_[トピック英語スネークケース].md`
- 保存先: `/Users/torukubota/ai-management/docs/pubmed/`
- 保存後、CLAUDE.mdの「今週の臨床疑問」セクションに `- [ ] [疑問の要約] → [ファイルへのリンク]` を追記

---

## 完了報告

全Step完了後、以下を報告:
1. 保存したファイルパス
2. 取得した文献数と主要文献トップ3
3. 臨床的結論の要約（3行以内）
CMDEOF'''

    add_code_block(doc, pico_content)

    doc.add_paragraph()

    # Command 2: evidence-update.md
    doc.add_heading("コマンド2: evidence-update.md（既存文献サマリーの更新）", level=3)
    doc.add_paragraph("以下のコマンドでファイルを作成します:")

    evidence_content = r'''cd ~/ai-management
cat > .claude/commands/evidence-update.md << 'CMDEOF'
# 既存文献サマリーの更新スキル

あなたはオーケストレーターです。以下のワークフローを**全てsubagent/Taskに委託**して実行してください。自分では実装しないこと。

## 対象トピックまたはファイルパス
$ARGUMENTS

---

## ワークフロー

### Step 1: 対象ファイルの特定（subagentに委託）

`/Users/torukubota/ai-management/docs/pubmed/` ディレクトリを確認し:
- 引数がファイルパスの場合: そのファイルを使用
- 引数がトピック名の場合: ファイル名に部分一致するファイルを探す
- 引数が空の場合: 最新の.mdファイル（ファイル名の日付順）を使用

既存ファイルを読み込み、以下を抽出:
- 元の臨床疑問とPICO
- 既存の検索クエリ（ファイル末尾に記録されているもの）
- 既存文献のPMIDリスト

### Step 2: 差分検索（subagentに委託）

既存ファイルの検索クエリを使い、**既存PMIDを除外した新規文献**をPubMed MCPで検索:
- 検索期間: 既存ファイルの検索日以降
- フィルタ: 高エビデンスレベル優先（SR/RCT/ガイドライン）
- 新規文献が5件以上ある場合のみ更新を推奨し、ユーザーに確認する

### Step 3: 更新内容の作成（subagentに委託）

新規文献から以下を作成:
- 既存セクションへの追記内容（新規参考文献、更新されたエビデンスレベル表）
- 「更新履歴」セクションへの追記: `### YYYY-MM-DD 更新: [追加文献数]件追加`
- 診断精度・治療推奨に変更がある場合はその旨を明記

### Step 4: ファイル更新（subagentに委託）

既存ファイルに差分を追記:
- ファイル名は変更しない
- ファイル末尾の「検索クエリ記録」セクションも更新（最新検索日を追記）
- 更新後、変更サマリーをユーザーに報告

---

## 完了報告

1. 更新したファイルパス
2. 追加した文献数
3. 主要な新知見（あれば）
4. 次回更新推奨時期（エビデンスの動向による）
CMDEOF'''

    add_code_block(doc, evidence_content)

    doc.add_paragraph()

    # Command 3: clinical-docx.md
    doc.add_heading("コマンド3: clinical-docx.md（文献サマリーDOCX変換）", level=3)
    doc.add_paragraph("以下のコマンドでファイルを作成します:")

    clinical_content = r'''cd ~/ai-management
cat > .claude/commands/clinical-docx.md << 'CMDEOF'
# 文献サマリー DOCX変換スキル

あなたはオーケストレーターです。以下のワークフローを**全てsubagent/Taskに委託**して実行してください。自分では実装しないこと。

## 対象ファイルパス（省略時は最新ファイル）
$ARGUMENTS

---

## ワークフロー

### Step 1: 対象MDファイルの特定（subagentに委託）

- 引数がファイルパスの場合: そのファイルを使用
- 引数が空の場合: `/Users/torukubota/ai-management/docs/pubmed/` 内の最新の.mdファイルを使用（ファイル名の日付で判断）
- INPUT と OUTPUT パスを確定する（OUTPUTは同ディレクトリ・同名で拡張子を.docxに変更）

### Step 2: 汎用変換スクリプトの実行（subagentに委託）

以下のコマンドを実行:
```
node /Users/torukubota/ai-management/scripts/convert_md_to_docx.js "<INPUT_PATH>" "<OUTPUT_PATH>"
```

スクリプトが存在しない場合、またはエラーが発生した場合はユーザーに報告して停止。

### Step 3: 完了確認

- 生成されたDOCXファイルのパスを報告
- ファイルサイズを確認して正常生成を確認

---

## 完了報告

1. 生成したDOCXファイルパス
2. ファイルサイズ
3. 変換元のMDファイル
CMDEOF'''

    add_code_block(doc, clinical_content)

    add_separator(doc)

    # ========== Step 9 ==========
    doc.add_heading("Step 9: プロジェクトフォルダをClaude Codeで開く", level=1)

    doc.add_paragraph("セットアップが完了したら、Claude Codeでプロジェクトを開きます。")
    doc.add_paragraph()
    doc.add_paragraph("方法A: デスクトップアプリから")
    add_bullet(doc, 'Claude for Desktop アプリを開く')
    add_bullet(doc, '「Open Folder」（フォルダを開く）からai-managementフォルダを選択')

    doc.add_paragraph()
    doc.add_paragraph("方法B: ターミナルから")
    add_code_block(doc, """cd ~/ai-management
claude""")

    doc.add_paragraph()
    p = doc.add_paragraph()
    run = p.add_run("CLAUDE.md が自動的に読み込まれ、プロジェクトのコンテキスト（過去の作業履歴、フォルダ構成、ルールなど）が引き継がれます。これがClaude Codeの「記憶」として機能します。")
    run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

    add_separator(doc)

    # ========== Step 10 ==========
    doc.add_heading("Step 10: 動作確認", level=1)

    # 10-1
    doc.add_heading("10-1: シミュレーター起動テスト", level=2)
    add_code_block(doc, """cd ~/ai-management
source .venv/bin/activate
streamlit run scripts/bed_control_simulator_app.py""")
    add_bullet(doc, "ブラウザが自動的に開き、ベッドコントロールシミュレーターが表示されれば成功です")
    add_bullet(doc, "停止するには Ctrl+C を押してください")

    doc.add_paragraph()

    # 10-2
    doc.add_heading("10-2: PubMed検索テスト", level=2)
    doc.add_paragraph("Claude Codeで以下のカスタムコマンドを入力してみてください:")
    add_code_block(doc, "/pico-search 肺炎の診断")
    add_bullet(doc, "PubMed文献検索が実行され、結果がdocs/pubmed/フォルダに保存されれば成功です")

    doc.add_paragraph()

    # 10-3
    doc.add_heading("10-3: Word変換テスト", level=2)
    add_code_block(doc, """cd ~/ai-management
node scripts/convert_md_to_docx.js docs/pubmed/2026-03-20_hospital_universal_masking_post_COVID.md""")
    add_bullet(doc, "同じフォルダに .docx ファイルが生成されれば成功です")

    add_separator(doc)

    # ========== Step 11 ==========
    doc.add_heading("Step 11: Streamlit Cloud（オプション）", level=1)

    doc.add_paragraph("ベッドコントロールシミュレーターは Streamlit Cloud にデプロイ済みです。別PCからもブラウザでアクセスできます。")
    doc.add_paragraph()

    p = doc.add_paragraph()
    run = p.add_run("URL: ")
    run.bold = True
    p.add_run("https://omc-bed-sim.streamlit.app")
    add_note(doc, "注意: 再デプロイが必要な場合があります")

    doc.add_paragraph()
    doc.add_paragraph("新規デプロイ手順:")
    add_bullet(doc, "share.streamlit.io にアクセス")
    add_bullet(doc, "GitHubアカウントでログイン")
    add_bullet(doc, "torukubota2023/bed-control-simulator", bold_prefix="Repository: ")
    add_bullet(doc, "main", bold_prefix="Branch: ")
    add_bullet(doc, "scripts/bed_control_simulator_app.py", bold_prefix="Main file path: ")
    add_bullet(doc, "3.12", bold_prefix="Python version: ")
    add_bullet(doc, "「Deploy!」ボタンをクリック")

    add_separator(doc)
    doc.add_page_break()

    # ========== トラブルシューティング ==========
    doc.add_heading("トラブルシューティング", level=1)

    # Q1
    doc.add_heading("Q: python3.12 コマンドが見つからない", level=3)
    p = doc.add_paragraph()
    run = p.add_run("→ ")
    run.bold = True
    p.add_run("brew install python@3.12 を再実行してください。それでも解決しない場合、PATHの設定が必要です:")
    add_code_block(doc, """echo 'export PATH="/opt/homebrew/opt/python@3.12/bin:$PATH"' >> ~/.zshrc
source ~/.zshrc""")

    doc.add_paragraph()

    # Q2
    doc.add_heading("Q: npm コマンドが見つからない", level=3)
    p = doc.add_paragraph()
    run = p.add_run("→ ")
    run.bold = True
    p.add_run("brew install node を再実行してください。")

    doc.add_paragraph()

    # Q3
    doc.add_heading("Q: Streamlitが起動しない", level=3)
    p = doc.add_paragraph()
    run = p.add_run("→ ")
    run.bold = True
    p.add_run("仮想環境を有効化しているか確認してください:")
    add_code_block(doc, "source .venv/bin/activate")

    doc.add_paragraph()

    # Q4
    doc.add_heading("Q: PubMed MCPが動かない", level=3)
    p = doc.add_paragraph()
    run = p.add_run("→ ")
    run.bold = True
    p.add_run(".mcp.json ファイルがプロジェクトルート（ai-managementフォルダ直下）にあるか確認してください:")
    add_code_block(doc, """cd ~/ai-management
ls -la .mcp.json""")

    doc.add_paragraph()

    # Q5
    doc.add_heading("Q: グラフの日本語が文字化けする", level=3)
    p = doc.add_paragraph()
    run = p.add_run("→ ")
    run.bold = True
    p.add_run("macOSでは通常問題ありません。Streamlit Cloud上で文字化けする場合は、packages.txt に fonts-noto-cjk が記載されていることを確認してください。")

    doc.add_paragraph()

    # Q6
    doc.add_heading('Q: 「permission denied」エラーが出る', level=3)
    p = doc.add_paragraph()
    run = p.add_run("→ ")
    run.bold = True
    p.add_run("Claude Code の設定画面で、必要なツール（Bash、Read、Write、Edit等）の権限を許可してください。")

    add_separator(doc)

    # ========== フォルダ構成 ==========
    doc.add_heading("フォルダ構成（参考）", level=1)

    add_code_block(doc, """ai-management/
├── CLAUDE.md          ← プロジェクトメモリ（最重要）
├── .mcp.json          ← PubMed検索設定（要手動作成）
├── .claude/commands/  ← カスタムコマンド（要手動作成）
├── .streamlit/        ← Streamlit設定
├── .venv/             ← Python仮想環境（要再作成）
├── docs/              ← 教育資料・文献・経営文書
├── scripts/           ← シミュレーター・変換スクリプト
├── requirements.txt   ← Python依存パッケージ
├── runtime.txt        ← Python版指定
└── packages.txt       ← システムパッケージ""")

    add_separator(doc)

    # ========== 重要なポイント ==========
    doc.add_heading("重要なポイント", level=1)

    p = doc.add_paragraph()
    run = p.add_run("CLAUDE.md")
    run.bold = True
    run.font.size = Pt(11)
    p.add_run(" がプロジェクトの「記憶」です。これがあればClaudeは過去の文脈を理解します。")

    p = doc.add_paragraph()
    run = p.add_run("カスタムコマンド")
    run.bold = True
    run.font.size = Pt(11)
    p.add_run("（/pico-search, /evidence-update, /clinical-docx）は手動で再作成が必要です。")

    p = doc.add_paragraph()
    run = p.add_run("MCP設定")
    run.bold = True
    run.font.size = Pt(11)
    p.add_run("（.mcp.json）も手動で再作成が必要です。")

    p = doc.add_paragraph()
    run = p.add_run("GitHubにpushされているファイル")
    run.bold = True
    run.font.size = Pt(11)
    p.add_run("は git pull で常に最新状態に同期できます:")
    add_code_block(doc, """cd ~/ai-management
git pull origin main""")

    # Save
    doc.save(OUTPUT_PATH)
    print(f"マニュアルを生成しました: {OUTPUT_PATH}")
    print(f"ファイルサイズ: {os.path.getsize(OUTPUT_PATH):,} bytes")


if __name__ == "__main__":
    create_manual()
