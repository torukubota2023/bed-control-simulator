#!/usr/bin/env python3
"""
Claude Code 環境移行マニュアル v2 DOCX生成スクリプト
別PCで同じワークスペースを継続するための手順書を生成する
"""

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
import os

OUTPUT_PATH = os.path.join(
    os.path.dirname(os.path.dirname(os.path.abspath(__file__))),
    "docs", "admin", "claude_code_migration_manual_v2.docx"
)


def setup_styles(doc):
    """ドキュメントのスタイルを設定"""
    # Normal style
    style = doc.styles["Normal"]
    font = style.font
    font.name = "Yu Gothic"
    font.size = Pt(10.5)
    font.color.rgb = RGBColor(0x33, 0x33, 0x33)
    style.element.rPr.rFonts.set(qn("w:eastAsia"), "Yu Gothic")
    pf = style.paragraph_format
    pf.space_after = Pt(6)
    pf.line_spacing = 1.4

    # Heading styles
    for level in range(1, 4):
        h = doc.styles[f"Heading {level}"]
        h.font.name = "Yu Gothic"
        h.element.rPr.rFonts.set(qn("w:eastAsia"), "Yu Gothic")
        h.font.color.rgb = RGBColor(0x1A, 0x47, 0x8A)
        h.font.bold = True
        if level == 1:
            h.font.size = Pt(20)
            h.paragraph_format.space_before = Pt(24)
            h.paragraph_format.space_after = Pt(12)
        elif level == 2:
            h.font.size = Pt(16)
            h.paragraph_format.space_before = Pt(18)
            h.paragraph_format.space_after = Pt(8)
        else:
            h.font.size = Pt(13)
            h.paragraph_format.space_before = Pt(12)
            h.paragraph_format.space_after = Pt(6)

    # Code block style
    if "Code Block" not in [s.name for s in doc.styles]:
        code_style = doc.styles.add_style("Code Block", 1)  # paragraph style
        code_style.font.name = "Courier New"
        code_style.font.size = Pt(9.5)
        code_style.font.color.rgb = RGBColor(0x1A, 0x1A, 0x1A)
        code_style.paragraph_format.space_before = Pt(4)
        code_style.paragraph_format.space_after = Pt(4)
        code_style.paragraph_format.line_spacing = 1.2
        code_style.paragraph_format.left_indent = Inches(0.3)

    # List Bullet style
    list_style = doc.styles["List Bullet"]
    list_style.font.name = "Yu Gothic"
    list_style.font.size = Pt(10.5)
    list_style.element.rPr.rFonts.set(qn("w:eastAsia"), "Yu Gothic")


def set_margins(doc):
    """ページマージンを設定"""
    for section in doc.sections:
        section.top_margin = Inches(1.0)
        section.bottom_margin = Inches(0.8)
        section.left_margin = Inches(1.0)
        section.right_margin = Inches(1.0)


def add_code_block(doc, code_text):
    """コードブロックを追加（グレー背景付き）"""
    for line in code_text.strip().split("\n"):
        p = doc.add_paragraph(line, style="Code Block")
        # Add gray shading
        shading = parse_xml(
            f'<w:shd {nsdecls("w")} w:fill="F0F0F0" w:val="clear"/>'
        )
        p.paragraph_format.element.get_or_add_pPr().append(shading)


def add_bold_paragraph(doc, text):
    """太字パラグラフを追加"""
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    run.font.name = "Yu Gothic"
    run.font.size = Pt(10.5)
    run.element.rPr.rFonts.set(qn("w:eastAsia"), "Yu Gothic")
    return p


def add_bullet(doc, text, bold_prefix=None):
    """箇条書きを追加。bold_prefixがあればその部分を太字に"""
    p = doc.add_paragraph(style="List Bullet")
    if bold_prefix:
        run = p.add_run(bold_prefix)
        run.bold = True
        run.font.name = "Yu Gothic"
        run.font.size = Pt(10.5)
        run.element.rPr.rFonts.set(qn("w:eastAsia"), "Yu Gothic")
        run = p.add_run(text)
        run.font.name = "Yu Gothic"
        run.font.size = Pt(10.5)
        run.element.rPr.rFonts.set(qn("w:eastAsia"), "Yu Gothic")
    else:
        run = p.add_run(text)
        run.font.name = "Yu Gothic"
        run.font.size = Pt(10.5)
        run.element.rPr.rFonts.set(qn("w:eastAsia"), "Yu Gothic")
    return p


def add_normal(doc, text):
    """通常テキストを追加"""
    p = doc.add_paragraph(text)
    return p


def add_page_break(doc):
    """改ページを追加"""
    doc.add_page_break()


def build_title_page(doc):
    """タイトルページを作成"""
    # Spacer
    for _ in range(4):
        doc.add_paragraph()

    # Title
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Claude Code 環境移行マニュアル")
    run.font.size = Pt(28)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0x1A, 0x47, 0x8A)
    run.font.name = "Yu Gothic"
    run.element.rPr.rFonts.set(qn("w:eastAsia"), "Yu Gothic")

    # Subtitle
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("〜別のPCで同じワークスペースを継続する方法〜")
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
    run.font.name = "Yu Gothic"
    run.element.rPr.rFonts.set(qn("w:eastAsia"), "Yu Gothic")

    doc.add_paragraph()

    # Separator line
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("━" * 40)
    run.font.color.rgb = RGBColor(0x1A, 0x47, 0x8A)
    run.font.size = Pt(10)

    doc.add_paragraph()

    # Author
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("おもろまちメディカルセンター AI管理ワークスペース")
    run.font.size = Pt(12)
    run.font.name = "Yu Gothic"
    run.element.rPr.rFonts.set(qn("w:eastAsia"), "Yu Gothic")

    # Date
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("2026年3月31日")
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
    run.font.name = "Yu Gothic"
    run.element.rPr.rFonts.set(qn("w:eastAsia"), "Yu Gothic")


def build_introduction(doc):
    """はじめにセクション"""
    doc.add_heading("はじめに", level=1)

    add_normal(doc,
        "このマニュアルは、現在お使いのMacで構築したClaude Code環境"
        "（ai-managementワークスペース）を、別のPC（Mac/Windows）に移行し、"
        "同じように作業を継続するための手順書です。"
    )
    add_normal(doc,
        "初心者の方でもコピペだけで進められるよう、全コマンドをそのまま記載しています。"
    )

    add_bold_paragraph(doc, "所要時間の目安: 約30〜45分")

    add_bold_paragraph(doc, "必要なもの:")
    add_bullet(doc, "移行先のPC（Mac または Windows）")
    add_bullet(doc, "インターネット接続")
    add_bullet(doc, "Anthropicアカウント（Claude Proプラン）")
    add_bullet(doc, "GitHubアカウント（torukubota2023）")


def build_toc(doc):
    """全体の流れ（目次）セクション"""
    doc.add_heading("全体の流れ", level=1)

    steps = [
        "ステップ1: Claude Codeをインストールする",
        "ステップ2: GitHubからプロジェクトを取得する",
        "ステップ3: Python環境をセットアップする",
        "ステップ4: Node.js環境をセットアップする",
        "ステップ5: Claude Codeの設定を復元する",
        "ステップ6: MCP（PubMed検索）を設定する",
        "ステップ7: 動作確認する",
    ]
    for step in steps:
        add_bullet(doc, step)


def build_step1(doc):
    """ステップ1: Claude Codeをインストールする"""
    add_page_break(doc)
    doc.add_heading("ステップ1: Claude Codeをインストールする", level=2)

    # Mac desktop
    doc.add_heading("Macの場合 — デスクトップ版アプリ（推奨）", level=3)

    add_normal(doc, "1. 以下のURLにアクセス:")
    add_normal(doc, "   https://claude.ai/download")
    add_normal(doc, '2.「Mac版をダウンロード」をクリック')
    add_normal(doc, "3. ダウンロードした .dmg ファイルを開き、Claudeをアプリケーションフォルダにドラッグ")
    add_normal(doc, "4. アプリケーションからClaudeを起動")
    add_normal(doc, "5. Anthropicアカウントでログイン")
    add_normal(doc, "6. Claude Codeの機能を有効化:")
    add_bullet(doc, 'Claudeアプリの設定（⌘+,）→「Claude Code」タブ →「Enable Claude Code」をオン')

    # Mac CLI
    doc.add_heading("Macの場合 — CLI版（ターミナル）", level=3)
    add_normal(doc, "ターミナル.appを開いて、以下をコピペして実行:")
    add_code_block(doc, "npm install -g @anthropic-ai/claude-code")
    add_normal(doc, "※ Node.jsが必要（ステップ4参照）")

    add_normal(doc, "CLIでのログイン:")
    add_code_block(doc, "claude login")

    # Windows
    doc.add_heading("Windowsの場合", level=3)
    add_normal(doc, "1. https://claude.ai/download からWindows版をダウンロード")
    add_normal(doc, "2. インストーラーを実行")
    add_normal(doc, "3. Anthropicアカウントでログイン")
    add_normal(doc, "4. 設定 → Claude Code → Enable Claude Code をオン")


def build_step2(doc):
    """ステップ2: GitHubからプロジェクトを取得する"""
    add_page_break(doc)
    doc.add_heading("ステップ2: GitHubからプロジェクトを取得する", level=2)

    doc.add_heading("2-1. Gitのインストール確認", level=3)
    add_code_block(doc, "git --version")
    add_normal(doc, "→ バージョンが表示されればOK")
    add_normal(doc, '→「command not found」の場合:')
    add_bullet(doc, 'Mac: Xcodeコマンドラインツールが自動インストール →「インストール」を押す')
    add_bullet(doc, "Windows: https://git-scm.com からダウンロード・インストール")

    doc.add_heading("2-2. プロジェクトをクローン（ダウンロード）", level=3)
    add_code_block(doc, "cd ~\ngit clone https://github.com/torukubota2023/bed-control-simulator.git ai-management")

    doc.add_heading("2-3. フォルダに移動して確認", level=3)
    add_code_block(doc, "cd ~/ai-management\nls -la")
    add_normal(doc, "以下のファイルが見えればOK:")
    add_bullet(doc, "CLAUDE.md")
    add_bullet(doc, "scripts/")
    add_bullet(doc, "docs/")
    add_bullet(doc, "data/")
    add_bullet(doc, "requirements.txt")
    add_bullet(doc, ".streamlit/")


def build_step3(doc):
    """ステップ3: Python環境をセットアップする"""
    add_page_break(doc)
    doc.add_heading("ステップ3: Python環境をセットアップする", level=2)

    doc.add_heading("3-1. Pythonのインストール確認", level=3)
    add_code_block(doc, "python3 --version")
    add_normal(doc, "→ Python 3.12.x が表示されればOK")

    add_bold_paragraph(doc, "Macでインストール（Homebrew経由）:")
    add_code_block(doc, "brew install python@3.12")
    add_normal(doc, "※ Homebrewが入っていない場合:")
    add_code_block(doc, '/bin/bash -c "$(curl -fsSL https://raw.githubusercontent.com/Homebrew/install/HEAD/install.sh)"')

    add_bold_paragraph(doc, "Windowsでインストール:")
    add_normal(doc, "https://www.python.org/downloads/ から Python 3.12 をダウンロード")
    add_normal(doc, '※「Add Python to PATH」に必ずチェック')

    doc.add_heading("3-2. 仮想環境（venv）を作成", level=3)
    add_code_block(doc, "cd ~/ai-management\npython3 -m venv .venv")

    doc.add_heading("3-3. 仮想環境を有効化", level=3)
    add_normal(doc, "Mac:")
    add_code_block(doc, "source .venv/bin/activate")
    add_normal(doc, "Windows:")
    add_code_block(doc, r".venv\Scripts\activate")
    add_normal(doc, "→ プロンプトの先頭に (.venv) が表示されればOK")

    doc.add_heading("3-4. 必要なパッケージをインストール", level=3)
    add_code_block(doc, "pip install -r requirements.txt")
    add_normal(doc, "インストールされるもの:")
    add_bullet(doc, "pandas（データ処理）")
    add_bullet(doc, "matplotlib（グラフ描画）")
    add_bullet(doc, "numpy（数値計算）")
    add_bullet(doc, "streamlit（Webアプリ）")


def build_step4(doc):
    """ステップ4: Node.js環境をセットアップする"""
    add_page_break(doc)
    doc.add_heading("ステップ4: Node.js環境をセットアップする", level=2)

    add_normal(doc, "PubMed MCPサーバーとdocx変換にNode.jsが必要です。")

    doc.add_heading("4-1. Node.jsのインストール確認", level=3)
    add_code_block(doc, "node --version")
    add_normal(doc, "→ v20以上が表示されればOK")

    add_bold_paragraph(doc, "Mac:")
    add_code_block(doc, "brew install node")
    add_bold_paragraph(doc, "Windows:")
    add_normal(doc, "https://nodejs.org/ から LTS版をダウンロード")

    doc.add_heading("4-2. Node.jsパッケージをインストール", level=3)
    add_code_block(doc, "cd ~/ai-management\nnpm init -y\nnpm install docx")


def build_step5(doc):
    """ステップ5: Claude Codeの設定を復元する"""
    add_page_break(doc)
    doc.add_heading("ステップ5: Claude Codeの設定を復元する", level=2)

    doc.add_heading("5-1. CLAUDE.md（自動）", level=3)
    add_normal(doc,
        "CLAUDE.mdはプロジェクトルートにあり、Claude Codeが自動的に読み込みます。特に操作不要。"
    )

    doc.add_heading("5-2. グローバル設定ファイルを作成", level=3)
    add_normal(doc, "以下をターミナルにコピペ:")
    add_code_block(doc, """mkdir -p ~/.claude

cat > ~/.claude/settings.json << 'SETTINGS'
{
  "env": {
    "CLAUDE_CODE_EXPERIMENTAL_AGENT_TEAMS": "1"
  },
  "teammateMode": "auto",
  "autoApprove": {
    "enabled": true
  }
}
SETTINGS""")

    doc.add_heading("5-3. プロジェクトメモリの復元（推奨）", level=3)
    add_normal(doc, "会話間で保持される記憶を復元します:")
    add_code_block(doc, """mkdir -p ~/.claude/projects/-Users-$(whoami)-ai-management/memory

cat > ~/.claude/projects/-Users-$(whoami)-ai-management/memory/user_profile.md << 'MEMO'
---
name: user_profile
description: 副院長・内科/呼吸器内科医、臨床・教育・経営業務
type: user
---

おもろまちメディカルセンター副院長、内科・呼吸器内科医。
臨床・教育・経営の3軸で業務を行う。
総病床数94床、月間入院数約150件。
MEMO

cat > ~/.claude/projects/-Users-$(whoami)-ai-management/memory/MEMORY.md << 'INDEX'
# Memory Index

## User
- [user_profile.md](user_profile.md) — 副院長・内科/呼吸器内科医、臨床・教育・経営業務
INDEX""")

    doc.add_heading("5-4. カスタムコマンドの復元（任意）", level=3)
    add_code_block(doc, """mkdir -p ~/ai-management/.claude/commands

cat > ~/ai-management/.claude/commands/pico-search.md << 'CMD'
臨床疑問を受け取り、PICO形式に分解してPubMed MCPで文献検索を実行する。
結果はdocs/pubmed/YYYY-MM-DD_topic.md に出力し、PMID・DOIを必ず記録する。
CMD""")


def build_step6(doc):
    """ステップ6: MCP（PubMed検索）を設定する"""
    add_page_break(doc)
    doc.add_heading("ステップ6: MCP（PubMed検索）を設定する", level=2)

    doc.add_heading("6-1. .mcp.json を作成", level=3)
    add_code_block(doc, """cd ~/ai-management

cat > .mcp.json << 'MCP'
{
  "mcpServers": {
    "pubmed-mcp": {
      "command": "npx",
      "args": [
        "-y",
        "@ncukondo/pubmed-mcp",
        "--email",
        "t.kubochan@gmail.com",
        "--cache-dir",
        "./pubmed-cache"
      ]
    }
  }
}
MCP""")

    doc.add_heading("6-2. キャッシュディレクトリを作成", level=3)
    add_code_block(doc, "mkdir -p ~/ai-management/pubmed-cache")


def build_step7(doc):
    """ステップ7: 動作確認"""
    add_page_break(doc)
    doc.add_heading("ステップ7: 動作確認する", level=2)

    doc.add_heading("7-1. シミュレーターの起動テスト", level=3)
    add_code_block(doc, "cd ~/ai-management\nsource .venv/bin/activate\nstreamlit run scripts/bed_control_simulator_app.py")
    add_normal(doc, "→ ブラウザが開きアプリ表示 = OK")
    add_normal(doc, "→ 停止は Ctrl+C")

    doc.add_heading("7-2. Claude Codeの起動テスト", level=3)
    add_bold_paragraph(doc, "デスクトップ版:")
    add_normal(doc, "1. Claudeアプリを開く")
    add_normal(doc, "2. Claude Code画面で ~/ai-management を開く")
    add_normal(doc, '3.「今の環境を確認して」と入力')
    add_normal(doc, "4. CLAUDE.mdの内容が認識されればOK")

    add_bold_paragraph(doc, "CLI版:")
    add_code_block(doc, "cd ~/ai-management\nclaude")
    add_normal(doc, '→「今の環境を確認して」と入力')

    doc.add_heading("7-3. PubMed MCPの動作テスト", level=3)
    add_normal(doc, 'Claude Codeで:「肺炎の最新のガイドラインをPubMedで検索して」')
    add_normal(doc, "→ 検索結果が返ればOK")


def build_troubleshooting(doc):
    """トラブルシューティング"""
    add_page_break(doc)
    doc.add_heading("トラブルシューティング", level=1)

    qa_items = [
        (
            'Q1:「command not found: claude」と出る',
            "→ デスクトップ版: アプリ設定でClaude Codeを有効化\n"
            "→ CLI版: npm install -g @anthropic-ai/claude-code を実行"
        ),
        (
            'Q2:「ModuleNotFoundError: No module named \'streamlit\'」',
            "→ 仮想環境が有効化されていない\n"
            "→ source .venv/bin/activate を実行"
        ),
        (
            'Q3:「pip: externally-managed-environment」',
            "→ venvの外でpipを実行している\n"
            "→ python3 -m venv .venv でvenvを作成→有効化"
        ),
        (
            "Q4: グラフの日本語が□□□になる",
            "→ Mac: brew install font-noto-sans-cjk-jp\n"
            "→ Streamlit Cloud: packages.txt で対応済み"
        ),
        (
            "Q5: CLAUDE.mdが認識されない",
            "→ ai-managementフォルダ内で起動しているか確認"
        ),
        (
            "Q6: PubMed MCPが動かない",
            "→ .mcp.json がプロジェクトルートにあるか確認\n"
            "→ node --version でNode.jsがあるか確認"
        ),
    ]

    for question, answer in qa_items:
        add_bold_paragraph(doc, question)
        for line in answer.split("\n"):
            add_normal(doc, line)
        doc.add_paragraph()  # spacing


def build_appendix_streamlit(doc):
    """補足: Streamlit Cloud"""
    add_page_break(doc)
    doc.add_heading("補足: Streamlit Cloudについて", level=1)

    add_normal(doc, "Streamlit CloudはGitHubに紐づいているため、PC移行の影響なし。")

    add_bold_paragraph(doc, "現在のデプロイ設定:")
    add_bullet(doc, "リポジトリ: torukubota2023/bed-control-simulator")
    add_bullet(doc, "ブランチ: main")
    add_bullet(doc, "メインファイル: scripts/bed_control_simulator_app.py")
    add_bullet(doc, "Python: 3.12")

    add_normal(doc, "新PCからgit pushすれば自動反映されます。")


def build_appendix_file_copy(doc):
    """補足: 直接ファイルコピーで移行する方法"""
    add_page_break(doc)
    doc.add_heading("補足: 直接ファイルコピーで移行する方法（上級者向け）", level=1)

    add_normal(doc, "GitHubに含まれないファイルをUSBメモリ等で転送:")

    add_bold_paragraph(doc, "コピーするファイル:")
    add_bullet(doc, "~/.claude/settings.json（グローバル設定）")
    add_bullet(doc, "~/.claude/settings.local.json（権限設定）")
    add_bullet(doc, "~/ai-management/.mcp.json（MCP設定）")
    add_bullet(doc, "~/ai-management/.claude/commands/（カスタムコマンド）")
    add_bullet(doc, "~/.claude/projects/…/memory/（メモリ）")

    p = doc.add_paragraph()
    run = p.add_run("⚠ .credentials.json は絶対にコピーしない（新PCで再ログイン）")
    run.bold = True
    run.font.color.rgb = RGBColor(0xCC, 0x00, 0x00)
    run.font.name = "Yu Gothic"
    run.font.size = Pt(10.5)
    run.element.rPr.rFonts.set(qn("w:eastAsia"), "Yu Gothic")


def build_closing(doc):
    """おわりに"""
    add_page_break(doc)
    doc.add_heading("おわりに", level=1)

    add_normal(doc,
        "以上で移行完了です。Claude CodeはCLAUDE.mdを自動で読み込むため、"
        "新PCでも同じルールで動作します。"
    )
    add_normal(doc,
        "GitHubのコード・ドキュメントはすべて新PCで利用可能。"
        "メモリは手動復元が必要ですが、使い続けると自動蓄積されます。"
    )


def main():
    doc = Document()
    setup_styles(doc)
    set_margins(doc)

    # Title page
    build_title_page(doc)
    add_page_break(doc)

    # Introduction
    build_introduction(doc)

    # Table of contents
    build_toc(doc)

    # Steps 1-7
    build_step1(doc)
    build_step2(doc)
    build_step3(doc)
    build_step4(doc)
    build_step5(doc)
    build_step6(doc)
    build_step7(doc)

    # Troubleshooting
    build_troubleshooting(doc)

    # Appendices
    build_appendix_streamlit(doc)
    build_appendix_file_copy(doc)

    # Closing
    build_closing(doc)

    # Save
    os.makedirs(os.path.dirname(OUTPUT_PATH), exist_ok=True)
    doc.save(OUTPUT_PATH)
    print(f"生成完了: {OUTPUT_PATH}")


if __name__ == "__main__":
    main()
