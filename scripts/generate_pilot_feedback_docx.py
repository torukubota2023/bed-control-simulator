"""試験運用フィードバックアンケート DOCX を生成.

副院長要望 (2026-04-27): docs/admin/pilot_feedback_survey_2026.md の内容を
紙印刷可能な DOCX に変換、Word MCP で PDF 化して院内 LAN 配布する。

出力:
    docs/admin/試験運用フィードバックアンケート_2026.docx

CLI:
    .venv/bin/python scripts/generate_pilot_feedback_docx.py
"""

from __future__ import annotations

from io import BytesIO
from pathlib import Path

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm, Pt


OUT_DIR = Path(__file__).resolve().parent.parent / "docs" / "admin"
DOCX_PATH = OUT_DIR / "試験運用フィードバックアンケート_2026.docx"


def _add_scale(doc: Document, labels: list[str]) -> None:
    """リッカートスケール行をテーブルで追加."""
    table = doc.add_table(rows=2, cols=len(labels))
    table.style = "Light Grid"
    for i, label in enumerate(labels):
        cell = table.rows[0].cells[i]
        cell.text = label
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in p.runs:
                run.font.size = Pt(10)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    for i in range(len(labels)):
        cell = table.rows[1].cells[i]
        cell.text = "□"
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in p.runs:
                run.font.size = Pt(20)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def _add_free_text(doc: Document, lines: int = 3) -> None:
    """自由記述用の罫線 (下線アンダースコア) を lines 行追加."""
    for _ in range(lines):
        p = doc.add_paragraph("_" * 60)
        p.paragraph_format.space_after = Pt(6)


def _add_checkboxes(doc: Document, options: list[str]) -> None:
    """チェックボックス選択肢を 1 行に追加."""
    p = doc.add_paragraph()
    for i, opt in enumerate(options):
        if i > 0:
            p.add_run("  ")
        p.add_run("□ " + opt)


def build() -> bytes:
    doc = Document()

    # 既定スタイル
    style = doc.styles["Normal"]
    style.font.name = "Yu Gothic"
    style.font.size = Pt(11)

    # ページ余白 (A4 縦)
    for section in doc.sections:
        section.top_margin = Cm(1.8)
        section.bottom_margin = Cm(1.8)
        section.left_margin = Cm(1.8)
        section.right_margin = Cm(1.8)

    # タイトル
    h = doc.add_heading("🩺 ベッドコントロールアプリ + 看護必要度マニュアル 試験運用フィードバック", level=0)

    # メタ情報
    p = doc.add_paragraph()
    p.add_run("実施期間: ").bold = True
    p.add_run("2026-______ 〜 2026-______ （1 週間）")
    p = doc.add_paragraph()
    p.add_run("所要時間: ").bold = True
    p.add_run("約 5 分")
    p = doc.add_paragraph()
    p.add_run("回収方法: ").bold = True
    p.add_run("ナースステーション ⃝⃝ ボックス、または副院長 (久保田) へ直接")
    p = doc.add_paragraph()
    p.add_run("目的: ").bold = True
    p.add_run("2026-06-01 の地域包括医療病棟入院料1 新基準達成に向け、現場で本当に役立つ形に改善するため")

    # 回答者プロフィール
    doc.add_heading("回答者プロフィール（任意・匿名可）", level=1)
    _add_checkboxes(doc, ["医師", "看護師", "看護師長", "その他（　　　　　）"])
    p = doc.add_paragraph()
    p.add_run("経験年数: ").bold = True
    _add_checkboxes(doc, ["5年未満", "5–10年", "10–20年", "20年以上"])
    p = doc.add_paragraph()
    p.add_run("所属: ").bold = True
    _add_checkboxes(doc, ["5F", "6F", "両方", "その他"])

    # Q1
    doc.add_heading("⭐ Q1. 「今日の一手」が分かったか", level=1)
    doc.add_paragraph(
        "アプリ各セクション最上段の「今日あと何をすればいいか」カードを見て、"
        "自分が今日すべきことが明確に分かりましたか?"
    )
    _add_scale(doc, ["1: まったく分からない", "2", "3: 普通", "4", "5: とても明確"])
    p = doc.add_paragraph()
    p.add_run("分かりにくかった部分があれば具体的に:").bold = True
    _add_free_text(doc, 3)

    # Q2
    doc.add_heading("⭐ Q2. 「数字の意味」が分かったか", level=1)
    doc.add_paragraph(
        "ヘッダーの稼働率・空床・在院患者数、6F 必要度Ⅱ ギャップ、救急患者応需係数 などの数字が、"
        "何を示しているか理解できましたか?"
    )
    _add_scale(doc, ["1: まったく分からない", "2", "3: 普通", "4", "5: とても明確"])
    p = doc.add_paragraph()
    p.add_run("意味が分からなかった指標があれば:").bold = True
    _add_free_text(doc, 3)

    # Q3 (重要・逆指標)
    doc.add_heading("⚠️ Q3. 「危ない処置を促された」と感じたか", level=1)
    p = doc.add_paragraph()
    p.add_run(
        "アプリやマニュアルが、医学的に適応外の処置をするように促していると感じた場面はありましたか?"
    )
    p = doc.add_paragraph()
    p.add_run("（これは「ない」が望ましい指標です。少しでも気になった点を率直にお書きください）").italic = True
    _add_scale(doc, ["1: 全く感じなかった", "2", "3: 少し気になった", "4", "5: 強く感じた・警戒"])
    p = doc.add_paragraph()
    run = p.add_run("「気になった」「強く感じた」場合、どの画面・どの記述でしたか?:")
    run.bold = True
    _add_free_text(doc, 5)

    doc.add_page_break()

    # Q4
    doc.add_heading("Q4. 「疾患別マニュアル (17 疾患)」を使いましたか?", level=1)
    p = doc.add_paragraph()
    p.add_run("□ 使った（何回くらい:　　　回）　□ 使っていない　□ 知らなかった")

    p = doc.add_paragraph()
    p.add_run("どの疾患で使いましたか? (複数回答可):").bold = True
    diseases_lines = [
        "□ 急性膵炎　□ 急性胆管炎　□ 消化管出血　□ 早期消化管腫瘍　□ 嚥下障害（PEG）",
        "□ 複雑性腹腔内感染　□ 重症肺炎　□ 気胸/膿胸　□ 高リスク PE　□ 重症 Af",
        "□ 心不全急性増悪　□ 敗血症　□ FN　□ 重症腎盂腎炎　□ 蜂窩織炎",
        "□ 糖尿病性足病変　□ PHN/CRPS",
    ]
    for line in diseases_lines:
        doc.add_paragraph(line)

    p = doc.add_paragraph()
    p.add_run("役立った場面・改善希望:").bold = True
    _add_free_text(doc, 3)

    # Q5
    doc.add_heading("Q5. 「DOCX/PDF 配布版」を使いましたか?", level=1)
    doc.add_paragraph("□ 使った（紙で印刷・タブレット閲覧）　□ 使っていない　□ 存在を知らなかった")

    # Q6
    doc.add_heading("Q6. このアプリ・マニュアルを「続けて使いたい」ですか?", level=1)
    doc.add_paragraph("□ はい　□ いいえ　□ どちらでもない")
    p = doc.add_paragraph()
    p.add_run("理由 (任意):").bold = True
    _add_free_text(doc, 2)

    # Q7
    doc.add_heading("Q7. 自由記述：改善してほしい点・現場で困った点", level=1)
    _add_free_text(doc, 6)

    # 集計後のアクション
    doc.add_heading("集計後のアクション (副院長より)", level=1)
    doc.add_paragraph("集計結果は以下の判断材料に使用します:")
    actions = [
        "Q1・Q2 が低い (中央値 ≤ 3) → UI 文言・指標説明の改善 PR",
        "Q3 が「3 以上」が 1 名でもいる → 該当箇所を緊急修正、必要なら表示中断",
        "Q4 で使用率 < 30% → マニュアルの導線・告知方法の見直し",
        "Q7 の自由記述 → 試験運用 2 週目の改善計画に反映",
    ]
    for a in actions:
        doc.add_paragraph(a, style="List Number")
    doc.add_paragraph(
        "集計結果は 2 週目開始前 (2026-______) に病棟会議で共有予定。"
    )

    # 提出について
    doc.add_heading("提出について", level=1)
    doc.add_paragraph("提出締切: 2026-______ （1 週間後）")
    p = doc.add_paragraph()
    p.add_run("提出方法:").bold = True
    submissions = [
        "紙: ナースステーション ⃝⃝ ボックス",
        "直接: 副院長 久保田 (内線 ____)",
        "電子: 副院長メール ____ または院内メッセンジャー",
    ]
    for s in submissions:
        doc.add_paragraph(s, style="List Bullet")

    # 締め
    doc.add_paragraph("")
    p = doc.add_paragraph(
        "ご協力ありがとうございます。皆さまの声が、6/1 新基準クリアと安全な運用の両立に直結します。"
    )
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p = doc.add_paragraph("— 副院長 久保田 徹")
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    buf = BytesIO()
    doc.save(buf)
    return buf.getvalue()


def main() -> int:
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    DOCX_PATH.write_bytes(build())
    print(f"✅ DOCX saved: {DOCX_PATH}")
    print(f"   bytes: {DOCX_PATH.stat().st_size:,}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
