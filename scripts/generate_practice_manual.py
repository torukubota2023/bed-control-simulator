#!/usr/bin/env python3
"""
ベッドコントロールシミュレーター 実践マニュアル生成スクリプト

python-docx を使用して、超詳細な実践マニュアルを Word 文書として生成する。
出力先: docs/admin/bed_control_simulator_practice_manual.docx
"""

from __future__ import annotations

import os
from pathlib import Path

from docx import Document
from docx.shared import Pt, Cm, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

# ---------------------------------------------------------------------------
# パス設定
# ---------------------------------------------------------------------------
PROJECT_ROOT = Path(__file__).resolve().parent.parent
OUTPUT_PATH = PROJECT_ROOT / "docs" / "admin" / "bed_control_simulator_practice_manual.docx"

# ---------------------------------------------------------------------------
# スタイルヘルパー
# ---------------------------------------------------------------------------

def set_cell_shading(cell, color_hex: str):
    """テーブルセルの背景色を設定する。"""
    shading_elm = parse_xml(
        f'<w:shd {nsdecls("w")} w:fill="{color_hex}" w:val="clear"/>'
    )
    cell._tc.get_or_add_tcPr().append(shading_elm)


def add_page_number(doc):
    """フッターにページ番号を追加する。"""
    section = doc.sections[0]
    footer = section.footer
    footer.is_linked_to_previous = False
    paragraph = footer.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    run = paragraph.add_run()
    fldChar1 = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="begin"/>')
    run._r.append(fldChar1)

    run2 = paragraph.add_run()
    instrText = parse_xml(f'<w:instrText {nsdecls("w")} xml:space="preserve"> PAGE </w:instrText>')
    run2._r.append(instrText)

    run3 = paragraph.add_run()
    fldChar2 = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="end"/>')
    run3._r.append(fldChar2)


def set_table_border(table):
    """テーブルに罫線を設定する。"""
    tbl = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else parse_xml(f'<w:tblPr {nsdecls("w")}/>')
    borders = parse_xml(
        f'<w:tblBorders {nsdecls("w")}>'
        '  <w:top w:val="single" w:sz="4" w:space="0" w:color="999999"/>'
        '  <w:left w:val="single" w:sz="4" w:space="0" w:color="999999"/>'
        '  <w:bottom w:val="single" w:sz="4" w:space="0" w:color="999999"/>'
        '  <w:right w:val="single" w:sz="4" w:space="0" w:color="999999"/>'
        '  <w:insideH w:val="single" w:sz="4" w:space="0" w:color="999999"/>'
        '  <w:insideV w:val="single" w:sz="4" w:space="0" w:color="999999"/>'
        '</w:tblBorders>'
    )
    tblPr.append(borders)


def add_styled_table(doc, headers: list[str], rows: list[list[str]], header_color: str = "2C3E50"):
    """ヘッダー行に背景色を設定したテーブルを追加する。"""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_border(table)

    # ヘッダー行
    header_row = table.rows[0]
    for i, header_text in enumerate(headers):
        cell = header_row.cells[i]
        cell.text = header_text
        set_cell_shading(cell, header_color)
        for paragraph in cell.paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in paragraph.runs:
                run.font.bold = True
                run.font.size = Pt(9.5)
                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

    # データ行
    for row_idx, row_data in enumerate(rows):
        row = table.rows[row_idx + 1]
        for col_idx, cell_text in enumerate(row_data):
            cell = row.cells[col_idx]
            cell.text = cell_text
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(9.5)
                # 数値列は右寄せ
                if col_idx > 0 and any(c.isdigit() for c in cell_text.replace(",", "").replace(".", "").replace("-", "").replace("%", "").replace("円", "").replace("〜", "").replace("日", "")):
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    return table


def add_body(doc, text: str, bold: bool = False, font_name: str = None):
    """本文段落を追加する。"""
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(10.5)
    run.bold = bold
    if font_name:
        run.font.name = font_name
    return p


def add_body_list(doc, items: list[str], bold_prefix: bool = False):
    """箇条書きリストを追加する。"""
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        if bold_prefix and "：" in item:
            parts = item.split("：", 1)
            run_bold = p.add_run(parts[0] + "：")
            run_bold.bold = True
            run_bold.font.size = Pt(10.5)
            run_normal = p.add_run(parts[1])
            run_normal.font.size = Pt(10.5)
        else:
            # clear default text and re-add
            p.clear()
            run = p.add_run(item)
            run.font.size = Pt(10.5)


def add_code_block(doc, code: str):
    """コードブロックを追加する（Courier New フォント、グレー背景）。"""
    p = doc.add_paragraph()
    # 段落にグレー背景を設定
    pPr = p._p.get_or_add_pPr()
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="F0F0F0" w:val="clear"/>')
    pPr.append(shading)
    run = p.add_run(code)
    run.font.name = "Courier New"
    run.font.size = Pt(9)
    return p


def add_indent_block(doc, text: str, indent_cm: float = 1.0):
    """インデント付き段落を追加する。"""
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(indent_cm)
    run = p.add_run(text)
    run.font.size = Pt(10.5)
    return p


def add_page_break(doc):
    """改ページを追加する。"""
    doc.add_page_break()


# ---------------------------------------------------------------------------
# メイン生成関数
# ---------------------------------------------------------------------------

def generate_manual():
    """マニュアル全体を生成する。"""
    doc = Document()

    # ページ設定
    section = doc.sections[0]
    section.page_height = Cm(29.7)
    section.page_width = Cm(21.0)
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

    # デフォルトフォントスタイル設定
    style = doc.styles["Normal"]
    font = style.font
    font.size = Pt(10.5)

    # ページ番号追加
    add_page_number(doc)

    # ================================================================
    # 表紙
    # ================================================================
    for _ in range(6):
        doc.add_paragraph()

    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title_p.add_run("ベッドコントロールシミュレーター")
    run.font.size = Pt(28)
    run.bold = True
    run.font.color.rgb = RGBColor(0x2C, 0x3E, 0x50)

    title_p2 = doc.add_paragraph()
    title_p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run2 = title_p2.add_run("実践マニュアル")
    run2.font.size = Pt(28)
    run2.bold = True
    run2.font.color.rgb = RGBColor(0x2C, 0x3E, 0x50)

    doc.add_paragraph()

    subtitle_p = doc.add_paragraph()
    subtitle_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run3 = subtitle_p.add_run("地域包括医療病棟の病棟運営の最適化のために")
    run3.font.size = Pt(16)
    run3.font.color.rgb = RGBColor(0x7F, 0x8C, 0x8D)

    for _ in range(4):
        doc.add_paragraph()

    org_p = doc.add_paragraph()
    org_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run4 = org_p.add_run("おもろまちメディカルセンター")
    run4.font.size = Pt(14)

    date_p = doc.add_paragraph()
    date_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run5 = date_p.add_run("2026年3月25日")
    run5.font.size = Pt(12)
    run5.font.color.rgb = RGBColor(0x7F, 0x8C, 0x8D)

    add_page_break(doc)

    # ================================================================
    # 目次
    # ================================================================
    doc.add_heading("目次", level=1)
    toc_items = [
        "第1章　シミュレーターとは何か",
        "　1-1. このツールの目的",
        "　1-2. なぜベッドコントロールが重要なのか",
        "　1-3. A/B/Cフェーズとは",
        "　1-4. 包括払いの前提",
        "　1-5. このシミュレーターでできること",
        "　1-6. このシミュレーターの限界",
        "第2章　3つの戦略を理解する",
        "　2-1. 回転重視戦略（rotation）",
        "　2-2. 安定維持戦略（stable）",
        "　2-3. バランス戦略（balanced）",
        "　2-4. どの戦略を選ぶべきか",
        "第3章　画面の見方（Streamlit版）",
        "　3-1. 全体構成",
        "　3-2. サイドバーの入力項目",
        "　3-3. メインエリアのタブ",
        "第4章　実践シナリオ",
        "　シナリオ1〜5",
        "第5章　結果の読み方と病棟運営への活用",
        "　5-1. 最も重要な指標",
        "　5-2. グラフの読み方のコツ",
        "　5-3. 運営改善アラートの活用",
        "　5-4. 月次レポートへの活用例",
        "第6章　よくある質問",
        "第7章　用語集",
        "付録A　CLI版の使い方",
        "付録B　パラメータ一覧表",
    ]
    for item in toc_items:
        p = doc.add_paragraph()
        run = p.add_run(item)
        run.font.size = Pt(10.5)

    add_page_break(doc)

    # ================================================================
    # 第1章：シミュレーターとは何か
    # ================================================================
    doc.add_heading("第1章　シミュレーターとは何か", level=1)

    # 1-1
    doc.add_heading("1-1. このツールの目的", level=2)
    add_body(doc,
        "地域包括医療病棟（94床）において、病床稼働率90〜95%を維持しながら、"
        "収益とコストのバランスを最適化するための病棟運営支援ツールです。")
    add_body(doc,
        "病棟運営において「どのタイミングで退院を促進すべきか」「新規入院をどこまで受け入れるべきか」"
        "という日々の判断を、データに基づいてシミュレーションすることで、"
        "より精度の高い病棟運営判断を支援します。")

    # 1-2
    doc.add_heading("1-2. なぜベッドコントロールが重要なのか", level=2)
    add_body(doc, "病床管理は病院経営の根幹です。以下の問題が常に発生します。")
    add_body_list(doc, [
        "稼働率が低い → 収益減（空床はゼロ収益だがコストは発生）",
        "稼働率が高すぎる → 新規入院を断る未活用病床コスト、スタッフ負担増",
        "「ただ埋める」「ただ長く入院させる」では最適化できない",
        "患者の入院フェーズごとの収益性が異なるため、フェーズ構成比のコントロールが鍵",
    ])
    add_body(doc,
        "つまり、単純に「満床にする」ことが最適解ではなく、"
        "入院患者の構成（どのフェーズの患者がどれだけいるか）を意識した管理が必要です。", bold=True)

    # 1-3
    doc.add_heading("1-3. A/B/Cフェーズとは", level=2)
    add_body(doc,
        "このシミュレーターの核となる概念です。患者を入院日数で3グループに分類します。")

    # A群
    doc.add_heading("A群（入院1〜5日目）＝入院初期", level=3)
    add_styled_table(doc,
        ["項目", "値"],
        [
            ["日次診療報酬", "36,000円"],
            ["日次変動費", "12,000円"],
            ["日次運営貢献額", "24,000円"],
        ],
        header_color="E74C3C")
    add_body(doc, "")
    add_body_list(doc, [
        "特徴：検査・初期治療で変動費が高く、運営貢献額は3群中最小",
        "運営面での意味：必要なフェーズだが、A群ばかりだと変動費増で運営貢献額が相対的に低下",
    ])

    # B群
    doc.add_heading("B群（入院6〜14日目）＝回復期", level=3)
    add_styled_table(doc,
        ["項目", "値"],
        [
            ["日次診療報酬", "36,000円"],
            ["日次変動費", "6,000円"],
            ["日次運営貢献額", "30,000円"],
        ],
        header_color="27AE60")
    add_body(doc, "")
    add_body_list(doc, [
        "特徴：治療が安定し、変動費が大幅に下がるが収益は維持。最も運営貢献額を生むフェーズ",
        "運営面での意味：B群を厚くすることが運営貢献額の最大化の鍵",
    ])

    # C群
    doc.add_heading("C群（入院15日目以降）＝安定期", level=3)
    add_styled_table(doc,
        ["項目", "値"],
        [
            ["日次診療報酬", "33,400円"],
            ["日次変動費", "4,500円"],
            ["日次運営貢献額", "28,900円"],
        ],
        header_color="2980B9")
    add_body(doc, "")
    add_body_list(doc, [
        "特徴：変動費最小で運営貢献額は高い。ただし長期化すると新規入院を阻害（未活用病床コスト）",
        "運営面での意味：需給の調整弁。閑散期は長めに保持、繁忙期は早めに退院促進",
    ])

    # フェーズ比較表
    doc.add_heading("フェーズ比較一覧", level=3)
    add_styled_table(doc,
        ["フェーズ", "入院日数", "日次診療報酬", "日次変動費", "日次運営貢献額", "運営面での役割"],
        [
            ["A群", "1〜5日目", "36,000円", "12,000円", "24,000円", "初期投資（必要コスト）"],
            ["B群", "6〜14日目", "36,000円", "6,000円", "30,000円", "運営貢献の源泉（最重要）"],
            ["C群", "15日目〜", "33,400円", "4,500円", "28,900円", "退院調整の柔軟性が高い層"],
        ])

    # 1-4
    doc.add_heading("1-4. 包括払いの前提", level=2)
    add_body(doc,
        "地域包括医療病棟は包括払い制度です。日次診療報酬は入院14日目まで約36,000円/日、15日目以降は約33,400円/日です。")
    add_body(doc,
        "変動費は初期（検査・処置）に集中するため、入院後半ほど運営貢献額が大きくなります。"
        "A群（初期）は変動費が高く運営貢献額は3群中最小ですが、B群・C群になるにつれて運営貢献額が拡大します。")
    add_body(doc,
        "ただし、C群を長期保持すると新規入院の未活用病床コストが発生するため、"
        "「いつ退院させるか」が病棟運営判断の核心となります。", bold=True)

    # 1-5
    doc.add_heading("1-5. このシミュレーターでできること", level=2)
    add_body_list(doc, [
        "3つの戦略（回転重視・安定維持・バランス）で病棟運営をシミュレーション",
        "30日間の日次推移を可視化（稼働率・在院患者数・入退院数・運営貢献額）",
        "月次の収益・変動費・運営貢献額を算出",
        "A/B/Cフェーズ構成比の時系列変化を確認",
        "運営改善アラート（稼働率低下、A群過多、C群滞留など）を自動検知",
        "3戦略の一括比較で最適戦略を特定",
        "パラメータを変更して「もし○○だったら」のシナリオ分析",
    ])

    # 1-6
    doc.add_heading("1-6. このシミュレーターの限界（知っておくべきこと）", level=2)
    add_body(doc, "本ツールは病棟運営判断の参考情報であり、以下の限界があります。", bold=True)
    add_body_list(doc, [
        "診療報酬の厳密な計算ではない（病棟運営判断のための簡易モデル）",
        "個別患者の重症度・疾患は考慮していない",
        "看護必要度の算定への影響は未反映",
        "緊急入院と予定入院の区別なし",
        "乱数を使用しているため、毎回微妙に異なる結果が出る（シード値を固定すれば再現可能）",
    ])

    add_page_break(doc)

    # ================================================================
    # 第2章：3つの戦略を理解する
    # ================================================================
    doc.add_heading("第2章　3つの戦略を理解する", level=1)
    add_body(doc,
        "シミュレーターには3つの戦略が実装されています。"
        "それぞれ退院確率の調整方法と新規入院の制御ロジックが異なります。")

    # 2-1 回転重視戦略
    doc.add_heading("2-1. 回転重視戦略（rotation）", level=2)
    add_body(doc, "コンセプト：「早く退院 → 早く入院 → B群を厚くする」", bold=True)

    doc.add_heading("戦略の特徴", level=3)
    add_body_list(doc, [
        "C群（15日目以降）の退院確率を2.2倍に引き上げ、積極的に退院を促進",
        "B群後半（12日目以降）も退院確率を1.6倍に引き上げ",
        "B群中盤（10日目以降）は退院確率を1.2倍に微増",
        "空床があれば積極的に新規入院を受け入れ（稼働率97%まで）",
        "稼働率85%未満では退院確率を0.3倍に抑制して下限を守る",
        "稼働率85〜87%では退院確率を0.5倍に抑制",
    ])

    doc.add_heading("向いている状況", level=3)
    add_body_list(doc, [
        "入院需要が旺盛な繁忙期",
        "新規入院の待ちが多い時期",
        "B群比率を高めて運営貢献構成を改善したい場合",
    ])

    doc.add_heading("注意点", level=3)
    add_body_list(doc, [
        "平均在院日数が短くなる（13〜15日程度）",
        "A群比率が高くなりやすい（35%前後） → 初期コスト増",
        "稼働率が不安定になりやすい",
    ])

    doc.add_heading("シミュレーション結果の目安（初期パラメータ）", level=3)
    add_styled_table(doc,
        ["指標", "目安値"],
        [
            ["平均稼働率", "約90%"],
            ["月次運営貢献額", "約3,062万円"],
            ["B群比率", "約49%（最高）"],
            ["目標レンジ日数", "9/30日"],
        ])

    add_page_break(doc)

    # 2-2 安定維持戦略
    doc.add_heading("2-2. 安定維持戦略（stable）", level=2)
    add_body(doc, "コンセプト：「C群を長めに保持して稼働率を安定させる」", bold=True)

    doc.add_heading("戦略の特徴", level=3)
    add_body_list(doc, [
        "C群の退院確率を0.6倍に低く設定し、退院を抑制",
        "稼働率90%未満ではさらに全体の退院確率を0.4倍に抑制",
        "稼働率85%未満ではさらに0.4倍を追加適用（0.6 × 0.4 × 0.4 = 約0.1倍）",
        "新規入院は稼働率93%で抑制開始（95%までに制限）",
        "稼働率96%超ではC群の退院確率を1.5倍に引き上げ（過密防止）",
    ])

    doc.add_heading("向いている状況", level=3)
    add_body_list(doc, [
        "入院需要が少ない閑散期",
        "稼働率の安定を最優先したい場合",
        "空床を出したくない時期",
    ])

    doc.add_heading("注意点", level=3)
    add_body_list(doc, [
        "C群比率が高くなる（26%前後） → 空床リスク",
        "新規入院を抑制しすぎるとB群が不足する可能性",
    ])

    doc.add_heading("シミュレーション結果の目安（初期パラメータ）", level=3)
    add_styled_table(doc,
        ["指標", "目安値"],
        [
            ["平均稼働率", "約94.5%（最高）"],
            ["月次運営貢献額", "約3,445万円（最高）"],
            ["C群比率", "約27%（最高）"],
            ["目標レンジ日数", "16/30日"],
        ])

    add_page_break(doc)

    # 2-3 バランス戦略
    doc.add_heading("2-3. バランス戦略（balanced）", level=2)
    add_body(doc, "コンセプト：「稼働率90〜95%を最優先、構成比も調整」", bold=True)

    doc.add_heading("戦略の特徴", level=3)
    add_body(doc, "稼働率ゾーン別の自動制御を行います。", bold=True)
    add_styled_table(doc,
        ["稼働率ゾーン", "退院制御", "入院制御"],
        [
            ["90%未満", "退院を強く抑制（0.3倍）", "空床全てに入院受入"],
            ["90〜95%", "通常運用（構成比微調整のみ）", "目標上限までの入院受入（A群35%超で抑制）"],
            ["95%超", "C群退院促進（1.8倍）、B群も1.2倍", "95%までに制限"],
        ])
    add_body(doc, "")
    add_body_list(doc, [
        "B群構成比が25%未満の場合、B群の退院確率を0.7倍に抑制",
        "C群構成比が35%超の場合、C群の退院確率を1.3倍に引き上げ",
        "A群比率35%超で新規入院を半減（空床の半分のみ受入）",
    ])

    doc.add_heading("向いている状況", level=3)
    add_body_list(doc, [
        "通常月の標準運用",
        "稼働率の安定と運営貢献額のバランスを取りたい場合",
        "実務上のデフォルト戦略",
    ])

    doc.add_heading("シミュレーション結果の目安（初期パラメータ）", level=3)
    add_styled_table(doc,
        ["指標", "目安値"],
        [
            ["平均稼働率", "約91.2%"],
            ["月次運営貢献額", "約3,360万円"],
            ["目標レンジ日数", "17/30日（最多＝最も安定）"],
        ])

    # 戦略比較一覧表
    doc.add_heading("3戦略の比較一覧", level=3)
    add_styled_table(doc,
        ["指標", "回転重視", "安定維持", "バランス"],
        [
            ["平均稼働率", "約90%", "約94.5%", "約91.2%"],
            ["月次運営貢献額", "約3,062万円", "約3,445万円", "約3,360万円"],
            ["B群比率", "約49%", "約33%", "約40%"],
            ["C群比率", "約18%", "約27%", "約23%"],
            ["目標レンジ日数", "9/30日", "16/30日", "17/30日"],
            ["安定性", "低い", "中程度", "高い"],
            ["推奨時期", "繁忙期", "閑散期", "通常月"],
        ])

    add_page_break(doc)

    # 2-4 判断フローチャート
    doc.add_heading("2-4. どの戦略を選ぶべきか（判断フロー）", level=2)
    add_body(doc, "以下の判断フローに従って、今月の最適戦略を選択してください。", bold=True)

    add_body(doc, "")
    add_body(doc, "Q1: 今月の入院需要は？", bold=True)
    add_indent_block(doc, "→ 多い（繁忙期） → 回転重視戦略")
    add_indent_block(doc, "→ 少ない（閑散期） → 安定維持戦略")
    add_indent_block(doc, "→ 通常 → Q2へ")

    add_body(doc, "")
    add_body(doc, "Q2: 現在の稼働率は？", bold=True)
    add_indent_block(doc, "→ 85%未満 → 安定維持戦略")
    add_indent_block(doc, "→ 85〜90% → バランス戦略")
    add_indent_block(doc, "→ 90〜95% → バランス戦略")
    add_indent_block(doc, "→ 95%超 → 回転重視戦略")

    add_body(doc, "")
    add_body(doc, "Q3: 最も重視する指標は？", bold=True)
    add_indent_block(doc, "→ 運営貢献額の最大化 → 安定維持戦略")
    add_indent_block(doc, "→ 稼働率安定 → バランス戦略")
    add_indent_block(doc, "→ B群比率最大 → 回転重視戦略")

    add_page_break(doc)

    # ================================================================
    # 第3章：画面の見方（Streamlit版）
    # ================================================================
    doc.add_heading("第3章　画面の見方（Streamlit版）", level=1)

    # 3-1
    doc.add_heading("3-1. 全体構成", level=2)
    add_body(doc,
        "画面は左サイドバー（入力エリア）とメインエリア（結果表示）に分かれています。"
        "サイドバーでパラメータを設定し、「シミュレーション実行」ボタンをクリックすると、"
        "メインエリアに結果がタブ形式で表示されます。")

    # 3-2
    doc.add_heading("3-2. サイドバーの入力項目", level=2)

    doc.add_heading("病棟基本条件", level=3)
    add_styled_table(doc,
        ["項目", "初期値", "説明"],
        [
            ["病床数", "94", "病棟の総ベッド数。変更することで他病棟のシミュレーションも可能"],
            ["目標稼働率下限", "0.90", "これを下回ると「稼働率低下」フラグが立つ"],
            ["目標稼働率上限", "0.95", "これを超えると「稼働率超過」フラグが立つ"],
            ["月の日数", "30", "シミュレーション期間（28〜31日）"],
            ["月間新規入院数", "150", "1ヶ月の新規入院見込み。当院は約150件/月"],
            ["平均在院日数", "19", "退院タイミングの基準。16〜21日で調整可能"],
            ["退院調整日数", "2", "退院準備に要する日数。この日数以内は退院不可"],
            ["入院流入変動係数", "1.0", "1.0=通常、1.2=繁忙（+20%）、0.8=閑散（-20%）"],
        ])

    doc.add_heading("患者フェーズ別パラメータ（折りたたみ）", level=3)
    add_body(doc, "各群の日次診療報酬・日次コストを個別に設定できます。初期値は当院の概算値です。")
    add_styled_table(doc,
        ["パラメータ", "A群", "B群", "C群"],
        [
            ["日次診療報酬", "36,000円", "36,000円", "33,400円"],
            ["日次変動費", "12,000円", "6,000円", "4,500円"],
        ])

    doc.add_heading("追加パラメータ（折りたたみ）", level=3)
    add_styled_table(doc,
        ["項目", "初期値", "説明"],
        [
            ["初日加算", "0円", "入院初日に加算される収益"],
            ["14日以内加算", "0円/日", "14日以内の入院に加算される収益"],
            ["リハビリ出来高", "0円/日", "リハビリによる追加収益（B群・C群に適用）"],
            ["未活用病床コスト", "10,000円", "満床で入院を断った場合の空床1日あたり逸失診療報酬"],
            ["退院促進閾値", "0.95", "この稼働率を超えるとC群退院確率を1.3倍に増加"],
            ["新規入院抑制閾値", "0.97", "回転重視戦略でこの稼働率を超えると新規入院を抑制"],
            ["乱数シード", "42", "同じ値なら同じ結果を再現できる"],
        ])

    doc.add_heading("戦略選択", level=3)
    add_body_list(doc, [
        "バランス戦略 / 回転重視戦略 / 安定維持戦略 から1つ選択",
        "「全戦略比較」にチェックを入れると、3戦略の比較タブが表示される",
    ])

    add_page_break(doc)

    # 3-3
    doc.add_heading("3-3. メインエリアのタブ", level=2)

    doc.add_heading("タブ1: 日次推移", level=3)
    add_body_list(doc, [
        "稼働率推移グラフ：青い線が日々の稼働率。緑の帯が目標レンジ（90〜95%）。帯に入っている日が多いほど良い",
        "在院患者数推移：日々の在院患者数の変化",
        "新規入院・退院数：緑棒が入院、赤棒が退院。バランスが取れていることが重要",
        "日次運営貢献額：棒グラフで日々の運営貢献額。安定して高いのが理想",
    ])

    doc.add_heading("タブ2: フェーズ構成", level=3)
    add_body_list(doc, [
        "A/B/C積み上げ面グラフ：赤(A群)・緑(B群)・青(C群)の構成比の時系列変化。緑（B群）が厚いほど運営貢献額が良い",
        "平均円グラフ：月間平均の構成比。一目で全体バランスを把握",
        "フェーズ別患者数：各群の実数推移",
    ])

    doc.add_heading("タブ3: 運営分析", level=3)
    add_body_list(doc, [
        "メトリクスカード：月次収益・変動費・運営貢献額・平均稼働率などの主要指標",
        "日次診療報酬・変動費推移：収益と変動費の乖離が大きいほど運営貢献額が大きい",
        "累積運営貢献額推移：右肩上がりの角度が急なほど効率が良い",
    ])

    doc.add_heading("タブ4: 運営改善アラート", level=3)
    add_body_list(doc, [
        "各日のフラグ一覧（赤=警告、緑=正常）",
        "フラグ集計：「稼働率90%未満が何日あったか」等のサマリー",
        "推奨アクション：フラグに基づく具体的なアクション提案",
    ])

    doc.add_heading("タブ5: 戦略比較（全戦略比較チェック時のみ）", level=3)
    add_body_list(doc, [
        "3戦略の主要指標を横並びで比較",
        "棒グラフで視覚的に比較",
        "最適戦略のハイライト",
    ])

    doc.add_heading("タブ6: データ", level=3)
    add_body_list(doc, [
        "日次データの全テーブル表示",
        "CSVダウンロードボタン：Excelで開いて詳細分析可能",
    ])

    add_page_break(doc)

    # ================================================================
    # 第4章：実践シナリオ
    # ================================================================
    doc.add_heading("第4章　実践シナリオ", level=1)
    add_body(doc,
        "以下のシナリオに沿って操作することで、シミュレーターの活用方法を体験できます。")

    # シナリオ1
    doc.add_heading("シナリオ1：通常月の最適戦略を見つける", level=2)
    add_body(doc, "目的：初期パラメータで3戦略を比較し、通常月の最適戦略を特定する", bold=True)
    add_body(doc, "手順：")
    p = doc.add_paragraph(style="List Number")
    p.clear()
    p.add_run("サイドバーのパラメータを初期値のまま保持する").font.size = Pt(10.5)
    p = doc.add_paragraph(style="List Number")
    p.clear()
    p.add_run("「全戦略比較」にチェックを入れる").font.size = Pt(10.5)
    p = doc.add_paragraph(style="List Number")
    p.clear()
    p.add_run("「シミュレーション実行」をクリック").font.size = Pt(10.5)
    p = doc.add_paragraph(style="List Number")
    p.clear()
    p.add_run("タブ5「戦略比較」で3戦略を比較する").font.size = Pt(10.5)
    p = doc.add_paragraph(style="List Number")
    p.clear()
    p.add_run("月次運営貢献額と目標レンジ日数を確認する").font.size = Pt(10.5)
    add_body(doc, "結論：バランス戦略が目標レンジ17日で最も安定しており、通常月のデフォルト戦略として推奨される。", bold=True)

    # シナリオ2
    doc.add_heading("シナリオ2：繁忙月（入院需要+20%）のシミュレーション", level=2)
    add_body(doc, "目的：入院需要が増加した場合に最適な戦略を特定する", bold=True)
    add_body(doc, "手順：")
    p = doc.add_paragraph(style="List Number")
    p.clear()
    p.add_run("入院流入変動係数を1.2に変更する").font.size = Pt(10.5)
    p = doc.add_paragraph(style="List Number")
    p.clear()
    p.add_run("月間新規入院数を180に変更する").font.size = Pt(10.5)
    p = doc.add_paragraph(style="List Number")
    p.clear()
    p.add_run("「全戦略比較」にチェックを入れる").font.size = Pt(10.5)
    p = doc.add_paragraph(style="List Number")
    p.clear()
    p.add_run("実行して比較する").font.size = Pt(10.5)
    add_body(doc, "注目点：回転重視戦略の未活用病床コストが最も少ないか確認する。")
    add_body(doc, "結論：繁忙月は回転重視でB群比率を維持しつつ新規入院を受け入れる。", bold=True)

    # シナリオ3
    doc.add_heading("シナリオ3：閑散月（入院需要-20%）のシミュレーション", level=2)
    add_body(doc, "目的：入院需要が減少した場合に稼働率を維持する戦略を特定する", bold=True)
    add_body(doc, "手順：")
    p = doc.add_paragraph(style="List Number")
    p.clear()
    p.add_run("入院流入変動係数を0.8に変更する").font.size = Pt(10.5)
    p = doc.add_paragraph(style="List Number")
    p.clear()
    p.add_run("月間新規入院数を120に変更する").font.size = Pt(10.5)
    p = doc.add_paragraph(style="List Number")
    p.clear()
    p.add_run("「全戦略比較」にチェックを入れて実行する").font.size = Pt(10.5)
    add_body(doc, "注目点：安定維持戦略の稼働率が90%を維持できているか確認する。")
    add_body(doc, "結論：閑散月は安定維持でC群を長めに保持し稼働率を確保する。", bold=True)

    # シナリオ4
    doc.add_heading("シナリオ4：平均在院日数を変えた場合の影響", level=2)
    add_body(doc, "目的：在院日数の変化がフェーズ構成比にどう影響するか確認する", bold=True)
    add_body(doc, "手順：")
    p = doc.add_paragraph(style="List Number")
    p.clear()
    p.add_run("平均在院日数を16日に変更（短縮） → 実行").font.size = Pt(10.5)
    p = doc.add_paragraph(style="List Number")
    p.clear()
    p.add_run("平均在院日数を21日に変更（延長） → 実行").font.size = Pt(10.5)
    p = doc.add_paragraph(style="List Number")
    p.clear()
    p.add_run("結果を比較する").font.size = Pt(10.5)
    add_body(doc, "観察ポイント：")
    add_body_list(doc, [
        "在院日数短縮 → A群比率増・回転率アップ",
        "在院日数延長 → C群比率増・安定性向上",
    ])
    add_body(doc, "結論：在院日数そのものではなく、フェーズ構成比の変化に注目することが重要。", bold=True)

    # シナリオ5
    doc.add_heading("シナリオ5：初日加算・14日以内加算を設定した場合", level=2)
    add_body(doc, "目的：加算設定が収支に与える影響を確認する", bold=True)
    add_body(doc, "手順：")
    p = doc.add_paragraph(style="List Number")
    p.clear()
    p.add_run("追加パラメータを展開する").font.size = Pt(10.5)
    p = doc.add_paragraph(style="List Number")
    p.clear()
    p.add_run("初日加算を15,000円、14日以内加算を5,000円に設定する").font.size = Pt(10.5)
    p = doc.add_paragraph(style="List Number")
    p.clear()
    p.add_run("実行して運営分析タブを確認する").font.size = Pt(10.5)
    add_body(doc, "注目点：回転重視戦略の運営貢献額が改善されるか確認する（入院初期の収益が増えるため）。")
    add_body(doc, "結論：加算が大きい場合、回転重視戦略の相対的な優位性が高まる。", bold=True)

    add_page_break(doc)

    # ================================================================
    # 第5章：結果の読み方と病棟運営への活用
    # ================================================================
    doc.add_heading("第5章　結果の読み方と病棟運営への活用", level=1)

    # 5-1
    doc.add_heading("5-1. 最も重要な指標", level=2)
    add_body(doc, "シミュレーション結果を確認する際、以下の優先順位で指標を確認してください。", bold=True)
    add_styled_table(doc,
        ["優先順位", "指標", "意味", "目標"],
        [
            ["1", "月次運営貢献額", "運営の最終目標", "3,000万円以上"],
            ["2", "目標レンジ日数", "運営の安定性", "20日以上/30日"],
            ["3", "B群構成比", "運営貢献の源泉", "40〜50%"],
            ["4", "未活用病床コスト", "見えないコスト", "できるだけ少なく"],
        ])

    # 5-2
    doc.add_heading("5-2. グラフの読み方のコツ", level=2)

    doc.add_heading("稼働率グラフ", level=3)
    add_body(doc,
        "緑帯（90〜95%）に入っている日が多いほど良い運営です。"
        "帯から大きく外れる日は要因分析が必要です。"
        "連続して帯を外れる場合は戦略の見直しを検討してください。")

    doc.add_heading("フェーズ構成比", level=3)
    add_body(doc,
        "B群（緑）が40〜50%が理想です。A群（赤）が35%を超えている場合は新規入院の過剰を疑い、"
        "一時的に入院を抑制する判断が必要かもしれません。")

    doc.add_heading("日次運営貢献額", level=3)
    add_body(doc,
        "安定して100万円以上が目標です。大きな落ち込みがあれば、"
        "その日のフェーズ構成比や入退院数を確認して原因を探ります。")

    doc.add_heading("新規入院・退院数", level=3)
    add_body(doc,
        "入退院のバランスが崩れると稼働率が不安定になります。"
        "特に退院が集中する日（月曜日等）のパターンに注意してください。")

    # 5-3
    doc.add_heading("5-3. 運営改善アラートの活用", level=2)
    add_body(doc, "各フラグが5日以上出現した場合のアクション指針：", bold=True)
    add_styled_table(doc,
        ["フラグ", "条件", "推奨アクション"],
        [
            ["稼働率低下", "稼働率90%未満", "退院抑制・入院促進を検討"],
            ["稼働率超過", "稼働率95%超", "退院促進・入院抑制を検討"],
            ["A群過多", "A群比率35%超", "一時的に入院を抑制、退院支援を強化"],
            ["B群不足", "B群比率25%未満", "A群からB群への遷移を待つ（=退院を抑制）"],
            ["C群滞留", "C群比率30%超", "退院支援カンファレンスの実施"],
            ["日次目標未達", "日次運営貢献額がマイナス", "コスト構造の見直し"],
        ])

    # 5-4
    doc.add_heading("5-4. 月次レポートへの活用例", level=2)
    add_body(doc, "毎月の病棟運営会議で以下を報告することを推奨します。", bold=True)
    p = doc.add_paragraph(style="List Number")
    p.clear()
    p.add_run("実績稼働率 vs シミュレーション結果").font.size = Pt(10.5)
    p = doc.add_paragraph(style="List Number")
    p.clear()
    p.add_run("採用した戦略とその効果").font.size = Pt(10.5)
    p = doc.add_paragraph(style="List Number")
    p.clear()
    p.add_run("来月の推奨戦略（入院需要予測に基づく）").font.size = Pt(10.5)
    p = doc.add_paragraph(style="List Number")
    p.clear()
    p.add_run("フェーズ構成比の改善ポイント").font.size = Pt(10.5)

    add_page_break(doc)

    # ================================================================
    # 第6章：よくある質問
    # ================================================================
    doc.add_heading("第6章　よくある質問", level=1)

    qa_list = [
        ("Q: シミュレーション結果は毎回変わりますか？",
         "A: 乱数シードが同じであれば同じ結果になります。異なるシードを試すことで、結果のばらつきを確認できます。デフォルトのシード値は42です。"),
        ("Q: 実際の診療報酬と金額が違うのですが？",
         "A: このシミュレーターは病棟運営用の簡易モデルです。実際の診療報酬を入力したい場合は、「患者フェーズ別パラメータ」の収益・コストを実態に合わせて調整してください。"),
        ("Q: 94床以外の病棟でも使えますか？",
         "A: はい。サイドバーの「病床数」を変更すれば他の病棟規模でもシミュレーション可能です（10〜200床の範囲）。"),
        ("Q: 結果をExcelで分析したい場合は？",
         "A: メインエリアの「データ」タブにある「CSVダウンロード」ボタンでCSVファイルを保存し、Excelで開いてください。日次データの全項目が含まれています。"),
        ("Q: 2ヶ月連続のシミュレーションはできますか？",
         "A: 現バージョンでは1ヶ月単位です。月をまたぐ分析は、前月末の稼働率を参考に初期パラメータ（入院流入変動係数など）を調整して実行してください。"),
        ("Q: 看護必要度への影響は反映されていますか？",
         "A: 現バージョンでは未対応です。今後の改善で看護必要度連動機能を追加予定です。"),
        ("Q: 初期状態の患者配置はどうなっていますか？",
         "A: シミュレーション開始時、稼働率約90%でA群15%・B群45%・C群40%の比率で患者を配置します。C群は退院ラッシュ防止のため15〜18日目に集中配置されます。"),
        ("Q: 退院確率はどのように計算されていますか？",
         "A: 在院日数に応じた基本退院確率（A群0.5%〜C群後期35%）に、戦略別の倍率と稼働率別の補正を適用しています。最終確率は0〜80%にクリップされます。"),
    ]

    for q, a in qa_list:
        add_body(doc, q, bold=True)
        add_body(doc, a)
        add_body(doc, "")

    add_page_break(doc)

    # ================================================================
    # 第7章：用語集
    # ================================================================
    doc.add_heading("第7章　用語集", level=1)
    add_styled_table(doc,
        ["用語", "定義"],
        [
            ["稼働率", "在院患者数 ÷ 病床数 × 100%。目標は90〜95%"],
            ["運営貢献額", "収益 - 変動費。人件費等の固定費は空床でも発生するため除外した利益"],
            ["未活用病床コスト", "空床や満床で入院を受け入れられなかった場合の逸失診療報酬。約25,000円/空床/日"],
            ["フェーズ構成比", "A/B/C各群の患者数が全体に占める割合。B群40〜50%が理想"],
            ["包括払い", "入院期間を通じて1日あたりの報酬が定額で支払われる制度"],
            ["在院日数", "入院してからの経過日数。フェーズ分類の基準"],
            ["退院調整の柔軟性が高い層", "C群患者を退院タイミングの調整に活用する考え方。閑散期は保持、繁忙期は退院促進"],
            ["シード値", "乱数の初期値。同じ値なら同じシミュレーション結果を再現できる（デフォルト42）"],
            ["退院確率", "各日に患者が退院する確率。在院日数・フェーズ・戦略・稼働率で動的に変化"],
            ["ポアソン分布", "新規入院数の日々のばらつきを表現するための確率分布"],
            ["シグモイド関数", "退院確率の計算に使用される0〜1の範囲の滑らかな関数"],
            ["目標レンジ日数", "30日間のうち稼働率が目標範囲（90〜95%）に収まった日数"],
            ["入院流入変動係数", "入院需要の季節変動を表す倍率（1.0=通常、1.2=繁忙、0.8=閑散）"],
        ])

    add_page_break(doc)

    # ================================================================
    # 付録A：CLI版の使い方
    # ================================================================
    doc.add_heading("付録A　CLI版の使い方", level=1)
    add_body(doc, "Streamlit版のほかに、ターミナルから直接実行できるCLI版も用意されています。")

    doc.add_heading("実行方法", level=2)
    add_code_block(doc, "cd ~/ai-management")
    add_code_block(doc, "source .venv/bin/activate")
    add_code_block(doc, "python scripts/bed_control_simulator.py")

    doc.add_heading("出力先", level=2)
    add_body(doc, "output/ フォルダにCSVファイルとPNG画像が保存されます。")
    add_body_list(doc, [
        "CSVファイル：日次シミュレーション結果の全データ",
        "PNGファイル：稼働率推移・フェーズ構成比・運営分析のグラフ",
    ])

    doc.add_heading("Streamlit版の起動方法", level=2)
    add_code_block(doc, "cd ~/ai-management")
    add_code_block(doc, "source .venv/bin/activate")
    add_code_block(doc, "streamlit run scripts/bed_control_simulator_app.py")
    add_body(doc, "ブラウザが自動的に開き、Streamlit版のインターフェースが表示されます。")

    doc.add_heading("主な関数（CLI版）", level=2)
    add_styled_table(doc,
        ["関数名", "説明"],
        [
            ["create_default_params()", "デフォルトパラメータ辞書を生成"],
            ["simulate_bed_control(params, strategy)", "日次シミュレーションを実行しDataFrameを返す"],
            ["summarize_results(df)", "月次サマリーを辞書で返す"],
            ["compare_strategies(params)", "3戦略の一括比較表を生成"],
        ])

    add_page_break(doc)

    # ================================================================
    # 付録B：パラメータ一覧表
    # ================================================================
    doc.add_heading("付録B　パラメータ一覧表", level=1)
    add_body(doc, "シミュレーターで使用する全パラメータの一覧です。", bold=True)

    doc.add_heading("病棟基本設定", level=2)
    add_styled_table(doc,
        ["変数名", "説明", "初期値", "範囲", "単位"],
        [
            ["num_beds", "総病床数", "94", "10〜200", "床"],
            ["target_occupancy_lower", "目標稼働率下限", "0.90", "0.80〜1.00", "（比率）"],
            ["target_occupancy_upper", "目標稼働率上限", "0.95", "0.80〜1.00", "（比率）"],
            ["days_in_month", "シミュレーション日数", "30", "28〜31", "日"],
            ["monthly_admissions", "月間入院件数", "150", "50〜300", "件"],
            ["avg_length_of_stay", "平均在院日数", "19", "10〜30", "日"],
            ["discharge_adjustment_days", "退院調整ラグ日数", "2", "0〜5", "日"],
            ["admission_variation_coeff", "入院数変動係数", "1.0", "0.50〜1.50", "（倍率）"],
        ])

    doc.add_heading("収益・コスト設定", level=2)
    add_styled_table(doc,
        ["変数名", "説明", "初期値", "単位"],
        [
            ["phase_a_revenue", "A群 日次診療報酬", "30,000", "円/日/人"],
            ["phase_a_cost", "A群 日次コスト", "28,000", "円/日/人"],
            ["phase_b_revenue", "B群 日次診療報酬", "30,000", "円/日/人"],
            ["phase_b_cost", "B群 日次コスト", "13,000", "円/日/人"],
            ["phase_c_revenue", "C群 日次診療報酬", "30,000", "円/日/人"],
            ["phase_c_cost", "C群 日次コスト", "11,000", "円/日/人"],
        ])

    doc.add_heading("加算設定", level=2)
    add_styled_table(doc,
        ["変数名", "説明", "初期値", "適用条件", "単位"],
        [
            ["first_day_bonus", "入院初日加算", "0", "新規入院時", "円/人"],
            ["within_14days_bonus", "14日以内退院加算", "0", "14日以内退院時", "円/人"],
            ["rehab_fee", "リハビリ加算", "0", "B群・C群患者", "円/日/人"],
        ])

    doc.add_heading("閾値・制御パラメータ", level=2)
    add_styled_table(doc,
        ["変数名", "説明", "初期値", "範囲", "単位"],
        [
            ["opportunity_cost", "未活用病床コスト", "10,000", "0〜50,000", "円/空床/日"],
            ["discharge_promotion_threshold", "退院促進閾値", "0.95", "0.80〜1.00", "（比率）"],
            ["admission_suppression_threshold", "入院抑制閾値", "0.97", "0.80〜1.00", "（比率）"],
            ["random_seed", "乱数シード", "42", "任意の整数", "−"],
        ])

    doc.add_heading("戦略別の内部制御パラメータ", level=2)
    add_body(doc, "以下は各戦略で自動的に適用される内部パラメータです（ユーザーが直接変更する項目ではありません）。")
    add_styled_table(doc,
        ["パラメータ", "回転重視", "安定維持", "バランス"],
        [
            ["C群退院倍率（通常時）", "2.2倍", "0.6倍", "1.0倍（状況依存）"],
            ["C群退院倍率（95%超）", "2.2倍", "1.5倍", "1.8倍"],
            ["B群後半退院倍率", "1.6倍", "1.0倍", "1.0倍"],
            ["低稼働時抑制", "0.3倍（85%未満）", "0.4倍×0.4倍", "0.3倍（90%未満）"],
            ["新規入院上限稼働率", "97%", "95%（93%から抑制）", "95%"],
            ["A群比率チェック", "なし", "なし", "35%超で入院半減"],
        ])

    # ================================================================
    # 最終ページ
    # ================================================================
    add_page_break(doc)
    for _ in range(8):
        doc.add_paragraph()
    end_p = doc.add_paragraph()
    end_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = end_p.add_run("ベッドコントロールシミュレーター 実践マニュアル")
    run.font.size = Pt(14)
    run.bold = True
    run.font.color.rgb = RGBColor(0x2C, 0x3E, 0x50)

    end_p2 = doc.add_paragraph()
    end_p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run2 = end_p2.add_run("おもろまちメディカルセンター")
    run2.font.size = Pt(12)

    end_p3 = doc.add_paragraph()
    end_p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run3 = end_p3.add_run("2026年3月25日 初版")
    run3.font.size = Pt(10)
    run3.font.color.rgb = RGBColor(0x7F, 0x8C, 0x8D)

    # ================================================================
    # 保存
    # ================================================================
    OUTPUT_PATH.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(OUTPUT_PATH))
    print(f"マニュアルを生成しました: {OUTPUT_PATH}")
    print(f"ファイルサイズ: {OUTPUT_PATH.stat().st_size / 1024:.1f} KB")


if __name__ == "__main__":
    generate_manual()
