#!/usr/bin/env python3
"""シナリオ台本 Markdown → DOCX 一括変換ツール.

4 本のシナリオ台本 .md を .docx に変換する。
フォントは Hiragino Kaku Gothic ProN（macOS 標準の日本語ゴシック）に統一。

対応する Markdown 要素:
    - 見出し (#, ##, ###, ####)
    - 段落（空行区切り）
    - 太字 (**text**)
    - 斜体 (*text*) ただし箇条書き記号 `*` とは区別
    - インラインコード (`code`)
    - リンク ([text](url)) → text (url) 形式で展開
    - 箇条書き (- , * ) ネスト 2 段階
    - 番号リスト (1. 2. ) ネスト 2 段階
    - テーブル (| col | col |)
    - 引用 (> text)
    - 水平線 (---)
    - コードブロック (``` で囲まれた部分、10pt 固定)

実行:
    python3 scripts/generate_scenario_docx.py
"""
from __future__ import annotations

import re
import subprocess
from dataclasses import dataclass
from datetime import datetime
from pathlib import Path
from typing import Iterator

from docx import Document
from docx.document import Document as DocumentType
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor
from docx.text.paragraph import Paragraph
from docx.text.run import Run

# ---------- 設定 ----------

FONT_NAME = "Hiragino Kaku Gothic ProN"
BODY_PT = 11
CODE_PT = 10
APP_VERSION = "v4"

REPO_ROOT = Path(__file__).resolve().parent.parent
OUTPUT_DIR = REPO_ROOT / "docs" / "admin" / "scenario_scripts_v4"

# 変換対象: (入力 md, 出力ファイル名、v4)
TARGETS: list[tuple[Path, str]] = [
    (REPO_ROOT / "docs" / "admin" / "demo_scenario_v4.md",
     "demo_scenario_v4.docx"),
    (REPO_ROOT / "docs" / "admin" / "presentation_script_bedcontrol_v4.md",
     "presentation_script_bedcontrol_v4.docx"),
    (REPO_ROOT / "docs" / "admin" / "carnf_scenario_v4.md",
     "carnf_scenario_v4.docx"),
    (REPO_ROOT / "docs" / "admin" / "BedControl_Manual_v4.md",
     "BedControl_Manual_v4.docx"),
    (REPO_ROOT / "docs" / "admin" / "slides" / "weekend_holiday_kpi" / "script.md",
     "weekend_holiday_kpi_script.docx"),
]


# ---------- フォント制御 ----------

def _apply_east_asian_font(run_or_style_element, font_name: str = FONT_NAME) -> None:
    """Run / Style の rPr に East Asian font を設定."""
    rPr = run_or_style_element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.append(rFonts)
    rFonts.set(qn("w:eastAsia"), font_name)
    rFonts.set(qn("w:ascii"), font_name)
    rFonts.set(qn("w:hAnsi"), font_name)
    rFonts.set(qn("w:cs"), font_name)


def _set_run_font(
    run: Run,
    size: int = BODY_PT,
    bold: bool = False,
    italic: bool = False,
    color: str | None = None,
    mono: bool = False,
) -> None:
    """Run にフォント設定を適用."""
    run.font.name = FONT_NAME
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    if color:
        r, g, b = int(color[0:2], 16), int(color[2:4], 16), int(color[4:6], 16)
        run.font.color.rgb = RGBColor(r, g, b)
    _apply_east_asian_font(run._element)
    # mono は使わず単純にサイズで差別化（Hiraginoを全体統一するため）


def _apply_style_font(style_element, size_pt: int, color_hex: str | None = None) -> None:
    """Style 要素にフォント設定を適用."""
    rPr = style_element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.append(rFonts)
    rFonts.set(qn("w:eastAsia"), FONT_NAME)
    rFonts.set(qn("w:ascii"), FONT_NAME)
    rFonts.set(qn("w:hAnsi"), FONT_NAME)
    rFonts.set(qn("w:cs"), FONT_NAME)


# ---------- インライン要素パーサ ----------

# [text](url) → "text (url)" へ平文化（DOCX のクリック可能リンクは複雑なので省略）
_RE_LINK = re.compile(r"\[([^\]]+)\]\(([^)]+)\)")
# **bold**
_RE_BOLD = re.compile(r"\*\*([^*]+)\*\*")
# `inline code`
_RE_INLINE_CODE = re.compile(r"`([^`]+)`")
# *italic* — 行頭の `* ` （リスト）と区別するため呼び出し側で除外する
_RE_ITALIC = re.compile(r"(?<![\*\w])\*([^*\s][^*]*[^*\s]|\S)\*(?!\*)")


@dataclass
class Span:
    text: str
    bold: bool = False
    italic: bool = False
    code: bool = False


def _parse_inline(text: str) -> list[Span]:
    """インライン Markdown を Span のリストに分解.

    優先順位: インラインコード > 太字 > 斜体 > リンク展開
    """
    # まずリンクを平文化
    text = _RE_LINK.sub(lambda m: f"{m.group(1)} ({m.group(2)})", text)

    spans: list[Span] = []

    # インラインコードから順に処理するため、トークナイザを手書き
    i = 0
    n = len(text)
    while i < n:
        ch = text[i]

        # inline code `...`
        if ch == "`":
            j = text.find("`", i + 1)
            if j != -1:
                spans.append(Span(text[i + 1:j], code=True))
                i = j + 1
                continue

        # **bold**
        if ch == "*" and i + 1 < n and text[i + 1] == "*":
            j = text.find("**", i + 2)
            if j != -1:
                inner = text[i + 2:j]
                # ネストしない（簡易実装）
                # 内部にさらに italic / code はあり得るが、スクリプト向け簡易版として平文化
                spans.append(Span(inner, bold=True))
                i = j + 2
                continue

        # *italic* （直前が英数字・アスタリスクでない場合のみ）
        if ch == "*" and (i == 0 or text[i - 1] not in "*\\w"):
            j = text.find("*", i + 1)
            if j != -1 and j != i + 1:
                inner = text[i + 1:j]
                # 空白のみ / 記号のみは italic 扱いしない
                if inner.strip() and "\n" not in inner:
                    spans.append(Span(inner, italic=True))
                    i = j + 1
                    continue

        # プレーン文字列を読み込む（次の特殊文字まで）
        next_special = n
        for marker in ("`", "**", "*"):
            k = text.find(marker, i + 1)
            if k != -1 and k < next_special:
                next_special = k
        if next_special == n:
            spans.append(Span(text[i:]))
            break
        else:
            if next_special > i:
                spans.append(Span(text[i:next_special]))
            i = next_special

    # 連続する通常 Span をマージ
    merged: list[Span] = []
    for s in spans:
        if merged and not (s.bold or s.italic or s.code) \
                and not (merged[-1].bold or merged[-1].italic or merged[-1].code):
            merged[-1] = Span(merged[-1].text + s.text)
        else:
            merged.append(s)
    return merged


def _add_inline_runs(paragraph: Paragraph, text: str, base_size: int = BODY_PT) -> None:
    """段落にインライン書式付き runs を追加."""
    spans = _parse_inline(text)
    for span in spans:
        run = paragraph.add_run(span.text)
        _set_run_font(
            run,
            size=CODE_PT if span.code else base_size,
            bold=span.bold,
            italic=span.italic and not span.code,
        )


# ---------- ドキュメントスタイル初期化 ----------

def _init_document_styles(doc: DocumentType) -> None:
    """ドキュメント全体のデフォルトスタイルを Hiragino に設定."""
    # Normal
    normal = doc.styles["Normal"]
    normal.font.name = FONT_NAME
    normal.font.size = Pt(BODY_PT)
    _apply_style_font(normal.element, BODY_PT)
    normal.paragraph_format.space_after = Pt(4)
    normal.paragraph_format.line_spacing = 1.35

    # Heading 1..4
    for level, size, color in [
        (1, 24, "1A478A"),
        (2, 18, "1A478A"),
        (3, 14, "2C5282"),
        (4, 12, "2C5282"),
    ]:
        try:
            h = doc.styles[f"Heading {level}"]
        except KeyError:
            continue
        h.font.name = FONT_NAME
        h.font.size = Pt(size)
        h.font.bold = True
        r, g, b = int(color[0:2], 16), int(color[2:4], 16), int(color[4:6], 16)
        h.font.color.rgb = RGBColor(r, g, b)
        _apply_style_font(h.element, size, color)
        h.paragraph_format.space_before = Pt(size * 0.7)
        h.paragraph_format.space_after = Pt(size * 0.3)
        h.paragraph_format.keep_with_next = True

    # List Bullet / List Number
    for style_name in ("List Bullet", "List Number"):
        try:
            s = doc.styles[style_name]
        except KeyError:
            continue
        s.font.name = FONT_NAME
        s.font.size = Pt(BODY_PT)
        _apply_style_font(s.element, BODY_PT)

    # 余白
    for section in doc.sections:
        section.top_margin = Cm(2.0)
        section.bottom_margin = Cm(2.0)
        section.left_margin = Cm(2.0)
        section.right_margin = Cm(2.0)


# ---------- ブロック要素の処理 ----------

_RE_HEADING = re.compile(r"^(#{1,6})\s+(.*)$")
_RE_TABLE_ROW = re.compile(r"^\|(.+)\|\s*$")
_RE_TABLE_SEP = re.compile(r"^\|[\s\-:|]+\|\s*$")
_RE_BULLET = re.compile(r"^(\s*)[-*+]\s+(.*)$")
_RE_NUMBERED = re.compile(r"^(\s*)(\d+)\.\s+(.*)$")
_RE_QUOTE = re.compile(r"^>\s?(.*)$")
_RE_HR = re.compile(r"^-{3,}\s*$")
_RE_CODE_FENCE = re.compile(r"^```(.*)$")


def _parse_table_row(line: str) -> list[str]:
    """`| a | b | c |` → ['a', 'b', 'c']."""
    inner = line.strip()
    if inner.startswith("|"):
        inner = inner[1:]
    if inner.endswith("|"):
        inner = inner[:-1]
    return [cell.strip() for cell in inner.split("|")]


def _add_table(doc: DocumentType, rows: list[list[str]]) -> None:
    """Word テーブルを追加."""
    if not rows:
        return
    n_cols = max(len(r) for r in rows)
    table = doc.add_table(rows=len(rows), cols=n_cols)
    table.style = "Light Grid Accent 1"
    table.alignment = WD_ALIGN_PARAGRAPH.LEFT

    for i, row in enumerate(rows):
        for j in range(n_cols):
            cell = table.cell(i, j)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            # 既存の空段落をクリア
            cell.paragraphs[0].text = ""
            p = cell.paragraphs[0]
            text = row[j] if j < len(row) else ""
            _add_inline_runs(p, text, base_size=BODY_PT - 1)
            # ヘッダ行は太字に
            if i == 0:
                for run in p.runs:
                    run.font.bold = True


def _add_code_block(doc: DocumentType, lines: list[str]) -> None:
    """コードブロックを段落として追加（10pt、インデント）."""
    for line in lines:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Cm(0.5)
        p.paragraph_format.space_after = Pt(0)
        run = p.add_run(line if line else " ")
        _set_run_font(run, size=CODE_PT)


def _add_quote(doc: DocumentType, text: str) -> None:
    """引用段落を追加."""
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.7)
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after = Pt(3)
    _add_inline_runs(p, text, base_size=BODY_PT)
    # 左側に薄いマーカー用に斜体＋グレーで差別化
    for run in p.runs:
        run.font.italic = True
        run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)


def _add_hr(doc: DocumentType) -> None:
    """水平線（段落の下罫線）を追加."""
    p = doc.add_paragraph()
    pPr = p._element.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "999999")
    pBdr.append(bottom)
    pPr.append(pBdr)


def _add_list_item(
    doc: DocumentType,
    text: str,
    style_name: str,
    indent_level: int = 0,
) -> None:
    """箇条書き・番号リスト項目を追加."""
    try:
        p = doc.add_paragraph(style=style_name)
    except KeyError:
        p = doc.add_paragraph()
    if indent_level:
        p.paragraph_format.left_indent = Cm(0.7 + 0.7 * indent_level)
    _add_inline_runs(p, text, base_size=BODY_PT)


# ---------- Markdown → DOCX 変換 ----------

def md_to_docx(md_path: Path, docx_path: Path, commit_hash: str) -> None:
    """Markdown を DOCX に変換."""
    doc = Document()
    _init_document_styles(doc)

    md_text = md_path.read_text(encoding="utf-8")
    lines = md_text.split("\n")

    # メタデータ（先頭に）
    meta_p = doc.add_paragraph()
    meta_run = meta_p.add_run(
        f"生成日: {datetime.now().strftime('%Y-%m-%d')}    "
        f"対応アプリ: {APP_VERSION}    コミット: {commit_hash[:7]}"
    )
    _set_run_font(meta_run, size=9, color="888888")
    _add_hr(doc)

    i = 0
    n = len(lines)
    while i < n:
        line = lines[i]
        stripped = line.rstrip()

        # 空行 → 何もしない（段落区切り）
        if not stripped:
            i += 1
            continue

        # コードフェンス
        m_fence = _RE_CODE_FENCE.match(stripped)
        if m_fence:
            code_lines: list[str] = []
            i += 1
            while i < n and not _RE_CODE_FENCE.match(lines[i].rstrip()):
                code_lines.append(lines[i])
                i += 1
            if i < n:
                i += 1  # 閉じフェンスをスキップ
            _add_code_block(doc, code_lines)
            continue

        # 水平線
        if _RE_HR.match(stripped):
            _add_hr(doc)
            i += 1
            continue

        # 見出し
        m_h = _RE_HEADING.match(stripped)
        if m_h:
            level = min(len(m_h.group(1)), 4)
            heading_text = m_h.group(2).strip()
            # インライン太字などが含まれる場合もタイトル内に適用
            try:
                h = doc.add_heading(level=level)
            except Exception:
                h = doc.add_paragraph()
            # 見出しはフォント揃えるため手動で run を追加
            _add_inline_runs(h, heading_text, base_size={1: 24, 2: 18, 3: 14, 4: 12}.get(level, 12))
            # bold 強制
            for run in h.runs:
                run.font.bold = True
                _apply_east_asian_font(run._element)
            i += 1
            continue

        # テーブル
        if _RE_TABLE_ROW.match(stripped):
            table_lines: list[str] = []
            while i < n and _RE_TABLE_ROW.match(lines[i].rstrip()):
                table_lines.append(lines[i].rstrip())
                i += 1
            # セパレータ行（---）を除外してセルデータだけ抽出
            rows = [
                _parse_table_row(tl)
                for tl in table_lines
                if not _RE_TABLE_SEP.match(tl)
            ]
            _add_table(doc, rows)
            continue

        # 引用
        m_q = _RE_QUOTE.match(stripped)
        if m_q:
            # 連続する > を1段落にまとめる
            quote_parts: list[str] = [m_q.group(1)]
            i += 1
            while i < n:
                mnext = _RE_QUOTE.match(lines[i].rstrip())
                if mnext:
                    quote_parts.append(mnext.group(1))
                    i += 1
                else:
                    break
            _add_quote(doc, " ".join(quote_parts))
            continue

        # 箇条書き
        m_b = _RE_BULLET.match(line)
        if m_b:
            leading_ws = len(m_b.group(1).replace("\t", "    "))
            level = min(leading_ws // 2, 2)
            _add_list_item(doc, m_b.group(2), "List Bullet", indent_level=level)
            i += 1
            continue

        # 番号リスト
        m_n = _RE_NUMBERED.match(line)
        if m_n:
            leading_ws = len(m_n.group(1).replace("\t", "    "))
            level = min(leading_ws // 2, 2)
            _add_list_item(doc, m_n.group(3), "List Number", indent_level=level)
            i += 1
            continue

        # 通常段落（連続行を連結）
        para_lines = [stripped]
        i += 1
        while i < n:
            nxt = lines[i].rstrip()
            if not nxt:
                break
            if (_RE_HEADING.match(nxt)
                    or _RE_TABLE_ROW.match(nxt)
                    or _RE_BULLET.match(lines[i])
                    or _RE_NUMBERED.match(lines[i])
                    or _RE_QUOTE.match(nxt)
                    or _RE_HR.match(nxt)
                    or _RE_CODE_FENCE.match(nxt)):
                break
            para_lines.append(nxt)
            i += 1
        para_text = " ".join(para_lines)
        p = doc.add_paragraph()
        _add_inline_runs(p, para_text, base_size=BODY_PT)

    docx_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(docx_path))


# ---------- エントリポイント ----------

def _get_commit_hash() -> str:
    """現在の HEAD コミットハッシュを取得."""
    try:
        result = subprocess.run(
            ["git", "rev-parse", "HEAD"],
            capture_output=True,
            text=True,
            cwd=str(REPO_ROOT),
            check=True,
        )
        return result.stdout.strip()
    except Exception:
        return "unknown"


def main() -> int:
    commit_hash = _get_commit_hash()
    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)

    print(f"[generate_scenario_docx] output dir: {OUTPUT_DIR}")
    print(f"[generate_scenario_docx] commit:     {commit_hash[:7]}")
    print(f"[generate_scenario_docx] font:       {FONT_NAME}")
    print("-" * 60)

    errors = 0
    for md_path, docx_name in TARGETS:
        if not md_path.exists():
            print(f"  [ERROR] not found: {md_path}")
            errors += 1
            continue
        docx_path = OUTPUT_DIR / docx_name
        try:
            md_to_docx(md_path, docx_path, commit_hash)
            size_kb = docx_path.stat().st_size / 1024
            print(f"  [OK] {md_path.name} -> {docx_name} ({size_kb:.1f} KB)")
        except Exception as e:
            print(f"  [ERROR] {md_path.name}: {e}")
            errors += 1

    return 0 if errors == 0 else 1


if __name__ == "__main__":
    raise SystemExit(main())
