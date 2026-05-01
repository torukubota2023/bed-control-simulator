"""院内 SE 向け ベッドコントロールシミュレーター 設置マニュアル DOCX 生成スクリプト.

副院長指示 (2026-05-01):
- 配布: Microsoft Teams で SE に渡す
- ブラウザ: Chrome Portable（DLリンクをマニュアル内に記載）
- マニュアル形式: DOCX 一本
- SE スキル: Python は扱える
- ネット環境: GitHub アクセスは外部 PC で実施 → zip で院内に持ち込み

使い方:
    .venv/bin/python scripts/generate_se_install_manual.py
    → docs/admin/SE_install_manual_2026-05-01.docx
"""
from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT


JP_FONT = "Hiragino Mincho ProN"
MONO_FONT = "Menlo"

# カラー定数
C_TITLE = RGBColor(0x1F, 0x29, 0x37)
C_ACCENT = RGBColor(0x25, 0x63, 0xEB)
C_DANGER = RGBColor(0xDC, 0x26, 0x26)
C_WARN = RGBColor(0xF5, 0x9E, 0x0B)
C_OK = RGBColor(0x10, 0xB9, 0x81)
C_MUTED = RGBColor(0x6B, 0x72, 0x80)
C_FAINT = RGBColor(0x9C, 0xA3, 0xAF)


def _set_jp_font(run, size_pt=11, bold=False, color=None):
    run.font.name = JP_FONT
    run.font.size = Pt(size_pt)
    run.font.bold = bold
    if color is not None:
        run.font.color.rgb = color
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.insert(0, rFonts)
    rFonts.set(qn("w:eastAsia"), JP_FONT)
    rFonts.set(qn("w:ascii"), JP_FONT)
    rFonts.set(qn("w:hAnsi"), JP_FONT)


def _set_mono_font(run, size_pt=10, color=None):
    run.font.name = MONO_FONT
    run.font.size = Pt(size_pt)
    if color is not None:
        run.font.color.rgb = color
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.insert(0, rFonts)
    rFonts.set(qn("w:ascii"), MONO_FONT)
    rFonts.set(qn("w:hAnsi"), MONO_FONT)


def add_para(doc, text, size=11, bold=False, align=None, color=None):
    p = doc.add_paragraph()
    if align is not None:
        p.alignment = align
    run = p.add_run(text)
    _set_jp_font(run, size_pt=size, bold=bold, color=color)
    return p


def add_heading(doc, text, level=1):
    sizes = {0: 26, 1: 20, 2: 16, 3: 13}
    h = doc.add_heading(level=level)
    run = h.add_run(text)
    _set_jp_font(run, size_pt=sizes.get(level, 12), bold=True, color=C_TITLE)


def add_code(doc, code):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(1.0)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(code)
    _set_mono_font(run, size_pt=10)


def add_callout(doc, label, body, color=None):
    if color is None:
        color = C_ACCENT
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.5)
    run_label = p.add_run(label + " ")
    _set_jp_font(run_label, size_pt=11, bold=True, color=color)
    run_body = p.add_run(body)
    _set_jp_font(run_body, size_pt=11)


def add_step(doc, num, title, body_lines, code=None):
    p = doc.add_paragraph()
    run_num = p.add_run(f"ステップ {num}：")
    _set_jp_font(run_num, size_pt=13, bold=True, color=C_DANGER)
    run_title = p.add_run(title)
    _set_jp_font(run_title, size_pt=13, bold=True)
    for line in body_lines:
        add_para(doc, line, size=11)
    if code:
        add_code(doc, code)


def add_table(doc, header, rows, col_widths_cm=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(header))
    table.style = "Light Grid Accent 1"
    table.alignment = WD_TABLE_ALIGNMENT.LEFT

    if col_widths_cm:
        for col, width in zip(table.columns, col_widths_cm):
            for cell in col.cells:
                cell.width = Cm(width)

    # ヘッダー
    hdr = table.rows[0].cells
    for i, label in enumerate(header):
        hdr[i].text = ""
        run = hdr[i].paragraphs[0].add_run(label)
        _set_jp_font(run, size_pt=10, bold=True, color=C_TITLE)

    # 本体
    for r_i, row_data in enumerate(rows, start=1):
        cells = table.rows[r_i].cells
        for c_i, val in enumerate(row_data):
            cells[c_i].text = ""
            run = cells[c_i].paragraphs[0].add_run(str(val))
            _set_jp_font(run, size_pt=10)
    return table


def add_separator(doc):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("─" * 30)
    _set_jp_font(run, size_pt=10, color=C_FAINT)


# ======================================================================
# マニュアル本体
# ======================================================================

def build():
    doc = Document()

    # 既定スタイル
    style = doc.styles["Normal"]
    style.font.name = JP_FONT
    style.font.size = Pt(11)

    # ページ余白
    for section in doc.sections:
        section.top_margin = Cm(2.2)
        section.bottom_margin = Cm(2.2)
        section.left_margin = Cm(2.2)
        section.right_margin = Cm(2.2)

    # ====================================================================
    # 表紙
    # ====================================================================
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(60)
    run = p.add_run("ベッドコントロール")
    _set_jp_font(run, size_pt=32, bold=True, color=C_TITLE)

    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run2 = p2.add_run("シミュレーター")
    _set_jp_font(run2, size_pt=32, bold=True, color=C_TITLE)

    p3 = doc.add_paragraph()
    p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p3.paragraph_format.space_before = Pt(20)
    run3 = p3.add_run("院内 LAN 設置マニュアル")
    _set_jp_font(run3, size_pt=24, bold=True, color=C_ACCENT)

    p4 = doc.add_paragraph()
    p4.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p4.paragraph_format.space_before = Pt(36)
    run4 = p4.add_run("〜 院内 SE 様向け 〜")
    _set_jp_font(run4, size_pt=14, color=C_MUTED)

    p5 = doc.add_paragraph()
    p5.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p5.paragraph_format.space_before = Pt(100)
    run5 = p5.add_run("おもろまちメディカルセンター")
    _set_jp_font(run5, size_pt=14)

    p6 = doc.add_paragraph()
    p6.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run6 = p6.add_run("副院長 久保田 透")
    _set_jp_font(run6, size_pt=12)

    p7 = doc.add_paragraph()
    p7.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p7.paragraph_format.space_before = Pt(8)
    run7 = p7.add_run("作成日：2026年5月1日")
    _set_jp_font(run7, size_pt=11, color=C_FAINT)

    doc.add_page_break()

    # ====================================================================
    # 0. はじめに
    # ====================================================================
    add_heading(doc, "0. はじめに", level=1)
    add_para(doc, "SE ご担当者様")
    add_para(doc, "")
    add_para(
        doc,
        "お疲れさまです。副院長の久保田です。"
        "病棟のベッドコントロール業務を支援する Web アプリケーションを開発しました。"
        "院内 LAN 上のサーバー PC（常時稼働）に設置し、病棟スタッフがブラウザから "
        "アクセスできるようにしていただきたく、お願いいたします。"
    )
    add_para(doc, "")

    add_heading(doc, "本アプリの特徴", level=2)
    add_para(doc, "・病棟の入退院データを日次で入力し、稼働率・在院日数・収益を可視化する Streamlit アプリ")
    add_para(doc, "・患者個人情報は一切含まない（匿名集計データのみ／医師コードのみ）")
    add_para(doc, "・院内 LAN のみでアクセス可能（外部公開しない）")
    add_para(doc, "・Python 3.11+ で動作（pandas / streamlit / plotly 等）")
    add_para(doc, "")

    add_heading(doc, "本パッケージの構成", level=2)
    add_para(
        doc,
        "Microsoft Teams で本マニュアル（DOCX）と一緒に、"
        "オフラインインストール用の zip ファイルが配布されます。zip の中身は以下です。",
    )

    add_code(
        doc,
        "bed_control_simulator_offline_2026-05-01.zip\n"
        "├─ scripts/                   ← Python アプリ本体\n"
        "├─ data/                      ← データ（空スキーマ）\n"
        "├─ settings/                  ← 設定 YAML（手動シード入力先）\n"
        "├─ docs/                      ← ドキュメント\n"
        "├─ wheels/                    ← pip オフラインインストール用 .whl ファイル一式\n"
        "├─ deploy/\n"
        "│   ├─ launch_bed_control.bat              ← サーバー起動\n"
        "│   ├─ launch_portable_chrome.bat          ← クライアント launcher（Chrome 版）\n"
        "│   ├─ launch_portable_firefox.bat         ← クライアント launcher（Firefox 版・予備）\n"
        "│   └─ open_bed_control.ps1                ← PowerShell 版 launcher\n"
        "├─ tools/browser_probe.html   ← ブラウザ互換チェッカー\n"
        "├─ requirements.txt           ← Python 依存リスト\n"
        "├─ requirements-edge90.txt    ← 古い Edge 互換確認用（参考）\n"
        "└─ README_FIRST.txt           ← まずこのファイルを読んでください",
    )
    add_para(doc, "")

    add_callout(
        doc,
        "💡 ヒント",
        "ネット接続のある PC で zip を展開しています。pip パッケージも wheels/ にまとめてあるため、"
        "院内 LAN から外に出ずにインストールが完結します。",
        color=C_OK,
    )

    doc.add_page_break()

    # ====================================================================
    # 1. 全体構成図
    # ====================================================================
    add_heading(doc, "1. 全体構成図", level=1)
    add_para(doc, "システム全体は次のようになります。")
    add_para(doc, "")

    add_code(
        doc,
        "┌─────────────────────────────────────┐\n"
        "│ サーバー PC（常時稼働）                │\n"
        "│   - Python 3.11+                      │\n"
        "│   - Streamlit サーバー（ポート 8501）  │\n"
        "│   - データ: data/admission_details.csv │\n"
        "│   - 設定: settings/*.yaml              │\n"
        "└─────────────────┬───────────────────┘\n"
        "                  │ 院内 LAN（HTTP）\n"
        "                  ▼\n"
        "┌─────────────────────────────────────┐\n"
        "│ 病棟 PC（複数台）                      │\n"
        "│   - Chrome Portable（電子カルテ用      │\n"
        "│     Edge 90 とは別ブラウザ）          │\n"
        "│   - URL: http://(サーバーIP):8501     │\n"
        "└─────────────────────────────────────┘",
    )

    add_para(doc, "")
    add_callout(
        doc,
        "⚠️ 重要",
        "電子カルテ用の Edge 90 には一切干渉しません。Chrome Portable は USB / 共有フォルダ"
        "から起動するため、レジストリも変更しません。",
        color=C_WARN,
    )

    doc.add_page_break()

    # ====================================================================
    # 2. 必要環境
    # ====================================================================
    add_heading(doc, "2. 必要環境", level=1)

    add_heading(doc, "サーバー PC 側", level=2)
    add_table(
        doc,
        ["項目", "要件"],
        [
            ["OS", "Windows 10 / 11、macOS、Linux のいずれか"],
            ["Python", "3.11 以上（python --version で確認）"],
            ["ディスク空き", "5 GB 以上推奨（依存パッケージ + データ蓄積用）"],
            ["メモリ", "4 GB 以上"],
            ["常時稼働", "業務時間中（または 24h）電源 ON のまま"],
            ["ネットワーク", "院内 LAN に固定 IP で接続"],
            ["ポート", "8501（TCP）を院内 LAN 内で受信可能にする"],
        ],
        col_widths_cm=[4.0, 11.0],
    )

    add_para(doc, "")
    add_heading(doc, "クライアント PC 側（病棟 PC）", level=2)
    add_table(
        doc,
        ["項目", "要件"],
        [
            ["OS", "Windows 10 / 11"],
            ["ブラウザ", "Chrome Portable（マニュアル中で配布手順を案内）"],
            ["既存環境", "電子カルテ用 Edge 90 はそのまま維持（影響なし）"],
            ["ネットワーク", "サーバー PC に院内 LAN 経由で到達可能"],
        ],
        col_widths_cm=[4.0, 11.0],
    )

    doc.add_page_break()

    # ====================================================================
    # 3. サーバーセットアップ（オフライン手順）
    # ====================================================================
    add_heading(doc, "3. サーバーセットアップ（オフライン手順）", level=1)

    add_callout(
        doc,
        "📌 前提",
        "このマニュアルは「院内 LAN は GitHub に直接アクセスできない」前提で書かれています。"
        "副院長（または SE ご自身）が外部ネット接続のある PC で zip を取得済みです。",
        color=C_ACCENT,
    )
    add_para(doc, "")

    add_step(
        doc, 1,
        "zip の展開",
        ["Teams から受け取った `bed_control_simulator_offline_2026-05-01.zip` を、",
         "サーバー PC の任意の場所（例: `C:\\BedControl\\` または `~/bed-control-simulator/`）に展開します。"],
        code="C:\\BedControl\\bed_control_simulator_offline_2026-05-01\\\n"
             "  ├─ scripts/\n"
             "  ├─ wheels/\n"
             "  └─ ...",
    )
    add_para(doc, "")

    add_step(
        doc, 2,
        "Python 仮想環境の作成",
        ["展開フォルダに移動し、Python の仮想環境を作成します。"],
        code="cd C:\\BedControl\\bed_control_simulator_offline_2026-05-01\n"
             "python -m venv .venv\n"
             "\n"
             "# Windows の場合\n"
             ".venv\\Scripts\\activate\n"
             "\n"
             "# macOS / Linux の場合\n"
             "source .venv/bin/activate",
    )
    add_para(doc, "")

    add_step(
        doc, 3,
        "依存パッケージのオフラインインストール",
        ["wheels/ フォルダに同梱された .whl ファイルから、外部ネットを使わずインストールします。"],
        code="pip install --no-index --find-links=wheels/ -r requirements.txt",
    )
    add_callout(
        doc,
        "💡 補足",
        "wheels/ には requirements.txt の全依存（pandas, streamlit, plotly 等）が含まれています。"
        "外部ネットへの DNS 解決すら発生しません。",
        color=C_OK,
    )
    add_para(doc, "")

    add_step(
        doc, 4,
        "動作確認（ローカルのみ）",
        ["まずサーバー PC 自身からアクセスできるかを確認します。"],
        code="streamlit run scripts/bed_control_simulator_app.py\n"
             "\n"
             "# ブラウザで http://localhost:8501 が開けば成功",
    )
    add_callout(
        doc,
        "✅ チェック",
        "「📋 日次データ入力」タブの最上部に「📭 空（未入力）」または「⚠️ デモデータ」"
        "のバナーが表示されていれば OK です。これは v3.5m で追加されたデータ純度チェック機能です。",
        color=C_OK,
    )
    add_para(doc, "")

    doc.add_page_break()

    # ====================================================================
    # 4. 院内 LAN 公開設定
    # ====================================================================
    add_heading(doc, "4. 院内 LAN 公開設定", level=1)

    add_step(
        doc, 1,
        "サーバー PC の IP アドレス確認",
        ["Windows: コマンドプロンプトで以下を実行。"],
        code="ipconfig\n"
             "# 「IPv4 アドレス」を確認（例: 192.168.1.100）",
    )
    add_para(doc, "（macOS / Linux の場合）")
    add_code(doc, "ifconfig | grep \"inet \"\n# またはip addr")
    add_para(doc, "")

    add_step(
        doc, 2,
        "院内 LAN 公開モードで起動",
        ["`--server.address 0.0.0.0` を付けて起動すると、院内 LAN 内のすべての PC からアクセス可能になります。"],
        code="streamlit run scripts/bed_control_simulator_app.py \\\n"
             "  --server.address 0.0.0.0 \\\n"
             "  --server.port 8501",
    )
    add_callout(
        doc,
        "💡 簡略化",
        "deploy/launch_bed_control.bat（Windows用）をダブルクリックすれば上記が自動実行されます。",
        color=C_OK,
    )
    add_para(doc, "")

    add_step(
        doc, 3,
        "ファイアウォール開放",
        ["サーバー PC の Windows Defender ファイアウォールで、TCP ポート 8501 を院内 LAN からの受信許可にします。"],
        code="新しい規則 → 受信 → ポート → TCP 8501 → 接続を許可\n"
             "プロファイル: ドメイン / プライベート（パブリックは外す）\n"
             "名前: BedControlSimulator",
    )
    add_para(doc, "")

    add_step(
        doc, 4,
        "病棟 PC からのアクセス確認",
        ["別の PC のブラウザから以下にアクセス。"],
        code="http://192.168.1.100:8501\n"
             "# ※ 192.168.1.100 はサーバー PC の実 IP に置き換える",
    )
    add_callout(
        doc,
        "⚠️ 注意",
        "アクセスできない場合は、ファイアウォール（受信規則）と --server.address 0.0.0.0 が"
        "ついているかをまず確認してください。",
        color=C_WARN,
    )

    doc.add_page_break()

    # ====================================================================
    # 5. 自動起動設定
    # ====================================================================
    add_heading(doc, "5. 自動起動設定", level=1)
    add_para(
        doc,
        "サーバー PC の電源を入れたら自動でアプリが起動するよう設定します。",
    )
    add_para(doc, "")

    add_heading(doc, "Windows: タスクスケジューラ", level=2)
    add_para(doc, "1. 「タスク スケジューラ」を起動")
    add_para(doc, "2. 「タスクの作成」をクリック")
    add_para(doc, "3. 「全般」タブ：名前 = BedControlSimulator、「ユーザーがログオンしているかどうかにかかわらず実行する」をチェック")
    add_para(doc, "4. 「トリガー」タブ：「スタートアップ時」を追加")
    add_para(doc, "5. 「操作」タブ：「プログラムの開始」")
    add_code(
        doc,
        "プログラム/スクリプト: cmd.exe\n"
        "引数の追加: /c \"C:\\BedControl\\bed_control_simulator_offline_2026-05-01\\deploy\\launch_bed_control.bat\"\n"
        "開始: C:\\BedControl\\bed_control_simulator_offline_2026-05-01\\",
    )
    add_para(doc, "")

    add_heading(doc, "macOS: launchd", level=2)
    add_para(doc, "deploy/com.omoromachi.claude-allow-watcher.example.plist を参考に、")
    add_para(doc, "~/Library/LaunchAgents/com.omc.bed-control.plist を作成し:")
    add_code(
        doc,
        "launchctl load ~/Library/LaunchAgents/com.omc.bed-control.plist",
    )
    add_para(doc, "")

    add_heading(doc, "Linux: systemd", level=2)
    add_para(doc, "/etc/systemd/system/bed-control.service を作成し:")
    add_code(
        doc,
        "[Unit]\n"
        "Description=Bed Control Simulator\n"
        "After=network.target\n"
        "\n"
        "[Service]\n"
        "Type=simple\n"
        "User=bedctrl\n"
        "WorkingDirectory=/home/bedctrl/bed-control-simulator\n"
        "ExecStart=/home/bedctrl/bed-control-simulator/.venv/bin/streamlit run scripts/bed_control_simulator_app.py --server.address 0.0.0.0 --server.port 8501\n"
        "Restart=always\n"
        "\n"
        "[Install]\n"
        "WantedBy=multi-user.target",
    )
    add_code(
        doc,
        "sudo systemctl enable bed-control\n"
        "sudo systemctl start bed-control",
    )

    doc.add_page_break()

    # ====================================================================
    # 6. Chrome Portable の配布（重要）
    # ====================================================================
    add_heading(doc, "6. Chrome Portable の配布（重要）", level=1)

    add_callout(
        doc,
        "🎯 目的",
        "院内端末の Edge 90（電子カルテ用）にはパッチを当てず、Streamlit アプリ専用に Chrome Portable を"
        "別途配布します。レジストリ変更不要・インストール不要・USB / 共有フォルダから起動可能です。",
        color=C_ACCENT,
    )
    add_para(doc, "")

    add_heading(doc, "6.1 Chrome Portable のダウンロード", level=2)
    add_para(doc, "外部ネット接続のある PC で以下にアクセスしてください:")
    add_code(
        doc,
        "https://portableapps.com/apps/internet/google_chrome_portable",
    )
    add_para(doc, "")
    add_para(doc, "「Download from PortableApps.com」ボタンからインストーラ（.paf.exe）を取得します。")
    add_callout(
        doc,
        "💡 補足",
        "PortableApps.com は MIT ライセンスのオープンソースランチャーで、"
        "Chrome 公式ビルドを「インストールせず」フォルダに展開する仕組みを提供しています。",
        color=C_OK,
    )
    add_para(doc, "")

    add_heading(doc, "6.2 インストーラの展開", level=2)
    add_para(doc, "ダウンロードした .paf.exe を実行し、展開先を指定します。")
    add_table(
        doc,
        ["展開方式", "展開先パス", "用途"],
        [
            ["共有フォルダ展開（推奨）", "\\\\server\\shared\\PortableApps\\GoogleChromePortable\\", "全端末から共通利用（更新も 1 箇所）"],
            ["端末ローカル展開", "C:\\PortableApps\\GoogleChromePortable\\", "ネットワーク障害時のフォールバック"],
            ["デスクトップ展開", "%USERPROFILE%\\Desktop\\GoogleChromePortable\\", "個別端末の試用"],
        ],
        col_widths_cm=[4.0, 7.0, 4.0],
    )
    add_para(doc, "")
    add_callout(
        doc,
        "✅ 推奨",
        "共有フォルダ + 端末ローカルの両方に置くと最も堅牢です。"
        "deploy/launch_portable_chrome.bat は共有 → ローカル → デスクトップの順で自動探索します。",
        color=C_OK,
    )
    add_para(doc, "")

    add_heading(doc, "6.3 Edge との関係", level=2)
    add_para(doc, "・電子カルテ用 Edge 90 はそのまま維持されます")
    add_para(doc, "・Chrome Portable はレジストリを書きません（プロファイルは Chrome のフォルダ内に保存）")
    add_para(doc, "・両ブラウザを同時に開いて使用しても問題ありません")

    doc.add_page_break()

    # ====================================================================
    # 7. クライアント端末セットアップ
    # ====================================================================
    add_heading(doc, "7. クライアント端末セットアップ", level=1)

    add_step(
        doc, 1,
        "launch_portable_chrome.bat の編集",
        ["zip 内の deploy/launch_portable_chrome.bat をテキストエディタで開き、",
         "下記の SERVER_URL の IP をサーバー PC の実際の IP に書き換えます。"],
        code="SET SERVER_URL=http://192.168.1.100:8501\n"
             "                  ↑ ここをサーバー PC の実 IP に変更",
    )
    add_para(doc, "")

    add_step(
        doc, 2,
        "ショートカットの作成",
        ["編集後の launch_portable_chrome.bat を端末のデスクトップにコピー、",
         "または右クリック → 「ショートカットの作成」を選択。",
         "アイコン名を「ベッドコントロール」に変更してください。"],
    )
    add_para(doc, "")

    add_step(
        doc, 3,
        "起動確認",
        ["デスクトップの「ベッドコントロール」ショートカットをダブルクリックします。",
         "Chrome Portable が起動し、自動的にアプリの URL が開けば成功です。"],
    )
    add_callout(
        doc,
        "⚠️ 起こりうるエラー",
        "「Chrome Portable が見つかりません」と表示された場合、bat 内のパスを実際のフォルダ位置に書き換えてください。",
        color=C_WARN,
    )

    doc.add_page_break()

    # ====================================================================
    # 8. ブラウザ互換チェック
    # ====================================================================
    add_heading(doc, "8. ブラウザ互換チェック", level=1)
    add_para(
        doc,
        "tools/browser_probe.html は、現在使用しているブラウザが Streamlit に必要な"
        "ECMAScript / WebAssembly / fetch API などをサポートしているか自動判定する HTML ツールです。",
    )
    add_para(doc, "使い方：")
    add_para(doc, "1. tools/browser_probe.html を端末上で右クリック → Chrome Portable で開く")
    add_para(doc, "2. 各項目に ✓（緑）が表示されていれば Streamlit が動作します")
    add_para(doc, "3. ✗（赤）が表示された項目はブラウザを別バージョンに変更してください")
    add_para(doc, "")
    add_callout(
        doc,
        "💡 ヒント",
        "電子カルテ用 Edge 90 でこの HTML を開くと、いくつかの項目で ✗ が出ます。"
        "それは想定どおりで、その端末ではアプリ用に Chrome Portable を使用してください。",
        color=C_ACCENT,
    )

    doc.add_page_break()

    # ====================================================================
    # 9. データ初期化（運用開始時）
    # ====================================================================
    add_heading(doc, "9. データ初期化（運用開始時）", level=1)
    add_para(
        doc,
        "本パッケージにはデモデータ（A医師〜J医師の架空データ）が含まれている可能性があります。"
        "本番運用開始前に必ず初期化してください。アプリ内 UI から安全に実行できます。",
    )
    add_para(doc, "")

    add_step(
        doc, 1,
        "「📋 日次データ入力」タブを開く",
        ["サイドバーから ⚙️ データ・設定 → 📋 日次データ入力 を選択します。"],
    )
    add_para(doc, "")

    add_step(
        doc, 2,
        "データ純度バナーを確認",
        ["タブ最上部に下記のいずれかが表示されます：",
         "・🔴 デモと実の混在 → 初期化が必要",
         "・⚠️ デモデータ → 初期化が必要",
         "・📭 空（未入力） → 初期化済み（OK）",
         "・✅ 実データ → 既に運用中（OK）"],
    )
    add_para(doc, "")

    add_step(
        doc, 3,
        "本番初期化を実行",
        ["「🛠 本番初期化（デモデータを退避 → 空スキーマを作成）」エキスパンダーを展開。",
         "チェックボックスで同意 → 「📦 本番初期化を実行」ボタンをクリック。",
         "退避ファイル data/archive/admission_details_demo_YYYYMMDD_HHMMSS.csv が作成されます。"],
    )
    add_callout(
        doc,
        "✅ 完了",
        "実行後、バナーが「📭 空（未入力）」に変わります。これで本番運用の開始準備が整いました。",
        color=C_OK,
    )

    doc.add_page_break()

    # ====================================================================
    # 10. バックアップ運用
    # ====================================================================
    add_heading(doc, "10. バックアップ運用", level=1)
    add_para(doc, "下記の 3 フォルダを定期的に院内バックアップサーバーへコピーしてください。")
    add_para(doc, "")
    add_table(
        doc,
        ["フォルダ", "内容", "重要度"],
        [
            ["data/", "日次入院データ（admission_details.csv）など", "★★★ 最重要"],
            ["settings/", "手動シード YAML（救急15% / 看護必要度）", "★★★ 副院長手入力データ"],
            ["data/archive/", "デモ退避ファイル / 過去のデータスナップショット", "★★ 履歴用"],
        ],
        col_widths_cm=[4.0, 8.0, 3.0],
    )
    add_para(doc, "")

    add_heading(doc, "Windows タスクスケジューラ例（毎日 22:00 自動コピー）", level=2)
    add_code(
        doc,
        "@echo off\n"
        "SET SRC=C:\\BedControl\\bed_control_simulator_offline_2026-05-01\n"
        "SET DST=\\\\backup-server\\share\\BedControl\\%date:~0,4%%date:~5,2%%date:~8,2%\n"
        "\n"
        "robocopy \"%SRC%\\data\"     \"%DST%\\data\"     /E\n"
        "robocopy \"%SRC%\\settings\" \"%DST%\\settings\" /E",
    )

    add_heading(doc, "macOS / Linux cron 例", level=2)
    add_code(
        doc,
        "0 22 * * * /bin/cp -r ~/bed-control-simulator/data /var/backup/bedcontrol/$(date +%Y%m%d)/\n"
        "5 22 * * * /bin/cp -r ~/bed-control-simulator/settings /var/backup/bedcontrol/$(date +%Y%m%d)/",
    )

    doc.add_page_break()

    # ====================================================================
    # 11. アップデート方法
    # ====================================================================
    add_heading(doc, "11. アップデート方法", level=1)
    add_para(
        doc,
        "院内 LAN は GitHub に直接アクセスできない前提のため、",
        bold=True,
    )
    add_para(
        doc,
        "副院長が外部ネット PC で更新版 zip を作成し、Teams で SE に再配布する運用とします。",
    )
    add_para(doc, "")

    add_heading(doc, "更新時の手順", level=2)
    add_para(doc, "1. 副院長が新しい zip（例: bed_control_simulator_offline_2026-MM-DD.zip）を Teams で送付")
    add_para(doc, "2. SE は **データを必ずバックアップ**（10. 参照）")
    add_para(doc, "3. 新 zip を別フォルダに展開")
    add_para(doc, "4. 旧フォルダの data/ と settings/ を新フォルダにコピー（データ引き継ぎ）")
    add_para(doc, "5. 新フォルダで venv 再作成 + pip install --no-index --find-links=wheels/")
    add_para(doc, "6. Streamlit を再起動して動作確認")
    add_para(doc, "")
    add_callout(
        doc,
        "⚠️ 重要",
        "data/ と settings/ を必ず引き継いでください。これを忘れると入院データと手動シードが失われます。",
        color=C_WARN,
    )

    doc.add_page_break()

    # ====================================================================
    # 12. トラブルシューティング
    # ====================================================================
    add_heading(doc, "12. トラブルシューティング", level=1)

    add_table(
        doc,
        ["症状", "対処"],
        [
            ["pip install で「No matching distribution」エラー",
             "wheels/ 内の .whl ファイルが requirements.txt と一致しているか確認。\n"
             "pip install --no-index --find-links=wheels/ -r requirements.txt の構文を確認。"],
            ["streamlit コマンドが見つからない",
             "venv の activate を実行しているか確認（source .venv/bin/activate or .venv\\Scripts\\activate）。"],
            ["http://localhost:8501 は開けるが、他 PC から開けない",
             "--server.address 0.0.0.0 が起動コマンドにあるか確認。\n"
             "Windows Defender ファイアウォールでポート 8501 が許可されているか確認。"],
            ["アプリは開くがアクセスが遅い",
             "Chrome Portable を共有フォルダではなく端末ローカルに展開すると改善することが多い。"],
            ["データ純度バナーが「🔴 デモと実の混在」",
             "9. データ初期化を実行 → デモデータを退避してから運用開始。"],
            ["ポート 8501 が他で使われている",
             "--server.port 8502 等で別ポートを指定（クライアント側 bat の URL も合わせて変更）。"],
            ["Python 3.10 以下しか入っていない",
             "Python 3.11 以上をインストール（python.org から取得 or winget install Python.Python.3.11）。"],
            ["Chrome Portable で日本語が文字化け",
             "Chrome Portable のフォルダ内に日本語フォントが認識されているか確認。"
             "OS 標準フォント（メイリオ等）があれば通常問題なし。"],
        ],
        col_widths_cm=[5.5, 9.5],
    )

    doc.add_page_break()

    # ====================================================================
    # 13. 副院長への引渡し前チェックリスト
    # ====================================================================
    add_heading(doc, "13. 副院長への引渡し前チェックリスト", level=1)
    add_para(doc, "下記 5 項目をすべて確認してから副院長へ「設置完了」のご連絡をお願いします。")
    add_para(doc, "")

    add_heading(doc, "✅ チェック 1: サーバー PC でアプリが起動する", level=2)
    add_code(doc, "streamlit run scripts/bed_control_simulator_app.py")
    add_para(doc, "→ ローカルから http://localhost:8501 が開ける")
    add_para(doc, "")

    add_heading(doc, "✅ チェック 2: 院内 LAN 公開モードでも起動する", level=2)
    add_code(doc, "streamlit run scripts/bed_control_simulator_app.py --server.address 0.0.0.0 --server.port 8501")
    add_para(doc, "→ 別 PC のブラウザから http://(サーバーIP):8501 が開ける")
    add_para(doc, "")

    add_heading(doc, "✅ チェック 3: Chrome Portable で開ける", level=2)
    add_para(doc, "→ deploy/launch_portable_chrome.bat ダブルクリックで Chrome Portable が起動し、アプリ画面が表示される")
    add_para(doc, "")

    add_heading(doc, "✅ チェック 4: データ純度バナーが「📭 空」または「✅ 実データ」", level=2)
    add_para(doc, "→ アプリの「⚙️ データ・設定 > 📋 日次データ入力」タブの最上部を確認")
    add_para(doc, "→ デモデータが残っていれば 9. の本番初期化を実行")
    add_para(doc, "")

    add_heading(doc, "✅ チェック 5: 自動起動設定が完了", level=2)
    add_para(doc, "→ サーバー PC を再起動 → ログイン後 1 分以内に Streamlit サーバーが起動している")
    add_para(doc, "")

    doc.add_page_break()

    # ====================================================================
    # 14. 連絡先 / FAQ
    # ====================================================================
    add_heading(doc, "14. 連絡先 / FAQ", level=1)
    add_para(doc, "ご不明点があれば下記まで遠慮なくご連絡ください。")
    add_para(doc, "")
    add_table(
        doc,
        ["項目", "内容"],
        [
            ["責任者", "副院長 久保田 透（おもろまちメディカルセンター）"],
            ["連絡方法", "Microsoft Teams（直接メッセージ可）"],
            ["緊急時", "院内 PHS"],
            ["GitHub リポジトリ", "torukubota2023/bed-control-simulator (private)"],
            ["参考ドキュメント", "本パッケージ内 docs/admin/pre_lan_deployment_checklist.md"],
        ],
        col_widths_cm=[4.0, 11.0],
    )
    add_para(doc, "")

    add_separator(doc)
    add_para(doc, "")
    add_para(
        doc,
        "本マニュアルは scripts/generate_se_install_manual.py から自動生成されています。"
        "副院長による更新時は同スクリプトを編集 → 再実行してください。",
        size=10,
        color=C_FAINT,
        align=WD_ALIGN_PARAGRAPH.CENTER,
    )

    # 保存
    output_dir = Path(__file__).resolve().parent.parent / "docs" / "admin"
    output_dir.mkdir(parents=True, exist_ok=True)
    output_path = output_dir / "SE_install_manual_2026-05-01.docx"
    doc.save(output_path)
    print(f"✅ Generated: {output_path}")
    return output_path


if __name__ == "__main__":
    build()
