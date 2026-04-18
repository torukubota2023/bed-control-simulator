#!/usr/bin/env python3
"""
ベッドコントロールシミュレーター Streamlit Community Cloud デプロイマニュアル生成スクリプト
出力: /Users/torukubota/ai-management/docs/admin/streamlit_cloud_deploy_manual.docx
"""

from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

def set_cell_shading(cell, color):
    shading_elm = OxmlElement('w:shd')
    shading_elm.set(qn('w:fill'), color)
    shading_elm.set(qn('w:val'), 'clear')
    cell._tc.get_or_add_tcPr().append(shading_elm)

def add_page_number(doc):
    """フッターにページ番号を追加"""
    for section in doc.sections:
        footer = section.footer
        footer.is_linked_to_previous = False
        paragraph = footer.paragraphs[0]
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = paragraph.add_run()
        fldChar1 = OxmlElement('w:fldChar')
        fldChar1.set(qn('w:fldCharType'), 'begin')
        run._r.append(fldChar1)
        run2 = paragraph.add_run()
        instrText = OxmlElement('w:instrText')
        instrText.set(qn('xml:space'), 'preserve')
        instrText.text = ' PAGE '
        run2._r.append(instrText)
        run3 = paragraph.add_run()
        fldChar2 = OxmlElement('w:fldChar')
        fldChar2.set(qn('w:fldCharType'), 'end')
        run3._r.append(fldChar2)
        for r in paragraph.runs:
            r.font.size = Pt(9)
            r.font.name = 'Yu Gothic'

def set_run_font(run, font_name='Yu Gothic', size=None, bold=False, italic=False, color=None):
    run.font.name = font_name
    r = run._element
    r.rPr.rFonts.set(qn('w:eastAsia'), font_name)
    if size:
        run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    if color:
        run.font.color.rgb = RGBColor(*color)

def add_heading_styled(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.name = 'Yu Gothic'
        r = run._element
        r.rPr.rFonts.set(qn('w:eastAsia'), 'Yu Gothic')
    return h

def add_para(doc, text, font_name='Yu Gothic', size=10.5, bold=False, alignment=None, space_after=6, first_line_indent=None):
    p = doc.add_paragraph()
    if alignment:
        p.alignment = alignment
    p.paragraph_format.space_after = Pt(space_after)
    if first_line_indent:
        p.paragraph_format.first_line_indent = Cm(first_line_indent)
    run = p.add_run(text)
    set_run_font(run, font_name, size, bold)
    return p

def add_para_mixed(doc, parts, size=10.5, space_after=6):
    """parts: list of (text, bold, font_name_or_None, color_or_None)"""
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(space_after)
    for text, bold, font_override, color in parts:
        run = p.add_run(text)
        fn = font_override if font_override else 'Yu Gothic'
        set_run_font(run, fn, size, bold, color=color)
    return p

def add_code_block(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.left_indent = Cm(1)
    # Add shading via XML
    pPr = p._element.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:fill'), 'F0F0F0')
    pPr.append(shd)
    run = p.add_run(text)
    set_run_font(run, 'Courier New', 9.5)
    return p

def add_bullet(doc, text, level=0, bold_prefix=None):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.space_after = Pt(3)
    if level > 0:
        p.paragraph_format.left_indent = Cm(1.5 * (level + 1))
    if bold_prefix:
        run_b = p.add_run(bold_prefix)
        set_run_font(run_b, 'Yu Gothic', 10.5, bold=True)
        run_t = p.add_run(text)
        set_run_font(run_t, 'Yu Gothic', 10.5)
    else:
        # Clear default run and add styled one
        p.clear()
        run = p.add_run(text)
        set_run_font(run, 'Yu Gothic', 10.5)
    return p

def add_numbered_item(doc, number, text, sub_items=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.left_indent = Cm(1)
    p.paragraph_format.hanging_indent = Cm(0.8)
    run = p.add_run(f"{number}. ")
    set_run_font(run, 'Yu Gothic', 10.5, bold=True)
    run2 = p.add_run(text)
    set_run_font(run2, 'Yu Gothic', 10.5)
    if sub_items:
        for item in sub_items:
            ps = doc.add_paragraph()
            ps.paragraph_format.space_after = Pt(2)
            ps.paragraph_format.left_indent = Cm(2)
            rs = ps.add_run(f"・{item}")
            set_run_font(rs, 'Yu Gothic', 10)
    return p

def create_table(doc, headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    # Header row
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ''
        p = cell.paragraphs[0]
        run = p.add_run(header)
        set_run_font(run, 'Yu Gothic', 10, bold=True, color=(255, 255, 255))
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_cell_shading(cell, '4472C4')
    # Data rows
    for r_idx, row in enumerate(rows):
        for c_idx, val in enumerate(row):
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = ''
            p = cell.paragraphs[0]
            run = p.add_run(val)
            set_run_font(run, 'Yu Gothic', 10)
            if r_idx % 2 == 1:
                set_cell_shading(cell, 'D9E2F3')
    return table


def main():
    doc = Document()

    # --- ページ設定 ---
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

    # --- デフォルトフォント ---
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Yu Gothic'
    font.size = Pt(10.5)
    style.element.rPr.rFonts.set(qn('w:eastAsia'), 'Yu Gothic')

    # Heading styles
    for i in range(1, 4):
        hs = doc.styles[f'Heading {i}']
        hs.font.name = 'Yu Gothic'
        hs.font.bold = True
        hs.element.rPr.rFonts.set(qn('w:eastAsia'), 'Yu Gothic')
        if i == 1:
            hs.font.size = Pt(18)
            hs.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)
        elif i == 2:
            hs.font.size = Pt(14)
            hs.font.color.rgb = RGBColor(0x2E, 0x74, 0xB5)
        else:
            hs.font.size = Pt(12)
            hs.font.color.rgb = RGBColor(0x2E, 0x74, 0xB5)

    # ===========================
    # 表紙
    # ===========================
    for _ in range(6):
        doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('ベッドコントロールシミュレーター')
    set_run_font(run, 'Yu Gothic', 26, bold=True, color=(0x1F, 0x49, 0x7D))

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('Streamlit Community Cloud\nデプロイマニュアル')
    set_run_font(run, 'Yu Gothic', 22, bold=True, color=(0x2E, 0x74, 0xB5))

    doc.add_paragraph()
    doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('URLひとつで院内外の誰でもアクセスできるようにする方法')
    set_run_font(run, 'Yu Gothic', 14, color=(0x59, 0x56, 0x59))

    for _ in range(4):
        doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('おもろまちメディカルセンター')
    set_run_font(run, 'Yu Gothic', 12)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('作成日：2026年3月25日')
    set_run_font(run, 'Yu Gothic', 12)

    doc.add_page_break()

    # ===========================
    # 目次（簡易）
    # ===========================
    add_heading_styled(doc, '目次', level=1)
    toc_items = [
        '1. はじめに',
        '   1-1. Streamlit Community Cloud とは',
        '   1-2. LAN共有との違い',
        '2. 全体の流れ（概要）',
        '3. ステップ1：GitHubアカウントの作成',
        '   3-1. GitHubとは',
        '   3-2. アカウント作成手順',
        '4. ステップ2：GitHubにコードをアップロード',
        '   4-1. 方法A：GitHub上で直接作成（初心者向け）',
        '   4-2. 方法B：ターミナルからgit push（経験者向け）',
        '5. ステップ3：Streamlit Community Cloudでデプロイ',
        '   5-1. Streamlit Cloudにアクセス',
        '   5-2. 新しいアプリをデプロイ',
        '   5-3. デプロイ完了を待つ',
        '   5-4. URLを共有',
        '6. 重要：ファイル配置の注意点',
        '7. アプリの更新方法',
        '8. トラブルシューティング',
        '9. セキュリティに関する注意',
        '10. よくある質問（FAQ）',
        '付録：クイックリファレンスカード',
    ]
    for item in toc_items:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(2)
        run = p.add_run(item)
        set_run_font(run, 'Yu Gothic', 10.5)

    doc.add_page_break()

    # ===========================
    # 1. はじめに
    # ===========================
    add_heading_styled(doc, '1. はじめに', level=1)

    add_heading_styled(doc, '1-1. Streamlit Community Cloud とは', level=2)
    add_para(doc, 'Streamlit Community Cloud は、Pythonで作られたアプリケーションをインターネット上に無料で公開できるクラウドサービスです。')
    add_para(doc, 'このマニュアルに従って作業を進めると、以下のことが実現できます：')

    bullets = [
        'https://xxxxx.streamlit.app のようなURLが発行される',
        'インターネットに接続できる誰でも、ブラウザでベッドコントロールシミュレーターにアクセスできるようになる',
        'ホストPCの電源が入っていなくても24時間アクセス可能',
        '完全無料（2026年3月時点）',
    ]
    for b in bullets:
        add_bullet(doc, b)

    doc.add_paragraph()
    p = add_para(doc, '', space_after=8)
    run = p.add_run('【重要】患者個人情報は絶対に含めないでください。')
    set_run_font(run, 'Yu Gothic', 10.5, bold=True, color=(0xC0, 0x00, 0x00))
    run2 = p.add_run('本シミュレーターはパラメータのみを扱うため、個人情報の問題はありません。')
    set_run_font(run2, 'Yu Gothic', 10.5)

    # --- LAN共有との違い ---
    add_heading_styled(doc, '1-2. LAN共有との違い', level=2)
    add_para(doc, '社内LAN共有（streamlit run コマンドでの共有）と Streamlit Community Cloud の違いを以下にまとめます。')

    headers = ['項目', '社内LAN共有', 'Streamlit Community Cloud']
    rows = [
        ['アクセス範囲', '院内ネットワークのみ', 'インターネット上どこでも'],
        ['ホストPC', '必要（起動中のみ）', '不要（クラウドが24時間稼働）'],
        ['コスト', '無料', '無料'],
        ['準備の手間', '少ない', 'GitHubアカウント作成が必要'],
        ['セキュリティ', '高い（院内限定）', 'URLを知っている人はアクセス可'],
    ]
    create_table(doc, headers, rows)

    doc.add_page_break()

    # ===========================
    # 2. 全体の流れ
    # ===========================
    add_heading_styled(doc, '2. 全体の流れ（概要）', level=1)
    add_para(doc, 'デプロイは以下の3ステップで完了します。合計約20分の作業です。')
    doc.add_paragraph()

    # ステップ図（テーブルで表現）
    step_table = doc.add_table(rows=1, cols=5)
    step_table.alignment = WD_TABLE_ALIGNMENT.CENTER

    steps_data = [
        ('ステップ1\n\nGitHubアカウント\nを作る\n\n【約5分】', '4472C4'),
        ('→', None),
        ('ステップ2\n\nGitHubにコード\nをアップロード\n\n【約10分】', '2E74B5'),
        ('→', None),
        ('ステップ3\n\nStreamlit Cloud\nでデプロイ\n\n【約5分】', '1F497D'),
    ]
    for i, (text, color) in enumerate(steps_data):
        cell = step_table.rows[0].cells[i]
        cell.text = ''
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(text)
        if color:
            set_run_font(run, 'Yu Gothic', 11, bold=True, color=(255, 255, 255))
            set_cell_shading(cell, color)
        else:
            set_run_font(run, 'Yu Gothic', 16, bold=True, color=(0x2E, 0x74, 0xB5))

    doc.add_paragraph()
    add_para(doc, '各ステップの詳細を次章以降で説明します。すでにGitHubアカウントをお持ちの場合は、ステップ1をスキップできます。')

    doc.add_page_break()

    # ===========================
    # 3. ステップ1: GitHub
    # ===========================
    add_heading_styled(doc, '3. ステップ1：GitHubアカウントの作成', level=1)

    add_heading_styled(doc, '3-1. GitHubとは', level=2)
    add_para(doc, 'GitHub（ギットハブ）は、プログラムのコードを保管・管理するための無料サービスです。世界中の開発者が使っているクラウド上の「コード倉庫」と考えてください。')
    add_para(doc, 'Streamlit Community Cloud にアプリを公開するには、まずコードをGitHubにアップロードする必要があります。そのため、GitHubアカウントが必要です。')

    p = add_para(doc, '')
    run = p.add_run('※すでにGitHubアカウントをお持ちの場合は、このステップをスキップして「4. ステップ2」に進んでください。')
    set_run_font(run, 'Yu Gothic', 10, italic=True, color=(0x59, 0x56, 0x59))

    add_heading_styled(doc, '3-2. アカウント作成手順', level=2)

    github_steps = [
        ('ブラウザで https://github.com にアクセスします。', 'Safari、Google Chrome など、お使いのブラウザで開いてください。'),
        ('画面右上の「Sign up」ボタンをクリックします。', '「Sign in」ではなく「Sign up」（新規登録）です。'),
        ('メールアドレスを入力し、「Continue」をクリックします。', '普段お使いのメールアドレスを入力してください。確認メールが届きます。'),
        ('パスワードを設定し、「Continue」をクリックします。', '15文字以上、または8文字以上で数字と小文字を含む必要があります。忘れないようにメモしてください。'),
        ('ユーザー名を入力し、「Continue」をクリックします。', '英数字とハイフン（-）のみ使用可能です。例：omoromachi-med'),
        ('メール通知の設定が表示されます。「n」を入力してスキップし、「Continue」をクリックします。', '後から設定を変更することも可能です。'),
        ('パズル認証（CAPTCHA）が表示されるので、指示に従って完了させます。', 'ロボットでないことを確認するための簡単なパズルです。'),
        ('「Create account」をクリックします。', 'アカウントの作成が開始されます。'),
        ('入力したメールアドレスに確認コード（6桁の数字）が届くので、入力します。', 'メールが届かない場合は、迷惑メールフォルダも確認してください。'),
        ('簡単なアンケートが表示されます（スキップ可能）。', '右下の「Skip personalization」でスキップできます。'),
        ('プランの選択画面で「Free」（無料）を選択し、「Continue」をクリックします。', '無料プランで十分です。有料プランを選ぶ必要はありません。'),
        ('ダッシュボード画面が表示されれば、アカウント作成完了です。', 'おめでとうございます！次のステップに進みましょう。'),
    ]

    for i, (step, note) in enumerate(github_steps, 1):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(2)
        p.paragraph_format.left_indent = Cm(1)
        p.paragraph_format.hanging_indent = Cm(0.8)
        run = p.add_run(f'{i}. ')
        set_run_font(run, 'Yu Gothic', 10.5, bold=True)
        run2 = p.add_run(step)
        set_run_font(run2, 'Yu Gothic', 10.5)

        if note:
            pn = doc.add_paragraph()
            pn.paragraph_format.space_after = Pt(6)
            pn.paragraph_format.left_indent = Cm(1.8)
            run_n = pn.add_run(f'💡 {note}')
            set_run_font(run_n, 'Yu Gothic', 9.5, color=(0x59, 0x56, 0x59))

    doc.add_page_break()

    # ===========================
    # 4. ステップ2: コードアップロード
    # ===========================
    add_heading_styled(doc, '4. ステップ2：GitHubにコードをアップロード', level=1)
    add_para(doc, 'コードをGitHubにアップロードする方法は2通りあります。ITに詳しくない方は「方法A」をおすすめします。')

    # --- 方法A ---
    add_heading_styled(doc, '4-1. 方法A：GitHub上で直接作成（初心者向け・おすすめ）', level=2)
    add_para(doc, 'ターミナル（黒い画面）を使わずに、ブラウザ上の操作だけでコードをアップロードする方法です。')

    add_heading_styled(doc, 'リポジトリ（保管場所）の作成', level=3)
    repo_steps = [
        'GitHubにログインした状態で、画面右上の「+」ボタンをクリックし、「New repository」を選択します。',
        'Repository name 欄に bed-control-simulator と入力します。',
        'Description 欄に「ベッドコントロールシミュレーター」と入力します（任意）。',
        '「Public」が選択されていることを確認します。',
        '「Add a README file」にチェックを入れます。',
        '「Create repository」をクリックします。',
    ]
    for i, step in enumerate(repo_steps, 1):
        add_numbered_item(doc, i, step)

    p = add_para(doc, '')
    run = p.add_run('【重要】')
    set_run_font(run, 'Yu Gothic', 10.5, bold=True, color=(0xC0, 0x00, 0x00))
    run2 = p.add_run('「Public」を選択してください。Streamlit Community Cloud の無料プランは Public リポジトリのみに対応しています。')
    set_run_font(run2, 'Yu Gothic', 10.5)

    add_heading_styled(doc, 'ファイルのアップロード', level=3)
    upload_steps = [
        '作成されたリポジトリのページで、「Add file」ボタンをクリックし、「Upload files」を選択します。',
        'Finder（ファイル管理アプリ）を開き、以下の3つのファイルを見つけます：',
        'Finderからブラウザの「Drag additional files here」エリアに、3つのファイルをドラッグ＆ドロップします。',
        '3つのファイルが表示されていることを確認します。',
        '画面下部の「Commit changes」ボタンをクリックします。',
        'リポジトリのページに戻り、3つのファイルが表示されていれば成功です。',
    ]
    for i, step in enumerate(upload_steps, 7):
        add_numbered_item(doc, i, step)
        if i == 8:
            # ファイル一覧
            files_info = [
                ('bed_control_simulator.py', '~/ai-management/scripts/bed_control_simulator.py'),
                ('bed_control_simulator_app.py', '~/ai-management/scripts/bed_control_simulator_app.py'),
                ('requirements.txt', '~/ai-management/requirements.txt'),
            ]
            for fname, fpath in files_info:
                p = doc.add_paragraph()
                p.paragraph_format.space_after = Pt(2)
                p.paragraph_format.left_indent = Cm(2.5)
                run = p.add_run(f'・{fname}')
                set_run_font(run, 'Courier New', 10, bold=True)
                p2 = doc.add_paragraph()
                p2.paragraph_format.space_after = Pt(4)
                p2.paragraph_format.left_indent = Cm(3)
                run2 = p2.add_run(f'場所：{fpath}')
                set_run_font(run2, 'Yu Gothic', 9.5, color=(0x59, 0x56, 0x59))

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    pPr = p._element.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:fill'), 'FFF3CD')
    pPr.append(shd)
    run = p.add_run('💡 Finderでファイルの場所がわからない場合：Finderのメニューバーで「移動」→「フォルダへ移動」を選択し、~/ai-management/scripts と入力してEnterキーを押すと、該当フォルダが開きます。')
    set_run_font(run, 'Yu Gothic', 10)

    doc.add_page_break()

    # --- 方法B ---
    add_heading_styled(doc, '4-2. 方法B：ターミナルからgit pushする方法（経験者向け）', level=2)
    add_para(doc, 'ターミナル（コマンドライン）の操作に慣れている方向けの方法です。初心者の方は方法Aをご利用ください。')

    add_heading_styled(doc, '手順', level=3)
    add_para(doc, 'ターミナルを開き（Launchpad →「ターミナル」、またはSpotlightで「ターミナル」と検索）、以下のコマンドを1行ずつ入力してEnterキーを押します。')

    commands = [
        ('cd ~/ai-management', 'ai-managementフォルダに移動します。'),
        ('git init', 'Gitの管理を開始します（初回のみ必要）。'),
        ('git add scripts/bed_control_simulator.py scripts/bed_control_simulator_app.py requirements.txt',
         'アップロードするファイルを選択します。'),
        ('git commit -m "ベッドコントロールシミュレーター初回コミット"', '変更内容を記録します。'),
        ('git remote add origin https://github.com/ユーザー名/bed-control-simulator.git',
         'GitHubのリポジトリと紐づけます。「ユーザー名」の部分はご自身のGitHubユーザー名に置き換えてください。'),
        ('git branch -M main', 'ブランチ名をmainに設定します。'),
        ('git push -u origin main', 'GitHubにコードをアップロードします。初回はGitHubのログイン情報を求められます。'),
    ]

    for cmd, explanation in commands:
        add_code_block(doc, cmd)
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(6)
        p.paragraph_format.left_indent = Cm(1)
        run = p.add_run(f'↑ {explanation}')
        set_run_font(run, 'Yu Gothic', 9.5, color=(0x59, 0x56, 0x59))

    p = add_para(doc, '')
    run = p.add_run('※注意：')
    set_run_font(run, 'Yu Gothic', 10.5, bold=True)
    run2 = p.add_run('方法Bではファイルが scripts/ フォルダ内に配置されます。ステップ3でMain file pathを設定する際に scripts/bed_control_simulator_app.py と指定する必要があります。また、requirements.txt をリポジトリのルートにも配置する必要がある場合があります（詳細は「6. ファイル配置の注意点」を参照）。')
    set_run_font(run2, 'Yu Gothic', 10.5)

    doc.add_page_break()

    # ===========================
    # 5. ステップ3: デプロイ
    # ===========================
    add_heading_styled(doc, '5. ステップ3：Streamlit Community Cloudでデプロイ', level=1)

    add_heading_styled(doc, '5-1. Streamlit Cloudにアクセス', level=2)
    deploy_access = [
        ('ブラウザで https://share.streamlit.io にアクセスします。', None),
        ('「Continue with GitHub」をクリックします。', 'GitHubアカウントでログインする方法です。'),
        ('GitHubアカウントのメールアドレスとパスワードを入力してログインします。', '初回はGitHubとの連携許可が求められます。「Authorize streamlit」をクリックして許可してください。'),
    ]
    for i, (step, note) in enumerate(deploy_access, 1):
        add_numbered_item(doc, i, step)
        if note:
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(6)
            p.paragraph_format.left_indent = Cm(1.8)
            run = p.add_run(f'💡 {note}')
            set_run_font(run, 'Yu Gothic', 9.5, color=(0x59, 0x56, 0x59))

    add_heading_styled(doc, '5-2. 新しいアプリをデプロイ', level=2)
    deploy_steps = [
        '画面右上の「New app」ボタンをクリックします。',
        '以下の項目を設定します：',
    ]
    for i, step in enumerate(deploy_steps, 4):
        add_numbered_item(doc, i, step)

    # 設定項目テーブル
    settings_headers = ['設定項目', '入力内容']
    settings_rows = [
        ['Repository', 'ユーザー名/bed-control-simulator を選択'],
        ['Branch', 'main'],
        ['Main file path', 'bed_control_simulator_app.py\n（方法Bの場合: scripts/bed_control_simulator_app.py）'],
    ]
    create_table(doc, settings_headers, settings_rows)
    doc.add_paragraph()

    more_deploy = [
        '「Advanced settings」をクリックし、Python version を 3.11（または最新版）に設定します。その他の設定はデフォルトのままで問題ありません。',
        '「Deploy!」ボタンをクリックします。',
    ]
    for i, step in enumerate(more_deploy, 6):
        add_numbered_item(doc, i, step)

    add_heading_styled(doc, '5-3. デプロイ完了を待つ', level=2)
    add_para(doc, '「Deploy!」をクリックすると、以下の流れでデプロイが進みます：')
    wait_items = [
        '画面に「Your app is in the oven」と表示され、デプロイ処理が開始されます。',
        '通常2〜5分で完了します。そのまま画面を閉じずにお待ちください。',
        'デプロイが完了すると、自動的にアプリの画面に遷移します。',
        'ベッドコントロールシミュレーターの画面が表示されれば成功です！',
    ]
    for item in wait_items:
        add_bullet(doc, item)

    doc.add_paragraph()
    add_para(doc, '発行されるURLの形式：', bold=True)
    add_code_block(doc, 'https://ユーザー名-bed-control-simulator-ファイル名-xxxxxx.streamlit.app')

    add_heading_styled(doc, '5-4. URLを共有する', level=2)
    share_steps = [
        'ブラウザのアドレスバーに表示されているURLをコピーします（アドレスバーをクリック → Command+A → Command+C）。',
        'メール、チャット、LINEなどで共有します。',
        '共有されたURLにアクセスするだけで、誰でもブラウザ上でシミュレーターを利用できます。',
    ]
    for item in share_steps:
        add_bullet(doc, item)

    doc.add_page_break()

    # ===========================
    # 6. ファイル配置の注意点
    # ===========================
    add_heading_styled(doc, '6. 重要：ファイル配置の注意点', level=1)

    p = add_para(doc, '')
    run = p.add_run('この章は重要です。')
    set_run_font(run, 'Yu Gothic', 10.5, bold=True, color=(0xC0, 0x00, 0x00))
    run2 = p.add_run('デプロイ時にエラーが出る場合、ファイル配置が原因であることが多いです。')
    set_run_font(run2, 'Yu Gothic', 10.5)

    add_heading_styled(doc, 'requirements.txt の配置', level=2)
    add_para(doc, 'Streamlit Community Cloud は、リポジトリのルート（最上位フォルダ）にある requirements.txt を自動的に読み込みます。')
    add_para(doc, '方法Aでアップロードした場合は、3ファイルともルートに配置されるので問題ありません。')

    add_heading_styled(doc, 'インポートの関係', level=2)
    add_para(doc, 'bed_control_simulator_app.py は bed_control_simulator.py をインポートして使用しています。そのため、この2つのファイルは必ず同じディレクトリに配置する必要があります。')

    add_heading_styled(doc, '推奨されるファイル配置', level=2)
    add_para(doc, '最も簡単で確実な配置は以下のとおりです（方法Aの場合）：')

    add_code_block(doc, 'bed-control-simulator/          ← リポジトリのルート')
    add_code_block(doc, '├── README.md')
    add_code_block(doc, '├── requirements.txt            ← ここにあることが必須')
    add_code_block(doc, '├── bed_control_simulator.py')
    add_code_block(doc, '└── bed_control_simulator_app.py ← Main file path に指定')

    doc.add_page_break()

    # ===========================
    # 7. アプリの更新方法
    # ===========================
    add_heading_styled(doc, '7. アプリの更新方法', level=1)
    add_para(doc, 'GitHubリポジトリのファイルを更新すると、Streamlit Community Cloud が自動的に変更を検知し、再デプロイが行われます。通常1〜2分で反映されます。')

    add_heading_styled(doc, '更新手順（ブラウザで操作）', level=2)
    update_steps = [
        'GitHubにログインし、bed-control-simulator リポジトリを開きます。',
        '更新したいファイルをクリックして開きます。',
        '画面右上のペンのアイコン（Edit this file）をクリックします。',
        'ファイルの内容を編集します。',
        '画面右上の「Commit changes...」をクリックします。',
        '確認画面で「Commit changes」をクリックします。',
        '1〜2分後、Streamlit Cloudのアプリに自動的に反映されます。',
    ]
    for i, step in enumerate(update_steps, 1):
        add_numbered_item(doc, i, step)

    add_heading_styled(doc, 'ファイルの差し替え（新しいファイルで上書き）', level=2)
    replace_steps = [
        'リポジトリのページで「Add file」→「Upload files」をクリックします。',
        '新しいファイルをドラッグ＆ドロップします（同名ファイルは上書きされます）。',
        '「Commit changes」をクリックします。',
    ]
    for i, step in enumerate(replace_steps, 1):
        add_numbered_item(doc, i, step)

    doc.add_page_break()

    # ===========================
    # 8. トラブルシューティング
    # ===========================
    add_heading_styled(doc, '8. トラブルシューティング', level=1)
    add_para(doc, '問題が発生した場合は、以下の症状から対処法を確認してください。')

    troubles = [
        {
            'title': '症状1：デプロイ時にエラーが出る',
            'checks': [
                'requirements.txt がリポジトリのルート（最上位）にあるか確認してください。',
                'requirements.txt の中身が正しいか確認してください。以下の4行が含まれている必要があります：\npandas\nmatplotlib\nnumpy\nstreamlit',
                'Main file path が正しいか確認してください（方法Aの場合：bed_control_simulator_app.py）。',
            ]
        },
        {
            'title': '症状2：アプリは表示されるが、画面にエラーメッセージが出る',
            'checks': [
                'bed_control_simulator.py と bed_control_simulator_app.py が同じディレクトリにあるか確認してください。',
                'import文のパスが正しいか確認してください。',
                'Streamlit Cloudの「Manage app」→「Logs」でエラーの詳細を確認できます。',
            ]
        },
        {
            'title': '症状3：アプリにアクセスするとスリープ中と表示される',
            'checks': [
                'これは正常な動作です。Streamlit Community Cloud の無料プランでは、一定期間アクセスがないとアプリが自動的にスリープします。',
                'URLにアクセスすると自動的に復帰します。数秒〜1分程度お待ちください。',
                '頻繁にスリープが問題になる場合は、定期的にURLにアクセスすることで回避できます。',
            ]
        },
        {
            'title': '症状4：GitHubとの連携許可を求められる',
            'checks': [
                '初回のみ、GitHubとStreamlit Cloudの連携許可が必要です。',
                '「Authorize streamlit」をクリックして許可してください。',
                'これはGitHubのリポジトリにStreamlit Cloudがアクセスするための正常な手順です。',
            ]
        },
    ]

    for trouble in troubles:
        add_heading_styled(doc, trouble['title'], level=2)
        add_para(doc, '確認ポイント：', bold=True)
        for check in trouble['checks']:
            add_bullet(doc, check)
        doc.add_paragraph()

    doc.add_page_break()

    # ===========================
    # 9. セキュリティ
    # ===========================
    add_heading_styled(doc, '9. セキュリティに関する注意', level=1)

    security_items = [
        ('Publicリポジトリについて', 'GitHubのPublicリポジトリは、コードが誰でも閲覧可能です。ただし、本シミュレーターにはパラメータのみが含まれ、患者個人情報は一切含まれないため、Publicで問題ありません。'),
        ('URLの管理', 'デプロイされたアプリのURLを知っている人は誰でもアクセスできます。URLを共有する範囲を適切に管理してください。URLを知っている人だけに共有する運用で十分です。'),
        ('アクセス制限をかけたい場合', 'Streamlit Cloudの設定で「Viewer authentication」を有効にすると、GitHubアカウントでの認証が必要になります。Settings → Sharing から設定できます。'),
        ('コードを非公開にしたい場合', 'GitHubの有料プラン（月額$4〜）でPrivateリポジトリにし、Streamlit Cloudの有料プランと連携することで、コードを非公開にできます。'),
    ]

    for title, desc in security_items:
        add_heading_styled(doc, title, level=2)
        add_para(doc, desc)

    p = add_para(doc, '', space_after=10)
    run = p.add_run('【最重要】患者個人情報は絶対にGitHubやStreamlit Cloudにアップロードしないでください。')
    set_run_font(run, 'Yu Gothic', 11, bold=True, color=(0xC0, 0x00, 0x00))

    doc.add_page_break()

    # ===========================
    # 10. FAQ
    # ===========================
    add_heading_styled(doc, '10. よくある質問（FAQ）', level=1)

    faqs = [
        ('本当に無料ですか？', 'はい。Streamlit Community Cloudは完全無料です（2026年3月時点）。クレジットカードの登録も不要です。'),
        ('いつでもアクセスできますか？', 'はい。24時間アクセス可能です。ただし、長時間アクセスがないとアプリがスリープ状態になり、復帰に数秒〜1分ほどかかる場合があります。'),
        ('スマートフォンからもアクセスできますか？', 'はい。iPhone、Androidともに、ブラウザ（Safari、Chromeなど）でURLにアクセスすればそのまま利用できます。アプリのインストールは不要です。'),
        ('何人まで同時にアクセスできますか？', '明確な上限は公開されていませんが、少人数（10人程度）の同時アクセスであれば問題ありません。病院内での利用には十分です。'),
        ('アプリを削除したい場合はどうすればよいですか？', 'Streamlit Cloudのダッシュボード（https://share.streamlit.io）からアプリを選択し、「Delete app」で削除できます。必要であればGitHubリポジトリも削除してください。'),
        ('GitHubのパスワードを忘れた場合は？', 'GitHubのログイン画面（https://github.com/login）で「Forgot password?」をクリックし、登録メールアドレスを入力するとリセット用のメールが届きます。'),
        ('コードを非公開にしたい場合は？', 'GitHubの有料プラン（月額$4〜）でPrivateリポジトリにし、Streamlit Cloudの有料プランと連携することで実現可能です。通常の院内利用ではPublicで問題ありません。'),
        ('アプリの見た目を変更できますか？', 'はい。Pythonのコード（bed_control_simulator_app.py）を編集してGitHubにアップロードすれば、自動的に反映されます。'),
    ]

    for q, a in faqs:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(2)
        run_q = p.add_run(f'Q. {q}')
        set_run_font(run_q, 'Yu Gothic', 10.5, bold=True, color=(0x1F, 0x49, 0x7D))

        p2 = doc.add_paragraph()
        p2.paragraph_format.space_after = Pt(10)
        p2.paragraph_format.left_indent = Cm(0.5)
        run_a = p2.add_run(f'A. {a}')
        set_run_font(run_a, 'Yu Gothic', 10.5)

    doc.add_page_break()

    # ===========================
    # 付録
    # ===========================
    add_heading_styled(doc, '付録：クイックリファレンスカード', level=1)
    add_para(doc, 'この1ページに手順の要点をまとめています。印刷して手元に置いておくと便利です。', size=10)

    doc.add_paragraph()

    # クイックリファレンス テーブル
    ref_table = doc.add_table(rows=8, cols=2)
    ref_table.style = 'Table Grid'
    ref_table.alignment = WD_TABLE_ALIGNMENT.CENTER

    ref_data = [
        ('アプリ名', 'ベッドコントロールシミュレーター'),
        ('デプロイ済みURL', '【ここにURLを貼り付けてください】'),
        ('GitHubリポジトリ', 'https://github.com/ユーザー名/bed-control-simulator'),
        ('Streamlit Cloud\nダッシュボード', 'https://share.streamlit.io'),
        ('アプリの更新方法', 'GitHubでファイルを編集・差し替え → 1〜2分で自動反映'),
        ('スリープからの復帰', 'URLにアクセスするだけ（数秒〜1分待つ）'),
        ('問い合わせ先', '【担当者名・連絡先を記入してください】'),
        ('作成日', '2026年3月25日'),
    ]

    for i, (label, value) in enumerate(ref_data):
        cell_l = ref_table.rows[i].cells[0]
        cell_l.text = ''
        p = cell_l.paragraphs[0]
        run = p.add_run(label)
        set_run_font(run, 'Yu Gothic', 10, bold=True)
        set_cell_shading(cell_l, 'D9E2F3')

        cell_r = ref_table.rows[i].cells[1]
        cell_r.text = ''
        p = cell_r.paragraphs[0]
        run = p.add_run(value)
        if 'URL' in label or 'GitHub' in label or 'Streamlit' in label or 'ダッシュボード' in label:
            set_run_font(run, 'Courier New', 10)
        else:
            set_run_font(run, 'Yu Gothic', 10)

    doc.add_paragraph()
    doc.add_paragraph()

    # 手順サマリー
    add_heading_styled(doc, '手順サマリー（3ステップ）', level=2)

    summary = [
        'ステップ1：https://github.com でアカウント作成（5分）',
        'ステップ2：リポジトリ作成 → 3ファイルをアップロード（10分）',
        'ステップ3：https://share.streamlit.io でデプロイ（5分）',
    ]
    for i, item in enumerate(summary, 1):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(4)
        run = p.add_run(f'{i}. {item}')
        set_run_font(run, 'Yu Gothic', 11, bold=True, color=(0x2E, 0x74, 0xB5))

    doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(20)
    # border box
    pPr = p._element.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    for border_name in ['top', 'left', 'bottom', 'right']:
        border = OxmlElement(f'w:{border_name}')
        border.set(qn('w:val'), 'single')
        border.set(qn('w:sz'), '6')
        border.set(qn('w:space'), '8')
        border.set(qn('w:color'), '2E74B5')
        pBdr.append(border)
    pPr.append(pBdr)
    run = p.add_run('合計約20分でデプロイ完了！')
    set_run_font(run, 'Yu Gothic', 14, bold=True, color=(0x2E, 0x74, 0xB5))

    # --- ページ番号 ---
    add_page_number(doc)

    # --- 保存 ---
    output_path = '/Users/torukubota/ai-management/docs/admin/streamlit_cloud_deploy_manual.docx'
    os.makedirs(os.path.dirname(output_path), exist_ok=True)
    doc.save(output_path)
    print(f'文書を保存しました: {output_path}')


if __name__ == '__main__':
    main()
