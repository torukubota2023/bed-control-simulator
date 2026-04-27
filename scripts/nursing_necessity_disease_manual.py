"""看護必要度 疾患別マニュアル — 医師担当部分の判断・処置リファレンス.

副院長要望 (2026-04-26):
    看護必要度ミニレクチャー (概念整理) と並列に、疾患入口型の
    実臨床リファレンスを提供する。「この患者を見たらどう判断するか」を
    17 疾患について網羅。

色分け (レクチャーと統一):
    🟠 A 項目 (オレンジ #F59E0B) — 入院中の処置・薬剤
    🔵 B 項目 (青 #2563EB)       — 患者の状況等 (新基準では評価不要)
    🟢 C 項目 (緑 #10B981)       — 手術・侵襲的処置 (4-5 日カウント)

Single source of truth:
    DISEASE_DATA (構造化リスト) を起点に、
    - DISEASE_MANUAL_MARKDOWN (Streamlit 表示用)
    - generate_docx() (院内 LAN 配布用 DOCX、Word で開いて PDF 化可能)
    の両方を自動生成する。

CLI:
    .venv/bin/python scripts/nursing_necessity_disease_manual.py
        → docs/admin/nursing_necessity_disease_manual.md を再生成
"""

from __future__ import annotations

from io import BytesIO
from pathlib import Path
from typing import Any


# ---------------------------------------------------------------------------
# 色トークン (markdown 内で span style として使用)
# ---------------------------------------------------------------------------

COLOR_A = "#F59E0B"  # オレンジ
COLOR_B = "#2563EB"  # 青
COLOR_C = "#10B981"  # 緑

A_TAG = f'<span style="color:{COLOR_A};font-weight:bold">'
C_TAG = f'<span style="color:{COLOR_C};font-weight:bold">'
END_TAG = "</span>"


def _a(text: str) -> str:
    """A 項目用のオレンジ強調."""
    return f"{A_TAG}🟠 {text}{END_TAG}"


def _c(text: str) -> str:
    """C 項目用の緑強調."""
    return f"{C_TAG}🟢 {text}{END_TAG}"


# ---------------------------------------------------------------------------
# DISEASE_DATA: single source of truth
# ---------------------------------------------------------------------------

# 構造:
# [
#   ("疾患群名 (絵文字付き)", [
#       {
#           "id": int,
#           "name": "疾患名",
#           "icon": "絵文字",
#           "entry": [str, ...],            # 入口 (こんな患者を見たら)
#           "severity_assessment": [        # 重症度評価スコア (Phase 2A 以降に追加)
#               {
#                   "name": "スコア名 (出典)",
#                   "items": [{"label": "項目", "criteria": "閾値", "points": "1点"}, ...],
#                   "interpretation": [{"range": "0-1点", "meaning": "軽症 ..."}, ...],
#                   "source": "PMID xxxx / GL ..."
#               }, ...
#           ],
#           "options": [                    # 第一選択 vs 看護必要度貢献選択肢
#               {"scene": "...", "conventional": "...", "necessity": "..."}
#           ],
#           "contributions": [str, ...],    # 看護必要度寄与
#           "clinical_examples": [          # 具体例 (Phase 2A 以降に追加)
#               {
#                   "title": "例 A: 典型的入口 (XXスコア X点)",
#                   "background": "年齢・性別・既往",
#                   "presentation": "主訴・経過",
#                   "vitals": "バイタル・身体所見",
#                   "labs": "検査所見",
#                   "score_calc": "スコア計算 (各項目)",
#                   "plan": "治療方針",
#                   "necessity_contrib": "看護必要度寄与 (A項目X点 + C項目Y点)"
#               }, ...
#           ],
#           "evidence": [                   # エビデンス
#               {"strength": "強/中/弱", "ref": "PMID xxx or 日循 GL", "summary": "..."}
#           ],
#           "guardrails": [str, ...],       # 倫理ガードレール
#           "orders": [str, ...],           # オーダー例
#       },
#       ...
#   ]),
#   ...
# ]
#
# 注: severity_assessment と clinical_examples は Phase 2A 以降に充実化中。
# 既存疾患には未追加のものもある (空 list または欠損で互換維持)。

DISEASE_DATA: list[tuple[str, list[dict[str, Any]]]] = [
    # ========================================================================
    # 🍽️ 消化器 (6 疾患)
    # ========================================================================
    ("🍽️ 消化器系疾患", [
        {
            "id": 1,
            "name": "急性膵炎（重症）",
            "icon": "🍽️",
            "entry": [
                "上腹部激痛 + アミラーゼ・リパーゼ高値、CT 重症度スコア高",
                "BISAP ≥ 3 / Ranson 高 / 持続性 SIRS",
                "経口摂取困難で 3 日以上の絶食 + 持続点滴が必要",
            ],
            "severity_assessment": [
                {
                    "name": "BISAP（来院後 24 時間以内、5 点満点）",
                    "items": [
                        {"label": "B: BUN", "criteria": "BUN > 25 mg/dL", "points": "1点"},
                        {"label": "I: Impaired mental status", "criteria": "意識障害 (GCS < 15)", "points": "1点"},
                        {"label": "S: SIRS", "criteria": "SIRS ≥ 2 項目（体温・心拍・呼吸・WBC）", "points": "1点"},
                        {"label": "A: Age", "criteria": "年齢 > 60 歳", "points": "1点"},
                        {"label": "P: Pleural effusion", "criteria": "胸水あり（画像確認）", "points": "1点"},
                    ],
                    "interpretation": [
                        {"range": "0-2 点", "meaning": "軽症 — 死亡率 < 2%"},
                        {"range": "**≥ 3 点**", "meaning": "**重症 — 臓器不全 OR 7.4、持続的臓器不全 OR 12.7、膵壊死 OR 3.8、死亡率 5-22%**"},
                    ],
                    "source": "Singh VK et al. Am J Gastroenterol 2009 ([PMID 19293787](https://pubmed.ncbi.nlm.nih.gov/19293787/))",
                },
                {
                    "name": "改訂 Atlanta 分類 2012（重症度 3 段階）",
                    "items": [
                        {"label": "軽症 (Mild)", "criteria": "臓器不全なし + 局所合併症・全身合併症なし", "points": "—"},
                        {"label": "中等症 (Moderately severe)", "criteria": "一過性臓器不全（< 48h）/ 局所合併症 / 既存基礎疾患の増悪", "points": "—"},
                        {"label": "重症 (Severe)", "criteria": "**持続性臓器不全（> 48h）**", "points": "—"},
                    ],
                    "interpretation": [
                        {"range": "軽症", "meaning": "死亡率 < 1% — 通常入院加療"},
                        {"range": "中等症", "meaning": "死亡率 1-2% — 入院、慎重なモニタリング"},
                        {"range": "**重症**", "meaning": "**死亡率 30-50% — 個室管理 + 集中治療（当院対応範囲）/ 進行・合併症で高度医療機関へ転院**"},
                    ],
                    "source": "Banks PA et al. Gut 2013 改訂 Atlanta 分類（[PMID 23100216](https://pubmed.ncbi.nlm.nih.gov/23100216/)）",
                },
            ],
            "options": [
                {
                    "scene": "鎮痛",
                    "conventional": "アセトアミノフェン IV / NSAIDs (腎機能注意)",
                    "necessity": _a("A6③ モルヒネ持続点滴") + " — 入院当日開始安全",
                },
                {
                    "scene": "持続点滴のルート",
                    "conventional": "末梢繰り返し穿刺",
                    "necessity": _c("C23 CV 挿入") + " — A 項目ハブ・C 項目両取り",
                },
            ],
            "contributions": [
                _a("A6③ 麻薬注射") + ": **3 点 / 日 × 3-5 日**",
                _a("A4 シリンジポンプ") + ": 1 点（同時計上可）",
                _c("C23 CV 挿入実施日") + ": 4 日カウント (該当患者扱い)",
            ],
            "clinical_examples": [
                {
                    "title": "例 A: 胆石性膵炎の中等症（BISAP 2 点 / Atlanta 中等症）",
                    "background": "65 歳女性、胆石症既往、肥満（BMI 30）",
                    "presentation": "前夜の脂質摂取後、上腹部激痛 + 嘔吐、12 時間後に来院",
                    "vitals": "BT 37.8℃、BP 128/72、HR 102、呼吸数 22、SpO2 96% RA、上腹部圧痛 +、Murphy −",
                    "labs": "Lipase 1,200、Amylase 850、WBC 12,800、CRP 8.5、BUN 18、Cre 0.9、AST/ALT 96/118、T-Bil 2.1、CT で膵腫大・周囲液貯留軽度",
                    "score_calc": "BISAP = SIRS(1) + 65歳超(1) = **2 点（軽症圏）** / Atlanta = 一過性臓器不全 + 局所合併症 → **中等症**",
                    "plan": "入院、絶食、輸液 (ラクテック 200 mL/h)、フェンタニル 25 μg/h 持続、PPI、ERCP は黄疸進行で検討",
                    "necessity_contrib": _a("A6③ 麻薬注射") + " **3 点** / " + _a("A4 シリンジポンプ") + " 1 点 / " + _a("A3 注射薬剤 3 種類以上") + " 1 点（フェンタニル + ラクテック + PPI）→ 入院初日 **A 計 5 点で該当**",
                },
                {
                    "title": "例 B: アルコール性重症膵炎（BISAP 4 点 / Atlanta 重症）",
                    "background": "58 歳男性、アルコール性肝障害、慢性膵炎の既往、独居",
                    "presentation": "1 週間の飲酒後、上腹部激痛で発症、24 時間後に救急搬送、徐々に意識朦朧",
                    "vitals": "GCS 13、BT 38.9℃、BP 92/58、HR 124、呼吸数 26、SpO2 92% RA、Cullen 徴候 +",
                    "labs": "Lipase 4,500、WBC 22,000、CRP 28、BUN 36、Cre 1.8、Hct 48 (血液濃縮)、Lactate 4.2、CT で広範な膵壊死 + 両側胸水",
                    "score_calc": "BISAP = BUN(1) + Imp(1) + SIRS(1) + 60超なし + Pleural(1) = **4 点（重症）** / Atlanta = 持続性臓器不全 → **重症**",
                    "plan": "**個室入室、看護師監視強化下で集中管理（当院対応範囲）**、CV 挿入、強力輸液（最初 6 時間で 5-10 mL/kg/h）、モルヒネ 1-2 mg/h 持続、必要なら HFNC/NPPV、CT 再評価で感染性膵壊死出現なら抗菌薬 + IVR 介入、栄養は早期経腸栄養 (< 48h)。ショック移行・呼吸不全進行で挿管必要なら高度医療機関へ転院",
                    "necessity_contrib": _a("A6③ 麻薬注射") + " **3 点 × 5 日** / " + _a("A4 シリンジポンプ") + " 1 点 / " + _a("A3 注射薬剤 3 種類以上") + " 1 点 / " + _c("C23 CV 挿入") + " **4 日カウント該当患者扱い** / 感染合併で " + _a("A6⑦ 昇圧剤") + " **3 点** 追加可能性 → 圧倒的な該当患者",
                },
            ],
            "evidence": [
                {
                    "strength": "強",
                    "ref": "[PMID 19293787](https://pubmed.ncbi.nlm.nih.gov/19293787/)",
                    "summary": "Singh VK et al. Am J Gastroenterol 2009 (397 例) — BISAP ≥ 3 で死亡・臓器不全・膵壊死リスク有意上昇",
                },
                {
                    "strength": "強",
                    "ref": "[PMID 23100216](https://pubmed.ncbi.nlm.nih.gov/23100216/)",
                    "summary": "Banks PA et al. Gut 2013 — 改訂 Atlanta 分類、急性膵炎の標準的重症度分類",
                },
                {
                    "strength": "強",
                    "ref": "[PMID 38439202](https://pubmed.ncbi.nlm.nih.gov/38439202/)",
                    "summary": "Pandanaboyana et al. UEG Journal 2024 (118 施設、27 国、1768 例) — 入院当日のオピオイド開始は安全、6 日以上で重症化リスク↑",
                },
            ],
            "guardrails": [
                "重症度評価に基づく適応判断（軽症膵炎では不要）",
                "6 日以上の継続は重症化リスクあり、48 時間ごとに継続要否を再評価",
                "経口移行可能になり次第切替（疼痛コントロール十分後）",
            ],
            "orders": [
                "「モルヒネ持続点滴 3 日間予定、症状次第で延長」(A6③ 3 点/日)",
                "「アミラーゼ・リパーゼ・CRP 1 日 1 回モニター」",
                "CV 挿入オーダー（持続点滴・栄養管理目的）",
            ],
        },
        {
            "id": 2,
            "name": "急性胆管炎・閉塞性黄疸",
            "icon": "🍽️",
            "entry": [
                "Charcot 三徴 (発熱・黄疸・右上腹部痛)",
                "胆道酵素上昇 + 画像で胆道拡張・結石・腫瘍",
                "東京ガイドライン Grade II/III の中等症〜重症",
            ],
            "options": [
                {
                    "scene": "胆道ドレナージ",
                    "conventional": "保存的管理 + 抗菌薬のみで経過観察",
                    "necessity": _c("C21③ ERCP + 胆管ステント / EBD") + " — Grade II 以上で 24-48 時間以内に",
                },
                {
                    "scene": "経乳頭的アプローチ困難例",
                    "conventional": "手術的胆道ドレナージを検討",
                    "necessity": _c("C21③ PTCD（経皮経肝胆道ドレナージ）") + " — 術後は " + _a("A6⑩ ドレナージ管理"),
                },
            ],
            "contributions": [
                _c("C21③ 侵襲的消化器治療") + ": **4 日カウント** / 各処置",
                _a("A6⑩ ドレナージ管理") + ": 2 点 / 留置中毎日 (PTCD 後)",
                _a("A4 シリンジポンプ") + " + " + _a("A3 注射薬剤 3 種類以上") + ": 2 点 (β-ラクタム持続併用時)",
            ],
            "evidence": [
                {
                    "strength": "強",
                    "ref": "東京ガイドライン (TG18)",
                    "summary": "Grade II 以上の急性胆管炎は早期胆道ドレナージ推奨、Grade III は緊急 (24h 以内)",
                },
            ],
            "guardrails": [
                "Grade I (軽症) では保存的治療優先、不必要な ERCP は行わない",
                "出血傾向・抗血栓薬使用例では PTCD 優先",
                "ERCP は消化器内科コンサルト、適応の医学的妥当性を確認",
            ],
            "orders": [
                "「ERCP + 胆管ステント留置依頼 (Grade II 急性胆管炎)」",
                "「PIPC/TAZ 4.5g IV q8h or 持続点滴で開始」",
                "「PTCD 留置後は排液量・性状を 6h ごと記録」",
            ],
        },
        {
            "id": 3,
            "name": "消化管出血",
            "icon": "🍽️",
            "entry": [
                "吐血・下血・タール便、Hb 低下 (≥ 2g/dL)",
                "ショック傾向 / 起立性低血圧 / 頻脈",
                "PPI 内服歴・NSAIDs・抗血栓薬使用歴",
            ],
            "options": [
                {
                    "scene": "出血源同定と止血",
                    "conventional": "保存的 (PPI IV) + 様子観察",
                    "necessity": _c("C21③ 内視鏡的止血術") + " — 上部・下部消化管出血で 24h 以内に",
                },
                {
                    "scene": "Hb 低下例",
                    "conventional": "経過観察 + 鉄剤",
                    "necessity": _a("A5 輸血や血液製剤の管理") + " — RBC 輸血実施で 2 点",
                },
            ],
            "contributions": [
                _c("C21③ 内視鏡的止血") + ": **4 日カウント**",
                _a("A5 輸血管理") + ": 2 点 / 輸血実施日 (該当患者扱い)",
                _a("A6⑨ 抗血栓持続") + " (UFH への切替時): 3 点 / 日",
            ],
            "evidence": [
                {
                    "strength": "強",
                    "ref": "ACG ガイドライン 2021 (Laine et al.)",
                    "summary": "高リスク出血潰瘍は早期内視鏡 (24h 以内) + 内視鏡的止血で再出血率減少",
                },
            ],
            "guardrails": [
                "輸血は適応基準 (Hb < 7 g/dL or 症状あり) に基づく、過剰輸血を避ける",
                "抗血栓薬は出血リスクと血栓リスクを天秤にかけ、安易な再開は避ける",
                "内視鏡は消化器内科判断、緊急度に応じて時間設定",
            ],
            "orders": [
                "「緊急上部消化管内視鏡 (吐血、Hb 8.5)」",
                "「RBC 2 単位輸血、輸血前後 Hb 測定」",
                "「PPI (オメプラゾール) 80mg ボーラス → 8mg/h 持続」",
            ],
        },
        {
            "id": 4,
            "name": "早期消化管腫瘍（食道・胃・大腸）",
            "icon": "🍽️",
            "entry": [
                "内視鏡で発見された早期癌 (T1a/cT1b)",
                "脈管侵襲なし、リンパ節転移リスク低",
                "ESD/EMR 適応 (絶対適応 / 拡大適応)",
            ],
            "options": [
                {
                    "scene": "早期癌の治療選択",
                    "conventional": "手術 (外科紹介) を即決",
                    "necessity": _c("C21③ ESD / EMR") + " — 適応症例で内視鏡的切除を積極化",
                },
            ],
            "contributions": [
                _c("C21③ 侵襲的消化器治療") + ": **4 日カウント**",
                _a("A1 創傷処置") + ": 1 点 / 日 (穿孔・出血合併時のドレナージ管理)",
            ],
            "evidence": [
                {
                    "strength": "強",
                    "ref": "胃癌治療ガイドライン (日本胃癌学会)",
                    "summary": "ESD 絶対適応病変は 5 年生存率が外科手術と同等、QOL 維持に優位",
                },
            ],
            "guardrails": [
                "ESD/EMR の適応病変か慎重に判断（消化器内科・病理コンサルト）",
                "穿孔・後出血リスクの説明と同意",
                "拡大適応病変は症例検討で慎重に",
            ],
            "orders": [
                "「ESD 予定 (早期胃癌、絶対適応)」",
                "「術後 24 時間絶食、ドレーン留置・排液モニター」",
                "「術後出血モニター (Hb 6h ごと、12 時間)」",
            ],
        },
        {
            "id": 5,
            "name": "嚥下障害（経口摂取困難）",
            "icon": "🍽️",
            "entry": [
                "脳卒中後・神経変性疾患・頭頸部癌術後の嚥下機能低下",
                "VF / VE で誤嚥確認、経口摂取で誤嚥性肺炎反復",
                "経鼻胃管 (NG) 長期化 (4 週間以上) で QOL 低下",
            ],
            "options": [
                {
                    "scene": "長期栄養経路",
                    "conventional": "NG 継続 / TPN（CV 経由）",
                    "necessity": _c("C23 PEG 造設") + " — 長期見込みで 4 週後を目処に",
                },
            ],
            "contributions": [
                _c("C23 別に定める手術・処置 (PEG)") + ": **5 日カウント** (該当患者扱い)",
                _a("A1 創傷処置") + ": 1 点 / 日 (PEG 創部管理)",
            ],
            "evidence": [
                {
                    "strength": "中",
                    "ref": "Lancet 2005 FOOD trial",
                    "summary": "脳卒中後の早期 PEG vs NG 継続は転帰差なし、しかし長期 (>4 週) は PEG が QOL 優位",
                },
            ],
            "guardrails": [
                "予後 4 週以上を見込む症例で適応 (短期見込みは NG で対応)",
                "本人・家族との十分な説明と同意 (terminal care 例では適応外)",
                "PEG 後の経口訓練継続を放棄しない",
            ],
            "orders": [
                "「PEG 造設依頼 (脳梗塞後嚥下障害、4 週経過)」",
                "「PEG 後 24 時間絶食、創部消毒 1 日 1 回」",
                "「PEG 経管栄養 開始 (200 kcal × 3 → 漸増)」",
            ],
        },
        {
            "id": 6,
            "name": "複雑性腹腔内感染 (cIAI)",
            "icon": "🍽️",
            "entry": [
                "穿孔性虫垂炎・憩室炎穿孔・術後腹腔内膿瘍",
                "発熱 + 腹膜刺激症状 + CT で膿瘍・遊離ガス",
                "重症度評価 (qSOFA、APACHE II)",
            ],
            "options": [
                {
                    "scene": "抗菌薬投与",
                    "conventional": "セフメタゾール / セフトリアキソン IV (間欠)",
                    "necessity": _a("A4 + A3 ペア") + " — PIPC/TAZ・MEPM 持続点滴 + 多剤併用",
                },
                {
                    "scene": "膿瘍ドレナージ",
                    "conventional": "保存的に抗菌薬のみで経過観察",
                    "necessity": _a("A6⑩ ドレナージ (PCD)") + " — IR ガイド下に経皮的に",
                },
            ],
            "contributions": [
                _a("A4 シリンジポンプ") + " + " + _a("A3 注射薬剤 3 種類以上") + ": **2 点** (該当患者)",
                _a("A6⑩ ドレナージ管理") + ": 2 点 / 留置中毎日",
            ],
            "evidence": [
                {
                    "strength": "強",
                    "ref": "IDSA cIAI ガイドライン 2010",
                    "summary": "5cm 以上の膿瘍はドレナージ必須、抗菌薬は適切なスペクトラムで 4-7 日",
                },
                {
                    "strength": "中",
                    "ref": "[PMID 38864155](https://pubmed.ncbi.nlm.nih.gov/38864155/)",
                    "summary": "BLING-III RCT — β-ラクタム持続点滴は 90 日死亡で有意差なしだがメタ解析では有意 (重症で合理的)",
                },
            ],
            "guardrails": [
                "ドレナージ可能な膿瘍は積極的に IR コンサルト",
                "抗菌薬持続点滴は重症例 (qSOFA ≥ 2) に限定、軽症は間欠投与",
                "培養結果で 48-72h 後に de-escalation",
            ],
            "orders": [
                "「PIPC/TAZ 4.5g 負荷 → 13.5g/24h 持続点滴」",
                "「腹腔内膿瘍 PCD 留置依頼 (IR)」",
                "「血液培養 2 セット + 膿瘍液培養」",
            ],
        },
    ]),
    # ========================================================================
    # 🫁 呼吸器 (3 疾患)
    # ========================================================================
    ("🫁 呼吸器系疾患", [
        {
            "id": 7,
            "name": "重症肺炎（CURB-65 高・ショックなし）",
            "icon": "🫁",
            "entry": [
                "CURB-65 ≥ 3、A-DROP ≥ 3、PSI Class IV/V",
                "SpO2 < 90% (room air)、呼吸数 ≥ 30",
                "意識障害なし、血圧維持できているがハイリスク",
            ],
            "severity_assessment": [
                {
                    "name": "CURB-65（英国胸部学会、5 点満点）",
                    "items": [
                        {"label": "C: Confusion", "criteria": "意識障害（AMTS ≤ 8 / 見当識障害）", "points": "1点"},
                        {"label": "U: Urea", "criteria": "BUN > 20 mg/dL（脱水・腎機能低下）", "points": "1点"},
                        {"label": "R: Respiratory rate", "criteria": "呼吸数 ≥ 30/min（頻呼吸）", "points": "1点"},
                        {"label": "B: Blood pressure", "criteria": "収縮期 < 90 mmHg または 拡張期 ≤ 60 mmHg", "points": "1点"},
                        {"label": "65: Age", "criteria": "年齢 ≥ 65 歳", "points": "1点"},
                    ],
                    "interpretation": [
                        {"range": "0-1 点", "meaning": "軽症（30 日死亡率 0.7-3%）— 外来治療可"},
                        {"range": "2 点", "meaning": "中等症（死亡率 9%）— **入院検討**"},
                        {"range": "3 点", "meaning": "**重症（死亡率 17%）— 入院、個室管理 + 看護師監視強化（NPPV/HFNC 含む呼吸ケア）、抗菌薬 IV**"},
                        {"range": "4-5 点", "meaning": "**超重症（死亡率 41-57%）— ショック・挿管必要なら高度医療機関へ転院**"},
                    ],
                    "source": "Lim WS et al. Thorax 2003 ([PMID 12728155](https://pubmed.ncbi.nlm.nih.gov/12728155/))",
                },
                {
                    "name": "A-DROP（日本呼吸器学会、5 点満点）",
                    "items": [
                        {"label": "A: Age", "criteria": "男性 ≥ 70 歳 / 女性 ≥ 75 歳", "points": "1点"},
                        {"label": "D: Dehydration", "criteria": "BUN ≥ 21 mg/dL", "points": "1点"},
                        {"label": "R: Respiratory failure", "criteria": "SpO2 ≤ 90% または PaO2 ≤ 60 mmHg", "points": "1点"},
                        {"label": "O: Orientation disorder", "criteria": "意識障害", "points": "1点"},
                        {"label": "P: Pressure", "criteria": "収縮期血圧 ≤ 90 mmHg", "points": "1点"},
                    ],
                    "interpretation": [
                        {"range": "0 点", "meaning": "軽症 — 外来治療"},
                        {"range": "1-2 点", "meaning": "中等症 — 入院検討"},
                        {"range": "3 点", "meaning": "**重症 — 入院、個室管理 + NPPV/HFNC 含む呼吸ケア**"},
                        {"range": "4-5 点", "meaning": "**超重症 — ショック・挿管必要なら高度医療機関へ転院**"},
                    ],
                    "source": "日本呼吸器学会 成人肺炎診療ガイドライン 2024",
                },
            ],
            "options": [
                {
                    "scene": "抗菌薬選択",
                    "conventional": "セフトリアキソン + アジスロマイシン (間欠)",
                    "necessity": _a("A4 + A3 ペア") + " — PIPC/TAZ 持続点滴 + シリンジポンプ",
                },
                {
                    "scene": "酸素管理",
                    "conventional": "鼻カニュラで様子観察",
                    "necessity": _a("A2 呼吸ケア (酸素 A2)") + " — 高流量鼻カニュラで確実に確保",
                },
            ],
            "contributions": [
                _a("A4 + A3") + ": **2 点** (該当患者)",
                _a("A2 呼吸ケア (酸素 A2)") + ": 1 点 / 日 (酸素必要日)",
            ],
            "clinical_examples": [
                {
                    "title": "例 A: 高齢の脱水＋意識障害（CURB-65 3 点 / A-DROP 4 点）",
                    "background": "78 歳女性、施設入所中、糖尿病・高血圧",
                    "presentation": "3 日前から食事摂取低下と発熱、来院前から意識朦朧",
                    "vitals": "GCS 13、BT 38.7℃、呼吸数 28/min、BP 102/65、HR 110、SpO2 89% RA",
                    "labs": "WBC 14,500、CRP 18.2、BUN 38 mg/dL、Cre 1.6、Na 134、CXR 右下葉浸潤影",
                    "score_calc": "CURB-65 = C(1) + U(1) + 65(1) = **3 点（重症）** / A-DROP = A(1)+D(1)+R(1)+O(1) = **4 点（超重症）**",
                    "plan": "緊急入院、6F 病棟、酸素 3L 鼻カニュラで SpO2 ≥ 94% 維持、PIPC/TAZ 4.5g 負荷 → 13.5g/24h 持続、輸液 + 電解質補正、血液培養 2 セット",
                    "necessity_contrib": _a("A2 呼吸ケア (酸素)") + " 1 点 / " + _a("A4 シリンジポンプ") + " 1 点 / " + _a("A3 注射薬剤 3 種類以上") + " 1 点 / " + _a("A7 救急搬送後") + " 2 点（入院日と翌日） → 初日 **A 計 5 点で該当**",
                },
                {
                    "title": "例 B: 高齢のショック移行例（CURB-65 5 点）",
                    "background": "82 歳男性、誤嚥性肺炎の既往、嚥下機能低下、寝たきり",
                    "presentation": "発熱と呼吸困難で家族搬送、来院時すでに低血圧",
                    "vitals": "意識混濁（GCS 11）、BT 38.5℃、呼吸数 32/min、BP 80/50、HR 124、SpO2 86% RA",
                    "labs": "WBC 18,200、CRP 22、BUN 45、Cre 1.9、Lactate 3.2、CXR 両側浸潤影",
                    "score_calc": "CURB-65 = C(1) + U(1) + R(1) + B(1) + 65(1) = **5 点（超重症、院内死亡率 ≥ 22%）**",
                    "plan": "**当院対応範囲超過 → 高度医療機関への転院推奨**。転院前安定化として: CV 挿入、ノルアドレナリン 0.05 μg/kg/min 開始、MEPM 1g IV → 3g/24h 持続、HFNC または NPPV 装着、MRSA カバー追加検討。安定後に転院搬送",
                    "necessity_contrib": "（転院前安定化中の当院での寄与）" + _a("A6⑦ 昇圧剤注射") + " **3 点** / " + _a("A2 呼吸ケア (HFNC/NPPV)") + " 1 点 / " + _a("A4 シリンジポンプ") + " 1 点 / " + _a("A3 注射薬剤 3 種類以上") + " 1 点 / " + _c("C23 CV 挿入") + " 4 日カウント / " + _a("A7 救急搬送後") + " 2 点 → **A 計 8 点 + C 該当 = 圧倒的な該当患者**",
                },
            ],
            "evidence": [
                {
                    "strength": "強",
                    "ref": "ATS/IDSA CAP ガイドライン 2019",
                    "summary": "重症 CAP は β-ラクタム + マクロライド or キノロン併用、高度医療機関への転院適応評価必須",
                },
                {
                    "strength": "強",
                    "ref": "[PMID 12728155](https://pubmed.ncbi.nlm.nih.gov/12728155/)",
                    "summary": "Lim WS et al. Thorax 2003 — CURB-65 オリジナル論文。1,068 例の国際多施設研究で 5 段階の死亡率階層化を確立",
                },
                {
                    "strength": "中",
                    "ref": "[PMID 38864155](https://pubmed.ncbi.nlm.nih.gov/38864155/)",
                    "summary": "BLING-III: β-ラクタム持続点滴は重症感染症で合理的、メタ解析では NNT 26",
                },
            ],
            "guardrails": [
                "持続点滴は CURB-65 ≥ 3 の重症例に限定",
                "経口移行・de-escalation は培養結果と臨床経過で判断",
                "高度医療機関への転院適応の見極め（呼吸不全進行で挿管が必要、ショック移行で集中治療が必要）を逃さない",
            ],
            "orders": [
                "「PIPC/TAZ 4.5g 負荷 → 13.5g/24h 持続」",
                "「アジスロマイシン 500mg IV q24h」",
                "「血液培養 2 セット + 喀痰培養 + 尿中肺炎球菌・レジオネラ抗原」",
            ],
        },
        {
            "id": 8,
            "name": "気胸・膿胸・血胸",
            "icon": "🫁",
            "entry": [
                "突然の胸痛・呼吸困難、画像で気胸・胸水・血腫確認",
                "緊張性気胸 (緊急)、大量胸水 (>1L)",
                "外傷後の血胸、膿胸 (CT で液貯留 + 胸壁肥厚)",
            ],
            "options": [
                {
                    "scene": "胸腔ドレナージ",
                    "conventional": "保存的に経過観察 (小さい気胸)",
                    "necessity": _a("A6⑩ 胸腔ドレナージ") + " — 中等度以上で挿入閾値を下げる",
                },
            ],
            "contributions": [
                _a("A6⑩ ドレナージの管理") + ": **2 点 / 留置中毎日** (該当患者)",
                _a("A1 創傷処置") + ": 1 点 / 日 (ドレーン創部管理)",
            ],
            "evidence": [
                {
                    "strength": "強",
                    "ref": "BTS ガイドライン 2010",
                    "summary": "症候性気胸 (≥ 2cm) は胸腔ドレナージ推奨、膿胸は早期外科介入も検討",
                },
            ],
            "guardrails": [
                "小気胸 (< 2cm、無症状) は経過観察が原則",
                "ドレーン抜去は持続的な脱気・排液停止 + 画像改善で判断",
                "膿胸は線維素性で胸膜癒着あり、外科コンサルトを早期に",
            ],
            "orders": [
                "「胸腔ドレーン挿入 (24Fr、第 5 肋間中腋窩線)」",
                "「ドレナージ排液量・性状 6h ごと記録」",
                "「胸部 X 線 翌朝 + ドレーン抜去前」",
            ],
        },
        {
            "id": 9,
            "name": "高リスク PE / CrCl<30 DVT",
            "icon": "🫁",
            "entry": [
                "造影 CT で肺動脈本幹〜葉動脈の血栓 (高リスク PE)",
                "右心負荷所見 (RV/LV ≥ 1.0、TAPSE 低下、トロポニン上昇)",
                "腎不全 (CrCl < 30) で DOAC・LMWH が使用しにくい",
            ],
            "options": [
                {
                    "scene": "抗凝固",
                    "conventional": "DOAC (アピキサバン・リバーロキサバン)",
                    "necessity": _a("A6⑨ UFH 持続点滴") + " — 高リスク PE / CrCl<30 で第一選択",
                },
            ],
            "contributions": [
                _a("A6⑨ 抗血栓塞栓薬の持続点滴") + ": **3 点 / 日 × 3-7 日**",
                _a("A4 シリンジポンプ") + ": 1 点 (UFH 持続併用)",
            ],
            "evidence": [
                {
                    "strength": "強",
                    "ref": "[日循 2025 GL](docs/admin/references/pdf/jcs2025_pte_dvt_guideline.pdf)",
                    "summary": "DOAC 第一選択だが、高リスク PE / CrCl<30 では UFH 持続が推奨",
                },
            ],
            "guardrails": [
                "出血リスク (HAS-BLED スコア等) を必ず評価",
                "APTT モニタリング (1.5-2.5 倍) で用量調整",
                "臨床的安定後 (24-48h) に DOAC 切替を検討",
            ],
            "orders": [
                "「UFH 80 単位/kg 負荷 → 18 単位/kg/h 持続」",
                "「APTT 6 時間ごと、目標 1.5-2.5 倍」",
                "「血小板数 3 日後 (HIT スクリーニング)」",
            ],
        },
    ]),
    # ========================================================================
    # ❤️ 循環器 (2 疾患)
    # ========================================================================
    ("❤️ 循環器系疾患", [
        {
            "id": 10,
            "name": "重症 Af / 心房粗動（血行動態不安定）",
            "icon": "❤️",
            "entry": [
                "心房細動・粗動 + 頻脈 (HR > 130) + 血圧低下",
                "胸痛・心不全症状の合併",
                "経口β遮断薬・Ca 拮抗薬で コントロール不十分",
            ],
            "options": [
                {
                    "scene": "心拍コントロール",
                    "conventional": "経口メトプロロール・ビソプロロール",
                    "necessity": _a("A6⑧ IV アミオダロン") + " — 血行動態不安定で第一選択",
                },
            ],
            "contributions": [
                _a("A6⑧ 抗不整脈剤の使用") + ": **3 点 / 日 × 1-3 日**",
                _a("A4 シリンジポンプ") + ": 1 点 (アミオダロン持続併用)",
            ],
            "evidence": [
                {
                    "strength": "強",
                    "ref": "[PMID 39210723](https://pubmed.ncbi.nlm.nih.gov/39210723/)",
                    "summary": "2024 ESC AF GL — 血行動態不安定例で IV アミオダロン推奨",
                },
            ],
            "guardrails": [
                "甲状腺機能・肝機能・QT 評価必須 (アミオダロン使用前)",
                "電気的除細動の適応 (重度血行動態不安定) も考慮",
                "心拍コントロール後は経口移行を 24-48h で判断",
            ],
            "orders": [
                "「アミオダロン 150mg ローディング → 1mg/min × 6h → 0.5mg/min」",
                "「心電図モニター継続、QT 8 時間ごと」",
                "「TSH・FT4・AST/ALT 投与前」",
            ],
        },
        {
            "id": 11,
            "name": "心不全急性増悪（hot/wet 型: 高血圧 + うっ血）",
            "icon": "❤️",
            "entry": [
                "急性肺水腫 + 起座呼吸 + SpO2 低下",
                "**SBP > 140 mmHg + 末梢温暖**（hot/wet 型、急性肺水腫の典型）",
                "BNP 高値、CXR で肺うっ血",
                "※ cold/wet (cardiogenic shock 移行例) は SBP < 90 + 末梢冷感、A6⑦ 昇圧剤主体に切替",
            ],
            "severity_assessment": [
                {
                    "name": "Forrester 分類（うっ血と末梢循環の組み合わせ、4 群）",
                    "items": [
                        {"label": "I 型 (dry/warm)", "criteria": "うっ血なし + 末梢温暖（PCWP < 18, CI > 2.2）", "points": "—"},
                        {"label": "II 型 (wet/warm) — **hot/wet**", "criteria": "**うっ血あり + 末梢温暖**（PCWP > 18, CI > 2.2）", "points": "—"},
                        {"label": "III 型 (dry/cold)", "criteria": "うっ血なし + 末梢冷感（PCWP < 18, CI < 2.2）", "points": "—"},
                        {"label": "IV 型 (wet/cold) — **cold/wet**", "criteria": "**うっ血あり + 末梢冷感（cardiogenic shock）**（PCWP > 18, CI < 2.2）", "points": "—"},
                    ],
                    "interpretation": [
                        {"range": "I 型", "meaning": "経過観察、必要なら経口治療強化"},
                        {"range": "**II 型（hot/wet）**", "meaning": "**血管拡張薬 (NTG IV) + 利尿剤 — 本疾患の典型**"},
                        {"range": "III 型", "meaning": "輸液負荷検討、原因精査"},
                        {"range": "**IV 型（cold/wet）**", "meaning": "**昇圧剤 (NE/DOB) + 血管拡張薬慎重に — IABP / 集中治療**"},
                    ],
                    "source": "Forrester JS et al. NEJM 1976（古典的分類）/ ESC HF GL 2021 で再確認",
                },
                {
                    "name": "Killip 分類（急性心筋梗塞時の心不全重症度、4 段階）",
                    "items": [
                        {"label": "Class I", "criteria": "心不全所見なし", "points": "—"},
                        {"label": "Class II", "criteria": "ラ音 < 1/2 肺野、S3、頸静脈怒張", "points": "—"},
                        {"label": "Class III", "criteria": "**重症心不全（肺水腫、ラ音 > 1/2 肺野）**", "points": "—"},
                        {"label": "Class IV", "criteria": "**心原性ショック（SBP < 90、末梢冷感）**", "points": "—"},
                    ],
                    "interpretation": [
                        {"range": "Class I", "meaning": "院内死亡率 < 5%"},
                        {"range": "Class II", "meaning": "院内死亡率 10-20%"},
                        {"range": "**Class III**", "meaning": "**院内死亡率 30-40% — 入院、集中治療**"},
                        {"range": "**Class IV**", "meaning": "**院内死亡率 60-80% — 当院対応範囲超過、IABP/Impella 適応で高次医療機関へ転送**"},
                    ],
                    "source": "Killip T 3rd, Kimball JT. Am J Cardiol 1967 / Mello et al. Arq Bras Cardiol 2014 ([PMID 25014060](https://pubmed.ncbi.nlm.nih.gov/25014060/)) で 1906 例で再検証",
                },
            ],
            "options": [
                {
                    "scene": "前負荷・後負荷軽減（hot/wet の主軸）",
                    "conventional": "利尿剤 (フロセミド IV ボーラス) のみ",
                    "necessity": _a("A4 + A3 ペア") + " — ニトログリセリン IV 持続点滴 (シリンジポンプ管理) + 利尿剤 + 補液",
                },
                {
                    "scene": "cold/wet 移行例の血行動態維持",
                    "conventional": "経過観察、利尿剤継続",
                    "necessity": _a("A6⑦ NE/DOB 持続点滴") + " — SBP < 90 で必要十分な期間維持",
                },
            ],
            "contributions": [
                _a("A4 シリンジポンプ管理") + ": **1 点 / 日** (NTG 持続)",
                _a("A3 注射薬剤 3 種類以上") + ": **1 点 / 日** (利尿剤 + NTG + 補液)",
                "→ " + _a("A4 + A3 = 2 点") + " で **該当患者** (NE なしの基本パターン)",
                _a("A6⑦ 昇圧剤") + ": 3 点 / 日 (cold/wet 移行例で追加)",
                _a("A2 呼吸ケア (酸素 A2)") + ": 1 点 / 日 (酸素必要日)",
            ],
            "clinical_examples": [
                {
                    "title": "例 A: 高血圧緊急症型 hot/wet (Forrester II / Killip III)",
                    "background": "75 歳男性、高血圧・慢性心不全（HFpEF）、最近内服中断",
                    "presentation": "夜間突然の呼吸困難、起座呼吸で救急搬送",
                    "vitals": "BP 198/112、HR 118、呼吸数 32、SpO2 84% RA → 高流量酸素 (15L マスク) で 92%、両肺野でラ音、頸静脈怒張、末梢温暖",
                    "labs": "BNP 1,850、トロポニン陰性、Cre 1.4、Na 138、CXR で両側肺野バタフライ陰影、CTR 60%",
                    "score_calc": "Forrester = うっ血(+) + 末梢温暖 → **II 型 (hot/wet)** / Killip = 肺水腫 → **Class III**",
                    "plan": "緊急入院、6F 病棟 (or HCU)、酸素 8-10L マスク、ニトログリセリン 5 μg/min IV → 増量 max 200 μg/min（SBP 100 維持）、フロセミド 40mg IV、A line、24-72h で経口 ACE-I/ARNI 移行",
                    "necessity_contrib": _a("A4 シリンジポンプ管理 (NTG 持続)") + " 1 点 / " + _a("A3 注射薬剤 3 種類以上 (NTG + フロセミド + 補液)") + " 1 点 / " + _a("A2 呼吸ケア (酸素)") + " 1 点 / " + _a("A7 救急搬送後") + " 2 点（入院日と翌日）→ 入院日 **A 計 5 点で該当**、NTG 継続中は連日 A4+A3=2 点で該当維持",
                },
                {
                    "title": "例 B: cold/wet 移行 (Forrester IV / Killip IV — 心原性ショック)",
                    "background": "82 歳女性、虚血性心疾患、CKD ステージ 4",
                    "presentation": "前日からの呼吸困難 + 嘔吐、来院時すでに低血圧と末梢冷感",
                    "vitals": "意識朦朧 (GCS 13)、BP 78/45、HR 132、呼吸数 36、SpO2 82% RA、末梢冷感・チアノーゼ、両肺野で湿性ラ音",
                    "labs": "BNP 3,200、トロポニン上昇、Lactate 4.5、Cre 2.8、ABG: pH 7.28 / PaO2 58 / PaCO2 42 / HCO3 18",
                    "score_calc": "Forrester = うっ血(+) + 末梢冷感 → **IV 型 (cold/wet)** / Killip = 心原性ショック → **Class IV（院内死亡率 60-80%）**",
                    "plan": "**個室入室、看護師監視強化下で集中管理**、CV 挿入、ノルアドレナリン 0.05 μg/kg/min 開始、ドブタミン併用検討、HFNC または NPPV 装着、フロセミド少量 IV (尿量見ながら)、A line、Foley。IABP/Impella 適応または挿管必要時は高次医療機関（CCU）へ転送",
                    "necessity_contrib": _a("A6⑦ 昇圧剤注射 (NE 持続)") + " **3 点** / " + _a("A4 シリンジポンプ") + " 1 点 / " + _a("A3 注射薬剤 3 種類以上") + " 1 点 / " + _a("A2 呼吸ケア (酸素・NPPV)") + " 1 点 / " + _c("C23 CV 挿入") + " 4 日カウント / " + _a("A7 救急搬送後") + " 2 点 → **A 計 8 点 + C 該当 = 圧倒的な該当患者**",
                },
            ],
            "evidence": [
                {
                    "strength": "強",
                    "ref": "[PMID 35379503](https://pubmed.ncbi.nlm.nih.gov/35379503/)",
                    "summary": "2022 AHA/ACC/HFSA HF GL — 急性肺水腫 + 高血圧で IV vasodilator (NTG/SNP) 推奨",
                },
                {
                    "strength": "強",
                    "ref": "[PMID 25014060](https://pubmed.ncbi.nlm.nih.gov/25014060/)",
                    "summary": "Mello et al. Arq Bras Cardiol 2014 (1906 例 5 年追跡) — Killip 分類は 60 年経過後も有意な独立予後因子",
                },
                {
                    "strength": "強",
                    "ref": "[PMID 35636830](https://pubmed.ncbi.nlm.nih.gov/35636830/)",
                    "summary": "2021 ESC HF GL — SBP > 110 mmHg の急性心不全で IV vasodilator (Class IIa)",
                },
                {
                    "strength": "中",
                    "ref": "[PMID 40834833](https://pubmed.ncbi.nlm.nih.gov/40834833/)",
                    "summary": "Henry AJEM 2025 (SCAPE 441 例) — IV NTG 高用量 (≥100 μg/min) は酸素 weaning 早い (2.7h vs 3.3h)、低血圧少ない",
                },
                {
                    "strength": "中",
                    "ref": "[PMID 40833345](https://pubmed.ncbi.nlm.nih.gov/40833345/)",
                    "summary": "Galluzzo JCP 2025 (国際 survey、170 医師、32 国) — AHF での IVV 第一選択は NTG 48%、ニトロプルシド 29%",
                },
            ],
            "guardrails": [
                "**SBP < 90 mmHg は禁忌** (NTG 減量・中止)、最低 SBP > 100 mmHg を維持",
                "24-72h で経口血管拡張薬 (ACE-I/ARNI) への移行を検討、持続を惰性で続けない",
                "cold/wet 進行 (低血圧 + 末梢冷感) を見逃さない → A6⑦ NE/DOB 追加判断",
                "5 日以上の持続は通常不要 (A3 注射薬剤の最大 7 日カウントを意識)",
                "PDE5 阻害薬使用例では NTG は禁忌 (重度低血圧リスク)",
            ],
            "orders": [
                "「ニトログリセリン 5 μg/min 開始 → 5 μg/min ずつ 3-5 分ごと増量 (max 200 μg/min)」",
                "「フロセミド 20-40mg IV ボーラス + 適宜追加 (尿量モニター)」",
                "「シリンジポンプで NTG 持続、SBP > 100 mmHg 維持」",
                "「A line + NIBP モニター継続、尿量 1 時間ごと記録」",
                "「cold/wet 移行時は NE 0.05 μg/kg/min 追加開始」",
            ],
        },
    ]),
    # ========================================================================
    # 🦠 感染症 (3 疾患)
    # ========================================================================
    ("🦠 感染症（昇圧剤適応・非適応）", [
        {
            "id": 12,
            "name": "敗血症・敗血症性ショック",
            "icon": "🦠",
            "entry": [
                "感染巣同定 + qSOFA ≥ 2 / SOFA score 上昇",
                "MAP < 65 (輸液後も)、lactate > 2 mmol/L",
                "意識障害・乏尿・末梢冷感",
            ],
            "severity_assessment": [
                {
                    "name": "qSOFA（quick SOFA、ベッドサイド 3 項目）",
                    "items": [
                        {"label": "呼吸数", "criteria": "≥ 22/min", "points": "1点"},
                        {"label": "意識変容", "criteria": "GCS < 15（意識レベル低下）", "points": "1点"},
                        {"label": "収縮期血圧", "criteria": "≤ 100 mmHg", "points": "1点"},
                    ],
                    "interpretation": [
                        {"range": "0-1 点", "meaning": "敗血症リスク低、ただし疑いは継続"},
                        {"range": "**≥ 2 点**", "meaning": "**敗血症の可能性高 — SOFA で正式評価、輸液 + 抗菌薬 + 培養を 1 時間以内に開始**"},
                    ],
                    "source": "Singer M et al. JAMA 2016 Sepsis-3 ([PMID 26903338](https://pubmed.ncbi.nlm.nih.gov/26903338/))",
                },
                {
                    "name": "Sepsis-3 定義（Sepsis vs Septic shock）",
                    "items": [
                        {"label": "Sepsis", "criteria": "感染 + SOFA ≥ 2 点上昇（臓器障害）", "points": "—"},
                        {"label": "**Septic shock**", "criteria": "**Sepsis + 輸液負荷後も MAP ≥ 65 維持に昇圧剤必要 + Lactate > 2 mmol/L**", "points": "—"},
                    ],
                    "interpretation": [
                        {"range": "Sepsis", "meaning": "院内死亡率 > 10% — 集中治療"},
                        {"range": "**Septic shock**", "meaning": "**院内死亡率 > 40% — 個室管理 + 看護師監視強化（昇圧剤・CV/A line・source control）/ 進行・治療反応不良で高度医療機関へ転院**"},
                    ],
                    "source": "Singer M et al. JAMA 2016 Sepsis-3 ([PMID 26903338](https://pubmed.ncbi.nlm.nih.gov/26903338/))",
                },
            ],
            "options": [
                {
                    "scene": "昇圧剤",
                    "conventional": "輸液のみで経過観察",
                    "necessity": _a("A6⑦ NE 持続点滴") + " — Surviving Sepsis 2021 第一選択",
                },
                {
                    "scene": "抗菌薬",
                    "conventional": "セフトリアキソン IV (間欠)",
                    "necessity": _a("A4 + A3 ペア") + " — PIPC/TAZ・MEPM 持続点滴 (重症で合理的)",
                },
            ],
            "contributions": [
                _a("A6⑦ 昇圧剤") + ": **3 点 / 日**",
                _a("A4 + A3") + ": 2 点 (β-ラクタム持続併用時、該当患者ダブル達成)",
            ],
            "clinical_examples": [
                {
                    "title": "例 A: 尿路感染症由来の Sepsis（qSOFA 2 点 / Sepsis-3）",
                    "background": "76 歳女性、糖尿病、神経因性膀胱、尿道留置カテーテル",
                    "presentation": "1 日前から発熱と全身倦怠、腰背部痛、来院前に意識朦朧",
                    "vitals": "GCS 14、BT 38.9℃、BP 108/64、HR 116、呼吸数 24、SpO2 95% RA、CVA 圧痛 +",
                    "labs": "WBC 16,800、CRP 22、BUN 24、Cre 1.5（base 0.8）、Lactate 2.1、尿沈渣 WBC 50/HPF、尿培養提出、血液培養 2 セット提出",
                    "score_calc": "qSOFA = 呼吸数(1) + 意識(1) = **2 点（陽性）** → Sepsis-3 評価で SOFA ≥ 2 上昇 → **Sepsis**",
                    "plan": "緊急入院、輸液 30 mL/kg を 3 時間以内、CMZ 1g IV q8h（地域 ESBL 状況見て検討）、カテーテル交換、血糖管理、source control 評価",
                    "necessity_contrib": _a("A3 注射薬剤 3 種類以上 (抗菌薬 + 輸液 + 補正剤)") + " 1 点 / " + _a("A7 救急搬送後") + " 2 点 → 入院日 **A 計 3 点で該当**。昇圧剤不要のため A6⑦ なし",
                },
                {
                    "title": "例 B: Septic shock（qSOFA 3 点 / Sepsis-3 ショック）",
                    "background": "68 歳男性、肝硬変、糖尿病、自宅独居",
                    "presentation": "数日前からの発熱と全身倦怠、家族発見時に意識朦朧、救急搬送",
                    "vitals": "GCS 11、BT 39.2℃、BP 78/42 (輸液 1L 後も 88/52)、HR 138、呼吸数 30、SpO2 88% RA、末梢冷感、皮膚モットリング",
                    "labs": "WBC 22,500、CRP 28、PCT 35、BUN 48、Cre 2.4、AST/ALT 285/198、Lactate 4.8、ABG: pH 7.22 / HCO3 14、Procalcitonin 高値",
                    "score_calc": "qSOFA = 呼吸数(1) + 意識(1) + 血圧(1) = **3 点（陽性）** → SOFA 大幅上昇 + 輸液後も MAP < 65 + Lactate > 2 → **Septic shock（院内死亡率 > 40%）**",
                    "plan": "**個室入室、看護師監視強化下で集中管理**、CV 挿入、ノルアドレナリン 0.1 μg/kg/min から開始、MEPM 1g IV → 3g/24h 持続（重症 + 肝障害）、A line、Foley、HFNC または NPPV 装着、適応次第でハイドロコルチゾン 200 mg/日、source 探索 (CT)。挿管必要・治療反応不良なら高度医療機関へ転院",
                    "necessity_contrib": _a("A6⑦ 昇圧剤注射 (NE 持続)") + " **3 点** / " + _a("A4 シリンジポンプ (NE + MEPM 持続)") + " 1 点 / " + _a("A3 注射薬剤 3 種類以上") + " 1 点 / " + _c("C23 CV 挿入") + " **4 日カウント該当患者扱い** / " + _a("A7 救急搬送後") + " 2 点 → **A 計 7 点 + C 該当 = 圧倒的な該当患者**",
                },
            ],
            "evidence": [
                {
                    "strength": "強",
                    "ref": "[PMID 26903338](https://pubmed.ncbi.nlm.nih.gov/26903338/)",
                    "summary": "Singer M et al. JAMA 2016 — Sepsis-3 公式定義。qSOFA ≥ 2 / SOFA ≥ 2 上昇 / Septic shock の基準を確立",
                },
                {
                    "strength": "強",
                    "ref": "[PMID 34599691](https://pubmed.ncbi.nlm.nih.gov/34599691/)",
                    "summary": "Surviving Sepsis Campaign 2021 — NE 第一選択、心機能低下時は ± dobutamine",
                },
            ],
            "guardrails": [
                "感染源コントロール (drainage、source control) を最優先",
                "1 時間バンドル (培養 → 広域抗菌薬 → 輸液 → 昇圧剤) 遵守",
                "lactate clearance + SvO2 で輸液・昇圧剤調整",
            ],
            "orders": [
                "「NE 0.05 μg/kg/min 開始、MAP > 65 維持」",
                "「PIPC/TAZ 4.5g 負荷 → 13.5g/24h 持続」",
                "「血液培養 2 セット (抗菌薬投与前)」",
            ],
        },
        {
            "id": 13,
            "name": "発熱性好中球減少症 (FN)",
            "icon": "🦠",
            "entry": [
                "化学療法後 7-10 日 (好中球最低期)",
                "好中球数 < 500 /μL + 発熱 (38.3℃以上 or 38.0℃以上 1h 持続)",
                "MASCC スコアで高リスク (< 21)",
            ],
            "options": [
                {
                    "scene": "経験的抗菌薬",
                    "conventional": "セフェピム IV (間欠)",
                    "necessity": _a("A4 + A3 ペア") + " — PIPC/TAZ・MEPM 持続点滴 (免疫不全で trough 重要)",
                },
            ],
            "contributions": [
                _a("A4 + A3") + ": **2 点** (該当患者)",
                _a("A2 呼吸ケア (酸素 A2)") + ": 1 点 (呼吸器症状あれば)",
            ],
            "evidence": [
                {
                    "strength": "強",
                    "ref": "IDSA FN ガイドライン 2010",
                    "summary": "高リスク FN は早期広域抗菌薬 IV (1h 以内)、低リスクは経口切替検討",
                },
                {
                    "strength": "中",
                    "ref": "[PMID 38864155](https://pubmed.ncbi.nlm.nih.gov/38864155/)",
                    "summary": "BLING-III: 免疫不全例での持続点滴は trough 維持に有利",
                },
            ],
            "guardrails": [
                "MASCC スコアで高/低リスク層別、低リスクは外来管理も検討",
                "G-CSF 投与の適応評価 (高リスク・遷延)",
                "培養結果で 48-72h 以内に de-escalation",
            ],
            "orders": [
                "「PIPC/TAZ 4.5g 負荷 → 13.5g/24h 持続」",
                "「血液培養 2 セット (末梢 + CV)」",
                "「G-CSF (フィルグラスチム 75μg/日 SC) 適応評価」",
            ],
        },
        {
            "id": 14,
            "name": "重症腎盂腎炎（ESBL 等高 MIC）",
            "icon": "🦠",
            "entry": [
                "発熱・腰背部痛・CVA tenderness + 尿沈渣 WBC ≥ 10/HPF",
                "尿培養で ESBL 産生大腸菌・KP 検出 (or 既往あり)",
                "qSOFA ≥ 1、SIRS 合併",
            ],
            "options": [
                {
                    "scene": "抗菌薬",
                    "conventional": "セフトリアキソン IV (ESBL では無効)",
                    "necessity": _a("A4 + A3 ペア") + " — MEPM 持続点滴 + 多剤併用 (高 MIC で %T>MIC 重要)",
                },
            ],
            "contributions": [
                _a("A4 + A3") + ": **2 点** (該当患者)",
            ],
            "evidence": [
                {
                    "strength": "強",
                    "ref": "IDSA UTI ガイドライン",
                    "summary": "ESBL 産生菌による複雑性 UTI はカルバペネム第一選択",
                },
                {
                    "strength": "中",
                    "ref": "[PMID 38864155](https://pubmed.ncbi.nlm.nih.gov/38864155/)",
                    "summary": "高 MIC 菌では持続点滴で %T>MIC を確保するメリットあり",
                },
            ],
            "guardrails": [
                "閉塞性腎盂腎炎 (結石・腫瘍) は緊急ドレナージ (尿管ステント・PCN) 検討",
                "ESBL 既往 + 重症で empiric MEPM、培養結果で見直し",
                "腎機能評価で用量調整",
            ],
            "orders": [
                "「MEPM 1g 負荷 → 3g/24h 持続 (CrCl 80 想定)」",
                "「腹部 CT (閉塞・膿瘍評価)」",
                "「尿培養 + 血液培養 2 セット」",
            ],
        },
    ]),
    # ========================================================================
    # 🩹 創傷 (2 疾患)
    # ========================================================================
    ("🩹 創傷・皮膚軟部組織感染症", [
        {
            "id": 15,
            "name": "蜂窩織炎",
            "icon": "🩹",
            "entry": [
                "発赤・腫脹・熱感・圧痛 + 発熱",
                "境界不明瞭な紅斑、リンパ管炎所見あり",
                "重症化・壊死性筋膜炎との鑑別 (LRINEC スコア)",
            ],
            "options": [
                {
                    "scene": "創傷処置",
                    "conventional": "観察のみ、抗菌薬内服のみ",
                    "necessity": _a("A1 創傷処置 (褥瘡を除く)") + " + " + _a("A3 注射薬剤 3 種類") + " — 洗浄・被覆 + IV 抗菌薬 + 補液 + 鎮痛",
                },
            ],
            "contributions": [
                _a("A1 + A3") + ": **2 点** (該当患者)",
            ],
            "evidence": [
                {
                    "strength": "強",
                    "ref": "IDSA SSTI ガイドライン 2014",
                    "summary": "中等症以上の蜂窩織炎は IV β-ラクタム、外来治療失敗例は入院",
                },
            ],
            "guardrails": [
                "壊死性筋膜炎の見逃しに注意 (緊急外科コンサルト)",
                "MRSA リスク評価 (院内、ホームレス、IVDU)",
                "下肢・糖尿病例では深部感染・骨髄炎評価",
            ],
            "orders": [
                "「セファゾリン 1g IV q8h」",
                "「創部洗浄 1 日 1 回 + 滅菌ガーゼ被覆」",
                "「下肢挙上、安静」",
            ],
        },
        {
            "id": 16,
            "name": "糖尿病性足病変（感染合併）",
            "icon": "🩹",
            "entry": [
                "DM 患者の足部潰瘍・蜂窩織炎・壊疽",
                "Wagner 分類 / IWGDF 分類で評価",
                "X 線・MRI で骨髄炎の有無確認",
            ],
            "options": [
                {
                    "scene": "創傷処置",
                    "conventional": "外来管理を継続",
                    "necessity": _a("A1 創傷処置") + " + " + _a("A3 注射薬剤 3 種類") + " — 入院でデブリードマン + IV 抗菌薬 + 血糖管理 + 補液",
                },
            ],
            "contributions": [
                _a("A1 + A3") + ": **2 点** (該当患者)",
                _a("A6⑩ ドレナージ管理") + ": 2 点 (深部膿瘍・壊死組織のドレナージ後)",
            ],
            "evidence": [
                {
                    "strength": "強",
                    "ref": "IWGDF ガイドライン 2023",
                    "summary": "糖尿病性足感染症は早期評価 + 適切なデブリ + 多分野連携 (整形・形成・血管)",
                },
            ],
            "guardrails": [
                "血糖コントロール改善必須 (HbA1c、入院中血糖プロファイル)",
                "末梢動脈疾患 (PAD) 評価で血行再建も検討",
                "切断回避を最優先 (mortality 上昇の主因)",
            ],
            "orders": [
                "「足部 X 線 + MRI (骨髄炎評価)」",
                "「ABPC/SBT 3g IV q6h」",
                "「創傷洗浄 + デブリードマン (毎日)」",
                "「血糖プロファイル 4 検 (毎食前 + 寝る前)」",
            ],
        },
    ]),
    # ========================================================================
    # 💊 ペイン (1 疾患)
    # ========================================================================
    ("💊 ペイン領域（神経障害性疼痛）", [
        {
            "id": 17,
            "name": "PHN / CRPS / 難治性神経障害性疼痛",
            "icon": "💊",
            "entry": [
                "帯状疱疹後神経痛 (PHN) で経口薬無効 (3 種類以上で改善せず)",
                "CRPS Type I/II の急性期 (灼熱痛、アロディニア)",
                "他治療抵抗性の難治性神経障害性疼痛",
            ],
            "options": [
                {
                    "scene": "難治性神経痛のレスキュー",
                    "conventional": "プレガバリン・トラマドール継続 (経口)",
                    "necessity": _a("A6⑧ IV リドカイン") + " — ペイン科判断で症例選択（**弱エビデンス**）",
                },
            ],
            "contributions": [
                _a("A6⑧ 抗不整脈剤の使用") + ": **3 点 / 日 × 1-7 日**",
                _a("A4 シリンジポンプ") + ": 1 点 (リドカイン持続)",
            ],
            "evidence": [
                {
                    "strength": "弱",
                    "ref": "[PMID 16013891](https://pubmed.ncbi.nlm.nih.gov/16013891/)",
                    "summary": "Hempenstall et al. PLoS Med 2005 — PHN への IV リドカインは効果なしと結論",
                },
                {
                    "strength": "弱",
                    "ref": "[PMID 36288104](https://pubmed.ncbi.nlm.nih.gov/36288104/)",
                    "summary": "Lee et al. Clin J Pain 2022 — IV リドカイン慢性神経障害性疼痛のエビデンス不十分",
                },
            ],
            "guardrails": [
                "**エビデンスは弱い**ことを認識した上でペイン科判断で症例選択",
                "看護必要度のために適応外で投与することは **絶対にやらない**",
                "心電図モニター必須 (QT 延長・不整脈リスク)",
                "効果判定 24-48h で継続要否判断、無効なら速やかに中止",
            ],
            "orders": [
                "「リドカイン 1 mg/kg/h 開始、効果と副作用見て調整 (max 5 mg/kg/h)」",
                "「心電図モニター継続、リドカイン血中濃度 (利用可能なら)」",
                "「ペイン科コンサルト」",
            ],
        },
    ]),
]


# ---------------------------------------------------------------------------
# Markdown 生成 (DISEASE_DATA → DISEASE_MANUAL_MARKDOWN)
# ---------------------------------------------------------------------------

_HEADER = """# 看護必要度 疾患別マニュアル — 医師担当部分

**作成**: おもろまちメディカルセンター 副院長 久保田 徹
**対象**: 地域包括医療病棟入院料1（5F・6F）担当医
**基準日**: 2026-06-01 新基準対応
**位置づけ**: ミニレクチャー（行動変容の概念整理）に対する**疾患入口型 実臨床リファレンス**

---

## 🛡️ 倫理ガードレール（全疾患共通、最上段）

1. **適応外の処置は絶対に行わない**
2. **医学的判断は変わらない** — 適応判断・治療方針は従来通り、ガイドライン準拠
3. **「選び方」の問題として捉える** — 複数の選択肢がある中での選択を意識化する

→ 看護必要度のための過剰治療・適応外投与は医療倫理違反、施設基準取消・診療報酬遡及返還のリスク。

## 🎨 色分け凡例

| 区分 | 色 | 内容 |
|---|---|---|
| 🟠 A 項目 | オレンジ | 入院中の処置・薬剤（毎日評価） |
| 🔵 B 項目 | 青 | 患者の状況等（**新基準では評価不要**） |
| 🟢 C 項目 | 緑 | 手術・侵襲的処置（4-5 日カウント） |

---

## 📚 目次（17 疾患）

"""


def _render_disease_md(d: dict[str, Any]) -> str:
    """1 疾患の dict を markdown ブロックに変換.

    見出しの前に明示的な HTML アンカー (<a id="disease-N">) を置くことで、
    Streamlit の markdown スラッグ化が括弧・スラッシュ等で破綻する問題を回避。
    目次は (#disease-N) で参照する (シンプルで確実)。
    """
    lines: list[str] = []
    # 明示アンカー (Streamlit でも GitHub でも安定)
    lines.append(f'<a id="disease-{d["id"]}"></a>')
    lines.append(f"### {d['id']}. {d['icon']} {d['name']}")
    lines.append("")

    # 入口
    lines.append("#### 🩺 入口（こんな患者を見たら）")
    for e in d["entry"]:
        lines.append(f"- {e}")
    lines.append("")

    # 重症度評価スコア (Phase 2A 以降)
    if d.get("severity_assessment"):
        lines.append("#### 📊 重症度評価スコア")
        for score in d["severity_assessment"]:
            lines.append(f"##### {score['name']}")
            if score.get("items"):
                lines.append("")
                lines.append("| 項目 | 該当条件 | 点数 |")
                lines.append("|---|---|:-:|")
                for it in score["items"]:
                    lines.append(f"| {it['label']} | {it['criteria']} | {it.get('points','1点')} |")
                lines.append("")
            if score.get("interpretation"):
                lines.append("**解釈**:")
                for itp in score["interpretation"]:
                    lines.append(f"- **{itp['range']}**: {itp['meaning']}")
                lines.append("")
            if score.get("source"):
                lines.append(f"*出典: {score['source']}*")
                lines.append("")

    # 選択肢比較テーブル
    lines.append("#### 💊 第一選択 vs 看護必要度貢献選択肢（医学的に同等以上）")
    lines.append("")
    lines.append("| 場面 | 従来の選択肢 | 看護必要度に貢献する選択肢 |")
    lines.append("|---|---|---|")
    for opt in d["options"]:
        lines.append(f"| {opt['scene']} | {opt['conventional']} | {opt['necessity']} |")
    lines.append("")

    # 看護必要度寄与
    lines.append("#### 📋 看護必要度寄与")
    for c in d["contributions"]:
        lines.append(f"- {c}")
    lines.append("")

    # 具体例 (Phase 2A 以降)
    if d.get("clinical_examples"):
        lines.append("#### 🩺 具体例（典型的な患者像）")
        for ex in d["clinical_examples"]:
            lines.append(f"##### {ex['title']}")
            if ex.get("background"):
                lines.append(f"- **背景**: {ex['background']}")
            if ex.get("presentation"):
                lines.append(f"- **主訴・経過**: {ex['presentation']}")
            if ex.get("vitals"):
                lines.append(f"- **バイタル・身体所見**: {ex['vitals']}")
            if ex.get("labs"):
                lines.append(f"- **検査所見**: {ex['labs']}")
            if ex.get("score_calc"):
                lines.append(f"- **スコア計算**: {ex['score_calc']}")
            if ex.get("plan"):
                lines.append(f"- **治療方針**: {ex['plan']}")
            if ex.get("necessity_contrib"):
                lines.append(f"- **看護必要度寄与**: {ex['necessity_contrib']}")
            lines.append("")

    # エビデンス
    lines.append("#### 📚 エビデンス")
    for ev in d["evidence"]:
        lines.append(f"- **{ev['strength']}**: {ev['ref']} — {ev['summary']}")
    lines.append("")

    # 倫理ガードレール
    lines.append("#### 🛡️ 倫理ガードレール（疾患特異）")
    for g in d["guardrails"]:
        lines.append(f"- {g}")
    lines.append("")

    # オーダー例
    lines.append("#### 📝 オーダー例")
    for o in d["orders"]:
        lines.append(f"- {o}")
    lines.append("")
    lines.append("---")
    lines.append("")
    return "\n".join(lines)


def _build_markdown() -> str:
    """DISEASE_DATA から完全な markdown 文字列を生成."""
    out: list[str] = [_HEADER]

    # 目次 (明示アンカー #disease-N で参照)
    for group_name, diseases in DISEASE_DATA:
        out.append(f"### {group_name}")
        out.append("")
        for d in diseases:
            out.append(f"- [{d['id']}. {d['icon']} {d['name']}](#disease-{d['id']})")
        out.append("")
    out.append("---")
    out.append("")

    # 本文 (疾患群ごと)
    for group_name, diseases in DISEASE_DATA:
        out.append(f"## {group_name}")
        out.append("")
        for d in diseases:
            out.append(_render_disease_md(d))

    out.append("")
    out.append("---")
    out.append("")
    out.append("*このマニュアルは医師の臨床判断を支援する目的で作成。")
    out.append("医療倫理・診療報酬制度の最新情報は厚労省通知を必ず参照すること。*")

    return "\n".join(out)


DISEASE_MANUAL_MARKDOWN: str = _build_markdown()


# ---------------------------------------------------------------------------
# DOCX 生成 (院内 LAN 配布用、Word で開いて PDF 化可能)
# ---------------------------------------------------------------------------

# 簡易 markdown 除去 (DOCX には HTML span を入れない、絵文字 prefix のみ残す)
def _strip_html(s: str) -> str:
    """span タグを除去して絵文字 + 太字記法を残す."""
    import re
    # <span ...>X</span> → X
    s = re.sub(r'<span[^>]*>(.*?)</span>', r'\1', s, flags=re.DOTALL)
    return s


def generate_docx() -> bytes:
    """疾患別マニュアル DOCX を生成して bytes で返す.

    Streamlit の st.download_button に渡す用途。
    Word で開いて PDF 化することを想定 (高度な書式は使わない)。
    """
    from docx import Document
    from docx.shared import Pt, Cm

    doc = Document()

    # 既定スタイル (日本語フォント)
    style = doc.styles["Normal"]
    style.font.name = "Yu Gothic"
    style.font.size = Pt(11)

    # ページ余白 (A4 縦)
    for section in doc.sections:
        section.top_margin = Cm(2.0)
        section.bottom_margin = Cm(2.0)
        section.left_margin = Cm(2.0)
        section.right_margin = Cm(2.0)

    # タイトル
    doc.add_heading("看護必要度 疾患別マニュアル — 医師担当部分", level=0)
    doc.add_paragraph("作成: おもろまちメディカルセンター 副院長 久保田 徹")
    doc.add_paragraph("対象: 地域包括医療病棟入院料1（5F・6F）担当医")
    doc.add_paragraph("基準日: 2026-06-01 新基準対応")
    doc.add_paragraph("位置づけ: ミニレクチャー（行動変容の概念整理）に対する疾患入口型 実臨床リファレンス")

    # 倫理ガードレール
    doc.add_heading("🛡️ 倫理ガードレール（全疾患共通、最上段）", level=1)
    for line in [
        "1. 適応外の処置は絶対に行わない",
        "2. 医学的判断は変わらない — 適応判断・治療方針は従来通り、ガイドライン準拠",
        "3. 「選び方」の問題として捉える — 複数の選択肢がある中での選択を意識化する",
    ]:
        doc.add_paragraph(line)
    doc.add_paragraph(
        "→ 看護必要度のための過剰治療・適応外投与は医療倫理違反、"
        "施設基準取消・診療報酬遡及返還のリスク。"
    )

    # 色分け凡例
    doc.add_heading("🎨 色分け凡例", level=1)
    legend_table = doc.add_table(rows=4, cols=3)
    legend_table.style = "Light Grid"
    cells = legend_table.rows[0].cells
    cells[0].text = "区分"
    cells[1].text = "色"
    cells[2].text = "内容"
    legend_data = [
        ("🟠 A 項目", "オレンジ", "入院中の処置・薬剤（毎日評価）"),
        ("🔵 B 項目", "青", "患者の状況等（新基準では評価不要）"),
        ("🟢 C 項目", "緑", "手術・侵襲的処置（4-5 日カウント）"),
    ]
    for i, row_data in enumerate(legend_data, start=1):
        for j, val in enumerate(row_data):
            legend_table.rows[i].cells[j].text = val

    doc.add_page_break()

    # 17 疾患を疾患群ごとに章立て
    for group_name, diseases in DISEASE_DATA:
        doc.add_heading(group_name, level=1)
        for d in diseases:
            doc.add_heading(f"{d['id']}. {d['icon']} {d['name']}", level=2)

            # 入口
            doc.add_heading("🩺 入口（こんな患者を見たら）", level=3)
            for e in d["entry"]:
                doc.add_paragraph(_strip_html(e), style="List Bullet")

            # 重症度評価スコア (Phase 2A 以降)
            if d.get("severity_assessment"):
                doc.add_heading("📊 重症度評価スコア", level=3)
                for score in d["severity_assessment"]:
                    doc.add_heading(score["name"], level=4)
                    if score.get("items"):
                        sa_table = doc.add_table(rows=1 + len(score["items"]), cols=3)
                        sa_table.style = "Light Grid"
                        sa_table.rows[0].cells[0].text = "項目"
                        sa_table.rows[0].cells[1].text = "該当条件"
                        sa_table.rows[0].cells[2].text = "点数"
                        for i, it in enumerate(score["items"], start=1):
                            sa_table.rows[i].cells[0].text = _strip_html(it["label"])
                            sa_table.rows[i].cells[1].text = _strip_html(it["criteria"])
                            sa_table.rows[i].cells[2].text = _strip_html(it.get("points", "1点"))
                    if score.get("interpretation"):
                        doc.add_paragraph("解釈:")
                        for itp in score["interpretation"]:
                            doc.add_paragraph(
                                f"{itp['range']}: {itp['meaning']}",
                                style="List Bullet",
                            )
                    if score.get("source"):
                        p = doc.add_paragraph()
                        run = p.add_run(f"出典: {_strip_html(score['source'])}")
                        run.italic = True

            # 選択肢比較テーブル
            doc.add_heading("💊 第一選択 vs 看護必要度貢献選択肢", level=3)
            opt_table = doc.add_table(rows=1 + len(d["options"]), cols=3)
            opt_table.style = "Light Grid"
            opt_table.rows[0].cells[0].text = "場面"
            opt_table.rows[0].cells[1].text = "従来の選択肢"
            opt_table.rows[0].cells[2].text = "看護必要度に貢献する選択肢"
            for i, opt in enumerate(d["options"], start=1):
                opt_table.rows[i].cells[0].text = _strip_html(opt["scene"])
                opt_table.rows[i].cells[1].text = _strip_html(opt["conventional"])
                opt_table.rows[i].cells[2].text = _strip_html(opt["necessity"])

            # 看護必要度寄与
            doc.add_heading("📋 看護必要度寄与", level=3)
            for c in d["contributions"]:
                doc.add_paragraph(_strip_html(c), style="List Bullet")

            # 具体例 (Phase 2A 以降)
            if d.get("clinical_examples"):
                doc.add_heading("🩺 具体例（典型的な患者像）", level=3)
                for ex in d["clinical_examples"]:
                    doc.add_heading(ex["title"], level=4)
                    field_labels = [
                        ("background", "背景"),
                        ("presentation", "主訴・経過"),
                        ("vitals", "バイタル・身体所見"),
                        ("labs", "検査所見"),
                        ("score_calc", "スコア計算"),
                        ("plan", "治療方針"),
                        ("necessity_contrib", "看護必要度寄与"),
                    ]
                    for key, jp in field_labels:
                        if ex.get(key):
                            doc.add_paragraph(
                                f"{jp}: {_strip_html(ex[key])}",
                                style="List Bullet",
                            )

            # エビデンス
            doc.add_heading("📚 エビデンス", level=3)
            for ev in d["evidence"]:
                doc.add_paragraph(
                    f"[{ev['strength']}] {_strip_html(ev['ref'])} — {_strip_html(ev['summary'])}",
                    style="List Bullet",
                )

            # 倫理ガードレール
            doc.add_heading("🛡️ 倫理ガードレール（疾患特異）", level=3)
            for g in d["guardrails"]:
                doc.add_paragraph(_strip_html(g), style="List Bullet")

            # オーダー例
            doc.add_heading("📝 オーダー例", level=3)
            for o in d["orders"]:
                doc.add_paragraph(_strip_html(o), style="List Bullet")

            doc.add_paragraph("")  # 疾患間スペース

    # 締め
    doc.add_page_break()
    doc.add_paragraph(
        "このマニュアルは医師の臨床判断を支援する目的で作成。"
        "医療倫理・診療報酬制度の最新情報は厚労省通知を必ず参照すること。"
    )

    buf = BytesIO()
    doc.save(buf)
    return buf.getvalue()


# ---------------------------------------------------------------------------
# Streamlit 描画ヘルパー
# ---------------------------------------------------------------------------

def render_in_streamlit(streamlit_module) -> None:
    """Streamlit でマニュアルを描画 + DOCX ダウンロードボタンを提供.

    Args:
        streamlit_module: ``streamlit`` モジュール
    """
    st = streamlit_module
    st.markdown(DISEASE_MANUAL_MARKDOWN, unsafe_allow_html=True)
    st.markdown("---")
    try:
        docx_bytes = generate_docx()
        st.download_button(
            "📥 疾患別マニュアル DOCX をダウンロード（Word で開いて印刷・PDF 化可）",
            data=docx_bytes,
            file_name="看護必要度_疾患別マニュアル.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            key="nn_disease_manual_docx",
        )
    except Exception as e:
        st.warning(f"⚠️ DOCX 生成中にエラー: {e}（python-docx の依存関係を確認してください）")


# ---------------------------------------------------------------------------
# CLI: docs/admin/nursing_necessity_disease_manual.md を再生成
# ---------------------------------------------------------------------------

def _regenerate_md_file() -> Path:
    """git 管理用の md ファイルを docs/admin/ に書き出す."""
    out_path = Path(__file__).resolve().parent.parent / "docs" / "admin" / "nursing_necessity_disease_manual.md"
    out_path.write_text(DISEASE_MANUAL_MARKDOWN, encoding="utf-8")
    return out_path


if __name__ == "__main__":
    p = _regenerate_md_file()
    print(f"✅ Regenerated: {p}")
    print(f"   Total characters: {len(DISEASE_MANUAL_MARKDOWN):,}")
    n_diseases = sum(len(group_diseases) for _, group_diseases in DISEASE_DATA)
    print(f"   Total diseases: {n_diseases}")
