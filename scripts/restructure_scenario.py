"""
シナリオ台本の Scene 順序を再構成するスクリプト
朝のブリーフィング → 日次推移 → 稼働率の正体 → 運営分析 の流れに変更
"""
from docx import Document
from docx.shared import Pt
from copy import deepcopy
import lxml.etree as etree

src = Document('docs/admin/bed_control_demo_scenario.docx')

# --- paragraph ranges (inclusive) ---
HEADER   = (0, 3)      # タイトル
PREP     = (4, 7)      # 事前準備
SC1_OLD  = (8, 21)     # 旧Scene 1: 運営分析ギャップ
SC1B_OLD = (22, 49)    # 旧Scene 1-B: 稼働率の正体
SC2_OLD  = (50, 58)    # 旧Scene 2: 日次推移
SC2B     = (59, 96)    # Scene 2-B: フェーズ構成
SC2C     = (97, 127)   # Scene 2-C: What-if
SC2D     = (128, 141)  # Scene 2-D: 退院マネジメント
SC2E     = (142, 186)  # Scene 2-E: 夜勤安全ライン
SC3_OLD  = (187, 195)  # 旧Scene 3: LOS
SC4_OLD  = (196, 209)  # 旧Scene 4: クリア計画
SC5_OLD  = (210, 219)  # 旧Scene 5: 病棟別
SC6_OLD  = (220, 225)  # 旧Scene 6: 2026改定
SC7_OLD  = (226, 239)  # 旧Scene 7: 賞与
SC8_OLD  = (240, 278)  # 旧Scene 8: 助け合い
SC9_OLD  = (279, 427)  # 旧Scene 9: まとめ + デモ手順
SC_ADD   = (428, 469)  # Scene 追加: 短手3

def copy_para(src_para, dst_doc):
    """Copy a paragraph from src to dst document, preserving formatting."""
    new_p = deepcopy(src_para._element)
    dst_doc.element.body.append(new_p)

def add_text_para(dst_doc, text, bold=False, size_pt=None):
    """Add a simple text paragraph."""
    p = dst_doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    if size_pt:
        run.font.size = Pt(size_pt)
    return p

def copy_range(src_doc, dst_doc, start, end):
    """Copy paragraph range [start..end] inclusive."""
    for i in range(start, end + 1):
        copy_para(src_doc.paragraphs[i], dst_doc)

# --- Build new document ---
dst = Document()
# Remove default empty paragraph
if dst.paragraphs:
    dst.element.body.remove(dst.paragraphs[0]._element)

# Copy section properties (page size, margins) from source
src_sectPr = src.sections[0]._sectPr
dst_sectPr = dst.sections[0]._sectPr
for child in list(dst_sectPr):
    dst_sectPr.remove(child)
for child in src_sectPr:
    dst_sectPr.append(deepcopy(child))

# 1. Header (0-3)
copy_range(src, dst, 0, 3)

# 2. 事前準備 (4-7) — update opening hook
copy_range(src, dst, 4, 7)

# 3. NEW Scene 1: ☀️ 朝のブリーフィング
new_scene1_lines = [
    ("■ Scene 1　☀️ 朝のブリーフィング — 毎朝これを見ます　(2分)", True, 16.51),
    ("▶ ブリーフィング画面「☀️ 今日のブリーフィング」を表示", False, None),
    ("", False, None),
    ("💬 まず最初に、毎朝この画面を見るところから始めます。これが「今日のブリーフィング」です。", False, None),
    ("", False, None),
    ("▶ 稼働率ゲージを指しながら", False, None),
    ("", False, None),
    ("💬 左のゲージが「月平均稼働率」——今月の成績です。今88.8%。目標の90%まであと1.2ポイント。", False, None),
    ("", False, None),
    ("▶ 在院患者数メトリックを指しながら", False, None),
    ("", False, None),
    ("💬 その隣が「在院患者数（今の空き）」。今の空床数が一目でわかります。空床が多ければ「今日、入院を受けられる余裕がある」という即時判断ができます。", False, None),
    ("", False, None),
    ("▶ 💡インサイトコメントを指しながら", False, None),
    ("", False, None),
    ("💬 そして一番大事なのがこの一行コメントです。空床数と稼働率を組み合わせて、今の状況を4パターンで自動判定しています。", False, None),
    ("", False, None),
    ("💬 🟢「回転良好」—— 空床が少なく稼働率も高い。ベッドがよく回っている状態です。", False, None),
    ("💬 🔴「入院増が必要」—— 空床が多く稼働率も低い。入院受入れの強化が必要。", False, None),
    ("💬 ⚠️「詰まりの兆候」—— ベッドは埋まっているが回転していない。退院調整を確認。", False, None),
    ("💬 🟡「受入余地あり」—— 回転は良いが空床あり。新規入院でさらに伸ばせる。", False, None),
    ("", False, None),
    ("💬 ポイントは「空床数は今すぐの受入判断、稼働率は今月の運営評価」——この2つは別のものさしだということです。空床が少ないから安心とは限りません。稼働率が低ければ「詰まっている」。逆に空床があっても稼働率が高ければ、ベッドはよく回っている証拠です。", False, None),
    ("", False, None),
    ("💬 毎朝この画面を10秒見るだけで、今日の動き方が決まります。では次に、もう少し長いスパンで見てみましょう。", False, None),
    ("", False, None),
]

for text, bold, size in new_scene1_lines:
    add_text_para(dst, text, bold=bold, size_pt=size)

# 4. Scene 2 (new): 日次推移 = old Scene 2 (50-58), renumber
# First copy with title change
p_title = src.paragraphs[50]
add_text_para(dst,
    "■ Scene 2　稼働率の波を「見る」— 週末の谷に隠れたコスト　(2分)",
    bold=True, size_pt=16.51)
copy_range(src, dst, 51, 58)

# 5. Scene 3 (new): 稼働率の正体 = old Scene 1-B (22-49), renumber
add_text_para(dst,
    "■ Scene 3　稼働率の正体 — 100%を超えるのは、なぜ？　(2分)",
    bold=True, size_pt=16.51)
copy_range(src, dst, 23, 49)

# 6. Scene 4 (new): 運営分析 = old Scene 1 (8-21), renumber
# Remove paras 18-20 (briefing insight) since moved to new Scene 1
add_text_para(dst,
    "■ Scene 4　「今日は4月20日。あと10日で、いくら取り戻せますか？」　(2分)",
    bold=True, size_pt=16.51)
copy_range(src, dst, 9, 16)  # Skip 17 (empty), 18-20 (briefing), 21 (empty)

# 7. Scene 5 (new): フェーズ構成 = old Scene 2-B (59-96)
add_text_para(dst,
    "■ Scene 5　フェーズ構成の理想比率 — 数学で決める（Little法則）　(3分)",
    bold=True, size_pt=16.51)
copy_range(src, dst, 60, 96)

# 8. Scene 6 (new): What-if = old Scene 2-C (97-127)
add_text_para(dst,
    "■ Scene 6　What-if シミュレーション — 経営会議で「もし○○なら」を即答する　(2分)",
    bold=True, size_pt=16.51)
copy_range(src, dst, 98, 127)

# 9. Scene 7 (new): 退院マネジメント = old Scene 2-D (128-141)
add_text_para(dst,
    "■ Scene 7　退院マネジメント — 家族目線が運営を変える　(3分)",
    bold=True, size_pt=16.51)
copy_range(src, dst, 129, 141)

# 10. Scene 8 (new): 夜勤安全ライン = old Scene 2-E (142-186)
add_text_para(dst,
    "■ Scene 8　夜勤安全ライン — 夜勤を守りながら稼働率を上げる　(3分)",
    bold=True, size_pt=16.51)
copy_range(src, dst, 143, 186)

# 11. Scene 9 (new): LOS = old Scene 3 (187-195)
add_text_para(dst,
    "■ Scene 9　「6F平均在院日数 21.3日」— 病棟別に見ると基準超過　(3分)",
    bold=True, size_pt=16.51)
copy_range(src, dst, 188, 195)

# 12. Scene 10 (new): クリア計画 = old Scene 4 (196-209)
add_text_para(dst,
    "■ Scene 10　クリア計画 —「C群からあと何名退院させれば基準クリアか」　(3分)",
    bold=True, size_pt=16.51)
copy_range(src, dst, 197, 209)

# 13. Scene 11 (new): 病棟別 = old Scene 5 (210-219)
add_text_para(dst,
    "■ Scene 11　病棟別に見る — 5Fと6Fでは課題が違う　(2分)",
    bold=True, size_pt=16.51)
copy_range(src, dst, 211, 219)

# 14. Scene 12 (new): 2026改定 = old Scene 6 (220-225)
add_text_para(dst,
    "■ Scene 12　2026年度改定への備え　(1分)",
    bold=True, size_pt=16.51)
copy_range(src, dst, 221, 225)

# 15. Scene 13 (new): 賞与 = old Scene 7 (226-239)
add_text_para(dst,
    "■ Scene 13　「稼働率が上がると、私たちの賞与はどうなるか」　(2分)",
    bold=True, size_pt=16.51)
copy_range(src, dst, 227, 239)

# 16. Scene 14 (new): 助け合い = old Scene 8 (240-278)
add_text_para(dst,
    "■ Scene 14　助け合い方式（Equal Effort） —「同じだけ頑張る」から始める　(3分)",
    bold=True, size_pt=16.51)
copy_range(src, dst, 241, 278)

# 17. Scene 15 (new): まとめ = old Scene 9 (279-427)
# This includes まとめ + demo procedures
add_text_para(dst,
    "■ Scene 15　まとめ —「見える」が「動ける」に変わる　(1分)",
    bold=True, size_pt=16.51)
# Copy, but update scene references in demo procedure section
copy_range(src, dst, 280, 427)

# 18. Scene 追加: 短手3 (428-469)
copy_range(src, dst, 428, 469)

# --- Fix scene references in the まとめ/demo procedure section ---
# Update cross-references like "(Scene 1で使用)" to new numbering
ref_map = {
    "Scene 1で使用": "Scene 4で使用",
    "Scene 1・Scene 4で使用": "Scene 4・Scene 10で使用",
    "Scene 2・Scene 5・Scene 8で使用": "Scene 2・Scene 11・Scene 14で使用",
    "Scene 2-Bで使用": "Scene 5で使用",
    "Scene 2-Cで使用": "Scene 6で使用",
    "Scene 3で使用": "Scene 9で使用",
    "Scene 4で使用": "Scene 10で使用",
    "Scene 8・助け合いで使用": "Scene 14・助け合いで使用",
}

for p in dst.paragraphs:
    for run in p.runs:
        for old_ref, new_ref in ref_map.items():
            if old_ref in run.text:
                run.text = run.text.replace(old_ref, new_ref)

# Also update the demo checklist item for briefing
for p in dst.paragraphs:
    for run in p.runs:
        if "Scene 1で使用" in run.text:
            run.text = run.text.replace("Scene 1で使用", "Scene 1・Scene 4で使用")

# Save
dst.save('docs/admin/bed_control_demo_scenario.docx')
print("✅ シナリオ台本の再構成が完了しました")

# Verify new structure
doc2 = Document('docs/admin/bed_control_demo_scenario.docx')
print("\n--- 新しいScene構成 ---")
for i, p in enumerate(doc2.paragraphs):
    text = p.text.strip()
    if text.startswith('■ Scene') or text.startswith('■ 事前準備'):
        print(f'[{i}] {text}')
