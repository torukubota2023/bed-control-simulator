"""シナリオ台本 DOCX 生成のテスト.

4 本の .docx が期待通り出力されるかを検証する:
    - ファイルが生成される
    - ファイルサイズが 5KB 以上
    - python-docx で開けて段落 / テーブルが存在
    - フォント設定が Hiragino Kaku Gothic ProN
"""
from __future__ import annotations

from pathlib import Path

import pytest

from docx import Document
from docx.oxml.ns import qn

REPO_ROOT = Path(__file__).resolve().parent.parent
OUTPUT_DIR = REPO_ROOT / "docs" / "admin" / "scenario_scripts_v3.6"

EXPECTED_FILES = [
    "demo_scenario_v3.6.docx",
    "presentation_script_bedcontrol_v3.6.docx",
    "carnf_scenario_v1.docx",
    "weekend_holiday_kpi_script.docx",
]

FONT_NAME = "Hiragino Kaku Gothic ProN"


@pytest.fixture(scope="module", autouse=True)
def _ensure_generated():
    """出力ファイルが無ければスクリプトを実行して生成."""
    missing = [f for f in EXPECTED_FILES if not (OUTPUT_DIR / f).exists()]
    if missing:
        import subprocess
        subprocess.run(
            ["python3", str(REPO_ROOT / "scripts" / "generate_scenario_docx.py")],
            check=True,
            cwd=str(REPO_ROOT),
        )
    yield


@pytest.mark.parametrize("filename", EXPECTED_FILES)
def test_file_exists(filename: str) -> None:
    """4 つの .docx がすべて存在する."""
    assert (OUTPUT_DIR / filename).exists(), f"{filename} was not generated"


@pytest.mark.parametrize("filename", EXPECTED_FILES)
def test_file_size_over_5kb(filename: str) -> None:
    """各 .docx は 5KB 以上のサイズを持つ（空でない）."""
    path = OUTPUT_DIR / filename
    size_kb = path.stat().st_size / 1024
    assert size_kb >= 5.0, f"{filename} is too small: {size_kb:.1f} KB"


@pytest.mark.parametrize("filename", EXPECTED_FILES)
def test_docx_openable_with_content(filename: str) -> None:
    """python-docx で開けて、段落とテーブルがいずれも 0 でない."""
    doc = Document(str(OUTPUT_DIR / filename))
    n_paragraphs = len(doc.paragraphs)
    n_tables = len(doc.tables)
    assert n_paragraphs > 10, f"{filename}: paragraphs={n_paragraphs}"
    # テーブルは 0 以上（すべてのファイルにあるとは限らないが、期待する 4 本すべてに有り）
    assert n_tables > 0, f"{filename}: tables={n_tables}"


@pytest.mark.parametrize("filename", EXPECTED_FILES)
def test_font_hiragino_on_runs(filename: str) -> None:
    """段落の run のフォントが Hiragino Kaku Gothic ProN に設定されている."""
    doc = Document(str(OUTPUT_DIR / filename))
    checked = 0
    for para in doc.paragraphs:
        for run in para.runs:
            if not run.text.strip():
                continue
            # 明示設定（run.font.name）とEast Asian font両方を確認
            assert run.font.name == FONT_NAME, \
                f"{filename}: run.font.name={run.font.name!r}, text={run.text[:30]!r}"
            # rFonts の East Asian
            rPr = run._element.find(qn("w:rPr"))
            if rPr is not None:
                rFonts = rPr.find(qn("w:rFonts"))
                if rFonts is not None:
                    east_asia = rFonts.get(qn("w:eastAsia"))
                    assert east_asia == FONT_NAME, \
                        f"{filename}: eastAsia={east_asia!r}"
            checked += 1
            if checked >= 30:
                break
        if checked >= 30:
            break
    assert checked > 0, f"{filename}: no runs found to check"


@pytest.mark.parametrize("filename", EXPECTED_FILES)
def test_font_hiragino_on_tables(filename: str) -> None:
    """テーブルセルの run のフォントも Hiragino に設定されている."""
    doc = Document(str(OUTPUT_DIR / filename))
    checked = 0
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    for run in para.runs:
                        if not run.text.strip():
                            continue
                        assert run.font.name == FONT_NAME, \
                            f"{filename}: table run font={run.font.name!r}"
                        checked += 1
                        if checked >= 10:
                            return


@pytest.mark.parametrize("filename", EXPECTED_FILES)
def test_normal_style_font(filename: str) -> None:
    """Normal スタイルのフォント設定が Hiragino."""
    doc = Document(str(OUTPUT_DIR / filename))
    normal = doc.styles["Normal"]
    assert normal.font.name == FONT_NAME, \
        f"{filename}: Normal font={normal.font.name!r}"
    # East Asian
    rPr = normal.element.find(qn("w:rPr"))
    if rPr is not None:
        rFonts = rPr.find(qn("w:rFonts"))
        if rFonts is not None:
            east_asia = rFonts.get(qn("w:eastAsia"))
            assert east_asia == FONT_NAME


def test_all_four_files_generated() -> None:
    """期待される 4 ファイルがすべて生成されている."""
    actual = sorted(p.name for p in OUTPUT_DIR.glob("*.docx"))
    expected = sorted(EXPECTED_FILES)
    # 既存ファイルは上書きしないので、期待するファイルが含まれていることを確認
    for f in expected:
        assert f in actual, f"missing: {f}"


def test_existing_docx_not_modified() -> None:
    """既存の presentation_script_bedcontrol.docx (docs/admin 直下) は触っていない."""
    original = REPO_ROOT / "docs" / "admin" / "presentation_script_bedcontrol.docx"
    # 原本が生き残っているか（名前が v3.6 を付けずに残っている）
    assert original.exists(), "original presentation_script_bedcontrol.docx should remain"
    # 新しい v3.6 版は別ディレクトリに出力されていること
    new_version = OUTPUT_DIR / "presentation_script_bedcontrol_v3.6.docx"
    assert new_version.exists()
    assert original.parent != new_version.parent
